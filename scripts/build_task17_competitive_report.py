#!/usr/bin/env python3
"""Build a final competitive Task 17 DOCX report from existing outputs."""

from __future__ import annotations

import importlib.util
import json
import os
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parent.parent
RESULTS_DIR = PROJECT_ROOT / "results" / "gnnplus_task17_tiebreak_eval"
REPORT_DOCX = RESULTS_DIR / os.environ.get("GNNPLUS_REPORT_NAME", "GNNPLUS_TASK17_COMPETITIVE_FINAL_REPORT.docx")
SUMMARY_CSV = RESULTS_DIR / "packet_sdn_summary.csv"
FAILURE_CSV = RESULTS_DIR / "packet_sdn_failure.csv"
METRICS_CSV = RESULTS_DIR / "packet_sdn_sdn_metrics.csv"
PROPORTIONAL_CSV = RESULTS_DIR / "proportional_budget_summary.csv"
EXTREME_STRESS_CSV = RESULTS_DIR / "proportional_budget_extreme_stress.csv"
PLOTS_DIR = RESULTS_DIR / "plots"
TRAINING_SUMMARY_JSON = RESULTS_DIR / "training" / "training_summary.json"
INFERENCE_CALIBRATION_JSON = RESULTS_DIR / "training" / "inference_calibration.json"
MEMO_MD = RESULTS_DIR / "PROFESSOR_MEMO.md"

TASK10_FAILURE_CSV = PROJECT_ROOT / "results" / "gnnplus_3tracks_task10_clean" / "packet_sdn_failure.csv"
OVERALL_COMPARE_CSV = RESULTS_DIR / "comparison" / "overall_bundle_comparison.csv"
METHOD_RANKING_CSV = RESULTS_DIR / "comparison" / "new_run_method_ranking.csv"
HELPER_PATH = PROJECT_ROOT / "scripts" / "build_gnnplus_packet_sdn_report_fixed.py"

TOPOLOGY_LABELS = {
    "abilene": "Abilene",
    "cernet": "CERNET",
    "geant": "GEANT",
    "ebone": "Ebone",
    "sprintlink": "Sprintlink",
    "tiscali": "Tiscali",
    "germany50": "Germany50",
    "vtlwavenet2011": "VtlWavenet2011",
}
TOP_ORDER = [
    "abilene",
    "cernet",
    "geant",
    "ebone",
    "sprintlink",
    "tiscali",
    "germany50",
    "vtlwavenet2011",
]


def load_helper():
    spec = importlib.util.spec_from_file_location("task17_report_helper", HELPER_PATH)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Could not load helper from {HELPER_PATH}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def add_title_page(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GNN+ Competitive Packet-SDN Final Report")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x80)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "Scope: ECMP, Bottleneck, Original GNN, GNN+\n"
        "Study branch: standalone GNN+ zero-shot packet-SDN pipeline\n"
        "No MetaGate, No Stable MetaGate, No Bayesian Calibration"
    )
    run.font.size = Pt(11)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Important: packet-level SDN metrics in this report are model-based analytical metrics, not live Mininet measurements.")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    doc.add_page_break()


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_image(doc: Document, path: Path, caption: str, width: float = 6.2) -> None:
    if not path.exists():
        doc.add_paragraph(f"[Image missing: {path.name}]")
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def pct_text(delta_pct: float) -> str:
    return f"{delta_pct:+.2f}%"


def build_topology_results_table(summary_df: pd.DataFrame, failure_df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    summary_pivot = summary_df.pivot(index="topology", columns="method", values="mean_mlu")
    failure_pivot = failure_df.pivot_table(index=["topology", "scenario"], columns="method", values="mean_mlu")
    for topo in TOP_ORDER:
        top_label = TOPOLOGY_LABELS[topo]
        gnnplus = float(summary_pivot.loc[topo, "gnnplus"])
        bottleneck = float(summary_pivot.loc[topo, "bottleneck"])
        original_gnn_dist = float(summary_df[(summary_df["topology"] == topo) & (summary_df["method"] == "gnn")]["mean_disturbance"].iloc[0])
        gnnplus_dist = float(summary_df[(summary_df["topology"] == topo) & (summary_df["method"] == "gnnplus")]["mean_disturbance"].iloc[0])
        gnnplus_dt = float(summary_df[(summary_df["topology"] == topo) & (summary_df["method"] == "gnnplus")]["decision_time_ms"].iloc[0])
        bottleneck_dt = float(summary_df[(summary_df["topology"] == topo) & (summary_df["method"] == "bottleneck")]["decision_time_ms"].iloc[0])
        failure_win_tie = int(
            sum(
                float(failure_pivot.loc[(topo, scenario), "gnnplus"]) <= float(failure_pivot.loc[(topo, scenario), "bottleneck"]) + 1e-8
                for scenario in sorted(failure_df[failure_df["topology"] == topo]["scenario"].unique())
            )
        )
        delta_pct = ((gnnplus - bottleneck) / bottleneck * 100.0) if abs(bottleneck) > 1e-12 else 0.0
        if gnnplus <= bottleneck + 1e-8:
            normal_cell = f"Win/Tie ({pct_text(delta_pct)})"
        else:
            normal_cell = f"Worse ({pct_text(delta_pct)})"
        dist_delta = gnnplus_dist - original_gnn_dist
        dist_cell = ("Better" if dist_delta <= 1e-12 else "Worse") + f" ({dist_delta:+.4f})"
        rows.append(
            {
                "Topology": top_label,
                "Normal MLU vs Bottleneck": normal_cell,
                "Failure Win/Tie Count": f"{failure_win_tie}/5",
                "Disturbance vs Original GNN": dist_cell,
                "Decision Time Ratio": f"{gnnplus_dt / bottleneck_dt:.2f}x",
            }
        )
    return pd.DataFrame(rows)


def build_flagship_table(failure_df: pd.DataFrame, task10_failure_df: pd.DataFrame) -> pd.DataFrame:
    current = failure_df.pivot_table(index=["topology", "scenario"], columns="method", values="mean_mlu")
    old = task10_failure_df.pivot_table(index=["topology", "scenario"], columns="method", values="mean_mlu")
    idx = ("vtlwavenet2011", "random_link_failure_2")
    g17 = float(current.loc[idx, "gnnplus"])
    g10 = float(old.loc[idx, "gnnplus"])
    b17 = float(current.loc[idx, "bottleneck"])
    return pd.DataFrame(
        [
            {
                "Topology": "VtlWavenet2011",
                "Scenario": "Random Link Failure (2)",
                "Task17 GNN+ MLU": g17,
                "Task10 GNN+ MLU": g10,
                "Delta vs Task10": g17 - g10,
                "Task17 Bottleneck MLU": b17,
                "Gap vs Bottleneck": g17 - b17,
            }
        ]
    )


def build_scope_table(training_summary: dict) -> pd.DataFrame:
    return pd.DataFrame(
        [
            {"Item": "Methods", "Value": "ECMP, Bottleneck, Original GNN, GNN+"},
            {"Item": "Known topologies", "Value": "Abilene, CERNET, GEANT, Ebone, Sprintlink, Tiscali"},
            {"Item": "Unseen topologies", "Value": "Germany50, VtlWavenet2011"},
            {"Item": "Zero-shot protocol", "Value": "Yes; no unseen-topology retraining or calibration"},
            {"Item": "MetaGate / Stable MetaGate", "Value": "No"},
            {"Item": "Fixed K", "Value": "40"},
            {"Item": "Feature profile", "Value": training_summary.get("feature_variant", "section7_temporal")},
            {"Item": "Base checkpoint", "Value": training_summary.get("base_checkpoint")},
            {"Item": "Final checkpoint", "Value": training_summary.get("final_checkpoint")},
            {"Item": "Final model source", "Value": training_summary.get("final_model_source")},
        ]
    )


def build_method_changes_table() -> pd.DataFrame:
    return pd.DataFrame(
        [
            {"Area": "Architecture", "Change": "Gated MoE-style selector between Bottleneck prior/expert floor and learned GNN+ scorer"},
            {"Area": "Supervised training", "Change": "Plackett-Luce listwise loss with candidate-aware window = 180"},
            {"Area": "Supervised gate", "Change": "LP-gap gate replaced recall-only acceptance"},
            {"Area": "RL", "Change": "State-conditional KL anchor: beta_normal 0.10 -> 0.01, beta_failure 0.001"},
            {"Area": "RL", "Change": "Gumbel-Top-K exploration on failure rollouts only"},
            {"Area": "RL", "Change": "Running-mean baseline and 5-type failure coverage aligned to evaluation"},
            {"Area": "Inference", "Change": "Do-no-harm LP fallback: 1.02x known, 1.00x unseen"},
            {"Area": "Inference", "Change": "Disturbance-aware tie-break with p25 epsilon on GEANT, Tiscali, Germany50, VtlWavenet2011"},
            {"Area": "Latency", "Change": "Vectorized prefilter and graph-build path; VtlWavenet ratio reduced from 1.57x main reference to 1.22x"},
        ]
    )


def build_report() -> Path:
    helper = load_helper()
    helper.PLOTS_DIR = PLOTS_DIR

    summary_df = pd.read_csv(SUMMARY_CSV)
    failure_df = pd.read_csv(FAILURE_CSV)
    metrics_df = pd.read_csv(METRICS_CSV)
    proportional_df = pd.read_csv(PROPORTIONAL_CSV) if PROPORTIONAL_CSV.exists() else pd.DataFrame()
    extreme_df = pd.read_csv(EXTREME_STRESS_CSV) if EXTREME_STRESS_CSV.exists() else pd.DataFrame()
    overall_compare = pd.read_csv(OVERALL_COMPARE_CSV) if OVERALL_COMPARE_CSV.exists() else pd.DataFrame()
    method_ranking = pd.read_csv(METHOD_RANKING_CSV) if METHOD_RANKING_CSV.exists() else pd.DataFrame()
    training_summary = json.loads(TRAINING_SUMMARY_JSON.read_text(encoding="utf-8"))
    inference_calibration = json.loads(INFERENCE_CALIBRATION_JSON.read_text(encoding="utf-8"))
    task10_failure_df = pd.read_csv(TASK10_FAILURE_CSV)

    doc = Document()
    helper.set_default_style(doc)
    add_title_page(doc)

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This final Task 17 bundle shows that the standalone GNN+ packet-SDN branch is competitive with the Bottleneck heuristic "
        "and materially stronger than the prior main reference on hard failure recovery and latency. The strongest single result is "
        "VtlWavenet2011 under 2-link random failure, where GNN+ improves mean MLU by 2176.81 units relative to the Task 10 supervised-only reference. "
        "The final branch preserves normal-condition MLU at 6/8 win-or-tie against Bottleneck, wins or ties 27/40 failure scenarios, and keeps decision time within 1.28x Bottleneck. "
        "The unresolved metric is disturbance: GNN+ is lower than or equal to Original GNN on 4/8 topologies, so the correct claim is competitive, not dominant."
    )

    doc.add_heading("2. Scope and Protocol", level=1)
    helper.add_dataframe_table(doc, build_scope_table(training_summary), font_size=9)
    add_bullet(doc, "Packet-level SDN metrics in this report are model-based analytical metrics, not live Mininet measurements.")
    add_bullet(doc, "Germany50 and VtlWavenet2011 remain pure unseen evaluation topologies in this branch.")
    add_bullet(doc, "No MetaGate, no Stable MetaGate, and no Bayesian calibration are used in this report.")

    doc.add_heading("3. What Changed in This Final Branch", level=1)
    helper.add_dataframe_table(doc, build_method_changes_table(), font_size=9)

    doc.add_heading("4. Training and Checkpoint Summary", level=1)
    train_rows = pd.DataFrame(
        [
            {"Field": "Base checkpoint", "Value": training_summary["base_checkpoint"]},
            {"Field": "Final checkpoint", "Value": training_summary["final_checkpoint"]},
            {"Field": "Supervised best epoch", "Value": training_summary["supervised"]["best_epoch"]},
            {"Field": "Supervised best recall@40", "Value": training_summary["supervised"]["best_val_recall_at_40"]},
            {"Field": "Aggregate mean LP gap", "Value": training_summary["supervised"]["lp_gap_gate"]["aggregate"]["mean_relative_gap"]},
            {"Field": "Aggregate p95 LP gap", "Value": training_summary["supervised"]["lp_gap_gate"]["aggregate"]["p95_relative_gap"]},
            {"Field": "RL best epoch", "Value": training_summary["reinforce"]["best_epoch"]},
            {"Field": "RL best val failure MLU", "Value": training_summary["reinforce"]["best_val_failure_mlu"]},
            {"Field": "RL best val disturbance", "Value": training_summary["reinforce"]["best_val_disturbance"]},
            {"Field": "Inference calibration", "Value": str(INFERENCE_CALIBRATION_JSON.relative_to(PROJECT_ROOT))},
        ]
    )
    helper.add_dataframe_table(doc, train_rows, font_size=9)
    doc.add_paragraph(
        "The final Task 17 report reuses the best RL checkpoint from Task 16 and applies the Task 17 aggressive tie-break inference policy at report time. "
        "The aggressive tie-break uses per-topology p25 score-gap epsilon on GEANT and Tiscali, and a global aggressive p25 fallback for Germany50 and VtlWavenet2011 to avoid any unseen-topology validation leakage."
    )
    tie_rows = []
    for topo in ["geant", "tiscali", "germany50", "vtlwavenet2011"]:
        cfg = inference_calibration.get("per_topology", {}).get(topo, {})
        tie_rows.append(
            {
                "Topology": TOPOLOGY_LABELS[topo],
                "Gate Temperature": cfg.get("gate_temperature", inference_calibration.get("global_gate_temperature")),
                "Tie-Break p5": cfg.get("tie_break_epsilon_p5", inference_calibration.get("global_tie_break_epsilon")),
                "Tie-Break p25": cfg.get("tie_break_epsilon_p25", inference_calibration.get("global_tie_break_epsilon_aggressive")),
                "Source": "Per-topology" if topo in inference_calibration.get("per_topology", {}) else "Global aggressive fallback",
            }
        )
    helper.add_dataframe_table(doc, pd.DataFrame(tie_rows), font_size=8)

    doc.add_heading("5. Normal Results", level=1)
    add_image(doc, PLOTS_DIR / "mlu_comparison_normal.png", "Figure 1. Normal-condition mean MLU across the 8 topologies.")
    add_image(doc, PLOTS_DIR / "throughput_comparison_normal.png", "Figure 2. Normal-condition throughput across the 8 topologies.")
    add_image(doc, PLOTS_DIR / "disturbance_comparison.png", "Figure 3. Routing disturbance across the 8 topologies.")
    add_image(doc, PLOTS_DIR / "decision_time_comparison.png", "Figure 4. Decision time across the 8 topologies.")
    helper.add_dataframe_table(doc, build_topology_results_table(summary_df, failure_df), font_size=8)

    doc.add_heading("6. Failure Results", level=1)
    add_image(doc, PLOTS_DIR / "failure_recovery_gnnplus.png", "Figure 5. GNN+ failure recovery time by scenario.")
    flagship = build_flagship_table(failure_df, task10_failure_df)
    helper.add_dataframe_table(doc, flagship, title="Flagship failure result", font_size=8)
    doc.add_paragraph(
        "The flagship result is VtlWavenet2011 under 2-link random failure. In Task 17, GNN+ reaches 18513.28 mean MLU, compared with 20690.09 in the Task 10 supervised-only reference and 19585.37 for Bottleneck in the same Task 17 bundle. "
        "This is the largest absolute improvement observed in the project and it occurs on the hardest unseen topology."
    )
    failure_overview = []
    failure_pivot = failure_df.pivot_table(index=["topology", "scenario"], columns="method", values="mean_mlu")
    for topo in TOP_ORDER:
        wins = []
        for scenario in sorted(failure_df[failure_df["topology"] == topo]["scenario"].unique()):
            g = float(failure_pivot.loc[(topo, scenario), "gnnplus"])
            b = float(failure_pivot.loc[(topo, scenario), "bottleneck"])
            wins.append("Win/Tie" if g <= b + 1e-8 else "Worse")
        failure_overview.append({"Topology": TOPOLOGY_LABELS[topo], "Per-scenario status": ", ".join(wins)})
    helper.add_dataframe_table(doc, pd.DataFrame(failure_overview), font_size=8)

    doc.add_heading("7. Model-Based Packet-SDN Metrics", level=1)
    doc.add_paragraph(
        "The table below summarizes the packet-SDN metrics used in this branch. These are model-based analytical measurements built from the routing outputs; they are not live controller or Mininet measurements."
    )
    helper.add_dataframe_table(doc, metrics_df, font_size=8)

    doc.add_heading("8. Comparison vs Previous Bundles", level=1)
    if not overall_compare.empty:
        helper.add_dataframe_table(doc, overall_compare, font_size=9)
    if not method_ranking.empty:
        helper.add_dataframe_table(doc, method_ranking, title="Method ranking in the new run", font_size=8)
    doc.add_paragraph(
        "Relative to the prior main-reference branch, the final Task 17 branch is materially faster and more robust on the hardest failure cases. "
        "Relative to Task 10 supervised-only, it preserves the strong normal MLU behavior while adding the strongest hard-failure gain on VtlWavenet2011 random link failure (2)."
    )

    if not proportional_df.empty:
        doc.add_heading("9. Proportional-Budget Fairness Study", level=1)
        doc.add_paragraph(
            "This proportional-budget study is separate from the main fixed-K40 benchmark. Here the active budget is recomputed each timestep as K = int(ratio x active OD pairs), and the same ratio-defined budget is applied to ECMP, Bottleneck, and GNN+."
        )
        helper.add_dataframe_table(doc, proportional_df, font_size=8)

    if not extreme_df.empty:
        doc.add_heading("10. Extreme Stress Test (VtlWavenet2011)", level=1)
        doc.add_paragraph(
            "This stress test isolates the unseen VtlWavenet2011 topology under a 5% active-flow budget and cascading random failures. It is kept separate from the fixed-K40 benchmark because it is a fairness/scalability stress study rather than part of the main thesis table."
        )
        helper.add_dataframe_table(doc, extreme_df, font_size=8)

    doc.add_heading("11. What Did Not Work", level=1)
    add_bullet(doc, "Universal disturbance dominance was not achieved. After both Task 16 RL disturbance tuning and Task 17 aggressive tie-break tuning, GNN+ still remained at 4/8 topologies better-than-or-equal-to Original GNN.")
    add_bullet(doc, "CERNET and VtlWavenet2011 normal MLU remain close to Bottleneck but not clearly better: +0.39% and +0.18% respectively.")
    add_bullet(doc, "Traffic-spike remains the hardest failure type to improve consistently; earlier RL diagnostics showed the most negative per-type advantage on traffic_spike_2x.")

    doc.add_heading("12. Scientific Contributions Claimed", level=1)
    add_bullet(doc, "State-conditional KL anchoring for RL from a near-optimal supervised prior, which prevented the archfix-style RL collapse.")
    add_bullet(doc, "LP-gap as the correct supervised gate metric for top-K set-selection problems where many near-equivalent solutions exist.")
    add_bullet(doc, "Empirical evidence that candidate prefiltering, LP warm-starting, and vectorized graph construction can keep GNN-based traffic engineering latency close to heuristic baselines at 92-node scale.")

    doc.add_heading("13. Conclusions and Future Work", level=1)
    doc.add_paragraph(
        "The correct conclusion from this final Task 17 branch is that GNN+ is competitive with Bottleneck, not dominant over all baselines on every metric. "
        "The branch is strongest on hard failure recovery and on the larger or unseen topologies, while disturbance remains the unresolved architectural tradeoff. "
        "The next defensible research questions are temporal-memory alternatives for disturbance control, traffic-spike specialization, and evaluation on larger unseen topologies beyond the current 50- to 92-node range."
    )
    if MEMO_MD.exists():
        doc.add_paragraph(
            f"Professor memo reference: {MEMO_MD.relative_to(PROJECT_ROOT)}"
        )

    REPORT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(REPORT_DOCX))
    return REPORT_DOCX


def main() -> int:
    path = build_report()
    print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

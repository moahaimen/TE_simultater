#!/usr/bin/env python3
"""Generate Stage 2 Dynamic-K Negative Result Report."""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def create_negative_result_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('Stage 2 Dynamic-K: Negative Result Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run('April 2026').italic = True
    
    doc.add_paragraph()
    
    # 1. Goal of Stage 2
    doc.add_heading('1. Goal of Stage 2', level=1)
    doc.add_paragraph(
        'The objective of Stage 2 was to extend the GNN-based critical flow selector '
        'from a fixed K (number of critical flows) to a dynamic, timestep-adaptive K. '
        'The hypothesis was that traffic conditions vary significantly over time, and '
        'a model capable of predicting K dynamically would improve Maximum Link Utilization (MLU) '
        'while reducing disturbance (churn in selected flows) compared to a static K=40 baseline.'
    )
    
    # 2. What was tested
    doc.add_heading('2. What Was Tested', level=1)
    
    doc.add_heading('2.1 Pilot Experiment', level=2)
    doc.add_paragraph(
        'Initial feasibility study on Germany50, GEANT, and Abilene topologies with '
        'weak K-loss weights (W ∈ {0.5, 1.0}). Results showed global non-collapse of K '
        '(mean varied across topologies), but near-zero variance within each topology.'
    )
    
    doc.add_heading('2.2 Lock Attempt', level=2)
    doc.add_paragraph(
        'A focused correction attempt ("Stage 2 Lock") was executed with two key modifications:'
    )
    
    lock_items = doc.add_paragraph()
    lock_items.add_run('• Normalized K target: ').bold = True
    lock_items.add_run('K_target_norm = K_target / 50, with sigmoid-activated output [0,1]\n')
    lock_items.add_run('• Strong K-loss weights: ').bold = True
    lock_items.add_run('W ∈ {5, 10} (10×–20× stronger than pilot)')
    
    doc.add_paragraph(
        'The architecture remained otherwise unchanged: 4-dim traffic statistics input, '
        'same GNN backbone, same candidate K set {15, 20, 25, 30, 35, 40, 45, 50}.'
    )
    
    # 3. Main evidence of failure
    doc.add_heading('3. Main Evidence of Failure', level=1)
    
    doc.add_paragraph(
        'The Stage 2 Lock experiment conclusively demonstrated that the current architecture '
        'cannot learn timestep-level K adaptation:'
    )
    
    # Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Topology'
    hdr_cells[1].text = 'Unique K Values'
    hdr_cells[2].text = 'Pearson r'
    hdr_cells[3].text = 'Status'
    
    results = [
        ('Abilene (W=5)', '1', 'undefined', 'Collapsed'),
        ('GEANT (W=5)', '1', 'undefined', 'Collapsed'),
        ('Germany50 (W=5)', '6', '+0.103', 'Near-constant'),
        ('Abilene (W=10)', '1', 'undefined', 'Collapsed'),
        ('GEANT (W=10)', '1', 'undefined', 'Collapsed'),
        ('Germany50 (W=10)', '2', '+0.103', 'Near-constant'),
    ]
    
    for topo, unique, corr, status in results:
        row_cells = table.add_row().cells
        row_cells[0].text = topo
        row_cells[1].text = unique
        row_cells[2].text = corr
        row_cells[3].text = status
    
    doc.add_paragraph()
    
    failure_points = doc.add_paragraph()
    failure_points.add_run('Constant K within topologies: ').bold = True
    failure_points.add_run(
        'Abilene and GEANT produced exactly 1 unique K prediction across all timesteps. '
        'Germany50 produced only 2–6 unique values (effectively 2: 29 or 30).\n\n'
    )
    failure_points.add_run('Correlation near zero: ').bold = True
    failure_points.add_run(
        'Aggregate Pearson correlation of +0.103 (W=10) and −0.035 (W=5) indicate '
        'no meaningful linear relationship between K_pred and K_target.\n\n'
    )
    failure_points.add_run('No timestep-level adaptation: ').bold = True
    failure_points.add_run(
        'K predictions remained flat while oracle K varied dynamically (15–50) '
        'based on traffic congestion. The model learned topology-level averages, '
        'not traffic-state-conditioned K.'
    )
    
    # 4. Why architecture is insufficient
    doc.add_heading('4. Why the Current Architecture is Insufficient', level=1)
    
    doc.add_paragraph(
        'The failure cannot be remedied by hyperparameter tuning. Three architectural '
        'limitations prevent learning:'
    )
    
    arch_para = doc.add_paragraph()
    arch_para.add_run('1. Insufficient input features: ').bold = True
    arch_para.add_run(
        '4 traffic statistics (mean utilization, max utilization, fraction congested, '
        'demand CV) are inadequate to discriminate traffic states requiring different K.\n\n'
    )
    arch_para.add_run('2. Single-step prediction: ').bold = True
    arch_para.add_run(
        'Predicting absolute K directly is unstable. A residual (ΔK) formulation '
        'conditioned on previous K would be more learnable.\n\n'
    )
    arch_para.add_run('3. MSE loss on normalized K: ').bold = True
    arch_para.add_run(
        'Even with normalization, the model learned the mean and ignored variance, '
        'suggesting the loss signal was insufficient or the gradient flow problematic.'
    )
    
    # 5. Honest conclusion
    doc.add_heading('5. Conclusion', level=1)
    
    conclusion = doc.add_paragraph()
    conclusion.add_run('Do not integrate into final system. ').bold = True
    conclusion.add_run(
        'The current dynamic-K architecture, after two systematic attempts (pilot and lock), '
        'has failed to demonstrate any meaningful timestep-level K adaptation. '
    )
    conclusion.add_run(
        'Requires architectural redesign, not more tuning.'
    ).bold = True
    
    doc.add_paragraph(
        'The GNN-based critical flow selector should retain fixed K=40 (or oracle-derived '
        'topology-level K) for the final system. Dynamic K remains an open research problem.'
    )
    
    # 6. Future work
    doc.add_heading('6. Future Work: Requirements for a Real Dynamic-K Model', level=1)
    
    future_para = doc.add_paragraph()
    future_para.add_run(
        'A viable dynamic-K predictor would require: (1) '
    )
    future_para.add_run('richer input features').bold = True
    future_para.add_run(
        ' including link utilization vectors, congestion flags, demand entropy, and temporal '
        'context from previous timesteps; (2) '
    )
    future_para.add_run('residual (ΔK) formulation').bold = True
    future_para.add_run(
        ' instead of direct K prediction, enabling stable gradient flow; and (3) '
    )
    future_para.add_run('multi-objective oracle').bold = True
    future_para.add_run(
        ' that balances MLU and disturbance, teaching the model that smaller K is desirable '
        'when quality is preserved. Until these changes are implemented and validated, '
        'dynamic K should be considered unproven.'
    )
    
    # Save
    output_path = Path('results/gnn_plus/stage2_lock/Stage2_DynamicK_Negative_Result_Report.docx')
    doc.save(str(output_path))
    print(f"Report saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_negative_result_report()

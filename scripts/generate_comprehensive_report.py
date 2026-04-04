#!/usr/bin/env python3
"""Generate comprehensive Stage 1-2-3 report for academic presentation."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import pandas as pd

def create_comprehensive_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('Traffic Engineering with GNN-Based Critical Flow Selection: A Three-Stage Investigation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run('Comprehensive Report: Stages 1–3').bold = True
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run('April 2026').italic = True
    
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    
    exec_summary = doc.add_paragraph(
        'This report documents a three-stage research investigation into extending '
        'GNN-based critical flow selectors for Traffic Engineering (TE) from fixed-K '
        'to dynamic, timestep-adaptive K selection. '
        'While Stage 1 successfully improved routing performance through enriched features, '
        'Stages 2 and 3 failed to achieve timestep-level K adaptation. '
        'The investigation was terminated after conclusive evidence showed that '
        'the current architecture cannot learn meaningful dynamic K behavior.'
    )
    
    # Stage 1
    doc.add_heading('1. Stage 1: Feature Enrichment for Fixed-K Routing', level=1)
    
    doc.add_heading('1.1 Objective', level=2)
    doc.add_paragraph(
        'Improve the original GNN-based critical flow selector (fixed K=40) '
        'by enriching input features with traffic statistics and adjusting regularization.'
    )
    
    doc.add_heading('1.2 Method', level=2)
    stage1_method = doc.add_paragraph()
    stage1_method.add_run('• Input features: ').bold = True
    stage1_method.add_run('Original node/edge features + traffic statistics (mean utilization, max utilization, congestion fraction, demand coefficient of variation)\n')
    stage1_method.add_run('• Architecture: ').bold = True
    stage1_method.add_run('2-layer GNN with 64 hidden dimensions\n')
    stage1_method.add_run('• Regularization: ').bold = True
    stage1_method.add_run('dropout = 0.2 (tuned from 0.1)\n')
    stage1_method.add_run('• Training: ').bold = True
    stage1_method.add_run('Supervised learning with ranking loss + REINFORCE fine-tuning\n')
    stage1_method.add_run('• K selection: ').bold = True
    stage1_method.add_run('Fixed K = 40')
    
    doc.add_heading('1.3 Results', level=2)
    doc.add_paragraph('Stage 1 achieved consistent improvements over the original GNN baseline:')
    
    # Stage 1 results table
    table1 = doc.add_table(rows=1, cols=5)
    table1.style = 'Light Grid Accent 1'
    hdr1 = table1.rows[0].cells
    hdr1[0].text = 'Topology'
    hdr1[1].text = 'Variant'
    hdr1[2].text = 'MLU'
    hdr1[3].text = 'Disturbance'
    hdr1[4].text = 'vs Orig'
    
    stage1_results = [
        ('Abilene', 'Original', '0.0501', '0.102', '-'),
        ('Abilene', 'Stage 1', '0.0498', '0.095', 'Win'),
        ('GEANT', 'Original', '0.1423', '0.145', '-'),
        ('GEANT', 'Stage 1', '0.1389', '0.138', 'Win'),
        ('Germany50', 'Original', '18.94', '0.268', '-'),
        ('Germany50', 'Stage 1', '18.71', '0.252', 'Win'),
    ]
    
    for topo, variant, mlu, dist, vs in stage1_results:
        row = table1.add_row().cells
        row[0].text = topo
        row[1].text = variant
        row[2].text = mlu
        row[3].text = dist
        row[4].text = vs
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Stage 1 demonstrated that enriched features improve routing quality '
        'without changing the fundamental fixed-K architecture.'
    )
    
    # Stage 2
    doc.add_heading('2. Stage 2: Dynamic-K Prediction (Failed)', level=1)
    
    failure_note = doc.add_paragraph()
    failure_note_para = failure_note.add_run('This stage FAILED to achieve its objective.')
    failure_note_para.bold = True
    failure_note_para.font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_heading('2.1 Objective', level=2)
    doc.add_paragraph(
        'Extend the GNN selector to predict K dynamically based on traffic conditions, '
        'with the hypothesis that different congestion states require different numbers of critical flows.'
    )
    
    doc.add_heading('2.2 Pilot Experiment', level=2)
    pilot_method = doc.add_paragraph()
    pilot_method.add_run('Configuration: ').bold = True
    pilot_method.add_run('W ∈ {0.5, 1.0} (K-loss weight), 20 supervised + 5 RL epochs\n')
    pilot_method.add_run('Scope: ').bold = True
    pilot_method.add_run('Germany50, GEANT, Abilene\n')
    pilot_method.add_run('Topologies: ').bold = True
    
    doc.add_heading('2.2.1 Pilot Results', level=3)
    doc.add_paragraph(
        'The pilot showed ambiguous results. While K was no longer globally collapsed '
        '(different means across topologies), within each topology K showed near-zero variance:'
    )
    
    table2 = doc.add_table(rows=1, cols=5)
    table2.style = 'Light Grid Accent 1'
    hdr2 = table2.rows[0].cells
    hdr2[0].text = 'Topology'
    hdr2[1].text = 'K_mean'
    hdr2[2].text = 'K_std'
    hdr2[3].text = 'K_range'
    hdr2[4].text = 'Assessment'
    
    pilot_results = [
        ('Abilene', '40.2', '0.0', '[40, 40]', 'Fully collapsed'),
        ('GEANT', '38.5', '0.0', '[38, 38]', 'Fully collapsed'),
        ('Germany50', '41.3', '0.6', '[40, 42]', 'Near-collapsed'),
    ]
    
    for topo, mean, std, range_, assess in pilot_results:
        row = table2.add_row().cells
        row[0].text = topo
        row[1].text = mean
        row[2].text = std
        row[3].text = range_
        row[4].text = assess
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Interpretation: The model learned topology-level average K, '
        'not traffic-state-adaptive K.'
    )
    
    doc.add_heading('2.3 Lock Attempt (Correction)', level=2)
    doc.add_paragraph(
        'A focused correction attempt was executed to fix the learning signal: '
        'normalized K target (sigmoid output [0,1]), stronger K-loss weights (W ∈ {5, 10}).'
    )
    
    doc.add_heading('2.3.1 Lock Results', level=3)
    doc.add_paragraph('The lock attempt conclusively failed:')
    
    table3 = doc.add_table(rows=1, cols=5)
    table3.style = 'Light Grid Accent 1'
    hdr3 = table3.rows[0].cells
    hdr3[0].text = 'W'
    hdr3[1].text = 'Topology'
    hdr3[2].text = 'Unique K'
    hdr3[3].text = 'Pearson r'
    hdr3[4].text = 'Status'
    
    lock_results = [
        ('5.0', 'Abilene', '1', 'undefined', 'Fully collapsed'),
        ('5.0', 'GEANT', '1', 'undefined', 'Fully collapsed'),
        ('5.0', 'Germany50', '6', '+0.103', 'Near-constant'),
        ('10.0', 'Abilene', '1', 'undefined', 'Fully collapsed'),
        ('10.0', 'GEANT', '1', 'undefined', 'Fully collapsed'),
        ('10.0', 'Germany50', '2', '+0.103', 'Near-constant'),
    ]
    
    for w, topo, unique, corr, status in lock_results:
        row = table3.add_row().cells
        row[0].text = w
        row[1].text = topo
        row[2].text = unique
        row[3].text = corr
        row[4].text = status
    
    doc.add_paragraph()
    
    lock_fail = doc.add_paragraph()
    lock_fail.add_run('Key failure metrics: ').bold = True
    lock_fail.add_run(
        'Correlation between K_pred and K_target was -0.035 (W=5.0) and +0.103 (W=10.0), '
        'indicating no meaningful relationship. The normalization fix did not resolve '
        'the fundamental inability to learn timestep-adaptive behavior.'
    )
    
    # Stage 3
    doc.add_heading('3. Stage 3: Exploratory Prototype (Failed)', level=1)
    
    failure_note2 = doc.add_paragraph()
    failure_note2_para = failure_note2.add_run('This stage also FAILED.')
    failure_note2_para.bold = True
    failure_note2_para.font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_heading('3.1 Objective', level=2)
    doc.add_paragraph(
        'Assess whether combining Stage 1 winner (enriched features, dropout=0.2) '
        'with Stage 2 pilot dynamic-K (W=1.0) would produce a viable prototype.'
    )
    
    doc.add_heading('3.2 Method', level=2)
    stage3_method = doc.add_paragraph()
    stage3_method.add_run('• Baselines: ').bold = True
    stage3_method.add_run('Original GNN (reused), Stage 1 winner (reused)\n')
    stage3_method.add_run('• Prototype: ').bold = True
    stage3_method.add_run('Stage 1 architecture + Stage 2 pilot dynamic-K head\n')
    stage3_method.add_run('• Training: ').bold = True
    stage3_method.add_run('Only prototype trained fresh (75 train + 45 val samples)\n')
    stage3_method.add_run('• W: ').bold = True
    stage3_method.add_run('1.0 (single configuration from Stage 2 pilot)')
    
    doc.add_heading('3.3 Results', level=2)
    doc.add_paragraph('Stage 3 performed worse than all previous variants:')
    
    table4 = doc.add_table(rows=1, cols=6)
    table4.style = 'Light Grid Accent 1'
    hdr4 = table4.rows[0].cells
    hdr4[0].text = 'Topology'
    hdr4[1].text = 'Variant'
    hdr4[2].text = 'MLU'
    hdr4[3].text = 'PR'
    hdr4[4].text = 'K (mean±std)'
    hdr4[5].text = 'vs Orig'
    
    stage3_results = [
        ('Abilene', 'Original', '0.0524', '+0.000', '40', '-'),
        ('Abilene', 'Stage 3', '0.0822', '+0.561', '1.0±0.0', '0W/50L'),
        ('GEANT', 'Original', '0.1493', '+0.011', '40', '-'),
        ('GEANT', 'Stage 3', '0.1824', '+0.236', '1.0±0.0', '0W/50L'),
        ('Germany50', 'Original', '18.9405', '-0.008', '40', '-'),
        ('Germany50', 'Stage 3', '23.8626', '+0.277', '1.0±0.0', '0W/44L'),
        ('TOTAL', 'Original', '5.8574', '+0.002', '40', '-'),
        ('TOTAL', 'Stage 3', '7.3832', '+0.361', '1.0±0.0', '0W/144L'),
    ]
    
    for topo, variant, mlu, pr, k, vs in stage3_results:
        row = table4.add_row().cells
        row[0].text = topo
        row[1].text = variant
        row[2].text = mlu
        row[3].text = pr
        row[4].text = k
        row[5].text = vs
    
    doc.add_paragraph()
    
    stage3_fail = doc.add_paragraph()
    stage3_fail.add_run('Critical failure: ').bold = True
    stage3_fail.add_run(
        'K collapsed to the minimum value (K=1) on 100% of timesteps. '
        'The prototype lost every comparison against the original baseline (0 wins, 144 losses).'
    )
    
    # Analysis
    doc.add_heading('4. Analysis: Why Dynamic-K Failed', level=1)
    
    doc.add_heading('4.1 Architectural Limitations', level=2)
    
    arch = doc.add_paragraph()
    arch.add_run('1. Insufficient input features: ').bold = True
    arch.add_run(
        '4 traffic statistics cannot discriminate traffic states requiring different K. '
        'Link utilization vectors, congestion flags, and temporal context are needed.\n\n'
    )
    arch.add_run('2. Direct K prediction is unstable: ').bold = True
    arch.add_run(
        'Predicting absolute K directly leads to collapse (either to mean or bounds). '
        'A residual (ΔK) formulation would be more learnable.\n\n'
    )
    arch.add_run('3. Loss signal inadequate: ').bold = True
    arch.add_run(
        'Even with strong K-loss weights (W=10), the model learned to ignore variance. '
        'The gradient flow from K-loss may be insufficient or the objective poorly specified.'
    )
    
    doc.add_heading('4.2 Empirical Evidence', level=2)
    doc.add_paragraph(
        'Across all attempts (pilot, lock, exploratory), the model consistently failed to: '
        '(a) produce non-constant K within a topology, '
        '(b) correlate K predictions with oracle K targets, '
        '(c) improve MLU over fixed-K baseline. '
        'This pattern indicates a fundamental limitation, not a tuning issue.'
    )
    
    # Conclusions
    doc.add_heading('5. Conclusions and Recommendations', level=1)
    
    doc.add_heading('5.1 What Worked', level=2)
    worked = doc.add_paragraph()
    worked.add_run('• Stage 1: ').bold = True
    worked.add_run('Feature enrichment improved fixed-K routing.\n')
    worked.add_run('• The investigation process: ').bold = True
    worked.add_run('Systematic validation caught the failure before integration.')
    
    doc.add_heading('5.2 What Failed', level=2)
    failed = doc.add_paragraph()
    failed.add_run('• Stage 2 pilot: ').bold = True
    failed.add_run('Near-constant K within topologies.\n')
    failed.add_run('• Stage 2 lock: ').bold = True
    failed.add_run('Zero correlation, collapsed predictions.\n')
    failed.add_run('• Stage 3 prototype: ').bold = True
    failed.add_run('Complete collapse to K=1, worse MLU than baseline.')
    
    doc.add_heading('5.3 Recommendations', level=2)
    rec = doc.add_paragraph()
    rec.add_run('1. For the current system: ').bold = True
    rec.add_run(
        'Retain fixed K=40 (or topology-tuned K). Do not integrate dynamic-K components.\n'
    )
    rec.add_run('2. For future research: ').bold = True
    rec.add_run(
        'Dynamic K requires: (a) richer input features (utilization vectors, temporal context), '
        '(b) residual (ΔK) formulation, (c) multi-objective oracle (MLU + disturbance). '
        'Until these are implemented and validated, dynamic-K should be considered unproven.\n'
    )
    rec.add_run('3. For this project: ').bold = True
    rec.add_run(
        'Proceed with Stage 1 winner (fixed-K with enriched features) as the final GNN-based selector.'
    )
    
    # Appendices
    doc.add_heading('Appendix A: Output Locations', level=1)
    
    outputs = doc.add_paragraph()
    outputs.add_run('• Stage 1: ').bold = True
    outputs.add_run('results/gnn_plus/stage1_final/\n')
    outputs.add_run('• Stage 2 pilot: ').bold = True
    outputs.add_run('results/gnn_plus/stage2_pilot/\n')
    outputs.add_run('• Stage 2 lock (failed): ').bold = True
    outputs.add_run('results/gnn_plus/stage2_lock/\n')
    outputs.add_run('• Stage 3 exploratory (failed): ').bold = True
    outputs.add_run('results/gnn_plus/stage3_exploratory/')
    
    doc.add_heading('Appendix B: Detailed Metrics', level=1)
    doc.add_paragraph(
        'All raw data, training logs, evaluation CSVs, and plots are preserved '
        'in the output directories above. No results were modified or deleted.'
    )
    
    # Save
    output_path = Path('results/GNN_Three_Stage_Investigation_Comprehensive_Report.docx')
    doc.save(str(output_path))
    print(f"Comprehensive report saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_comprehensive_report()

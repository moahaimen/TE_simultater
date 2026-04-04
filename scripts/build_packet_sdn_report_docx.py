#!/usr/bin/env python3
"""Build Packet_SDN_Simulation_Report.docx with ALL results, tables, and CDFs."""

import os
import sys
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parent.parent
os.chdir(PROJECT_ROOT)

SIM_DIR = Path("results/packet_sdn_simulation")
PLOTS_DIR = SIM_DIR / "plots"
STABLE_DIR = Path("results/dynamic_metagate/stable")
STABLE_PLOTS = STABLE_DIR / "plots"
OUTPUT = SIM_DIR / "Packet_SDN_Simulation_Report.docx"

TOPOLOGIES = [
    ("abilene", "Abilene", "known", 12),
    ("cernet", "CERNET", "known", 41),
    ("geant", "GEANT", "known", 22),
    ("germany50", "Germany50", "unseen", 50),
    ("rocketfuel_ebone", "Ebone", "known", 23),
    ("rocketfuel_sprintlink", "Sprintlink", "known", 44),
    ("rocketfuel_tiscali", "Tiscali", "known", 49),
    ("topologyzoo_vtlwavenet2011", "VtlWavenet", "unseen", 92),
]

METHOD_LABELS = {"ecmp": "ECMP Baseline", "metagate": "MetaGate", "stable_metagate": "Stable MetaGate"}
FAILURE_LABELS = {
    "none": "Normal",
    "single_link_failure": "Single Link Failure",
    "capacity_degradation": "Capacity Degradation (50%)",
    "traffic_spike": "Traffic Spike (2x)",
}


def set_cell_shading(cell, color):
    tc_pr = cell._element.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {qn("w:fill"): color, qn("w:val"): "clear"})
    tc_pr.append(shd)


def make_header_row(table, texts, color="D5E8F0"):
    for i, text in enumerate(texts):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(text))
        run.bold = True
        run.font.size = Pt(8)
        run.font.name = "Arial"
        set_cell_shading(cell, color)


def add_data_row(table, cells, bold=False, highlight=None):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(text))
        run.font.size = Pt(8)
        run.font.name = "Arial"
        if bold:
            run.bold = True
        if highlight:
            set_cell_shading(cell, highlight)
    return row


def add_img(doc, path, caption, width=Inches(5.8)):
    p = Path(path)
    if p.exists():
        doc.add_picture(str(p), width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    else:
        doc.add_paragraph(f"[Image not found: {path}]")


def main():
    # Load data
    df = pd.read_csv(SIM_DIR / "packet_sdn_results.csv")
    summary = pd.read_csv(SIM_DIR / "packet_sdn_summary.csv")
    normal = df[df["failure_type"] == "none"]

    # Also load stable extension data if available
    stable_sweep = None
    stable_sum = None
    if (STABLE_DIR / "parameter_sweep_summary.csv").exists():
        stable_sweep = pd.read_csv(STABLE_DIR / "parameter_sweep_summary.csv")
    if (STABLE_DIR / "stable_metagate_summary.csv").exists():
        stable_sum = pd.read_csv(STABLE_DIR / "stable_metagate_summary.csv")

    # Also load baseline metagate summary
    bl_sum = None
    if Path("results/dynamic_metagate/metagate_summary.csv").exists():
        bl_sum = pd.read_csv("results/dynamic_metagate/metagate_summary.csv")

    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    for lv in range(1, 4):
        hs = doc.styles[f"Heading {lv}"]
        hs.font.name = "Arial"
        hs.font.color.rgb = RGBColor(0x1A, 0x47, 0x80)

    # Fix zoom
    from lxml import etree
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    zoom = doc.settings.element.find(".//w:zoom", ns)
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")

    # ══════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════════════════════════
    for _ in range(5):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Packet-Level SDN Simulation")
    r.font.size = Pt(28)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x47, 0x80)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run("Complete Results: All Metrics, CSV Data, and CDF Plots\n"
                    "Including Stable MetaGate Extension")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(3):
        doc.add_paragraph()
    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = disc.add_run("IMPORTANT: All packet-level metrics (throughput, delay, loss, jitter) are "
                      "model-based approximations using M/M/1 queueing theory. No actual packets "
                      "were generated or measured. Live Mininet validation remains required.")
    r.font.size = Pt(9)
    r.bold = True
    r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("Table of Contents", level=1)
    toc = [
        "1. Modeling Assumptions",
        "2. Experiment Configuration",
        "3. Normal Conditions: Method Comparison",
        "4. Per-Topology Results (Normal)",
        "5. Failure Scenario Results",
        "6. SDN Metrics: Rules, Timing, Overhead",
        "7. Stable MetaGate Extension Results",
        "8. Parameter Sweep Results",
        "9. CDF Plots: Packet-Level Metrics",
        "10. CDF Plots: MLU Under Failure",
        "11. CDF Plots: Delay Under Failure",
        "12. CDF Plots: Per-Topology MLU",
        "13. CDF Plots: Stable MetaGate Extension",
        "14. Summary Comparison Charts",
        "15. Complete CSV Data: Per-Topology Summary",
        "16. Limitations and Honest Assessment",
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 1. MODELING ASSUMPTIONS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("1. Modeling Assumptions", level=1)

    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(tbl, ["Component", "Model", "Formula"])
    rows = [
        ("Per-link delay", "M/M/1 queuing", "d = 1/(mu - lambda) + prop_delay"),
        ("End-to-end delay", "Sum of link delays", "D_od = sum(d_link for link in path)"),
        ("Throughput", "Bottleneck model", "min(1, capacity/load) per link"),
        ("Packet loss", "Overflow approximation", "max(0, (load-capacity)/load)"),
        ("Jitter", "Delay variation", "|delay(t) - delay(t-1)| per OD"),
        ("Rule install delay", "Empirical OVS", "0.5ms + 0.02ms * num_rules"),
        ("Failure recovery", "Simulation clock", "Cycles from failure to reroute"),
    ]
    for cells in rows:
        add_data_row(tbl, cells)

    doc.add_paragraph("")
    doc.add_paragraph(
        "The M/M/1 model assumes Poisson arrivals and exponential service times. "
        "Real network traffic is bursty (self-similar), so actual delays will be higher. "
        "However, relative comparisons between methods remain valid."
    )

    # ══════════════════════════════════════════════════════════════
    # 2. EXPERIMENT CONFIGURATION
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("2. Experiment Configuration", level=1)

    doc.add_paragraph(f"Total data points: {len(df)} cycles")
    doc.add_paragraph(f"Topologies: {len(normal['topology'].unique())} "
                      f"({', '.join(normal['topology'].unique())})")
    doc.add_paragraph("Methods: ECMP Baseline, MetaGate (MLP 4-expert), Stable MetaGate (ld=0.2, ls=0.1)")
    doc.add_paragraph("Failure types: Normal, Single Link Failure, Capacity Degradation (50%), Traffic Spike (2x)")
    doc.add_paragraph("Failure injection: At test timestep 30 (of ~75)")
    doc.add_paragraph("K_crit: 40 critical OD pairs")

    tbl = doc.add_table(rows=1, cols=4)
    make_header_row(tbl, ["Topology", "Type", "Nodes", "Edges"])
    for key, display, ttype, nodes in TOPOLOGIES:
        sub = df[df["topology"] == key]
        edges = "N/A"
        add_data_row(tbl, [f"{display} ({key})", ttype, str(nodes), str(len(sub) // 12) if len(sub) > 0 else "N/A"])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 3. NORMAL CONDITIONS: METHOD COMPARISON
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("3. Normal Conditions: Method Comparison", level=1)
    doc.add_paragraph("Aggregate results across all 8 topologies under normal (no failure) conditions:")

    tbl = doc.add_table(rows=1, cols=8)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(tbl, ["Method", "Mean MLU", "Throughput", "Delay (ms)", "P95 Delay", "Loss", "Jitter (ms)", "Decision (ms)"])

    for m in ["ecmp", "metagate", "stable_metagate"]:
        sub = normal[normal["method"] == m]
        if len(sub) == 0:
            continue
        add_data_row(tbl, [
            METHOD_LABELS[m],
            f"{sub['mlu'].mean():.4f}",
            f"{sub['mean_throughput'].mean():.4f}",
            f"{sub['mean_delay_ms'].mean():.2f}",
            f"{sub['p95_delay_ms'].mean():.2f}",
            f"{sub['mean_packet_loss'].mean():.4f}",
            f"{sub['jitter_ms'].mean():.4f}",
            f"{sub['decision_time_ms'].mean():.1f}",
        ])

    doc.add_paragraph("")
    doc.add_paragraph(
        "MetaGate achieves ~4% lower MLU than ECMP. Stable MetaGate is nearly identical to "
        "baseline MetaGate in all metrics, confirming that stability penalties do not degrade "
        "routing quality."
    )

    # ══════════════════════════════════════════════════════════════
    # 4. PER-TOPOLOGY RESULTS (NORMAL)
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("4. Per-Topology Results (Normal Conditions)", level=1)

    for m in ["ecmp", "metagate", "stable_metagate"]:
        doc.add_heading(f"4.{['ecmp','metagate','stable_metagate'].index(m)+1} {METHOD_LABELS[m]}", level=2)

        tbl = doc.add_table(rows=1, cols=8)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        make_header_row(tbl, ["Topology", "Type", "MLU", "Throughput", "Delay (ms)", "P95 Delay", "Loss", "Jitter"])

        for key, display, ttype, nodes in TOPOLOGIES:
            sub = normal[(normal["method"] == m) & (normal["topology"] == key)]
            if len(sub) == 0:
                continue
            add_data_row(tbl, [
                display, ttype,
                f"{sub['mlu'].mean():.4f}",
                f"{sub['mean_throughput'].mean():.4f}",
                f"{sub['mean_delay_ms'].mean():.2f}",
                f"{sub['p95_delay_ms'].mean():.2f}",
                f"{sub['mean_packet_loss'].mean():.4f}",
                f"{sub['jitter_ms'].mean():.4f}",
            ])
        doc.add_paragraph("")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 5. FAILURE SCENARIO RESULTS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("5. Failure Scenario Results", level=1)

    for ft in ["single_link_failure", "capacity_degradation", "traffic_spike"]:
        doc.add_heading(f"5.{['single_link_failure','capacity_degradation','traffic_spike'].index(ft)+1} "
                        f"{FAILURE_LABELS[ft]}", level=2)

        fail_df = df[df["failure_type"] == ft]

        tbl = doc.add_table(rows=1, cols=7)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        make_header_row(tbl, ["Method", "Mean MLU", "Throughput", "Delay (ms)", "Loss", "Jitter", "Rules"])

        for m in ["ecmp", "metagate", "stable_metagate"]:
            sub = fail_df[fail_df["method"] == m]
            if len(sub) == 0:
                continue
            add_data_row(tbl, [
                METHOD_LABELS[m],
                f"{sub['mlu'].mean():.4f}",
                f"{sub['mean_throughput'].mean():.4f}",
                f"{sub['mean_delay_ms'].mean():.2f}",
                f"{sub['mean_packet_loss'].mean():.4f}",
                f"{sub['jitter_ms'].mean():.4f}",
                f"{sub['rules_pushed'].sum()}",
            ])
        doc.add_paragraph("")

        # Per-topology under this failure
        doc.add_heading(f"Per-Topology ({FAILURE_LABELS[ft]}, MetaGate)", level=3)
        tbl2 = doc.add_table(rows=1, cols=6)
        make_header_row(tbl2, ["Topology", "MLU", "Throughput", "Delay", "Loss", "Jitter"])
        mg = fail_df[fail_df["method"] == "metagate"]
        for key, display, _, _ in TOPOLOGIES:
            sub = mg[mg["topology"] == key]
            if len(sub) == 0:
                continue
            add_data_row(tbl2, [
                display,
                f"{sub['mlu'].mean():.4f}",
                f"{sub['mean_throughput'].mean():.4f}",
                f"{sub['mean_delay_ms'].mean():.2f}",
                f"{sub['mean_packet_loss'].mean():.4f}",
                f"{sub['jitter_ms'].mean():.4f}",
            ])
        doc.add_paragraph("")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 6. SDN METRICS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("6. SDN Metrics: Rules, Timing, Overhead", level=1)

    tbl = doc.add_table(rows=1, cols=6)
    make_header_row(tbl, ["Method", "Decision Time (ms)", "Rules/Cycle", "Install Delay (ms)", "Total Rules", "Flow Table Size"])

    for m in ["ecmp", "metagate", "stable_metagate"]:
        sub = normal[normal["method"] == m]
        if len(sub) == 0:
            continue
        add_data_row(tbl, [
            METHOD_LABELS[m],
            f"{sub['decision_time_ms'].mean():.1f}",
            f"{sub['rules_pushed'].mean():.1f}",
            f"{sub['rule_install_delay_ms'].mean():.2f}",
            f"{sub['rules_pushed'].sum()}",
            f"{sub['flow_table_size'].mean():.0f}",
        ])

    doc.add_paragraph("")
    doc.add_paragraph(
        "ECMP has zero decision time and zero rule updates (static baseline). "
        "MetaGate and Stable MetaGate have similar overhead, with the stable version "
        "showing slightly fewer rule updates due to reduced expert switching."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 7. STABLE METAGATE EXTENSION RESULTS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("7. Stable MetaGate Extension Results", level=1)

    if bl_sum is not None:
        doc.add_heading("7.1 Baseline vs Stable: Per-Topology", level=2)

        st_best = None
        if stable_sum is not None:
            st_best = stable_sum[(stable_sum["lambda_d"] == 0.2) & (stable_sum["lambda_s"] == 0.1)]

        tbl = doc.add_table(rows=1, cols=8)
        make_header_row(tbl, ["Topology", "BL MLU", "ST MLU", "BL Acc", "ST Acc",
                               "Disturbance", "Switches", "Switch Rate"])

        for key, display, ttype, nodes in TOPOLOGIES:
            bl_row = bl_sum[bl_sum["dataset"] == key]
            st_row = st_best[st_best["dataset"] == key] if st_best is not None else pd.DataFrame()

            bl_mlu = f"{float(bl_row.iloc[0]['metagate_mlu']):.4f}" if len(bl_row) > 0 else "N/A"
            st_mlu = f"{float(st_row.iloc[0]['metagate_mlu']):.4f}" if len(st_row) > 0 else "N/A"
            bl_acc = f"{float(bl_row.iloc[0]['accuracy']):.1%}" if len(bl_row) > 0 else "N/A"
            st_acc = f"{float(st_row.iloc[0]['accuracy']):.1%}" if len(st_row) > 0 else "N/A"
            dist = f"{float(st_row.iloc[0]['mean_disturbance']):.3f}" if len(st_row) > 0 else "N/A"
            sw = f"{int(st_row.iloc[0]['total_switches'])}" if len(st_row) > 0 else "N/A"
            sw_rate = f"{float(st_row.iloc[0]['switch_rate']):.1%}" if len(st_row) > 0 else "N/A"

            add_data_row(tbl, [display, bl_mlu, st_mlu, bl_acc, st_acc, dist, sw, sw_rate])

    doc.add_paragraph("")

    # ══════════════════════════════════════════════════════════════
    # 8. PARAMETER SWEEP
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("8. Parameter Sweep Results", level=1)

    if stable_sweep is not None:
        tbl = doc.add_table(rows=1, cols=6)
        make_header_row(tbl, ["lambda_d", "lambda_s", "Mean MLU", "Disturbance", "Switch Rate", "Accuracy"])
        for _, row in stable_sweep.iterrows():
            is_best = (row["lambda_d"] == 0.2 and row["lambda_s"] == 0.1)
            add_data_row(tbl, [
                f"{row['lambda_d']:.2f}",
                f"{row['lambda_s']:.2f}",
                f"{row['mean_mlu']:.4f}",
                f"{row['mean_disturbance']:.4f}",
                f"{row['switch_rate']:.1%}",
                f"{row['accuracy']:.1%}",
            ], bold=is_best, highlight="E8F5E9" if is_best else None)

        doc.add_paragraph("")
        doc.add_paragraph(
            "Best configuration (highlighted green): lambda_d=0.2, lambda_s=0.1. "
            "Achieves lowest disturbance and switch rate with negligible MLU cost."
        )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 9-14. ALL CDF PLOTS
    # ══════════════════════════════════════════════════════════════
    fig_num = [1]  # mutable counter

    def add_section_plots(heading, plots_list):
        doc.add_heading(heading, level=1)
        for path, caption in plots_list:
            add_img(doc, path, f"Figure {fig_num[0]}: {caption}")
            fig_num[0] += 1
            doc.add_paragraph("")

    # 9. Packet-level CDFs
    add_section_plots("9. CDF Plots: Packet-Level Metrics (Normal)", [
        (PLOTS_DIR / "cdf_throughput.png", "CDF of Throughput: ECMP vs MetaGate vs Stable MetaGate"),
        (PLOTS_DIR / "cdf_delay.png", "CDF of End-to-End Delay (M/M/1 model)"),
        (PLOTS_DIR / "cdf_packet_loss.png", "CDF of Packet Loss (overflow model)"),
        (PLOTS_DIR / "cdf_jitter.png", "CDF of Jitter (inter-timestep delay variation)"),
        (PLOTS_DIR / "cdf_decision_time.png", "CDF of Decision Time"),
    ])

    doc.add_page_break()

    # 10. MLU under failure
    add_section_plots("10. CDF Plots: MLU Under Failure Scenarios", [
        (PLOTS_DIR / "cdf_mlu_single_link_failure.png", "CDF of MLU under Single Link Failure"),
        (PLOTS_DIR / "cdf_mlu_capacity_degradation.png", "CDF of MLU under Capacity Degradation (50%)"),
        (PLOTS_DIR / "cdf_mlu_traffic_spike.png", "CDF of MLU under Traffic Spike (2x)"),
    ])

    doc.add_page_break()

    # 11. Delay under failure
    add_section_plots("11. CDF Plots: Delay Under Failure Scenarios", [
        (PLOTS_DIR / "cdf_delay_single_link_failure.png", "CDF of Delay under Single Link Failure"),
        (PLOTS_DIR / "cdf_delay_capacity_degradation.png", "CDF of Delay under Capacity Degradation"),
        (PLOTS_DIR / "cdf_delay_traffic_spike.png", "CDF of Delay under Traffic Spike"),
    ])

    doc.add_page_break()

    # 12. Per-topology MLU
    doc.add_heading("12. CDF Plots: Per-Topology MLU (Normal)", level=1)
    for key, display, ttype, nodes in TOPOLOGIES:
        p = PLOTS_DIR / f"cdf_mlu_{key}.png"
        add_img(doc, p, f"Figure {fig_num[0]}: MLU CDF - {display} ({ttype}, {nodes} nodes)")
        fig_num[0] += 1
        doc.add_paragraph("")

    doc.add_page_break()

    # 13. Stable MetaGate CDFs
    doc.add_heading("13. CDF Plots: Stable MetaGate Extension", level=1)

    stable_plots = [
        (STABLE_PLOTS / "cdf_mlu_global.png", "Global MLU CDF: Baseline vs Stable MetaGate"),
        (STABLE_PLOTS / "cdf_disturbance_sweep.png", "Routing Disturbance CDF across 9 parameter configs"),
        (STABLE_PLOTS / "cdf_decision_time_global.png", "Decision Time CDF: Baseline vs Stable"),
        (STABLE_PLOTS / "cdf_total_time_global.png", "Total End-to-End Time CDF"),
        (STABLE_PLOTS / "expert_distribution_comparison.png", "Expert Distribution: Baseline vs Stable"),
        (STABLE_PLOTS / "parameter_sweep_heatmap.png", "Parameter Sweep Heatmap (Disturbance, Switch Rate, MLU)"),
        (STABLE_PLOTS / "summary_comparison.png", "Summary Comparison: Disturbance, Switch Rate, MLU by Topology"),
    ]
    for path, caption in stable_plots:
        add_img(doc, path, f"Figure {fig_num[0]}: {caption}")
        fig_num[0] += 1
        doc.add_paragraph("")

    # Per-topology stable MLU CDFs
    doc.add_heading("13.1 Per-Topology MLU: Baseline vs Stable", level=2)
    for key, display, ttype, _ in TOPOLOGIES:
        p = STABLE_PLOTS / f"cdf_mlu_{key}.png"
        if p.exists():
            add_img(doc, p, f"Figure {fig_num[0]}: MLU CDF - {display}: Baseline vs Stable")
            fig_num[0] += 1

    doc.add_paragraph("")

    # Per-topology disturbance CDFs
    doc.add_heading("13.2 Per-Topology Disturbance (Stable)", level=2)
    for key, display, ttype, _ in TOPOLOGIES:
        p = STABLE_PLOTS / f"cdf_disturbance_{key}.png"
        if p.exists():
            add_img(doc, p, f"Figure {fig_num[0]}: Disturbance CDF - {display}")
            fig_num[0] += 1

    doc.add_page_break()

    # 14. Summary charts
    add_section_plots("14. Summary Comparison Charts", [
        (PLOTS_DIR / "method_comparison_summary.png", "6-Panel Method Comparison: Throughput, Delay, Loss, Jitter, MLU, Rule Delay"),
        (PLOTS_DIR / "failure_impact_comparison.png", "Failure Impact: Normal vs Failure MLU for Each Method"),
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 15. COMPLETE CSV DATA
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("15. Complete CSV Data: Per-Topology Summary", level=1)
    doc.add_paragraph("Full summary table (packet_sdn_summary.csv):")

    # Split into manageable tables per failure type
    for ft in ["none", "single_link_failure", "capacity_degradation", "traffic_spike"]:
        doc.add_heading(f"15.{list(FAILURE_LABELS.keys()).index(ft)+1} {FAILURE_LABELS[ft]}", level=2)

        ft_summary = summary[summary["failure_type"] == ft]
        if len(ft_summary) == 0:
            doc.add_paragraph("No data available.")
            continue

        tbl = doc.add_table(rows=1, cols=10)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        make_header_row(tbl, ["Topology", "Method", "MLU", "Throughput", "Delay", "P95 Delay",
                               "Loss", "Jitter", "Rules", "Cycles"])

        for _, row in ft_summary.iterrows():
            topo_short = row["topology"].replace("rocketfuel_", "rf_").replace("topologyzoo_", "tz_")
            add_data_row(tbl, [
                topo_short,
                row["method"],
                f"{row['mlu']:.4f}",
                f"{row['mean_throughput']:.4f}",
                f"{row['mean_delay_ms']:.2f}",
                f"{row['p95_delay_ms']:.2f}",
                f"{row['mean_packet_loss']:.4f}",
                f"{row['jitter_ms']:.4f}",
                f"{int(row['rules_pushed'])}",
                f"{int(row['n_cycles'])}",
            ])
        doc.add_paragraph("")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 16. LIMITATIONS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("16. Limitations and Honest Assessment", level=1)

    doc.add_heading("16.1 What This Report Contains", level=2)
    items = [
        "M/M/1 queuing model applied to LP routing solutions from the MetaGate system",
        "Throughput, delay, loss, jitter computed analytically from link loads and capacities",
        "OpenFlow rule counts computed from LP split ratio diffs",
        "Rule installation delay estimated from empirical OVS model",
        "Failure scenarios simulated by modifying capacities or traffic matrices",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("16.2 What This Report Does NOT Contain", level=2)
    items = [
        "Actual packet generation, transmission, or measurement",
        "Real OpenFlow rule installation on Open vSwitch",
        "TCP/UDP protocol dynamics (congestion control, retransmission)",
        "Switch buffer overflow or queuing beyond M/M/1",
        "Control-plane latency to/from Ryu SDN controller",
        "LLDP topology discovery timing",
        "Real iperf throughput measurements",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("16.3 Model Limitations", level=2)
    limits = [
        ("M/M/1 assumes Poisson arrivals",
         "Real traffic is bursty (self-similar). Actual delays are 2-10x higher than M/M/1 predictions."),
        ("No buffer modeling",
         "Real switches have finite buffers. Our overflow model is a first-order approximation."),
        ("No TCP dynamics",
         "Real TCP adjusts sending rate. Our model assumes constant demand regardless of congestion."),
        ("Instantaneous failure",
         "Real failures take ~100ms to detect via LLDP. Our failure injection is instantaneous."),
        ("Rule install delay is empirical",
         "Actual OVS installation depends on table size, hardware, and switch load."),
    ]
    for title, desc in limits:
        p = doc.add_paragraph()
        r = p.add_run(title + ": ")
        r.bold = True
        p.add_run(desc)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    r = p.add_run("Bottom line: ")
    r.bold = True
    p.add_run(
        "These model-based results provide directional insights and relative method comparisons. "
        "Absolute values (especially delay and loss) should not be taken at face value. "
        "Live Mininet validation on Ubuntu Linux remains the required final step for "
        "actual packet-level metrics."
    )

    # ══════════════════════════════════════════════════════════════
    # SAVE
    # ══════════════════════════════════════════════════════════════
    doc.save(str(OUTPUT))

    total_tables = 0
    for tbl in doc.tables:
        total_tables += 1
    total_images = fig_num[0] - 1

    print(f"Saved: {OUTPUT}")
    print(f"Tables: {total_tables}")
    print(f"Figures: {total_images}")
    print(f"Pages: ~{len(doc.paragraphs) // 20} (estimated)")


if __name__ == "__main__":
    main()

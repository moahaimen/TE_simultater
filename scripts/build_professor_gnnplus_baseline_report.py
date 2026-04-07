#!/usr/bin/env python3
"""Build a Sarah-style GNN+ vs baselines report without MetaGate."""

from __future__ import annotations

import importlib.util
import json
import math
import os
import sys
from datetime import datetime
from pathlib import Path

import matplotlib

PROJECT_ROOT = Path(__file__).resolve().parent.parent
os.chdir(PROJECT_ROOT)
os.environ.setdefault("MPLCONFIGDIR", str(PROJECT_ROOT / ".mplconfig"))
os.environ.setdefault("XDG_CACHE_HOME", str(PROJECT_ROOT / ".cache"))
Path(os.environ["MPLCONFIGDIR"]).mkdir(parents=True, exist_ok=True)
Path(os.environ["XDG_CACHE_HOME"]).mkdir(parents=True, exist_ok=True)
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


HELPER_PATH = PROJECT_ROOT / "scripts" / "build_gnnplus_packet_sdn_report_fixed.py"
INPUT_DIR = Path(
    os.environ.get(
        "PROF_GNNPLUS_BASELINE_INPUT_DIR",
        str(PROJECT_ROOT / "results" / "professor_gnnplus_baselines_zeroshot"),
    )
).resolve()
PLOTS_DIR = INPUT_DIR / "plots"
REPORT_DOC = INPUT_DIR / os.environ.get(
    "PROF_GNNPLUS_BASELINE_REPORT_NAME",
    "GNNPLUS_SARAH_STYLE_BASELINE_REPORT.docx",
)
AUDIT_MD = INPUT_DIR / os.environ.get(
    "PROF_GNNPLUS_BASELINE_AUDIT_NAME",
    "report_audit_gnnplus_sarah_baselines.md",
)

SUMMARY_CSV = INPUT_DIR / "packet_sdn_summary.csv"
FAILURE_CSV = INPUT_DIR / "packet_sdn_failure.csv"
SDN_CSV = INPUT_DIR / "packet_sdn_sdn_metrics.csv"

CHECKPOINT_PATH = PROJECT_ROOT / "results" / "gnn_plus_retrained_fixedk40" / "gnn_plus_fixed_k40.pt"
SOURCE_SUMMARY_JSON = PROJECT_ROOT / "results" / "gnn_plus" / "stage1_regularization" / "training_d02" / "summary.json"

TOPOLOGY_ORDER = [
    "abilene",
    "cernet",
    "geant",
    "ebone",
    "sprintlink",
    "tiscali",
    "germany50",
    "vtlwavenet2011",
]
TOPOLOGY_DISPLAY = {
    "abilene": "Abilene",
    "cernet": "CERNET",
    "geant": "GEANT",
    "ebone": "Ebone",
    "sprintlink": "Sprintlink",
    "tiscali": "Tiscali",
    "germany50": "Germany50",
    "vtlwavenet2011": "VtlWavenet2011",
}
METHOD_ORDER = ["gnnplus", "bottleneck", "topk", "sensitivity", "ecmp", "ospf"]
METHOD_LABELS = {
    "gnnplus": "GNN+",
    "bottleneck": "Bottleneck",
    "topk": "TopK",
    "sensitivity": "Sensitivity",
    "ecmp": "ECMP",
    "ospf": "OSPF",
}
METHOD_COLORS = {
    "gnnplus": "#d62728",
    "bottleneck": "#2ca02c",
    "topk": "#9467bd",
    "sensitivity": "#8c564b",
    "ecmp": "#1f77b4",
    "ospf": "#17becf",
}
SCENARIO_ORDER = [
    "single_link_failure",
    "random_link_failure_1",
    "random_link_failure_2",
    "capacity_degradation_50",
    "traffic_spike_2x",
]
SCENARIO_LABELS = {
    "single_link_failure": "Single Link Failure",
    "random_link_failure_1": "Random Link Failure (1)",
    "random_link_failure_2": "Random Link Failure (2)",
    "capacity_degradation_50": "Capacity Degradation (50%)",
    "traffic_spike_2x": "Traffic Spike (2x)",
}
TITLE_MAIN = os.environ.get("PROF_GNNPLUS_BASELINE_TITLE", "GNN+ Baseline Comparison Report")
TITLE_SUB = os.environ.get(
    "PROF_GNNPLUS_BASELINE_SUBTITLE",
    "Sarah-Style Seminar Version\n"
    "Proposed Method: GNN+\n"
    "Baselines: Bottleneck, TopK, Sensitivity, ECMP, OSPF\n"
    "No MetaGate, No Stable MetaGate, No Original GNN in Main Comparison",
)
SYNTHETIC_WARNING = os.environ.get("PROF_GNNPLUS_BASELINE_SYNTHETIC_WARNING", "").strip()


def load_helper():
    spec = importlib.util.spec_from_file_location("build_gnnplus_packet_sdn_report_fixed", HELPER_PATH)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Could not load helper from {HELPER_PATH}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def add_section_table(helper, doc: Document, title: str, df: pd.DataFrame, font_size: int = 8) -> None:
    doc.add_paragraph(title)
    helper.add_dataframe_table(doc, df, font_size=font_size)
    doc.add_paragraph("")


def plot_grouped_metric(df: pd.DataFrame, metric: str, title: str, ylabel: str, output: Path, log_scale: bool = False):
    fig, ax = plt.subplots(figsize=(14, 6))
    x = np.arange(len(TOPOLOGY_ORDER))
    width = 0.12
    offsets = np.linspace(-width * 2.5, width * 2.5, len(METHOD_ORDER))

    for idx, method in enumerate(METHOD_ORDER):
        vals = []
        for topo in TOPOLOGY_ORDER:
            row = df[(df["topology"] == topo) & (df["method"] == method)]
            vals.append(float(row.iloc[0][metric]) if not row.empty else np.nan)
        ax.bar(x + offsets[idx], vals, width=width, label=METHOD_LABELS[method], color=METHOD_COLORS[method], edgecolor="white", linewidth=0.6)

    ax.set_title(title)
    ax.set_xlabel("Topology")
    ax.set_ylabel(ylabel)
    ax.set_xticks(x)
    ax.set_xticklabels([TOPOLOGY_DISPLAY[t] for t in TOPOLOGY_ORDER], rotation=35, ha="right")
    ax.grid(axis="y", alpha=0.25)
    ax.legend(ncol=3, fontsize=8)
    if log_scale:
        ax.set_yscale("log")
    fig.tight_layout()
    fig.savefig(output, dpi=160, bbox_inches="tight")
    plt.close(fig)


def plot_gnnplus_gap(summary_df: pd.DataFrame, output: Path):
    fig, ax = plt.subplots(figsize=(12, 5))
    gaps = []
    for topo in TOPOLOGY_ORDER:
        topo_df = summary_df[summary_df["topology"] == topo]
        gnnplus = float(topo_df[topo_df["method"] == "gnnplus"].iloc[0]["mean_mlu"])
        best = float(topo_df["mean_mlu"].min())
        gaps.append(((gnnplus - best) / max(abs(best), 1e-12)) * 100.0)
    x = np.arange(len(TOPOLOGY_ORDER))
    ax.bar(x, gaps, color="#d62728", edgecolor="white")
    ax.axhline(0.0, color="black", linewidth=1.0)
    ax.set_xticks(x)
    ax.set_xticklabels([TOPOLOGY_DISPLAY[t] for t in TOPOLOGY_ORDER], rotation=35, ha="right")
    ax.set_ylabel("GNN+ Gap to Best Method (%)")
    ax.set_title("GNN+ Gap to the Best Baseline per Topology")
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(output, dpi=160, bbox_inches="tight")
    plt.close(fig)


def plot_failure_recovery(failure_df: pd.DataFrame, output: Path):
    fig, ax = plt.subplots(figsize=(14, 6))
    x = np.arange(len(TOPOLOGY_ORDER))
    width = 0.14
    offsets = np.linspace(-width * 2, width * 2, len(SCENARIO_ORDER))
    gnnplus_df = failure_df[failure_df["method"] == "gnnplus"]
    colors = ["#4c78a8", "#f58518", "#54a24b", "#e45756", "#72b7b2"]
    for idx, scenario in enumerate(SCENARIO_ORDER):
        vals = []
        for topo in TOPOLOGY_ORDER:
            row = gnnplus_df[(gnnplus_df["topology"] == topo) & (gnnplus_df["scenario"] == scenario)]
            vals.append(float(row.iloc[0]["failure_recovery_ms"]) if not row.empty else np.nan)
        ax.bar(x + offsets[idx], vals, width=width, label=SCENARIO_LABELS[scenario], color=colors[idx], edgecolor="white", linewidth=0.6)
    ax.set_xticks(x)
    ax.set_xticklabels([TOPOLOGY_DISPLAY[t] for t in TOPOLOGY_ORDER], rotation=35, ha="right")
    ax.set_ylabel("Recovery Time (ms)")
    ax.set_title("GNN+ Failure Recovery Across Scenarios")
    ax.grid(axis="y", alpha=0.25)
    ax.legend(ncol=2, fontsize=8)
    fig.tight_layout()
    fig.savefig(output, dpi=160, bbox_inches="tight")
    plt.close(fig)


def plot_cdf(summary_df: pd.DataFrame, metric: str, title: str, output: Path):
    fig, ax = plt.subplots(figsize=(8, 5))
    for method in METHOD_ORDER:
        vals = np.sort(summary_df[summary_df["method"] == method][metric].to_numpy(dtype=float))
        y = np.arange(1, len(vals) + 1) / max(len(vals), 1)
        ax.plot(vals, y, label=METHOD_LABELS[method], color=METHOD_COLORS[method], linewidth=2)
    ax.set_title(title)
    ax.set_ylabel("CDF")
    ax.set_xlabel(metric.replace("_", " ").title())
    ax.grid(alpha=0.25)
    ax.legend(fontsize=8)
    fig.tight_layout()
    fig.savefig(output, dpi=160, bbox_inches="tight")
    plt.close(fig)


def build_report():
    helper = load_helper()
    PLOTS_DIR.mkdir(parents=True, exist_ok=True)
    if not SUMMARY_CSV.exists() or not FAILURE_CSV.exists() or not SDN_CSV.exists():
        missing = [str(p.relative_to(PROJECT_ROOT)) for p in [SUMMARY_CSV, FAILURE_CSV, SDN_CSV] if not p.exists()]
        raise FileNotFoundError("Missing required input files:\n" + "\n".join(missing))

    summary = pd.read_csv(SUMMARY_CSV)
    failure = pd.read_csv(FAILURE_CSV)
    sdn_metrics = pd.read_csv(SDN_CSV)
    stage1_summary = json.loads(SOURCE_SUMMARY_JSON.read_text(encoding="utf-8"))
    provenance = helper.load_checkpoint_provenance()

    plot_grouped_metric(summary, "mean_mlu", "GNN+ vs Baselines: Mean MLU", "Mean MLU", PLOTS_DIR / "mlu_comparison.png", log_scale=True)
    plot_grouped_metric(summary, "throughput", "GNN+ vs Baselines: Throughput", "Throughput", PLOTS_DIR / "throughput_comparison.png")
    plot_grouped_metric(summary, "mean_disturbance", "GNN+ vs Baselines: Routing Disturbance", "Mean Disturbance", PLOTS_DIR / "disturbance_comparison.png")
    plot_grouped_metric(summary, "decision_time_ms", "GNN+ vs Baselines: Decision Time", "Decision Time (ms)", PLOTS_DIR / "decision_time_comparison.png")
    plot_gnnplus_gap(summary, PLOTS_DIR / "gnnplus_gap.png")
    plot_failure_recovery(failure, PLOTS_DIR / "gnnplus_failure_recovery.png")
    plot_cdf(summary, "mean_mlu", "CDF of Mean MLU Across Methods", PLOTS_DIR / "cdf_mean_mlu.png")
    plot_cdf(summary, "decision_time_ms", "CDF of Decision Time Across Methods", PLOTS_DIR / "cdf_decision_time.png")

    doc = Document()
    helper.set_default_style(doc)

    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE_MAIN)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x80)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(TITLE_SUB)
    run.font.size = Pt(11)

    if SYNTHETIC_WARNING:
        warn = doc.add_paragraph()
        warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = warn.add_run(SYNTHETIC_WARNING)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xAA, 0x00, 0x00)

    stamp = doc.add_paragraph()
    stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stamp.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_page_break()

    doc.add_heading("Table of Contents", level=1)
    toc = [
        "1. Method Summary",
        "2. MLU Results (Maximum Link Utilization)",
        "3. Performance Ratio (PR)",
        "4. Network Disturbance (DB)",
        "5. Execution Time and Decision Latency",
        "6. Training Efficiency and Checkpoint Provenance",
        "7. Routing Update Overhead",
        "8. Failure Robustness",
        "9. CDF Plots",
        "10. SDN Deployment Validation",
        "11. Complexity Analysis",
        "12. Unseen-Topology Zero-Shot Deep Dive",
        "13. Seminar Positioning",
        "14. Limitations",
        "15. Live Mininet Testbed",
        "16. Exact Thesis Method Description",
    ]
    for item in toc:
        doc.add_paragraph(item)
    doc.add_page_break()

    doc.add_heading("1. Method Summary", level=1)
    doc.add_paragraph(
        "This report reframes the seminar story around GNN+ as the proposed learned method. "
        "The comparison set is GNN+ vs Bottleneck, TopK, Sensitivity, ECMP, and OSPF. "
        "MetaGate and Stable MetaGate are intentionally excluded from this report."
    )
    helper.add_bullet(doc, "This is a direct baseline-comparison study, not an adaptive expert-selection study.")
    helper.add_bullet(doc, "Unseen topologies (Germany50 and VtlWavenet2011) are evaluated directly with the trained GNN+ checkpoint and no calibration.")
    helper.add_bullet(doc, "Original GNN is intentionally omitted from the main visible comparison for seminar clarity.")
    helper.add_bullet(doc, "All SDN metrics in this report are model-based analytical control-loop metrics, not live Mininet measurements.")
    if SYNTHETIC_WARNING:
        helper.add_bullet(doc, SYNTHETIC_WARNING)
    doc.add_page_break()

    doc.add_heading("2. MLU Results (Maximum Link Utilization)", level=1)
    helper.add_image(doc, PLOTS_DIR / "mlu_comparison.png", "Figure 1. GNN+ compared with Bottleneck, TopK, Sensitivity, ECMP, and OSPF.", width=6.8)
    mlu_rows = []
    for topo in TOPOLOGY_ORDER:
        topo_df = summary[summary["topology"] == topo]
        row = {"Topology": TOPOLOGY_DISPLAY[topo], "Status": topo_df.iloc[0]["status"].title()}
        for method in METHOD_ORDER:
            row[METHOD_LABELS[method]] = float(topo_df[topo_df["method"] == method].iloc[0]["mean_mlu"])
        mlu_rows.append(row)
    helper.add_dataframe_table(doc, pd.DataFrame(mlu_rows), font_size=7)
    helper.add_image(doc, PLOTS_DIR / "gnnplus_gap.png", "Figure 2. GNN+ gap to the best baseline on each topology.", width=6.2)
    doc.add_page_break()

    doc.add_heading("3. Performance Ratio (PR)", level=1)
    pr_rows = []
    for topo in TOPOLOGY_ORDER:
        topo_df = summary[summary["topology"] == topo].copy()
        best = float(topo_df["mean_mlu"].min())
        row = {"Topology": TOPOLOGY_DISPLAY[topo]}
        for method in METHOD_ORDER:
            val = float(topo_df[topo_df["method"] == method].iloc[0]["mean_mlu"])
            row[METHOD_LABELS[method] + " PR"] = val / max(best, 1e-12)
        pr_rows.append(row)
    helper.add_dataframe_table(doc, pd.DataFrame(pr_rows), font_size=7)
    doc.add_page_break()

    doc.add_heading("4. Network Disturbance (DB)", level=1)
    helper.add_image(doc, PLOTS_DIR / "disturbance_comparison.png", "Figure 3. Routing disturbance comparison across all methods.", width=6.8)
    disturb_df = summary.copy()
    disturb_df["Method"] = disturb_df["method"].map(METHOD_LABELS)
    disturb_df["Topology"] = disturb_df["topology"].map(TOPOLOGY_DISPLAY)
    disturb_df["Status"] = disturb_df["status"].str.title()
    helper.add_dataframe_chunks(
        doc,
        disturb_df[["Method", "Topology", "Status", "mean_disturbance"]].rename(columns={"mean_disturbance": "Mean Disturbance"}),
        "Routing disturbance table",
        chunk_size=18,
    )
    doc.add_page_break()

    doc.add_heading("5. Execution Time and Decision Latency", level=1)
    helper.add_image(doc, PLOTS_DIR / "decision_time_comparison.png", "Figure 4. Decision-time comparison across all methods.", width=6.8)
    timing_df = summary.copy()
    timing_df["Method"] = timing_df["method"].map(METHOD_LABELS)
    timing_df["Topology"] = timing_df["topology"].map(TOPOLOGY_DISPLAY)
    helper.add_dataframe_chunks(
        doc,
        timing_df[["Method", "Topology", "decision_time_ms"]].rename(columns={"decision_time_ms": "Decision Time (ms)"}),
        "Decision-time table",
        chunk_size=18,
    )
    doc.add_page_break()

    doc.add_heading("6. Training Efficiency and Checkpoint Provenance", level=1)
    train_rows = pd.DataFrame(
        [
            {"Metric": "Supervised stage time (s)", "Value": stage1_summary.get("sup_time")},
            {"Metric": "Supervised best epoch", "Value": stage1_summary.get("sup_best_epoch")},
            {"Metric": "RL fine-tuning time (s)", "Value": stage1_summary.get("rl_time")},
            {"Metric": "RL best epoch", "Value": stage1_summary.get("rl_best_epoch")},
            {"Metric": "Total GNN+ training time (s)", "Value": stage1_summary.get("total_time")},
            {"Metric": "Checkpoint used", "Value": str(CHECKPOINT_PATH.relative_to(PROJECT_ROOT))},
        ]
    )
    helper.add_dataframe_table(doc, train_rows, font_size=8)
    prov_rows = pd.DataFrame(
        [
            {"Field": "Dropout", "Value": provenance["dropout"]},
            {"Field": "learn_k_crit", "Value": provenance["learn_k_crit"]},
            {"Field": "Fixed K", "Value": provenance["fixed_k"]},
            {"Field": "Source checkpoint", "Value": provenance["source_checkpoint_path"]},
        ]
    )
    helper.add_dataframe_table(doc, prov_rows, font_size=8)
    doc.add_page_break()

    doc.add_heading("7. Routing Update Overhead", level=1)
    overhead = sdn_metrics.copy()
    overhead["Method"] = overhead["method"].map(METHOD_LABELS)
    overhead["Topology"] = overhead["topology"].map(TOPOLOGY_DISPLAY)
    helper.add_dataframe_chunks(
        doc,
        overhead[["Method", "Topology", "flow_table_updates", "rule_install_delay_ms"]].rename(
            columns={"flow_table_updates": "Flow Table Updates", "rule_install_delay_ms": "Rule Install Delay (ms)"}
        ),
        "Routing-update overhead table",
        chunk_size=18,
    )
    doc.add_page_break()

    doc.add_heading("8. Failure Robustness", level=1)
    helper.add_image(doc, PLOTS_DIR / "gnnplus_failure_recovery.png", "Figure 5. GNN+ failure recovery across the five scenarios.", width=6.8)
    for idx, scenario in enumerate(SCENARIO_ORDER, start=1):
        doc.add_heading(f"8.{idx} {SCENARIO_LABELS[scenario]}", level=2)
        scenario_df = failure[failure["scenario"] == scenario].copy()
        scenario_df["Method"] = scenario_df["method"].map(METHOD_LABELS)
        scenario_df["Topology"] = scenario_df["topology"].map(TOPOLOGY_DISPLAY)
        scenario_df["Status"] = scenario_df["status"].str.title()
        helper.add_dataframe_chunks(
            doc,
            scenario_df[["Method", "Topology", "Status", "mean_mlu", "failure_recovery_ms"]].rename(
                columns={"mean_mlu": "Post-Failure MLU", "failure_recovery_ms": "Recovery (ms)"}
            ),
            f"{SCENARIO_LABELS[scenario]} table",
            chunk_size=18,
        )
    doc.add_page_break()

    doc.add_heading("9. CDF Plots", level=1)
    helper.add_image(doc, PLOTS_DIR / "cdf_mean_mlu.png", "Figure 6. CDF of mean MLU across methods.", width=6.3)
    helper.add_image(doc, PLOTS_DIR / "cdf_decision_time.png", "Figure 7. CDF of decision time across methods.", width=6.3)
    doc.add_page_break()

    doc.add_heading("10. SDN Deployment Validation", level=1)
    note = doc.add_paragraph()
    run = note.add_run(
        "HONESTY NOTE: the SDN metrics below are model-based analytical control-loop metrics, not live packet measurements from Mininet."
    )
    run.bold = True
    run.font.color.rgb = RGBColor(0xAA, 0x00, 0x00)
    sdn_df = sdn_metrics.copy()
    sdn_df["Method"] = sdn_df["method"].map(METHOD_LABELS)
    sdn_df["Topology"] = sdn_df["topology"].map(TOPOLOGY_DISPLAY)
    sdn_df["Status"] = sdn_df["status"].str.title()
    sdn_df = sdn_df.rename(
        columns={
            "mean_mlu": "MLU",
            "throughput": "Throughput",
            "mean_latency_au": "Mean Delay",
            "p95_latency_au": "P95 Delay",
            "packet_loss": "Packet Loss",
            "jitter_au": "Jitter",
            "decision_time_ms": "Decision Time (ms)",
            "flow_table_updates": "Flow Table Updates",
            "rule_install_delay_ms": "Rule Install Delay (ms)",
            "avg_failure_recovery_ms": "Failure Recovery (ms)",
        }
    )
    helper.add_dataframe_chunks(
        doc,
        sdn_df[["Method", "Topology", "Status", "MLU", "Throughput", "Mean Delay", "P95 Delay", "Packet Loss", "Jitter", "Decision Time (ms)", "Flow Table Updates", "Rule Install Delay (ms)", "Failure Recovery (ms)"]],
        "SDN metrics table",
        chunk_size=12,
    )
    doc.add_page_break()

    doc.add_heading("11. Complexity Analysis", level=1)
    complexity_rows = pd.DataFrame(
        [
            {"Method": "GNN+", "Controller Work": "Graph-neural inference + LP reroute"},
            {"Method": "Bottleneck", "Controller Work": "Heuristic critical-flow selection + LP reroute"},
            {"Method": "TopK", "Controller Work": "Demand ranking + LP reroute"},
            {"Method": "Sensitivity", "Controller Work": "Sensitivity scoring + LP reroute"},
            {"Method": "ECMP", "Controller Work": "Static equal-cost split"},
            {"Method": "OSPF", "Controller Work": "Shortest-path split without learned inference"},
        ]
    )
    helper.add_dataframe_table(doc, complexity_rows, font_size=8)
    comp_numeric = []
    for method in METHOD_ORDER:
        mdf = summary[summary["method"] == method]
        comp_numeric.append(
            {
                "Method": METHOD_LABELS[method],
                "Avg Decision Time (ms)": float(mdf["decision_time_ms"].mean()),
                "Max Decision Time (ms)": float(mdf["decision_time_ms"].max()),
                "Avg Rule Delay (ms)": float(mdf["rule_install_delay_ms"].mean()),
            }
        )
    helper.add_dataframe_table(doc, pd.DataFrame(comp_numeric), font_size=8)
    doc.add_page_break()

    doc.add_heading("12. Unseen-Topology Zero-Shot Deep Dive", level=1)
    unseen = summary[summary["status"] == "unseen"].copy()
    unseen["Method"] = unseen["method"].map(METHOD_LABELS)
    unseen["Topology"] = unseen["topology"].map(TOPOLOGY_DISPLAY)
    helper.add_dataframe_chunks(
        doc,
        unseen[["Method", "Topology", "mean_mlu", "throughput", "decision_time_ms"]].rename(
            columns={"mean_mlu": "Mean MLU", "decision_time_ms": "Decision Time (ms)"}
        ),
        "Unseen-topology zero-shot table",
        chunk_size=18,
    )
    doc.add_page_break()

    doc.add_heading("13. Seminar Positioning", level=1)
    doc.add_paragraph(
        "For seminar presentation, GNN+ is the proposed learned method. The remaining visible methods in this report are baselines: Bottleneck, TopK, Sensitivity, ECMP, and OSPF."
    )
    helper.add_bullet(doc, "Original GNN is intentionally removed from the main visible comparison for presentation clarity.")
    helper.add_bullet(doc, "No method labels are renamed: each row keeps its true method identity.")
    doc.add_page_break()

    doc.add_heading("14. Limitations", level=1)
    helper.add_bullet(doc, "This report excludes MetaGate and Stable MetaGate by design.")
    helper.add_bullet(doc, "Original GNN is not shown in the main report body, even though it exists elsewhere in the repo.")
    helper.add_bullet(doc, "Paper baselines like FlexDATE, FlexEntry, and ERODRL are not included here because they are not runnable in this branch.")
    helper.add_bullet(doc, "Packet-level SDN metrics remain model-based analytical metrics, not live Mininet measurements.")
    doc.add_page_break()

    doc.add_heading("15. Live Mininet Testbed", level=1)
    doc.add_paragraph(
        "This report does not claim live Mininet packet measurements. A separate Mininet deployment would still be required for actual packet latency and dataplane validation."
    )
    doc.add_page_break()

    doc.add_heading("16. Exact Thesis Method Description", level=1)
    doc.add_paragraph(
        "We evaluate GNN+ as the proposed learned critical-flow selector in a direct baseline-comparison setting. "
        "At each control cycle, GNN+ selects K=40 critical flows from the current traffic matrix and topology features. "
        "Those critical flows are then optimized by LP, while the remaining non-critical flows stay on the baseline routing policy."
    )
    doc.add_paragraph(
        "The seminar version of this report compares GNN+ against Bottleneck, TopK, Sensitivity, ECMP, and OSPF. "
        "No MetaGate or Stable MetaGate is involved in this report, and no Bayesian calibration is applied on unseen topologies."
    )

    REPORT_DOC.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_DOC)

    audit_lines = [
        "# GNN+ Sarah-Style Baseline Report Audit",
        "",
        f"- Report: `{REPORT_DOC.relative_to(PROJECT_ROOT)}`",
        f"- Results directory: `{INPUT_DIR.relative_to(PROJECT_ROOT)}`",
        "- Proposed method highlighted: GNN+",
        "- Visible baselines: Bottleneck, TopK, Sensitivity, ECMP, OSPF",
        "- Original GNN hidden from main body: yes",
        "- MetaGate included: no",
        f"- Synthetic warning active: {'yes' if bool(SYNTHETIC_WARNING) else 'no'}",
    ]
    AUDIT_MD.write_text("\n".join(audit_lines), encoding="utf-8")
    print(f"Report saved to {REPORT_DOC}")


if __name__ == "__main__":
    build_report()

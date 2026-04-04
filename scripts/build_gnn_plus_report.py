#!/usr/bin/env python3
"""Generate GNN_Plus_Screening_Report.docx using python-docx.

Output: results/gnn_plus/GNN_Plus_Screening_Report.docx
"""

import os
import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))
os.chdir(PROJECT_ROOT)

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

ROOT = Path("results/gnn_plus")
PLOTS = ROOT / "plots"
OUT = ROOT / "GNN_Plus_Screening_Report.docx"


def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    tcPr.append(shading)


def add_table(doc, headers, rows, col_widths=None, header_color="2E4057"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Arial"
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = "Arial"

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)

    return t


def add_image_safe(doc, img_path, width=Inches(6)):
    if img_path.exists():
        doc.add_picture(str(img_path), width=width)
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[Image not found: {img_path.name}]")


doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)

# ==============================
# TITLE
# ==============================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("GNN+ Screening Experiment Report")
run.bold = True
run.font.size = Pt(20)
run.font.name = "Arial"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AI-Driven Traffic Routing with MLP Meta-Gate")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Date: 2026-03-30  |  Branch: gnn-plus-extension  |  Type: Screening (not full ablation)")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# Disclaimer
p = doc.add_paragraph()
run = p.add_run("IMPORTANT: ")
run.bold = True
run.font.color.rgb = RGBColor(0x85, 0x64, 0x04)
run = p.add_run("This is a screening experiment testing the combined GNN+ upgrade (richer features + dynamic K). "
                "It does NOT isolate the individual effects of features vs. dynamic K. "
                "The original GNN baseline is unchanged.")
run.font.color.rgb = RGBColor(0x85, 0x64, 0x04)

# ==============================
# 1. Motivation
# ==============================
doc.add_heading("1. Motivation", level=1)
doc.add_paragraph(
    "The original GNN expert in our MetaGate pipeline uses 34 effective input features "
    "(with 4 placeholder zeros in node features). It operates with a fixed K=40 critical "
    "flow budget regardless of topology or traffic conditions. This screening experiment "
    "tests whether enriching the GNN input features and enabling dynamic K prediction can "
    "improve routing quality."
)

# ==============================
# 2. Old GNN Limitations
# ==============================
doc.add_heading("2. Original GNN Limitations", level=1)
doc.add_paragraph("Identified weaknesses in the current GNN expert:")

limitations = [
    "4 wasted node feature slots (zero padding at indices 12-15)",
    "No ECMP per-OD demand contribution to node features",
    "No path overlap / shared bottleneck features between OD pairs",
    "No temporal features (traffic change between t and t-1)",
    "No hop count / path length in OD features (only weighted cost)",
    "No source/destination local congestion in OD scoring",
    "K_pred head exists but is dead code (force_default_k=True at inference)",
]
for lim in limitations:
    doc.add_paragraph(lim, style="List Bullet")

# ==============================
# 3. New Feature Design
# ==============================
doc.add_heading("3. GNN+ Feature Design", level=1)
doc.add_paragraph("GNN+ enriches all three feature groups while maintaining backward compatibility:")

doc.add_heading("3.1 Node Features [V, 16]", level=2)
doc.add_paragraph("Indices 0-11 unchanged. Indices 12-15 replaced with real features:")
add_table(doc,
    ["Idx", "Feature", "Source", "Replaces"],
    [
        ["12", "ECMP demand through node", "TM + OD pairs", "zero"],
        ["13", "Congested neighbor fraction", "Adjacency + util", "zero"],
        ["14", "Max residual capacity", "cap - load", "zero"],
        ["15", "Clustering coefficient", "Topology triangles", "zero"],
    ],
    col_widths=[0.5, 2.5, 1.8, 1.7],
)

doc.add_heading("3.2 Edge Features [E, 8 -> 12]", level=2)
add_table(doc,
    ["Idx", "Feature", "Source", "Status"],
    [
        ["0-7", "Original 8 features", "Unchanged", "Kept"],
        ["8", "OD paths sharing edge", "Path library", "NEW"],
        ["9", "Residual capacity (abs)", "cap - load", "NEW"],
        ["10", "Load change ratio (t/t-1)", "Temporal", "NEW"],
        ["11", "Is bottleneck indicator", "Path analysis", "NEW"],
    ],
    col_widths=[0.5, 2.5, 1.8, 1.7],
)

doc.add_heading("3.3 OD Features [num_od, 10 -> 18]", level=2)
add_table(doc,
    ["Idx", "Feature", "Source", "Status"],
    [
        ["0-9", "Original 10 features", "Unchanged", "Kept"],
        ["10", "Hop count (normalized)", "Path library", "NEW"],
        ["11", "Demand change ratio", "Temporal (t/t-1)", "NEW"],
        ["12", "Source congestion", "Node max_util_out", "NEW"],
        ["13", "Destination congestion", "Node max_util_in", "NEW"],
        ["14", "Path overlap score", "Edge sharing count", "NEW"],
        ["15", "ECMP congestion contribution", "demand * bottleneck", "NEW"],
        ["16", "Alternative path headroom", "Non-best paths", "NEW"],
        ["17", "Demand x hop (normalized)", "Resource proxy", "NEW"],
    ],
    col_widths=[0.5, 2.5, 1.8, 1.7],
)

doc.add_paragraph("Total effective features: 16 + 12 + 18 = 46 (up from 34).")

# ==============================
# 4. Dynamic Bounded K
# ==============================
doc.add_heading("4. Dynamic Bounded K Design", level=1)
p = doc.add_paragraph()
run = p.add_run("Rule: K = max(K_min, min(K_pred, 40)) if K_pred is not None else 40")
run.bold = True
doc.add_paragraph(
    "K_min = 15. Justification: Smallest topology (Abilene) has 132 OD pairs; 15 is ~11% coverage, "
    "the minimum meaningful selection for LP optimization. K_max = 40 matches all prior experiments "
    "for fair comparison."
)
doc.add_paragraph("Implementation: GNN+ model uses force_default_k=False, with k_head output clamped to [15, 40].")

# ==============================
# 5. Training Summary
# ==============================
doc.add_heading("5. Training Summary", level=1)
add_table(doc,
    ["Parameter", "Value"],
    [
        ["Topologies", "Abilene, GEANT, Germany50"],
        ["Train / Val samples", "120 / 45"],
        ["Max epochs", "30 (early stop patience=8)"],
        ["Best epoch", "28"],
        ["Best val loss", "2.815"],
        ["Final alpha (residual weight)", "0.308"],
        ["Val selection overlap", "0.790 (79.0% Jaccard with oracle)"],
        ["Training time", "72.3 seconds"],
        ["Learning rate", "5e-4 (AdamW + cosine annealing)"],
        ["K_pred at convergence", "39 (nearly fixed at upper bound)"],
    ],
    col_widths=[3.0, 3.5],
)

# ==============================
# 6. Screening Results
# ==============================
doc.add_page_break()
doc.add_heading("6. Screening Results", level=1)

doc.add_heading("6.1 Summary Table", level=2)
add_table(doc,
    ["Topology", "MLU Orig", "MLU GNN+", "PR Orig", "PR GNN+", "Dist Orig", "Dist GNN+", "GNN+ Wins", "MLU Change"],
    [
        ["Abilene", "0.0546", "0.0546", "0.00%", "0.00%", "0.093", "0.091", "27/75", "~0%"],
        ["GEANT", "0.1615", "0.1630", "0.80%", "1.70%", "0.120", "0.098", "15/75", "-0.94%"],
        ["Germany50", "18.941", "19.270", "-1.65%", "0.25%", "0.268", "0.174", "2/44", "-1.74%"],
        ["AGGREGATE", "4.379", "4.455", "-0.07%", "0.71%", "0.143", "0.113", "44/194", "-1.72%"],
    ],
    col_widths=[1.0, 0.75, 0.75, 0.65, 0.65, 0.65, 0.65, 0.7, 0.7],
)

p = doc.add_paragraph()
run = p.add_run("Key finding: GNN+ has WORSE MLU on GEANT (-0.94%) and Germany50 (-1.74%), "
                "but BETTER disturbance on all topologies (21% aggregate improvement).")
run.bold = True

doc.add_heading("6.2 Dynamic K Distribution", level=2)
add_table(doc,
    ["Topology", "K Mean", "K Min", "K Max", "K Std", "Note"],
    [
        ["Abilene", "39", "39", "39", "0", "Constant"],
        ["GEANT", "39", "39", "39", "0", "Constant"],
        ["Germany50", "39", "39", "39", "0", "Constant"],
    ],
    col_widths=[1.2, 0.9, 0.9, 0.9, 0.9, 1.7],
)

p = doc.add_paragraph()
run = p.add_run("Observation: K_pred converged to 39 for all topologies. The k_head learned to output "
                "near-maximum, meaning the model found no benefit from reducing K below 40. "
                "The dynamic K mechanism is effectively inactive.")
run.bold = True

doc.add_heading("6.3 Execution Time", level=2)
add_table(doc,
    ["Topology", "Orig GNN (ms)", "GNN+ (ms)", "Overhead"],
    [
        ["Abilene", "2.04", "2.93", "+44%"],
        ["GEANT", "3.95", "6.85", "+73%"],
        ["Germany50", "16.45", "29.22", "+78%"],
    ],
    col_widths=[1.6, 1.6, 1.6, 1.7],
)
doc.add_paragraph(
    "GNN+ feature building is ~44-78% slower due to path overlap computation and clustering "
    "coefficient calculation. Both remain sub-30ms, well within real-time requirements."
)

# ==============================
# 7. Comparison Plots
# ==============================
doc.add_page_break()
doc.add_heading("7. Comparison Plots", level=1)

doc.add_heading("7.1 MLU CDF: Original GNN vs GNN+", level=2)
add_image_safe(doc, PLOTS / "mlu_cdf_comparison.png", width=Inches(6.2))
doc.add_paragraph("Figure 1: CDF of MLU across test timesteps. Original GNN dominates on GEANT and Germany50.")

doc.add_heading("7.2 Dynamic K Distribution", level=2)
add_image_safe(doc, PLOTS / "k_distribution.png", width=Inches(6.2))
doc.add_paragraph("Figure 2: K_pred is constant at 39 across all topologies. Dynamic K mechanism did not activate.")

doc.add_heading("7.3 Mean MLU Bar Chart", level=2)
add_image_safe(doc, PLOTS / "mean_mlu_bar_chart.png", width=Inches(5.0))
doc.add_paragraph("Figure 3: Per-topology mean MLU comparison. GNN+ is slightly worse on GEANT and Germany50.")

doc.add_heading("7.4 Disturbance CDF", level=2)
add_image_safe(doc, PLOTS / "disturbance_cdf_comparison.png", width=Inches(6.2))
doc.add_paragraph("Figure 4: GNN+ achieves lower disturbance on all topologies, especially Germany50 (0.174 vs 0.268).")

# ==============================
# 8. Comparison with Original GNN
# ==============================
doc.add_page_break()
doc.add_heading("8. Comparison with Original GNN", level=1)

comparisons = [
    ("MLU Quality", "Original GNN is BETTER. GNN+ is 0.94-1.74% worse on GEANT and Germany50. On Abilene, they are tied."),
    ("Disturbance", "GNN+ is BETTER. 21% lower aggregate disturbance (0.113 vs 0.143). Especially strong on Germany50: 35% reduction."),
    ("Dynamic K", "NOT ACTIVE. K_pred converges to 39 for all topologies. The model learned that maximum K is optimal."),
    ("Execution Time", "GNN+ is 44-78% slower due to richer feature computation. Still sub-30ms, not a practical concern."),
    ("Training", "Similar convergence: 28-29 epochs, 72s total. Val overlap 0.79 (comparable to original)."),
]

for aspect, detail in comparisons:
    p = doc.add_paragraph()
    run = p.add_run(f"{aspect}: ")
    run.bold = True
    p.add_run(detail)

# ==============================
# 9. Honest Limitations
# ==============================
doc.add_heading("9. Honest Limitations", level=1)

honest_items = [
    "This is a SCREENING experiment, not a full ablation. We cannot isolate whether the MLU "
    "regression is caused by the richer features, the dynamic K, or the retraining from scratch.",
    "GNN+ was trained from random initialization with only 120 samples across 3 topologies. "
    "The original GNN was trained on 240 samples across 6 topologies with REINFORCE fine-tuning. "
    "This is NOT a fair comparison of feature quality.",
    "The original GNN also had REINFORCE fine-tuning (LP-in-the-loop), which GNN+ did not receive. "
    "This stage can significantly improve MLU quality beyond oracle-label supervision.",
    "K_pred stuck at 39 suggests the k_head loss weight (0.01) may be too low, or that fixed "
    "K=40 is genuinely optimal for these topologies and traffic patterns.",
    "Only 3 topologies tested (2 known + 1 unseen). More generalization topologies would be "
    "needed for definitive conclusions.",
    "Disturbance improvement may simply be due to different learned score distributions, "
    "not necessarily better features.",
]
for item in honest_items:
    doc.add_paragraph(item, style="List Bullet")

# ==============================
# 10. Recommendation
# ==============================
doc.add_heading("10. Should GNN+ Be Inserted into MetaGate?", level=1)

p = doc.add_paragraph()
run = p.add_run("Verdict: NOT YET.")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph(
    "The screening shows that the combined GNN+ upgrade (richer features + dynamic K) does not "
    "improve MLU quality over the original GNN. The original GNN wins 63% of timesteps overall."
)
doc.add_paragraph("")
doc.add_paragraph("Before inserting GNN+ into MetaGate, the following steps are recommended:")

next_steps = [
    "Run a proper ablation: test features-only (fixed K=40) vs dynamic-K-only (original features) to isolate effects.",
    "Apply REINFORCE fine-tuning to GNN+ to match the original GNN training protocol.",
    "Train on all 6 known topologies (not just 3) for fair comparison.",
    "Investigate why K_pred converges to 39 and whether the k_head needs architectural changes.",
    "If features-only ablation shows improvement, integrate the enriched features into the original GNN without dynamic K.",
]
for i, step in enumerate(next_steps, 1):
    doc.add_paragraph(f"{step}", style="List Number")

# Fix zoom percent validation
import docx.oxml
for section in doc.sections:
    settings = doc.settings.element
    zoom_list = settings.findall(qn("w:zoom"))
    for z in zoom_list:
        if z.get(qn("w:percent")) is None:
            z.set(qn("w:percent"), "100")

# Save
doc.save(str(OUT))
print(f"Written: {OUT} ({OUT.stat().st_size} bytes)")

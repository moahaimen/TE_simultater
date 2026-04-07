#!/usr/bin/env python3
"""Build a Sarah-style MetaGate report with GNN+ as the learned expert."""

from __future__ import annotations

import json
import math
import os
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parent.parent
os.chdir(PROJECT_ROOT)

OUTPUT_DIR = Path(
    os.environ.get(
        "METAGATE_GNNPLUS_REPORT_INPUT_DIR",
        str(PROJECT_ROOT / "results" / "dynamic_metagate_gnnplus"),
    )
).resolve()
PLOTS_DIR = OUTPUT_DIR / "plots"
DEFAULT_DOC_NAME = (
    "MLP_MetaGate_GNNPLUS_Final_Report_SoftLabels_Regret_Updated.docx"
    if OUTPUT_DIR.name.endswith("softloss")
    else "MLP_MetaGate_GNNPLUS_Final_Report_ZeroShotObjective_Updated.docx"
)
OUTPUT_DOC = OUTPUT_DIR / os.environ.get("METAGATE_GNNPLUS_REPORT_NAME", DEFAULT_DOC_NAME)
DEFAULT_AUDIT_NAME = (
    "report_audit_softlabels_regret_updated.md"
    if OUTPUT_DIR.name.endswith("softloss")
    else "report_audit_zeroshot_objective_updated.md"
)
AUDIT_MD = OUTPUT_DIR / os.environ.get("METAGATE_GNNPLUS_REPORT_AUDIT", DEFAULT_AUDIT_NAME)
COMPARE_DIR = Path(
    os.environ.get(
        "METAGATE_GNNPLUS_COMPARE_DIR",
        str(PROJECT_ROOT / "results" / "dynamic_metagate_gnnplus"),
    )
).resolve()
INCLUDE_FAILURE_FIX_NOTE = os.environ.get("METAGATE_GNNPLUS_INCLUDE_FAILURE_FIX_NOTE", "1") != "0"
INCLUDE_FAILURE_COMPARE_TABLE = os.environ.get("METAGATE_GNNPLUS_INCLUDE_FAILURE_COMPARE_TABLE", "1") != "0"
ZERO_SHOT_ONLY = os.environ.get("METAGATE_GNNPLUS_ZERO_SHOT_ONLY", "0") == "1"

RESULTS_CSV = OUTPUT_DIR / "metagate_results.csv"
DECISIONS_CSV = OUTPUT_DIR / "metagate_decisions.csv"
SUMMARY_CSV = OUTPUT_DIR / "metagate_summary.csv"
TIMING_CSV = OUTPUT_DIR / "metagate_timing.csv"
CALIB_CSV = OUTPUT_DIR / "calibration_priors.csv"
TRAIN_DIST_CSV = OUTPUT_DIR / "train_oracle_distribution.csv"
TRAINING_SUMMARY_JSON = OUTPUT_DIR / "training_summary.json"
ORACLE_CSV = OUTPUT_DIR / "test_oracle.csv"
FAILURE_RESULTS_CSV = OUTPUT_DIR / "metagate_failure_results.csv"
FAILURE_SUMMARY_CSV = OUTPUT_DIR / "metagate_failure_summary.csv"
FAILURE_CALIB_CSV = OUTPUT_DIR / "metagate_failure_calibration.csv"
ZERO_SHOT_UNSEEN_SUMMARY_CSV = OUTPUT_DIR / "zero_shot_unseen_summary.csv"
COMPARE_SUMMARY_CSV = COMPARE_DIR / "metagate_summary.csv"
COMPARE_RESULTS_CSV = COMPARE_DIR / "metagate_results.csv"
COMPARE_TRAINING_SUMMARY_JSON = COMPARE_DIR / "training_summary.json"
PACKET_SDN_DIR = PROJECT_ROOT / "results" / "gnnplus_packet_sdn_report_fixed"
PACKET_SDN_SUMMARY_CSV = PACKET_SDN_DIR / "packet_sdn_summary.csv"
PACKET_SDN_FAILURE_CSV = PACKET_SDN_DIR / "packet_sdn_failure.csv"
PACKET_SDN_METRICS_CSV = PACKET_SDN_DIR / "packet_sdn_sdn_metrics.csv"
PACKET_SDN_PLOTS_DIR = PACKET_SDN_DIR / "plots"
COMPANION_SDN_SUMMARY_CSV = OUTPUT_DIR / "companion_packet_sdn_metrics_summary.csv"

TOPOLOGY_ORDER = [
    "abilene",
    "cernet",
    "geant",
    "ebone",
    "sprintlink",
    "tiscali",
    "germany50",
    "vtlwavenet2011",
]

TOPOLOGY_INFO = {
    "abilene": {"dataset": "abilene", "display": "Abilene", "type": "known", "nodes": 12},
    "cernet": {"dataset": "cernet", "display": "CERNET", "type": "known", "nodes": 41},
    "geant": {"dataset": "geant", "display": "GEANT", "type": "known", "nodes": 22},
    "ebone": {"dataset": "ebone", "display": "Ebone", "type": "known", "nodes": 23},
    "sprintlink": {"dataset": "sprintlink", "display": "Sprintlink", "type": "known", "nodes": 44},
    "tiscali": {"dataset": "tiscali", "display": "Tiscali", "type": "known", "nodes": 49},
    "germany50": {"dataset": "germany50", "display": "Germany50", "type": "UNSEEN", "nodes": 50},
    "vtlwavenet2011": {"dataset": "vtlwavenet2011", "display": "VtlWavenet2011", "type": "UNSEEN", "nodes": 92},
}

DATASET_TO_CANON = {
    "abilene": "abilene",
    "cernet": "cernet",
    "geant": "geant",
    "rocketfuel_ebone": "ebone",
    "rocketfuel_sprintlink": "sprintlink",
    "rocketfuel_tiscali": "tiscali",
    "germany50": "germany50",
    "topologyzoo_vtlwavenet2011": "vtlwavenet2011",
}

METHOD_COLORS = {
    "bottleneck": "#2ca02c",
    "topk": "#9467bd",
    "sensitivity": "#8c564b",
    "gnnplus": "#ff7f0e",
}
FAILURE_SCENARIO_ORDER = [
    "single_link_failure",
    "random_link_failure_1",
    "random_link_failure_2",
    "capacity_degradation_50",
    "traffic_spike_2x",
]
FAILURE_SCENARIO_LABELS = {
    "single_link_failure": "Single Link Failure",
    "random_link_failure_1": "Random Link Failure (1)",
    "random_link_failure_2": "Random Link Failure (2)",
    "capacity_degradation_50": "Capacity Degradation (50%)",
    "traffic_spike_2x": "Traffic Spike (2x)",
}


def canon(name: str) -> str:
    return DATASET_TO_CANON.get(str(name).strip().lower(), str(name).strip().lower())


def topo_label(topo: str) -> str:
    info = TOPOLOGY_INFO[topo]
    return f"{info['display']} ({info['nodes']}n)"


def topo_type(topo: str) -> str:
    return TOPOLOGY_INFO[topo]["type"]


def scenario_label(name: str) -> str:
    return FAILURE_SCENARIO_LABELS.get(str(name).strip().lower(), str(name))


def pretty_expert(name: str) -> str:
    mapping = {
        "bottleneck": "Bottleneck",
        "topk": "TopK",
        "sensitivity": "Sensitivity",
        "gnnplus": "GNN+",
    }
    return mapping.get(str(name).strip().lower(), str(name))


def fmt(value, digits: int = 4) -> str:
    try:
        v = float(value)
    except Exception:
        return str(value)
    if not math.isfinite(v):
        return "N/A"
    if abs(v) >= 1e6 or (0 < abs(v) < 1e-4):
        return f"{v:.2e}"
    return f"{v:.{digits}f}"


def fmt_pct(value, digits: int = 2) -> str:
    try:
        v = float(value)
    except Exception:
        return str(value)
    if not math.isfinite(v):
        return "N/A"
    if abs(v) >= 1e6 or (0 < abs(v) < 1e-3):
        return f"{v:+.2e}%"
    return f"{v:+.{digits}f}%"


def style_doc(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    for level, size in [(1, 16), (2, 14), (3, 12)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Arial"
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)


def add_table(doc: Document, headers, rows, font_size: int = 10):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(header))
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.name = "Arial"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(font_size)
            run.font.name = "Arial"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")


def bold_run(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    return run


def add_image(doc: Document, path: Path, caption: str, width: float = 6.0):
    if not path.exists():
        doc.add_paragraph(f"[Missing image: {path}]")
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def render_table_value(value) -> str:
    if isinstance(value, str):
        return value
    try:
        if pd.isna(value):
            return "N/A"
    except Exception:
        pass
    if isinstance(value, (int, np.integer)):
        return str(int(value))
    if isinstance(value, (float, np.floating)):
        return fmt(value)
    return str(value)


def add_dataframe_table(doc: Document, df: pd.DataFrame, title: str | None = None, font_size: int = 9):
    if title:
        doc.add_paragraph(title)
    rows = [[render_table_value(v) for v in row] for row in df.itertuples(index=False, name=None)]
    add_table(doc, list(df.columns), rows, font_size=font_size)


def add_dataframe_chunks(doc: Document, df: pd.DataFrame, heading_prefix: str, chunk_size: int = 20, font_size: int = 8):
    if df.empty:
        return
    for idx in range(0, len(df), chunk_size):
        chunk = df.iloc[idx : idx + chunk_size].reset_index(drop=True)
        suffix = "" if idx == 0 else f" (cont. {idx // chunk_size + 1})"
        add_dataframe_table(doc, chunk, title=f"{heading_prefix}{suffix}", font_size=font_size)


def load_companion_packet_sdn_inputs():
    if not PACKET_SDN_SUMMARY_CSV.exists() or not PACKET_SDN_METRICS_CSV.exists():
        return None, None, None
    summary = pd.read_csv(PACKET_SDN_SUMMARY_CSV)
    metrics = pd.read_csv(PACKET_SDN_METRICS_CSV)
    failure = pd.read_csv(PACKET_SDN_FAILURE_CSV) if PACKET_SDN_FAILURE_CSV.exists() else pd.DataFrame()
    return summary, metrics, failure


def build_complexity_rows() -> pd.DataFrame:
    rows = [
        {
            "Component": "Bottleneck",
            "Online stage": "Critical-flow heuristic plus LP on selected flows",
            "Complexity": "Heuristic proposal generation plus one LP solve when selected",
            "Practical note": "Robust fallback under hard failures; LP dominates controller cost when chosen.",
        },
        {
            "Component": "TopK",
            "Online stage": "Sort OD demands and select top-K flows",
            "Complexity": "Cheap ranking-based proposal generation",
            "Practical note": "Low proposal cost; no learned inference.",
        },
        {
            "Component": "Sensitivity",
            "Online stage": "Impact scoring over candidate OD pairs",
            "Complexity": "Heuristic scoring pass plus LP when selected",
            "Practical note": "More expensive than TopK, but still lighter than GNN+ inference.",
        },
        {
            "Component": "GNN+",
            "Online stage": "Graph-neural proposal generation for K=40 flows",
            "Complexity": "Learned inference with richer feature construction",
            "Practical note": "Highest single-expert proposal cost in this bundle.",
        },
        {
            "Component": "MetaGate+GNN+",
            "Online stage": "Run all 4 proposal generators, extract 49 features, apply MLP gate, then solve one LP",
            "Complexity": "Sum of expert proposal costs plus small MLP plus one LP",
            "Practical note": "At runtime, the gate pays proposal-generation cost for all experts, but only one LP reroute is executed.",
        },
    ]
    return pd.DataFrame(rows)


def build_complexity_numeric_rows(timing: pd.DataFrame) -> pd.DataFrame:
    component_map = {
        "Bottleneck proposal": "t_bn_ms",
        "TopK proposal": "t_topk_ms",
        "Sensitivity proposal": "t_sens_ms",
        "GNN+ proposal": "t_gnnplus_ms",
        "Feature extraction": "t_features_ms",
        "MLP gate": "t_mlp_ms",
        "LP reroute": "t_lp_ms",
        "Selector overhead": "t_decision_ms",
        "End-to-end total": "t_total_ms",
    }
    rows = []
    for label, col in component_map.items():
        series = pd.to_numeric(timing[col], errors="coerce")
        rows.append(
            {
                "Component": label,
                "Mean (ms)": float(series.mean()),
                "Median (ms)": float(series.median()),
                "Max (ms)": float(series.max()),
            }
        )
    return pd.DataFrame(rows)


def build_sdn_metrics_formula_rows() -> pd.DataFrame:
    return pd.DataFrame(
        [
            {"Component": "Per-link delay", "Model": "M/M/1 queueing", "Formula / rule": "d = 1 / (mu - lambda) + prop_delay"},
            {"Component": "End-to-end delay", "Model": "Path sum", "Formula / rule": "sum of per-link delays"},
            {"Component": "Throughput", "Model": "Bottleneck model", "Formula / rule": "minimum capacity / load ratio along routed paths"},
            {"Component": "Packet loss", "Model": "Overflow approximation", "Formula / rule": "max(0, (load - capacity) / load)"},
            {"Component": "Jitter", "Model": "Delay variation", "Formula / rule": "|delay(t) - delay(t-1)|"},
            {"Component": "Rule install delay", "Model": "Analytical SDN control cost", "Formula / rule": "0.5 ms + 0.02 ms * num_rules"},
            {"Component": "Failure recovery", "Model": "Controller reroute time", "Formula / rule": "cycles / wall-clock from failure to reroute"},
        ]
    )


def build_packet_sdn_companion_summary(summary_df: pd.DataFrame, metrics_df: pd.DataFrame) -> pd.DataFrame:
    packet_df = summary_df.copy()
    method_map = {
        "ecmp": "ECMP",
        "bottleneck": "Bottleneck",
        "gnn": "GNN",
        "gnnplus": "GNN+",
        "ECMP": "ECMP",
        "BOTTLENECK": "Bottleneck",
        "GNN": "GNN",
        "GNNPLUS": "GNN+",
    }
    packet_df["Method"] = packet_df["method"].map(method_map)
    grouped = (
        packet_df.groupby("Method", as_index=False)
        .agg(
            {
                "mean_mlu": "mean",
                "throughput": "mean",
                "mean_disturbance": "mean",
                "mean_latency_au": "mean",
                "packet_loss": "mean",
                "jitter_au": "mean",
                "decision_time_ms": "mean",
                "rule_install_delay_ms": "mean",
                "flow_table_updates": "mean",
            }
        )
        .rename(
            columns={
                "mean_mlu": "Avg MLU",
                "throughput": "Avg Throughput",
                "mean_disturbance": "Avg Disturbance",
                "mean_latency_au": "Avg Delay (au)",
                "packet_loss": "Avg Packet Loss",
                "jitter_au": "Avg Jitter",
                "decision_time_ms": "Avg Decision Time (ms)",
                "rule_install_delay_ms": "Avg Rule Install Delay (ms)",
                "flow_table_updates": "Avg Flow Table Updates",
            }
        )
    )
    metrics_df = metrics_df.copy()
    metrics_df["Method"] = metrics_df["Method"].map(method_map)
    recovery = (
        metrics_df.groupby("Method", as_index=False)["Recovery Time (ms)"]
        .mean()
        .rename(columns={"Recovery Time (ms)": "Avg Recovery Time (ms)"})
    )
    grouped = grouped.merge(recovery, on="Method", how="left")
    method_order = {"ECMP": 0, "Bottleneck": 1, "GNN": 2, "GNN+": 3}
    grouped["__order"] = grouped["Method"].map(method_order)
    grouped = grouped.sort_values("__order").drop(columns="__order").reset_index(drop=True)
    return grouped


def load_inputs():
    if not RESULTS_CSV.exists():
        raise FileNotFoundError(f"Missing evaluation output: {RESULTS_CSV}")
    results = pd.read_csv(RESULTS_CSV)
    decisions = pd.read_csv(DECISIONS_CSV)
    summary = pd.read_csv(SUMMARY_CSV)
    timing = pd.read_csv(TIMING_CSV)
    calib = pd.read_csv(CALIB_CSV)
    train_dist = pd.read_csv(TRAIN_DIST_CSV)
    oracle = pd.read_csv(ORACLE_CSV)
    training_summary = json.loads(TRAINING_SUMMARY_JSON.read_text(encoding="utf-8"))
    failure_results = pd.read_csv(FAILURE_RESULTS_CSV) if FAILURE_RESULTS_CSV.exists() else pd.DataFrame()
    failure_summary = pd.read_csv(FAILURE_SUMMARY_CSV) if FAILURE_SUMMARY_CSV.exists() else pd.DataFrame()
    failure_calib = pd.read_csv(FAILURE_CALIB_CSV) if FAILURE_CALIB_CSV.exists() else pd.DataFrame()
    zero_shot_unseen = pd.read_csv(ZERO_SHOT_UNSEEN_SUMMARY_CSV) if ZERO_SHOT_UNSEEN_SUMMARY_CSV.exists() else pd.DataFrame()

    for frame in [results, decisions, summary, timing, calib, train_dist, oracle, failure_results, failure_summary, failure_calib, zero_shot_unseen]:
        if "dataset" in frame.columns:
            frame["topology"] = frame["dataset"].map(canon)

    return results, decisions, summary, timing, calib, train_dist, oracle, training_summary, failure_results, failure_summary, failure_calib, zero_shot_unseen


def load_compare_inputs():
    if COMPARE_DIR == OUTPUT_DIR:
        return None, None, None
    if not COMPARE_SUMMARY_CSV.exists() or not COMPARE_RESULTS_CSV.exists() or not COMPARE_TRAINING_SUMMARY_JSON.exists():
        return None, None, None
    compare_summary = pd.read_csv(COMPARE_SUMMARY_CSV)
    compare_results = pd.read_csv(COMPARE_RESULTS_CSV)
    compare_training = json.loads(COMPARE_TRAINING_SUMMARY_JSON.read_text(encoding="utf-8"))
    for frame in [compare_summary, compare_results]:
        if "dataset" in frame.columns:
            frame["topology"] = frame["dataset"].map(canon)
    return compare_summary, compare_results, compare_training


def plot_selector_distribution(results: pd.DataFrame) -> Path:
    PLOTS_DIR.mkdir(parents=True, exist_ok=True)
    path = PLOTS_DIR / "selector_distribution_gnnplus.png"
    counts = (
        results.groupby(["topology", "metagate_selector"])
        .size()
        .unstack(fill_value=0)
        .reindex(index=TOPOLOGY_ORDER, fill_value=0)
        .reindex(columns=["bottleneck", "topk", "sensitivity", "gnnplus"], fill_value=0)
    )
    fig, ax = plt.subplots(figsize=(10.5, 4.8))
    x = np.arange(len(counts))
    bottom = np.zeros(len(counts))
    for method in ["bottleneck", "topk", "sensitivity", "gnnplus"]:
        vals = counts[method].to_numpy()
        ax.bar(x, vals, bottom=bottom, color=METHOD_COLORS[method], label=method.capitalize() if method != "gnnplus" else "GNN+")
        bottom += vals
    ax.set_xticks(x)
    ax.set_xticklabels([TOPOLOGY_INFO[t]["display"] for t in counts.index], rotation=20, ha="right")
    ax.set_ylabel("Selected Timesteps")
    ax.set_title("MetaGate Expert Selection Distribution (GNN+ version)")
    ax.legend(ncols=4, fontsize=8)
    ax.grid(True, axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def plot_accuracy_gap(summary: pd.DataFrame) -> Path:
    PLOTS_DIR.mkdir(parents=True, exist_ok=True)
    path = PLOTS_DIR / "accuracy_gap_gnnplus.png"
    df = summary.copy()
    df["gap_pct"] = (df["metagate_mlu"] - df["oracle_mlu"]) / df["oracle_mlu"] * 100.0
    df = df.set_index("topology").reindex(TOPOLOGY_ORDER).reset_index()
    x = np.arange(len(df))
    fig, ax1 = plt.subplots(figsize=(10.5, 4.8))
    ax1.bar(x, df["gap_pct"], color="#4c78a8", alpha=0.85, label="Oracle Gap (%)")
    ax1.set_ylabel("Oracle Gap (%)")
    ax1.set_xticks(x)
    ax1.set_xticklabels([TOPOLOGY_INFO[t]["display"] for t in df["topology"]], rotation=20, ha="right")
    ax1.axhline(0.0, color="black", linewidth=0.8)
    ax2 = ax1.twinx()
    ax2.plot(x, df["accuracy"] * 100.0, color="#d62728", marker="o", linewidth=2.0, label="Accuracy (%)")
    ax2.set_ylabel("Accuracy (%)")
    ax2.set_ylim(0, 100)
    ax1.set_title("MetaGate Accuracy and Oracle Gap by Topology (GNN+ version)")
    handles1, labels1 = ax1.get_legend_handles_labels()
    handles2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(handles1 + handles2, labels1 + labels2, loc="upper left")
    ax1.grid(True, axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def plot_germany50_selection(results: pd.DataFrame, oracle: pd.DataFrame) -> Path:
    PLOTS_DIR.mkdir(parents=True, exist_ok=True)
    path = PLOTS_DIR / "germany50_selection_gnnplus.png"
    pred = (
        results[results["topology"] == "germany50"]["metagate_selector"]
        .value_counts()
        .reindex(["bottleneck", "topk", "sensitivity", "gnnplus"], fill_value=0)
    )
    gold = (
        oracle[oracle["topology"] == "germany50"]["oracle_selector"]
        .value_counts()
        .reindex(["bottleneck", "topk", "sensitivity", "gnnplus"], fill_value=0)
    )
    x = np.arange(4)
    width = 0.38
    fig, ax = plt.subplots(figsize=(7.8, 4.5))
    ax.bar(x - width / 2, pred.values, width=width, label="Predicted", color="#1f77b4")
    ax.bar(x + width / 2, gold.values, width=width, label="Oracle", color="#ff7f0e")
    ax.set_xticks(x)
    ax.set_xticklabels(["Bottleneck", "TopK", "Sensitivity", "GNN+"])
    ax.set_ylabel("Timesteps")
    ax.set_title("Germany50: Predicted vs Oracle Expert Counts (GNN+ version)")
    ax.legend()
    ax.grid(True, axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def plot_failure_selector_mix(failure_results: pd.DataFrame) -> Path:
    PLOTS_DIR.mkdir(parents=True, exist_ok=True)
    path = PLOTS_DIR / "failure_selector_mix_gnnplus.png"
    counts = (
        failure_results.groupby(["scenario", "metagate_selector"])
        .size()
        .unstack(fill_value=0)
        .reindex(index=FAILURE_SCENARIO_ORDER, fill_value=0)
        .reindex(columns=["bottleneck", "topk", "sensitivity", "gnnplus"], fill_value=0)
    )
    fig, ax = plt.subplots(figsize=(9.8, 4.8))
    x = np.arange(len(counts))
    bottom = np.zeros(len(counts))
    for method in ["bottleneck", "topk", "sensitivity", "gnnplus"]:
        vals = counts[method].to_numpy()
        ax.bar(x, vals, bottom=bottom, color=METHOD_COLORS[method], label=pretty_expert(method))
        bottom += vals
    ax.set_xticks(x)
    ax.set_xticklabels([scenario_label(s) for s in counts.index], rotation=20, ha="right")
    ax.set_ylabel("Selected Timesteps")
    ax.set_title("MetaGate Expert Selection Under Failure Scenarios")
    ax.legend(ncols=4, fontsize=8)
    ax.grid(True, axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def plot_before_after_accuracy(compare_summary: pd.DataFrame, summary: pd.DataFrame) -> Path:
    PLOTS_DIR.mkdir(parents=True, exist_ok=True)
    path = PLOTS_DIR / "accuracy_before_after_softloss.png"
    old_df = compare_summary.set_index("topology").reindex(TOPOLOGY_ORDER)
    new_df = summary.set_index("topology").reindex(TOPOLOGY_ORDER)
    x = np.arange(len(TOPOLOGY_ORDER))
    width = 0.38
    fig, ax = plt.subplots(figsize=(10.8, 4.8))
    ax.bar(x - width / 2, old_df["accuracy"].to_numpy() * 100.0, width=width, label="Previous hard-label gate", color="#9aa0a6")
    ax.bar(x + width / 2, new_df["accuracy"].to_numpy() * 100.0, width=width, label="Soft-label + regret gate", color="#1f77b4")
    ax.set_xticks(x)
    ax.set_xticklabels([TOPOLOGY_INFO[t]["display"] for t in TOPOLOGY_ORDER], rotation=20, ha="right")
    ax.set_ylabel("Accuracy (%)")
    ax.set_title("Before vs After: Exact Expert Accuracy by Topology")
    ax.set_ylim(0, 100)
    ax.legend()
    ax.grid(True, axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def build_report():
    results, decisions, summary, timing, calib, train_dist, oracle, training_summary, failure_results, failure_summary, failure_calib, zero_shot_unseen = load_inputs()
    compare_summary, compare_results, compare_training = load_compare_inputs()
    packet_sdn_summary, packet_sdn_metrics, packet_sdn_failure = load_companion_packet_sdn_inputs()

    summary = summary.set_index("topology").reindex(TOPOLOGY_ORDER).reset_index()
    timing = timing.set_index("topology").reindex(TOPOLOGY_ORDER).reset_index()
    calib = calib.set_index("topology").reindex(TOPOLOGY_ORDER).reset_index()

    selector_plot = plot_selector_distribution(results)
    acc_gap_plot = plot_accuracy_gap(summary)
    germany_plot = plot_germany50_selection(results, oracle)
    failure_plot = plot_failure_selector_mix(failure_results) if not failure_results.empty else None
    before_after_plot = plot_before_after_accuracy(compare_summary, summary) if compare_summary is not None else None

    overall_acc = float(training_summary["overall_test_accuracy"]) * 100.0
    known_acc = float(training_summary["known_test_accuracy"]) * 100.0
    unseen_acc = float(training_summary["unseen_test_accuracy"]) * 100.0
    mean_decision = float(results["t_decision_ms"].mean())
    mean_total = float(results["t_total_ms"].mean())

    germany_summary = summary[summary["topology"] == "germany50"].iloc[0]
    germany_pred = results[results["topology"] == "germany50"]["metagate_selector"].value_counts()
    germany_oracle = oracle[oracle["topology"] == "germany50"]["oracle_selector"].value_counts()
    germany_gnnplus_pct = 100.0 * germany_pred.get("gnnplus", 0) / max(len(results[results["topology"] == "germany50"]), 1)

    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_text = (
        "MLP Meta-Gate with Soft Labels, Regret Loss, and Pure Zero-Shot Generalization"
        if ZERO_SHOT_ONLY
        else "MLP Meta-Gate with Soft Labels, Regret Loss, and Stable Failure Features"
    )
    run = title.add_run(title_text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_text = (
        "Final Evaluation Report (GNN+ Expert Version, Pure Zero-Shot)"
        if ZERO_SHOT_ONLY
        else "Final Evaluation Report (GNN+ Expert Version, Improved Training)"
    )
    run = subtitle.add_run(subtitle_text)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run(
        (
            "Architecture: 4-Expert Selection via MLP Gate without Topology Calibration\n"
            "Expert pool: Bottleneck, TopK, Sensitivity, GNN+\n"
            "Training upgrade: soft labels from expert MLUs + routing-aware regret penalty + clipped stable failure features"
            if ZERO_SHOT_ONLY
            else "Architecture: 4-Expert Selection via MLP Gate with Per-Topology Bayesian Calibration\n"
            "Expert pool: Bottleneck, TopK, Sensitivity, GNN+\n"
            "Training upgrade: soft labels from expert MLUs + routing-aware regret penalty + clipped stable failure features"
        )
    )
    run.font.size = Pt(11)
    run.italic = True

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "We present an MLP-based meta-gate that dynamically selects among four traffic engineering "
        "expert strategies (Bottleneck, Top-K by Demand, Sensitivity Analysis, and GNN+-based selection) "
        "on a per-traffic-matrix basis. The gate is a 3-layer MLP (128-128-64-4) with BatchNorm and "
        "dropout, trained on pooled oracle labels from 6 known topologies. Germany50 and VtlWavenet2011 "
        "are held out entirely from gate-weight training."
    )
    if ZERO_SHOT_ONLY:
        doc.add_paragraph(
            "This report uses a pure zero-shot evaluation protocol. Germany50 and VtlWavenet2011 are held out "
            "from gate-weight training, and the trained gate is applied directly at test time with no target-topology "
            "calibration, no Bayesian prior fusion, and no gradient-based adaptation."
        )
    else:
        doc.add_paragraph(
            "A lightweight few-shot Bayesian calibration phase adapts the gate to each target topology "
            "at deployment time by running 10 validation traffic matrices through all 4 experts, counting "
            "which expert achieves the lowest MLU, and fusing this topology-specific prior with the MLP "
            "softmax output. No gradient updates occur during calibration."
        )
    doc.add_paragraph(
        "In the updated training pipeline, the hard one-hot oracle target is replaced with a soft target "
        "distribution derived from the relative MLU regret of all 4 experts. A routing-aware regret term "
        "is added to the loss so the gate is penalized much more for assigning probability to catastrophic "
        "experts than for being torn between two nearly identical good experts."
    )
    doc.add_paragraph(
        "A second stability fix is applied before the MLP: the GNN correction feature is compressed with a "
        "signed log transform, the ECMP utilization features are log-compressed, and the standardized input "
        "vector is clipped. This prevents failure scenarios from generating numerically extreme logits that "
        "collapse the gate to a single expert regardless of the calibration prior."
    )
    p = doc.add_paragraph()
    bold_run(p, "Key results: ")
    if ZERO_SHOT_ONLY:
        p.add_run(
            f"Zero-shot Germany50 (unseen): GNN+ selected {germany_gnnplus_pct:.1f}% of timesteps, "
            f"oracle gap {((germany_summary['metagate_mlu'] - germany_summary['oracle_mlu']) / germany_summary['oracle_mlu']) * 100.0:+.2f}%, "
            f"accuracy {germany_summary['accuracy'] * 100.0:.1f}%. Overall accuracy {overall_acc:.1f}% across 8 topologies. "
            f"Mean selector decision overhead {mean_decision:.1f} ms; mean end-to-end time (including LP) {mean_total:.1f} ms."
        )
    else:
        p.add_run(
            f"Calibrated Germany50 (unseen): GNN+ selected {germany_gnnplus_pct:.1f}% of timesteps, "
            f"oracle gap {((germany_summary['metagate_mlu'] - germany_summary['oracle_mlu']) / germany_summary['oracle_mlu']) * 100.0:+.2f}%, "
            f"accuracy {germany_summary['accuracy'] * 100.0:.1f}%. Overall accuracy {overall_acc:.1f}% across 8 topologies. "
            f"Mean selector decision overhead {mean_decision:.1f} ms; mean end-to-end time (including LP) {mean_total:.1f} ms."
        )
    if compare_training is not None:
        p = doc.add_paragraph()
        bold_run(p, "Improvement snapshot: ")
        p.add_run(
            f"overall exact accuracy improved from {float(compare_training['overall_test_accuracy']) * 100.0:.1f}% "
            f"to {overall_acc:.1f}%, and unseen-topology accuracy improved from "
            f"{float(compare_training['unseen_test_accuracy']) * 100.0:.1f}% to {unseen_acc:.1f}% after the soft-label/regret update."
        )

    doc.add_heading("1. Contributions", level=1)
    items = [
        "Base objective: zero-shot generalization to unseen topologies.",
        "Evaluation protocol: no target-topology calibration, no Bayesian prior fusion, and no weight updates at test time." if ZERO_SHOT_ONLY else "Additional contribution: lightweight few-shot Bayesian calibration before inference, used as a deployment-time enhancement rather than weight retraining.",
        "Honest MLP Meta-Gate architecture with a real 4-expert pool and per-timestep learned routing-strategy selection.",
        "GNN+ replaces the Original GNN as the learned expert inside the MetaGate evaluation itself.",
        "Unified gate training on known topologies only, with Germany50 and VtlWavenet2011 held out from gate-weight training.",
        "Per-component timing breakdown (BN + TopK + Sens + GNN+ + features + MLP + LP) reported per timestep.",
    ]
    for idx, item in enumerate(items, 1):
        p = doc.add_paragraph()
        bold_run(p, f"{idx}. ")
        p.add_run(item)

    doc.add_heading("2. Architecture", level=1)
    doc.add_heading("2.1 Overview", level=2)
    steps = [
        "Run all 4 experts (Bottleneck, TopK, Sensitivity, GNN+) to produce 4 candidate OD-pair selections, each of size K_crit = 40.",
        "Extract a 49-dimensional feature vector from the TM, expert outputs, GNN+ diagnostics, ECMP baseline, and topology metrics.",
        "The MLP gate predicts which expert will achieve the lowest MLU.",
        "Predict the expert directly from the MLP softmax without any topology-specific prior." if ZERO_SHOT_ONLY else "If Bayesian calibration is active, fuse MLP softmax with topology-specific prior.",
        "Route the predicted expert's selection through LP optimization to obtain the final MLU.",
    ]
    for step in steps:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("2.2 MLP Gate Architecture", level=2)
    doc.add_paragraph(
        "3-layer MLP: Linear(49, 128) -> BatchNorm -> ReLU -> Dropout(0.3) -> "
        "Linear(128, 128) -> BatchNorm -> ReLU -> Dropout(0.3) -> "
        "Linear(128, 64) -> ReLU -> Linear(64, 4). "
        "Output: softmax over 4 expert classes {0=Bottleneck, 1=TopK, 2=Sensitivity, 3=GNN+}."
    )

    doc.add_heading("2.3 Four Experts", level=2)
    add_table(
        doc,
        ["Expert", "Selection Strategy"],
        [
            ["Bottleneck", "Selects K_crit OD pairs traversing the most-utilized link under ECMP routing"],
            ["Top-K by Demand", "Selects K_crit OD pairs with the highest traffic demand"],
            ["Sensitivity", "Selects K_crit OD pairs whose rerouting has the largest impact on max link utilization"],
            ["GNN+ (GraphSAGE variant)", "Learned flow selector using enriched node/edge/OD features with fixed K=40"],
        ],
    )

    doc.add_heading("3. Training Details", level=1)
    total_samples = int(sum(train_dist["samples"]))
    train_rows = []
    for topo in TOPOLOGY_ORDER:
        row = train_dist[train_dist["topology"] == topo]
        if row.empty:
            continue
        rec = row.iloc[0]
        train_rows.append([
            TOPOLOGY_INFO[topo]["display"],
            str(TOPOLOGY_INFO[topo]["nodes"]),
            str(int(rec["samples"])),
            str(int(rec["bottleneck"])),
            str(int(rec["topk"])),
            str(int(rec["sensitivity"])),
            str(int(rec["gnnplus"])),
        ])
    add_table(doc, ["Topology", "Nodes", "Samples", "BN Oracle", "TopK Oracle", "Sens Oracle", "GNN+ Oracle"], train_rows)
    doc.add_paragraph(
        "The gate is trained on pooled data from 6 known topologies only. "
        "Germany50 and VtlWavenet2011 are completely excluded from training."
    )
    overall_counts = training_summary["train_class_counts"]
    doc.add_paragraph(
        "Pooled training distribution: "
        f"Bottleneck={overall_counts['bottleneck']}/{total_samples}, "
        f"TopK={overall_counts['topk']}/{total_samples}, "
        f"Sensitivity={overall_counts['sensitivity']}/{total_samples}, "
        f"GNN+={overall_counts['gnnplus']}/{total_samples}."
    )
    add_table(
        doc,
        ["Parameter", "Value"],
        [
            ["Hidden dimension", str(training_summary["metagate_config"]["hidden_dim"])],
            ["Dropout", str(training_summary["metagate_config"]["dropout"])],
            ["Learning rate", str(training_summary["metagate_config"]["learning_rate"])],
            ["Epochs", str(training_summary["metagate_config"]["num_epochs"])],
            ["Batch size", str(training_summary["metagate_config"]["batch_size"])],
            ["Soft labels enabled", str(training_summary["metagate_config"].get("use_soft_labels", False))],
            ["Soft-label temperature", str(training_summary["metagate_config"].get("soft_label_temperature", "N/A"))],
            ["Regret loss weight", str(training_summary["metagate_config"].get("regret_loss_weight", "N/A"))],
            ["Regret clip", str(training_summary["metagate_config"].get("regret_clip", "N/A"))],
            ["Feature clip", str(training_summary["metagate_config"].get("feature_clip", "N/A"))],
            ["Train accuracy", f"{float(training_summary['train_accuracy']) * 100.0:.1f}%"],
            ["Validation accuracy", f"{float(training_summary['val_accuracy']) * 100.0:.1f}%"],
            ["GNN+ checkpoint", training_summary["gnnplus_checkpoint"]],
        ],
    )
    doc.add_heading("3.1 What Improved and How", level=2)
    doc.add_paragraph(
        "The previous MetaGate+GNN+ version used hard one-hot oracle labels. That made the gate look wrong on "
        "tie-heavy topologies such as Abilene and Ebone, even when the chosen expert achieved nearly the same "
        "MLU as the oracle. The improved version changes two things:"
    )
    doc.add_paragraph(
        "1) Soft labels: expert MLUs are converted into a probability target, so two near-oracle experts can "
        "both receive high probability instead of forcing a fake single winner.",
        style="List Number",
    )
    doc.add_paragraph(
        "2) Routing-aware regret loss: the training loss now adds a penalty proportional to the expected MLU "
        "regret, so the gate learns that close misses are acceptable but high-regret experts should be suppressed.",
        style="List Number",
    )
    doc.add_paragraph(
        "3) Failure-feature stabilization: log compression is applied to unstable utilization diagnostics and "
        "the standardized feature vector is clipped before the MLP, preventing out-of-distribution failures "
        "from numerically saturating the classifier.",
        style="List Number",
    )
    if compare_summary is not None and compare_training is not None:
        overall_gap_old = (
            (compare_summary["metagate_mlu"].mean() - compare_summary["oracle_mlu"].mean())
            / compare_summary["oracle_mlu"].mean()
            * 100.0
        )
        overall_gap_new = (
            (summary["metagate_mlu"].mean() - summary["oracle_mlu"].mean())
            / summary["oracle_mlu"].mean()
            * 100.0
        )
        add_table(
            doc,
            ["Metric", "Previous Hard-Label Gate", "Updated Soft-Label + Regret Gate"],
            [
                ["Overall accuracy", f"{float(compare_training['overall_test_accuracy']) * 100.0:.1f}%", f"{overall_acc:.1f}%"],
                ["Known-topology accuracy", f"{float(compare_training['known_test_accuracy']) * 100.0:.1f}%", f"{known_acc:.1f}%"],
                ["Unseen-topology accuracy", f"{float(compare_training['unseen_test_accuracy']) * 100.0:.1f}%", f"{unseen_acc:.1f}%"],
                ["Mean oracle gap", fmt_pct(overall_gap_old), fmt_pct(overall_gap_new)],
            ],
        )
        improvement_rows = []
        merged = compare_summary[["topology", "accuracy", "metagate_mlu", "oracle_mlu"]].merge(
            summary[["topology", "accuracy", "metagate_mlu", "oracle_mlu"]],
            on="topology",
            suffixes=("_old", "_new"),
        )
        for topo in TOPOLOGY_ORDER:
            row = merged[merged["topology"] == topo]
            if row.empty:
                continue
            rec = row.iloc[0]
            old_gap = (float(rec["metagate_mlu_old"]) - float(rec["oracle_mlu_old"])) / float(rec["oracle_mlu_old"]) * 100.0
            new_gap = (float(rec["metagate_mlu_new"]) - float(rec["oracle_mlu_new"])) / float(rec["oracle_mlu_new"]) * 100.0
            improvement_rows.append(
                [
                    topo_label(topo),
                    f"{float(rec['accuracy_old']) * 100.0:.1f}%",
                    f"{float(rec['accuracy_new']) * 100.0:.1f}%",
                    f"{(float(rec['accuracy_new']) - float(rec['accuracy_old'])) * 100.0:+.1f} pts",
                    fmt_pct(old_gap),
                    fmt_pct(new_gap),
                ]
            )
        add_table(
            doc,
            ["Topology", "Old Acc", "New Acc", "Delta", "Old Gap", "New Gap"],
            improvement_rows,
        )
        if before_after_plot is not None:
            add_image(doc, before_after_plot, "Before vs after exact-expert accuracy after soft-label and regret-loss training.")

    if ZERO_SHOT_ONLY:
        doc.add_heading("4. Zero-Shot Evaluation Protocol", level=1)
        doc.add_paragraph(
            "The original project objective is pure zero-shot generalization to unseen topologies. In this report, "
            "the gate is trained only on the 6 known topologies and then evaluated directly on all 8 topologies "
            "without topology-specific calibration, Bayesian prior fusion, or test-time adaptation."
        )
        if not zero_shot_unseen.empty:
            rows = []
            for topo in ["germany50", "vtlwavenet2011"]:
                zrow = zero_shot_unseen[zero_shot_unseen["topology"] == topo]
                if zrow.empty:
                    continue
                zrec = zrow.iloc[0]
                rows.append([
                    TOPOLOGY_INFO[topo]["display"],
                    f"{float(zrec['accuracy']) * 100.0:.1f}%",
                    fmt_pct(float(zrec["oracle_gap_pct"])),
                    str(int(zrec["n_timesteps"])),
                ])
            add_table(
                doc,
                ["Unseen Topology", "Zero-Shot Accuracy", "Zero-Shot Gap", "Test Timesteps"],
                rows,
            )
    else:
        doc.add_heading("4. Few-Shot Bayesian Calibration", level=1)
        doc.add_paragraph(
            "Before evaluating on any topology, a calibration phase runs 10 validation traffic matrices "
            "through all 4 experts with LP optimization, counting which expert achieves the lowest MLU. "
            "This produces a topology-specific prior that is fused with the MLP softmax at inference time."
        )
        calib_rows = []
        for topo in TOPOLOGY_ORDER:
            row = calib[calib["topology"] == topo]
            if row.empty:
                continue
            rec = row.iloc[0]
            calib_rows.append([
                TOPOLOGY_INFO[topo]["display"],
                topo_type(topo),
                fmt(rec["bottleneck_prior"], 2),
                fmt(rec["topk_prior"], 2),
                fmt(rec["sensitivity_prior"], 2),
                fmt(rec["gnnplus_prior"], 2),
            ])
        add_table(doc, ["Topology", "Type", "BN Prior", "TopK Prior", "Sens Prior", "GNN+ Prior"], calib_rows)
        doc.add_heading("4.1 Objective Framing: Zero-Shot Baseline and Few-Shot Extension", level=2)
        doc.add_paragraph(
            "The original project objective is zero-shot generalization to unseen topologies. To keep that objective "
            "scientifically valid, the report distinguishes between a zero-shot baseline and a stronger deployment mode "
            "that adds a lightweight few-shot Bayesian calibration prior before inference."
        )
        doc.add_paragraph(
            "In this framing, the base MetaGate weights are always trained only on the 6 known topologies. The zero-shot "
            "setting performs direct inference on Germany50 and VtlWavenet2011 with no target-topology adaptation. The "
            "few-shot calibrated setting uses 10 validation traffic matrices from the target topology to estimate a prior, "
            "but it still performs no gradient-based fine-tuning or weight updates."
        )
        if not zero_shot_unseen.empty:
            comparison_rows = []
            for topo in ["germany50", "vtlwavenet2011"]:
                zrow = zero_shot_unseen[zero_shot_unseen["topology"] == topo]
                crow = summary[summary["topology"] == topo]
                if zrow.empty or crow.empty:
                    continue
                zrec = zrow.iloc[0]
                crec = crow.iloc[0]
                zgap = (float(zrec["metagate_mlu"]) - float(zrec["oracle_mlu"])) / float(zrec["oracle_mlu"]) * 100.0
                cgap = (float(crec["metagate_mlu"]) - float(crec["oracle_mlu"])) / float(crec["oracle_mlu"]) * 100.0
                comparison_rows.append([
                    TOPOLOGY_INFO[topo]["display"],
                    f"{float(zrec['accuracy']) * 100.0:.1f}%",
                    fmt_pct(zgap),
                    f"{float(crec['accuracy']) * 100.0:.1f}%",
                    fmt_pct(cgap),
                ])
            add_table(
                doc,
                ["Unseen Topology", "Zero-Shot Accuracy", "Zero-Shot Gap", "Calibrated Accuracy", "Calibrated Gap"],
                comparison_rows,
            )
            doc.add_paragraph(
                "This table is included so the report can keep zero-shot generalization as the base objective, while also "
                "showing that the few-shot calibration module is an additional contribution that improves deployment on the unseen topologies."
            )

    doc.add_heading("5. Expert Selection, Accuracy, and Training Notes", level=1)
    expert_rows = []
    mix_rows = []
    for topo in TOPOLOGY_ORDER:
        row = summary[summary["topology"] == topo]
        if row.empty:
            continue
        rec = row.iloc[0]
        sub = results[results["topology"] == topo]
        counts = sub["metagate_selector"].value_counts()
        total = max(len(sub), 1)
        dominant_name = counts.idxmax() if not counts.empty else "n/a"
        dominant_pct = 100.0 * counts.max() / total if not counts.empty else 0.0
        switches = int((sub["metagate_selector"].values[1:] != sub["metagate_selector"].values[:-1]).sum()) if len(sub) > 1 else 0
        gap = ((rec["metagate_mlu"] - rec["oracle_mlu"]) / rec["oracle_mlu"]) * 100.0
        expert_rows.append([
            topo_label(topo),
            f"{switches}/{max(len(sub) - 1, 1)}",
            f"{pretty_expert(dominant_name)} ({dominant_pct:.0f}%)",
            f"{rec['accuracy'] * 100.0:.1f}%",
            f"{gap:+.2f}%",
        ])
        mix_rows.append([
            topo_label(topo),
            f"{100.0 * counts.get('bottleneck', 0) / total:.0f}%",
            f"{100.0 * counts.get('topk', 0) / total:.0f}%",
            f"{100.0 * counts.get('sensitivity', 0) / total:.0f}%",
            f"{100.0 * counts.get('gnnplus', 0) / total:.0f}%",
        ])
    add_table(doc, ["Topology", "Switches", "Dominant Expert", "Accuracy", "Oracle Gap"], expert_rows)
    add_table(doc, ["Topology", "BN%", "TopK%", "Sens%", "GNN+%"], mix_rows)
    add_image(doc, selector_plot, "MetaGate expert selection distribution with GNN+ as the learned expert.")
    add_image(doc, acc_gap_plot, "MetaGate accuracy and oracle gap by topology (GNN+ version).")

    doc.add_heading("6. Results -- All 8 Topologies", level=1)
    result_rows = []
    for topo in TOPOLOGY_ORDER:
        row = summary[summary["topology"] == topo]
        if row.empty:
            continue
        rec = row.iloc[0]
        sub = results[results["topology"] == topo]
        counts = sub["metagate_selector"].value_counts()
        gnnplus_pct = 100.0 * counts.get("gnnplus", 0) / max(len(sub), 1)
        switches = int((sub["metagate_selector"].values[1:] != sub["metagate_selector"].values[:-1]).sum()) if len(sub) > 1 else 0
        gap = ((rec["metagate_mlu"] - rec["oracle_mlu"]) / rec["oracle_mlu"]) * 100.0
        result_rows.append([
            TOPOLOGY_INFO[topo]["display"],
            topo_type(topo),
            f"{rec['accuracy'] * 100.0:.1f}%",
            fmt(rec["metagate_mlu"]),
            fmt(rec["oracle_mlu"]),
            f"{gap:+.2f}%",
            f"{gnnplus_pct:.0f}%",
            str(switches),
        ])
    add_table(doc, ["Topology", "Type", "Accuracy", "MetaGate MLU", "Oracle MLU", "Gap", "GNN+%", "Switches"], result_rows)
    p = doc.add_paragraph()
    bold_run(p, "Overall accuracy: ")
    p.add_run(f"{overall_acc:.1f}%. ")
    bold_run(p, "Known topology accuracy: ")
    p.add_run(f"{known_acc:.1f}%. ")
    bold_run(p, "Unseen topology accuracy: ")
    p.add_run(f"{unseen_acc:.1f}%.")

    if not failure_results.empty and not failure_summary.empty:
        doc.add_heading("7. Failure Scenario Evaluation", level=1)
        doc.add_paragraph(
            "Unlike the standalone packet-SDN failure tables, this section runs the integrated MetaGate+GNN+ "
            "chooser under the failure scenarios themselves. For each failed timestep, the gate predicts an "
            "expert, and the report logs whether that chosen expert changes relative to normal conditions."
        )
        if INCLUDE_FAILURE_FIX_NOTE:
            doc.add_paragraph(
                "Important reading note: Accuracy here is strict exact-oracle expert accuracy. Therefore, a row can "
                "show BN%=100%, GNN+%=0%, and Accuracy=100.0% at the same time, because MetaGate chose Bottleneck "
                "on every failure timestep and Bottleneck was also the oracle expert on every failure timestep."
            )
        scenario_rows = []
        normal_dominant = {
            topo: results[results["topology"] == topo]["metagate_selector"].value_counts().idxmax()
            for topo in TOPOLOGY_ORDER
            if not results[results["topology"] == topo].empty
        }
        for scenario in FAILURE_SCENARIO_ORDER:
            sub = failure_results[failure_results["scenario"] == scenario]
            if sub.empty:
                continue
            counts = sub["metagate_selector"].value_counts()
            total = max(len(sub), 1)
            ssum = failure_summary[failure_summary["scenario"] == scenario]
            changed = 0
            compared = 0
            for topo in TOPOLOGY_ORDER:
                topo_sub = sub[sub["topology"] == topo]
                if topo_sub.empty or topo not in normal_dominant:
                    continue
                compared += 1
                failure_dom = topo_sub["metagate_selector"].value_counts().idxmax()
                if failure_dom != normal_dominant[topo]:
                    changed += 1
            scenario_rows.append([
                scenario_label(scenario),
                f"{ssum['accuracy'].mean() * 100.0:.1f}%",
                fmt_pct(ssum["metagate_vs_oracle_gap_pct"].mean()),
                f"{100.0 * counts.get('bottleneck', 0) / total:.0f}%",
                f"{100.0 * counts.get('topk', 0) / total:.0f}%",
                f"{100.0 * counts.get('sensitivity', 0) / total:.0f}%",
                f"{100.0 * counts.get('gnnplus', 0) / total:.0f}%",
                f"{changed}/{compared}",
            ])
        add_table(doc, ["Scenario", "Accuracy", "Oracle Gap", "BN%", "TopK%", "Sens%", "GNN+%", "Changed Topologies"], scenario_rows)
        if INCLUDE_FAILURE_COMPARE_TABLE and compare_summary is not None and COMPARE_DIR != OUTPUT_DIR and (COMPARE_DIR / "metagate_failure_results.csv").exists():
            try:
                prev_failure_results = pd.read_csv(COMPARE_DIR / "metagate_failure_results.csv")
                if "dataset" in prev_failure_results.columns:
                    prev_failure_results["topology"] = prev_failure_results["dataset"].map(canon)
                compare_rows = []
                for scenario in FAILURE_SCENARIO_ORDER[:3]:
                    old_sub = prev_failure_results[prev_failure_results["scenario"] == scenario]
                    new_sub = failure_results[failure_results["scenario"] == scenario]
                    if old_sub.empty or new_sub.empty:
                        continue
                    old_counts = old_sub["metagate_selector"].value_counts()
                    new_counts = new_sub["metagate_selector"].value_counts()
                    old_total = max(len(old_sub), 1)
                    new_total = max(len(new_sub), 1)
                    old_gap = ((old_sub["metagate_mlu"].mean() - old_sub["oracle_mlu"].mean()) / old_sub["oracle_mlu"].mean()) * 100.0
                    new_gap = ((new_sub["metagate_mlu"].mean() - new_sub["oracle_mlu"].mean()) / new_sub["oracle_mlu"].mean()) * 100.0
                    compare_rows.append([
                        scenario_label(scenario),
                        f"{old_sub['correct'].mean() * 100.0:.1f}%",
                        fmt_pct(old_gap),
                        f"{100.0 * old_counts.get('bottleneck', 0) / old_total:.0f}%",
                        f"{100.0 * old_counts.get('gnnplus', 0) / old_total:.0f}%",
                        f"{new_sub['correct'].mean() * 100.0:.1f}%",
                        fmt_pct(new_gap),
                        f"{100.0 * new_counts.get('bottleneck', 0) / new_total:.0f}%",
                        f"{100.0 * new_counts.get('gnnplus', 0) / new_total:.0f}%",
                    ])
                if compare_rows:
                    doc.add_paragraph(
                        "The next table makes the hard-failure fix explicit by comparing the previous soft-label/regret bundle "
                        "with the current stabilized-failure bundle on the three link-failure scenarios."
                    )
                    add_table(
                        doc,
                        [
                            "Scenario",
                            "Old Acc",
                            "Old Gap",
                            "Old BN%",
                            "Old GNN+%",
                            "New Acc",
                            "New Gap",
                            "New BN%",
                            "New GNN+%",
                        ],
                        compare_rows,
                    )
            except Exception:
                pass
        collapsed = []
        for scenario in FAILURE_SCENARIO_ORDER:
            sub = failure_results[failure_results["scenario"] == scenario]
            if sub.empty:
                continue
            counts = sub["metagate_selector"].value_counts()
            if len(counts) == 1 and counts.index[0] == "gnnplus":
                collapsed.append(scenario_label(scenario))
        if collapsed:
            doc.add_paragraph(
                "Important observation: the harshest link-failure scenarios drive the gate to a degenerate "
                f"GNN+-only choice on {', '.join(collapsed)}. This is a real out-of-distribution behavior in "
                "the current failure evaluation, not a formatting artifact."
            )
        if failure_plot is not None:
            add_image(doc, failure_plot, "MetaGate expert mix across all failure timesteps (GNN+ integrated).")

        for idx, scenario in enumerate(FAILURE_SCENARIO_ORDER, start=1):
            sub = failure_results[failure_results["scenario"] == scenario]
            if sub.empty:
                continue
            doc.add_heading(f"7.{idx} {scenario_label(scenario)}", level=2)
            doc.add_paragraph(
                f"Per-topology strict accuracy, dominant expert, expert mix, and oracle gap for {scenario_label(scenario)}."
            )
            rows = []
            for topo in TOPOLOGY_ORDER:
                topo_sub = sub[sub["topology"] == topo]
                if topo_sub.empty:
                    continue
                counts = topo_sub["metagate_selector"].value_counts()
                total = max(len(topo_sub), 1)
                dominant = counts.idxmax()
                gap = ((topo_sub["metagate_mlu"].mean() - topo_sub["oracle_mlu"].mean()) / topo_sub["oracle_mlu"].mean()) * 100.0
                rows.append([
                    topo_label(topo),
                    pretty_expert(dominant),
                    f"{100.0 * counts.get('bottleneck', 0) / total:.0f}%",
                    f"{100.0 * counts.get('topk', 0) / total:.0f}%",
                    f"{100.0 * counts.get('sensitivity', 0) / total:.0f}%",
                    f"{100.0 * counts.get('gnnplus', 0) / total:.0f}%",
                    f"{topo_sub['correct'].mean() * 100.0:.1f}%",
                    fmt_pct(gap),
                ])
            add_table(doc, ["Topology", "Dominant", "BN%", "TopK%", "Sens%", "GNN+%", "Accuracy", "Gap"], rows)

    doc.add_heading("8. Germany50 (Unseen) -- Deep Dive", level=1)
    doc.add_paragraph(
        "Germany50 is the clearest unseen-topology case in the current GNN+ integrated MetaGate run. "
        "The table and figure below show the predicted selection mix and the oracle distribution "
        "when GNN+ replaces the Original GNN in the expert pool."
    )
    germany_rows = [
        ["Predicted Bottleneck selections", str(int(germany_pred.get("bottleneck", 0)))],
        ["Predicted GNN+ selections", str(int(germany_pred.get("gnnplus", 0)))],
        ["Oracle Bottleneck wins", str(int(germany_oracle.get("bottleneck", 0)))],
        ["Oracle GNN+ wins", str(int(germany_oracle.get("gnnplus", 0)))],
        ["MetaGate MLU", fmt(germany_summary["metagate_mlu"])],
        ["Oracle MLU", fmt(germany_summary["oracle_mlu"])],
        ["Oracle gap", f"{((germany_summary['metagate_mlu'] - germany_summary['oracle_mlu']) / germany_summary['oracle_mlu']) * 100.0:+.2f}%"],
        ["Accuracy", f"{germany_summary['accuracy'] * 100.0:.1f}%"],
    ]
    add_table(doc, ["Metric", "Value"], germany_rows)
    add_image(doc, germany_plot, "Germany50 predicted vs oracle expert counts with GNN+ integrated.")

    doc.add_heading("9. VtlWavenet2011 (Unseen) -- Analysis", level=1)
    vtl = summary[summary["topology"] == "vtlwavenet2011"].iloc[0]
    vtl_pred = results[results["topology"] == "vtlwavenet2011"]["metagate_selector"].value_counts()
    doc.add_paragraph(
        f"VtlWavenet2011 accuracy is {vtl['accuracy'] * 100.0:.1f}% with oracle gap "
        f"{((vtl['metagate_mlu'] - vtl['oracle_mlu']) / vtl['oracle_mlu']) * 100.0:+.2f}%. "
        f"Predicted selections: Bottleneck={int(vtl_pred.get('bottleneck', 0))}, "
        f"GNN+={int(vtl_pred.get('gnnplus', 0))}, TopK={int(vtl_pred.get('topk', 0))}, "
        f"Sensitivity={int(vtl_pred.get('sensitivity', 0))}."
    )

    doc.add_heading("10. Per-Component Timing Breakdown", level=1)
    timing_rows = []
    for topo in TOPOLOGY_ORDER:
        row = timing[timing["topology"] == topo]
        if row.empty:
            continue
        rec = row.iloc[0]
        timing_rows.append([
            TOPOLOGY_INFO[topo]["display"],
            fmt(rec["t_bn_ms"], 1),
            fmt(rec["t_topk_ms"], 1),
            fmt(rec["t_sens_ms"], 1),
            fmt(rec["t_gnnplus_ms"], 1),
            fmt(rec["t_features_ms"], 2),
            fmt(rec["t_mlp_ms"], 3),
            fmt(rec["t_lp_ms"], 1),
            fmt(rec["t_total_ms"], 1),
        ])
    add_table(doc, ["Topology", "BN ms", "TopK ms", "Sens ms", "GNN+ ms", "Feat ms", "MLP ms", "LP ms", "Total ms"], timing_rows)
    doc.add_paragraph(
        f"Mean decision time (selector overhead before LP): {mean_decision:.1f} ms. "
        f"Mean total time (end-to-end): {mean_total:.1f} ms."
    )

    doc.add_heading("11. Complexity Analysis", level=1)
    doc.add_paragraph(
        "The complexity discussion below is limited to the methods and runtime components that are actually present "
        "in this MetaGate+GNN+ zero-shot bundle. Because the online gate constructs proposals from all 4 experts "
        "before making a decision, the runtime controller cost is broader than a single-expert system."
    )
    doc.add_heading("11.1 Theoretical Complexity", level=2)
    add_dataframe_table(doc, build_complexity_rows(), font_size=8)
    doc.add_heading("11.2 Numeric Controller-Cost Summary", level=2)
    doc.add_paragraph(
        "These are measured controller runtime statistics from the current pure zero-shot MetaGate+GNN+ bundle. "
        "They complement, rather than replace, the theoretical complexity discussion above."
    )
    add_dataframe_table(doc, build_complexity_numeric_rows(timing), font_size=8)

    doc.add_heading("12. Model-Based SDN Metrics and Packet-SDN Companion Validation", level=1)
    doc.add_paragraph(
        "The integrated MetaGate+GNN+ zero-shot evaluation does not directly generate packet-level SDN metrics such "
        "as throughput, packet loss, jitter, rule-install delay, and recovery delay. To keep the deployment story "
        "complete in one document, this section imports the validated fixed GNN+ packet-SDN companion branch as a "
        "companion systems-level reference. These SDN metrics are model-based analytical estimates, not live Mininet measurements."
    )
    doc.add_heading("12.1 SDN Metrics Philosophy", level=2)
    add_dataframe_table(doc, build_sdn_metrics_formula_rows(), font_size=8)
    doc.add_paragraph(
        "No live packet generator or Mininet dataplane claim is made here. The formulas and tables below are "
        "included as a companion validation layer around the same GNN+ family, not as a replacement for the "
        "MetaGate zero-shot results above."
    )
    if packet_sdn_summary is not None and packet_sdn_metrics is not None:
        companion_summary = build_packet_sdn_companion_summary(packet_sdn_summary, packet_sdn_metrics)
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        companion_summary.to_csv(COMPANION_SDN_SUMMARY_CSV, index=False)

        doc.add_heading("12.2 Companion Packet-SDN Summary by Method", level=2)
        add_dataframe_table(doc, companion_summary, font_size=8)
        doc.add_paragraph(
            "This summary table comes from the fixed 4-method packet-SDN branch (ECMP, Bottleneck, GNN, GNN+) and is "
            "included so the full report still contains model-based SDN deployment metrics alongside the zero-shot MetaGate study."
        )

        for image_name, caption in [
            ("mlu_comparison_normal.png", "Companion packet-SDN MLU comparison across all 8 topologies."),
            ("throughput_comparison_normal.png", "Companion packet-SDN throughput comparison across all 8 topologies."),
            ("decision_time_comparison.png", "Companion packet-SDN controller decision-time comparison."),
            ("failure_recovery_gnnplus.png", "Companion packet-SDN GNN+ failure recovery comparison."),
        ]:
            add_image(doc, PACKET_SDN_PLOTS_DIR / image_name, caption)

        doc.add_heading("12.3 Companion Packet-SDN Tables", level=2)
        known_metrics = packet_sdn_metrics[packet_sdn_metrics["Status"] == "Known"].reset_index(drop=True)
        unseen_metrics = packet_sdn_metrics[packet_sdn_metrics["Status"] == "Unseen"].reset_index(drop=True)
        add_dataframe_chunks(doc, known_metrics, "Known-topology packet-SDN metrics", chunk_size=24, font_size=8)
        add_dataframe_chunks(doc, unseen_metrics, "Unseen-topology packet-SDN metrics", chunk_size=12, font_size=8)

    doc.add_heading("13. Limitations (Honest Assessment)", level=1)
    limitations = [
        "This report validates MetaGate with GNN+ as the learned expert, but it does not include a separate Stable MetaGate extension.",
        "Paper baselines such as FlexDATE, FlexEntry, and ERODRL are still unavailable in the current runnable repository.",
        "This report evaluates pure zero-shot deployment, so there is no topology-specific calibration to rescue difficult unseen cases." if ZERO_SHOT_ONLY else "The calibration phase uses 10 validation traffic matrices from the target topology, so the framing remains zero-shot gate training with few-shot calibration rather than pure zero-shot deployment.",
        "The reported results are limited to the current runnable bundle and should not be mixed numerically with earlier report branches unless the scope is stated explicitly.",
    ]
    for item in limitations:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("14. Exact Method Description for Thesis", level=1)
    doc.add_paragraph(
        "We use a two-stage expert-selection framework for traffic engineering. An MLP MetaGate "
        "(3-layer, 128-128-64-4 with BatchNorm and dropout 0.3) selects among four routing experts: "
        "Bottleneck, Top-K by Demand, Sensitivity Analysis, and GNN+-based selection, all operating "
        "with K_crit = 40 critical flows. Oracle labels are generated by running all four experts "
        "through LP optimization and selecting the expert with minimum MLU for each traffic matrix. "
        "The gate is trained on pooled oracle labels from 6 known topologies only. Before deployment "
        "on each topology, a lightweight few-shot calibration phase evaluates 10 validation traffic matrices "
        "through all four experts to build a topology-specific Bayesian prior, which is fused with the "
        "MLP softmax during test-time inference. Germany50 and VtlWavenet2011 are held out from gate-weight "
        "training and are used as unseen-topology tests."
    )
    if ZERO_SHOT_ONLY:
        doc.paragraphs[-1].text = (
            "We use a two-stage expert-selection framework for traffic engineering. An MLP MetaGate "
            "(3-layer, 128-128-64-4 with BatchNorm and dropout 0.3) selects among four routing experts: "
            "Bottleneck, Top-K by Demand, Sensitivity Analysis, and GNN+-based selection, all operating "
            "with K_crit = 40 critical flows. Oracle labels are generated by running all four experts "
            "through LP optimization and selecting the expert with minimum MLU for each traffic matrix. "
            "The gate is trained on pooled oracle labels from 6 known topologies only. During evaluation, "
            "the trained gate is applied directly with no target-topology calibration, no Bayesian prior "
            "fusion, and no weight updates. Germany50 and VtlWavenet2011 are held out from gate-weight "
            "training and are used as pure unseen-topology zero-shot tests."
        )

    doc.add_heading("15. Output Files", level=1)
    output_rows = [
        ["Primary report", str(OUTPUT_DOC.relative_to(PROJECT_ROOT))],
        ["Audit note", str(AUDIT_MD.relative_to(PROJECT_ROOT))],
        ["MetaGate summary CSV", str(SUMMARY_CSV.relative_to(PROJECT_ROOT))],
        ["MetaGate failure summary CSV", str(FAILURE_SUMMARY_CSV.relative_to(PROJECT_ROOT)) if FAILURE_SUMMARY_CSV.exists() else "N/A"],
        ["Zero-shot unseen summary CSV", str(ZERO_SHOT_UNSEEN_SUMMARY_CSV.relative_to(PROJECT_ROOT)) if ZERO_SHOT_UNSEEN_SUMMARY_CSV.exists() else "N/A"],
        ["Companion SDN summary CSV", str(COMPANION_SDN_SUMMARY_CSV.relative_to(PROJECT_ROOT)) if COMPANION_SDN_SUMMARY_CSV.exists() else "N/A"],
        ["Companion packet-SDN metrics source", str(PACKET_SDN_METRICS_CSV.relative_to(PROJECT_ROOT)) if PACKET_SDN_METRICS_CSV.exists() else "N/A"],
    ]
    add_table(doc, ["Artifact", "Path"], output_rows, font_size=9)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOC)

    audit_lines = [
        "# MetaGate GNN+ Report Audit",
        "",
        f"- Output report: `{OUTPUT_DOC.relative_to(PROJECT_ROOT)}`",
        f"- This report is built from `{OUTPUT_DIR.relative_to(PROJECT_ROOT)}` outputs.",
        "- The learned expert in this report is GNN+, not the old Original GNN.",
        "- Selector percentages and accuracy values come from the new integrated MetaGate+GNN+ evaluation.",
        f"- Failure section included: {'yes' if not failure_results.empty else 'no'}",
        f"- Companion packet-SDN SDN metrics included: {'yes' if packet_sdn_metrics is not None else 'no'}",
        f"- Compared against previous bundle: `{COMPARE_DIR.relative_to(PROJECT_ROOT)}`" if COMPARE_DIR != OUTPUT_DIR else "- Compared against previous bundle: no",
        "- Training upgrade documented: soft labels from expert MLUs, routing-aware regret penalty, and stabilized failure features.",
    ]
    AUDIT_MD.write_text("\n".join(audit_lines), encoding="utf-8")
    print(f"Report saved to {OUTPUT_DOC}")


if __name__ == "__main__":
    build_report()

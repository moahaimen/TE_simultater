#!/usr/bin/env python3
"""Generate Stage 3 Exploratory Prototype Report."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def create_stage3_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('Stage 3 Exploratory Prototype Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run('April 2026').italic = True
    
    # IMPORTANT DISCLAIMER
    disclaimer = doc.add_paragraph()
    disclaimer_para = disclaimer.add_run('IMPORTANT: This is an EXPLORATORY experiment only. ')
    disclaimer_para.bold = True
    disclaimer_para.font.color.rgb = RGBColor(255, 0, 0)
    
    disclaimer2 = doc.add_paragraph()
    disclaimer2_para = disclaimer2.add_run(
        'Stage 2 dynamic-K was NOT validated by the later lock test. '
        'This experiment does NOT upgrade the official final system. '
        'It answers only: does the combined prototype show promise for future work?'
    )
    disclaimer2_para.font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_paragraph()
    
    # 1. Goal
    doc.add_heading('1. Goal', level=1)
    doc.add_paragraph(
        'Assess whether combining Stage 1 winner features (enriched inputs, dropout=0.2) '
        'with Stage 2 pilot dynamic-K (W=1.0) produces a viable prototype. '
        'Compare against: (1) Original GNN baseline, (2) Stage 1 winner.'
    )
    
    # 2. Method
    doc.add_heading('2. Method', level=1)
    
    doc.add_heading('2.1 Scope', level=2)
    doc.add_paragraph('Germany50 (mandatory), GEANT (mandatory), Abilene (included if loaded).')
    
    doc.add_heading('2.2 Variants Tested', level=2)
    
    variants = doc.add_paragraph()
    variants.add_run('1. Original GNN: ').bold = True
    variants.add_run('Baseline with original features, fixed K=40. Reused existing checkpoint.\n')
    variants.add_run('2. Stage 1 Winner: ').bold = True
    variants.add_run('Enriched features, dropout=0.2, fixed K=40. (Not available in this run.)\n')
    variants.add_run('3. Stage 3 Prototype: ').bold = True
    variants.add_run(
        'Stage 1 winner architecture + Stage 2 pilot dynamic-K head. '
        'W=1.0 (single configuration). Trained fresh for this experiment.'
    )
    
    doc.add_heading('2.3 Stage 2 Pilot Configuration Used', level=2)
    pilot_config = doc.add_paragraph()
    pilot_config.add_run('• K-loss weight: ').bold = True
    pilot_config.add_run('W = 1.0 (raw K target, no normalization)\n')
    pilot_config.add_run('• Dropout: ').bold = True
    pilot_config.add_run('0.2 (from Stage 1 winner)\n')
    pilot_config.add_run('• Training: ').bold = True
    pilot_config.add_run('20 supervised epochs, early stopping patience=6')
    
    # 3. Results
    doc.add_heading('3. Results', level=1)
    
    doc.add_paragraph('Evaluation on test timesteps (max 50 per topology):')
    
    # Table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Topology'
    hdr_cells[1].text = 'Variant'
    hdr_cells[2].text = 'MLU'
    hdr_cells[3].text = 'PR'
    hdr_cells[4].text = 'K (mean±std)'
    hdr_cells[5].text = 'vs Orig'
    
    results = [
        ('Abilene', 'Original', '0.0524', '+0.000', '40', '-'),
        ('Abilene', 'Stage 3', '0.0822', '+0.561', '1.0±0.0', '0W/50L'),
        ('GEANT', 'Original', '0.1493', '+0.011', '40', '-'),
        ('GEANT', 'Stage 3', '0.1824', '+0.236', '1.0±0.0', '0W/50L'),
        ('Germany50', 'Original', '18.9405', '-0.008', '40', '-'),
        ('Germany50', 'Stage 3', '23.8626', '+0.277', '1.0±0.0', '0W/44L'),
        ('TOTAL', 'Original', '5.8574', '+0.002', '40', '-'),
        ('TOTAL', 'Stage 3', '7.3832', '+0.361', '1.0±0.0', '0W/144L'),
    ]
    
    for topo, variant, mlu, pr, k, vs in results:
        row_cells = table.add_row().cells
        row_cells[0].text = topo
        row_cells[1].text = variant
        row_cells[2].text = mlu
        row_cells[3].text = pr
        row_cells[4].text = k
        row_cells[5].text = vs
    
    doc.add_paragraph()
    
    # Key observations
    doc.add_heading('3.1 Key Observations', level=2)
    
    obs = doc.add_paragraph()
    obs.add_run('• Complete K collapse: ').bold = True
    obs.add_run('Stage 3 predicted K=1 (minimum) on 100% of timesteps across all topologies.\n\n')
    obs.add_run('• Zero wins: ').bold = True
    obs.add_run('Stage 3 won 0/144 comparisons vs Original GNN (144 losses).\n\n')
    obs.add_run('• Worse MLU: ').bold = True
    obs.add_run('Stage 3 MLU was 56% worse on Abilene, 22% worse on GEANT, 26% worse on Germany50.\n\n')
    obs.add_run('• Lower disturbance (trivial): ').bold = True
    obs.add_run('Zero disturbance because K=1 selected the same single flow every timestep.')
    
    # 4. Why it failed
    doc.add_heading('4. Why the Prototype Failed', level=1)
    
    doc.add_paragraph(
        'The Stage 3 prototype performed worse than even the basic Stage 2 pilot. '
        'Several factors contributed:'
    )
    
    why = doc.add_paragraph()
    why.add_run('1. K collapsed to minimum: ').bold = True
    why.add_run(
        'The model learned to always predict K=1, likely because the K-loss gradient '
        'dominated and drove predictions toward the minimum bound.\n\n'
    )
    why.add_run('2. No exploration of K space: ').bold = True
    why.add_run(
        'Unlike the Stage 2 pilot which showed some K variance (especially on Germany50), '
        'the combined prototype showed zero variance.\n\n'
    )
    why.add_run('3. Architecture mismatch: ').bold = True
    why.add_run(
        'The Stage 1 winner features (designed for fixed-K routing) may not provide '
        'the right signal for dynamic-K prediction.')
    
    # 5. Conclusion
    doc.add_heading('5. Conclusion', level=1)
    
    conclusion = doc.add_paragraph()
    conclusion_para = conclusion.add_run('NOT worth continuing')
    conclusion_para.bold = True
    conclusion_para.font.size = Pt(14)
    conclusion_para.font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        'The Stage 3 exploratory prototype combining Stage 1 winner features with '
        'Stage 2 pilot dynamic-K (W=1.0) showed no promise. '
        'K completely collapsed to the minimum value (K=1) on all topologies, '
        'producing worse MLU than the original fixed-K baseline on every single timestep.'
    )
    
    final = doc.add_paragraph()
    final.add_run('Recommendation: ').bold = True
    final.add_run(
        'Stop this line of investigation. The dynamic-K approach, even when combined '
        'with enriched features, does not show sufficient promise for future work. '
        'The official system should retain fixed K=40 (or topology-tuned K).'
    )
    
    # 6. Honesty check
    doc.add_heading('6. Honesty Check', level=1)
    
    honesty = doc.add_paragraph()
    honesty.add_run('Stage 2 validation status: ').bold = True
    honesty.add_run('NOT validated. The Stage 2 lock test showed near-zero correlation.\n\n')
    honesty.add_run('Stage 3 upgrade status: ').bold = True
    honesty.add_run('This experiment does NOT upgrade the official system.\n\n')
    honesty.add_run('Result reliability: ').bold = True
    honesty.add_run(
        'The Stage 3 failure is unambiguous (0 wins, K collapsed). '
        'No further tuning would salvage this architecture.'
    )
    
    # Save
    output_path = Path('results/gnn_plus/stage3_exploratory/Stage3_Exploratory_Prototype_Report.docx')
    doc.save(str(output_path))
    print(f"Report saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_stage3_report()

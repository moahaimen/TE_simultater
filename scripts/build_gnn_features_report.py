#!/usr/bin/env python3
"""Generate GNN_Features_Only_Report.docx."""

import os, sys
from pathlib import Path
PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))
os.chdir(PROJECT_ROOT)

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

ROOT = Path("results/gnn_plus/step1_features_only")
PLOTS = ROOT / "plots"
OUT = ROOT / "GNN_Features_Only_Report.docx"


def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.makeelement(qn("w:shd"), {qn("w:fill"): color_hex, qn("w:val"): "clear"})
    tcPr.append(shading)


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.name = "Arial"
        set_cell_shading(c, "2E4057")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]; c.text = ""
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val)); r.font.size = Pt(9); r.font.name = "Arial"
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


def add_img(doc, name, w=Inches(6)):
    fp = PLOTS / name
    if fp.exists():
        doc.add_picture(str(fp), width=w)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[Image not found: {name}]")


doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)
doc.styles["Normal"].font.name = "Arial"
doc.styles["Normal"].font.size = Pt(11)

# Title
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Step 1 Ablation: Features-Only GNN"); r.bold = True; r.font.size = Pt(20)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Isolating the effect of enriched input features (dynamic K disabled)")
r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Date: 2026-03-31  |  Branch: gnn-plus-extension  |  Type: Step 1 Ablation")
r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ---- 1. Experiment Design ----
doc.add_heading("1. Experiment Design", level=1)
doc.add_paragraph(
    "This ablation isolates the effect of enriched input features by keeping everything else "
    "identical to the original GNN training pipeline: fixed K=40, same supervised + REINFORCE "
    "two-stage training, same hyperparameters, same 6 known training topologies."
)
add_table(doc, ["Parameter", "Original GNN", "Features-only GNN"], [
    ["Input features", "34 effective (4 zeros)", "46 effective (all real)"],
    ["Node dim", "16 (12 real + 4 zero)", "16 (all 16 real)"],
    ["Edge dim", "8", "12 (+4 new)"],
    ["OD dim", "10", "18 (+8 new)"],
    ["K", "Fixed 40", "Fixed 40"],
    ["Dynamic K", "Disabled", "Disabled"],
    ["Training: supervised", "30 epochs, lr=5e-4", "30 epochs, lr=5e-4"],
    ["Training: REINFORCE", "10 epochs, lr=1e-4", "10 epochs, lr=1e-4"],
    ["Train topologies", "6 known", "6 known"],
    ["Train samples", "240 train, 120 val", "240 train, 120 val"],
], col_widths=[2.0, 2.3, 2.3])

# ---- 2. Training Results ----
doc.add_heading("2. Training Results", level=1)
add_table(doc, ["Metric", "Value"], [
    ["Supervised: best epoch", "30 (ran full)"],
    ["Supervised: best val loss", "2.869"],
    ["Supervised: val overlap", "0.665 (66.5% Jaccard with oracle)"],
    ["Supervised: time", "185 seconds"],
    ["REINFORCE: best epoch", "1"],
    ["REINFORCE: best val MLU", "691.5"],
    ["REINFORCE: early stop", "Epoch 5 (patience=4)"],
    ["REINFORCE: time", "99 seconds"],
    ["Total training time", "284 seconds (4.7 min)"],
    ["Final alpha", "0.617"],
], col_widths=[3.0, 3.5])

doc.add_paragraph(
    "Note: REINFORCE val_mlu is high because it averages across topologies with very different "
    "MLU scales (e.g., Abilene ~0.05 vs Tiscali ~800). This is expected and matches the original pipeline."
)

# ---- 3. Main Results ----
doc.add_page_break()
doc.add_heading("3. Main Results", level=1)

doc.add_heading("3.1 Per-Topology Summary", level=2)
add_table(doc,
    ["Topology", "Type", "MLU Orig", "MLU Feat", "Change%", "Feat Wins", "Orig Wins", "Dist Orig", "Dist Feat"],
    [
        ["Abilene", "known", "0.0546", "0.0546", "~0%", "34", "26", "0.093", "0.087"],
        ["CERNET", "known", "1738.3", "1726.6", "+0.67%", "59", "16", "0.307", "0.426"],
        ["GEANT", "known", "0.1615", "0.1624", "-0.57%", "26", "43", "0.120", "0.114"],
        ["Ebone", "known", "379.6", "379.6", "~0%", "35", "33", "0.156", "0.197"],
        ["Sprintlink", "known", "891.4", "843.1", "+5.42%", "73", "1", "0.241", "0.219"],
        ["Tiscali", "known", "843.6", "839.7", "+0.46%", "29", "19", "0.397", "0.425"],
        ["", "", "", "", "", "", "", "", ""],
        ["Germany50", "unseen", "18.94", "19.16", "-1.18%", "11", "33", "0.268", "0.470"],
        ["VtlWavenet", "unseen", "12252", "12266", "-0.11%", "9", "66", "0.438", "0.464"],
        ["", "", "", "", "", "", "", "", ""],
        ["KNOWN AGG", "-", "642.2", "631.5", "+1.66%", "256", "138", "0.219", "0.245"],
        ["UNSEEN AGG", "-", "7728.9", "7737.8", "-0.12%", "20", "99", "0.375", "0.466"],
        ["TOTAL AGG", "-", "2124.3", "2117.7", "+0.31%", "276", "237", "0.252", "0.291"],
    ],
    col_widths=[1.1, 0.6, 0.75, 0.75, 0.65, 0.6, 0.6, 0.6, 0.6],
)

# Key findings
doc.add_heading("3.2 Key Findings", level=2)

p = doc.add_paragraph()
r = p.add_run("MLU: Features-only GNN is BETTER on known topologies. "); r.bold = True
p.add_run("Known aggregate: +1.66% improvement (642.2 -> 631.5). "
          "Standout: Sprintlink +5.42% (73 wins vs 1 loss). CERNET +0.67% (59 wins vs 16).")

p = doc.add_paragraph()
r = p.add_run("MLU: Features-only GNN is WORSE on unseen topologies. "); r.bold = True
p.add_run("Germany50 -1.18%, VtlWavenet -0.11%. Unseen aggregate: -0.12%. "
          "The enriched features may overfit to known topology structures.")

p = doc.add_paragraph()
r = p.add_run("Disturbance: Mixed results. "); r.bold = True
p.add_run("Features-only is better on Abilene, GEANT, Sprintlink. "
          "But worse on CERNET, Ebone, Tiscali, and both unseen topologies. "
          "Overall aggregate slightly worse: 0.291 vs 0.252.")

p = doc.add_paragraph()
r = p.add_run("Total aggregate: Marginal improvement. "); r.bold = True
p.add_run("276 wins vs 237 losses (+0.31% MLU improvement). "
          "Net positive but driven mainly by Sprintlink.")

# ---- 4. Execution Time ----
doc.add_heading("4. Execution Time", level=1)
add_table(doc, ["Topology", "Orig (ms)", "Feat (ms)", "Overhead"], [
    ["Abilene", "2.0", "2.9", "+45%"],
    ["CERNET", "11.2", "22.2", "+98%"],
    ["GEANT", "4.0", "6.9", "+71%"],
    ["Ebone", "4.3", "7.7", "+80%"],
    ["Sprintlink", "12.6", "25.3", "+100%"],
    ["Tiscali", "14.9", "30.9", "+107%"],
    ["Germany50", "15.8", "32.2", "+104%"],
    ["VtlWavenet", "58.3", "240.8", "+313%"],
], col_widths=[1.6, 1.2, 1.2, 2.5])

doc.add_paragraph(
    "Feature building overhead is 45-313%, scaling with topology size. "
    "The clustering coefficient computation and path overlap scoring are O(V*deg^2) and O(OD*path_len). "
    "VtlWavenet (92 nodes, 8372 ODs) shows the worst overhead. "
    "Still sub-250ms for the largest topology."
)

# ---- 5. Plots ----
doc.add_page_break()
doc.add_heading("5. Comparison Plots", level=1)

doc.add_heading("5.1 MLU CDF — Known Topologies", level=2)
add_img(doc, "mlu_cdf_known.png", Inches(6.3))
doc.add_paragraph("Figure 1: Features-only GNN matches or beats original on most known topologies.")

doc.add_heading("5.2 MLU CDF — Unseen Topologies", level=2)
add_img(doc, "mlu_cdf_unseen.png", Inches(6.3))
doc.add_paragraph("Figure 2: Original GNN is slightly better on unseen topologies.")

doc.add_heading("5.3 Mean MLU Bar Chart", level=2)
add_img(doc, "mean_mlu_bar.png", Inches(5.5))
doc.add_paragraph("Figure 3: Per-topology mean MLU. Sprintlink shows the largest improvement.")

doc.add_heading("5.4 Disturbance CDF", level=2)
add_img(doc, "disturbance_cdf.png", Inches(6.3))
doc.add_paragraph("Figure 4: Mixed disturbance results across topologies.")

# ---- 6. Honest Limitations ----
doc.add_page_break()
doc.add_heading("6. Honest Limitations", level=1)
for item in [
    "This is Step 1 of a staged ablation. Features and fixed K=40 are tested together. "
    "We have not yet tested dynamic K in isolation.",
    "The original GNN checkpoint includes REINFORCE fine-tuning on 240 samples. "
    "Our features-only model was trained with the same pipeline but from scratch. "
    "Random initialization differences may account for some variance.",
    "REINFORCE early-stopped at epoch 5 (patience=4 after epoch 1 was best). "
    "The original may have had different REINFORCE convergence characteristics.",
    "Sprintlink dominates the known-aggregate improvement. Without Sprintlink, "
    "the known-aggregate improvement drops significantly.",
    "Unseen topology regression (Germany50 -1.18%) suggests the richer features may "
    "encode topology-specific patterns that don't generalize.",
    "Feature computation overhead (45-313%) is acceptable for real-time routing "
    "but should be optimized if deployed at scale.",
    "Only 8 topologies tested. Results may not generalize to arbitrary network topologies.",
]:
    doc.add_paragraph(item, style="List Bullet")

# ---- 7. Verdict ----
doc.add_heading("7. Verdict and Next Steps", level=1)

p = doc.add_paragraph()
r = p.add_run("Q1: Did enriched features improve MLU?"); r.bold = True
doc.add_paragraph(
    "YES on known topologies (+1.66% aggregate, driven by Sprintlink +5.42%). "
    "NO on unseen topologies (-0.12% aggregate). "
    "Overall: marginal net positive (+0.31%)."
)

p = doc.add_paragraph()
r = p.add_run("Q2: Did they improve disturbance/stability?"); r.bold = True
doc.add_paragraph(
    "MIXED. Better on 3 topologies, worse on 5. Aggregate slightly worse (0.291 vs 0.252). "
    "Enriched features change the score distribution, increasing selection variability on some topologies."
)

p = doc.add_paragraph()
r = p.add_run("Q3: Is Step 2 (dynamic-K-only) justified?"); r.bold = True
doc.add_paragraph(
    "YES, but with low priority. The screening showed K_pred converges to 39 anyway. "
    "The more promising direction is investigating WHY features help on Sprintlink "
    "and whether feature selection (choosing a subset of new features) can recover unseen performance."
)

p = doc.add_paragraph()
r = p.add_run("Q4: Should this variant be stopped or continued?"); r.bold = True
doc.add_paragraph(
    "CONTINUE with caution. The Sprintlink result (+5.42%, 73/75 wins) is too strong to ignore. "
    "The enriched features are clearly capturing something real for complex topologies. "
    "But the unseen regression needs to be addressed before integration into MetaGate."
)

doc.add_paragraph("")
p = doc.add_paragraph()
r = p.add_run("Recommended next steps:"); r.bold = True; r.font.size = Pt(12)
for step in [
    "Feature importance analysis: which of the 12 new features drive the Sprintlink gain?",
    "Selective feature integration: try adding only the top 3-4 most impactful features.",
    "Regularization: test dropout or feature noise during training to reduce overfitting to known topologies.",
    "If selective features recover unseen performance, integrate into the main GNN and retrain MetaGate.",
    "Dynamic K (Step 2) has low priority given K_pred=39 in screening.",
]:
    doc.add_paragraph(step, style="List Number")

# Fix zoom
for section in doc.sections:
    settings = doc.settings.element
    for z in settings.findall(qn("w:zoom")):
        if z.get(qn("w:percent")) is None:
            z.set(qn("w:percent"), "100")

doc.save(str(OUT))
print(f"Written: {OUT} ({OUT.stat().st_size} bytes)")

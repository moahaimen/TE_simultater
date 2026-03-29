#!/usr/bin/env python3
"""Build comprehensive professor-ready report with all metrics, CDFs, and SDN analysis."""

import os
import sys
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

PROJECT = Path(__file__).resolve().parent.parent
OUTPUT = PROJECT / "results" / "dynamic_metagate" / "Professor_Final_Report.docx"

# ── Load all data ──
mg = pd.read_csv(PROJECT / "results/dynamic_metagate/metagate_results.csv")
mg_summary = pd.read_csv(PROJECT / "results/dynamic_metagate/metagate_summary.csv")
ar = pd.read_csv(PROJECT / "results/requirements_compliant_eval/all_results.csv")
fr = pd.read_csv(PROJECT / "results/requirements_compliant_eval/failure_results.csv")
sdn = pd.read_csv(PROJECT / "results/sdn/sdn_scaled_summary.csv")
pr_csv = pd.read_csv(PROJECT / "results/requirements_compliant_eval/pr_summary.csv")
gnn_train = pd.read_csv(PROJECT / "results/requirements_compliant_eval/gnn_training_log.csv")
train_eff = pd.read_csv(PROJECT / "results/requirements_compliant_eval/table_training_efficiency.csv")

NODES_MAP = {
    "abilene": 12, "geant": 22, "rocketfuel_ebone": 23, "cernet": 41,
    "rocketfuel_sprintlink": 44, "rocketfuel_tiscali": 49, "germany50": 50,
    "topologyzoo_vtlwavenet2011": 92,
}
TOPO_DISPLAY = {
    "abilene": "Abilene (12n)", "geant": "GEANT (22n)", "rocketfuel_ebone": "Ebone (23n)",
    "cernet": "CERNET (41n)", "rocketfuel_sprintlink": "Sprintlink (44n)",
    "rocketfuel_tiscali": "Tiscali (49n)", "germany50": "Germany50 (50n)",
    "topologyzoo_vtlwavenet2011": "VtlWavenet (92n)",
}

doc = Document()

# ── Styles ──
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level, size in [(1, 15), (2, 13), (3, 11)]:
    h = doc.styles[f"Heading {level}"]
    h.font.name = "Arial"
    h.font.size = Pt(size)
    h.font.bold = True
    h.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)


def add_table(headers, rows, col_widths=None, font_size=9):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.name = "Arial"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(font_size)
            run.font.name = "Arial"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Inches(w)
    doc.add_paragraph("")
    return t


def bold_run(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    return run


def add_image_if_exists(path, width_inches=5.5):
    p = Path(path)
    if p.exists():
        doc.add_picture(str(p), width=Inches(width_inches))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    else:
        doc.add_paragraph(f"[Figure not found: {p.name}]").italic = True
        return False


# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(4):
    doc.add_paragraph("")
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("AI-Driven Traffic Engineering\nwith MLP Meta-Gate Expert Selection")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

doc.add_paragraph("")
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Comprehensive Evaluation Report\nfor Professor Review")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph("")
method = doc.add_paragraph()
method.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = method.add_run(
    "Method: MLP Meta-Gate with Few-Shot Bayesian Calibration\n"
    "4 Experts: Bottleneck, TopK, Sensitivity, GNN\n"
    "Framing: Zero-Shot Gate Training + Few-Shot Calibration\n"
    "This is NOT a pure zero-shot evaluation"
)
run.font.size = Pt(11)
run.italic = True

doc.add_paragraph("")
toc_note = doc.add_paragraph()
toc_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = toc_note.add_run(
    "Covers: MLU, Performance Ratio, Disturbance, Execution Time,\n"
    "Training Efficiency, Failure Robustness, SDN Deployment,\n"
    "Complexity Analysis, CDF Plots, and Live Mininet Status"
)
run.font.size = Pt(10)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Method Summary",
    "2. MLU Results (Maximum Link Utilization)",
    "3. Performance Ratio (PR)",
    "4. Network Disturbance (DB)",
    "5. Execution Time and Decision Latency",
    "6. Training Efficiency and Stability",
    "7. Routing Update Overhead",
    "8. Failure Robustness",
    "9. CDF Plots",
    "10. SDN Deployment Validation",
    "11. Complexity Analysis",
    "12. Germany50 Unseen Topology Deep Dive",
    "13. CERNET Correction",
    "14. Limitations (Honest Assessment)",
    "15. Live Mininet Testbed: Required Final Step",
    "16. Exact Thesis Method Description",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. METHOD SUMMARY
# ============================================================
doc.add_heading("1. Method Summary", level=1)
doc.add_paragraph(
    "An MLP-based meta-gate (3-layer, 128-128-64-4, BatchNorm, dropout 0.3) dynamically selects "
    "among four expert flow-selection strategies per traffic matrix. Gate weights are trained on "
    "pooled oracle labels from 6 known topologies (~2,098 samples). Germany50 and VtlWavenet2011 "
    "are held out entirely from gate-weight training."
)
doc.add_paragraph(
    "At deployment, a lightweight few-shot calibration phase runs 10 validation TMs through all "
    "4 experts with LP, building a topology-specific Bayesian prior. This prior is fused with "
    "MLP softmax: P_final(i) = P_MLP(i) * prior(i)^5 / Z. No gradient updates during calibration."
)
p = doc.add_paragraph()
bold_run(p, "Framing: ")
p.add_run("Zero-shot gate training + few-shot calibration. This is NOT pure zero-shot.")

experts_tbl = [
    ("Bottleneck", "Selects K=40 OD pairs traversing the most-utilized ECMP link"),
    ("Top-K by Demand", "Selects K=40 OD pairs with highest traffic demand"),
    ("Sensitivity", "Selects K=40 OD pairs whose rerouting maximally reduces max link utilization"),
    ("GNN (GraphSAGE)", "Learned selector using graph neural network embeddings, K=40"),
]
add_table(["Expert", "Strategy (all K_crit=40)"], experts_tbl, col_widths=[1.5, 5.0])

# ============================================================
# 2. MLU RESULTS
# ============================================================
doc.add_heading("2. MLU Results (Maximum Link Utilization)", level=1)
doc.add_heading("2.1 MetaGate vs Oracle vs Individual Experts", level=2)

mlu_rows = []
for _, r in mg_summary.sort_values("dataset").iterrows():
    ds = r["dataset"]
    tp = "UNSEEN" if r["topology_type"] == "unseen" else "known"
    mlu_rows.append((
        TOPO_DISPLAY.get(ds, ds), tp,
        f"{r['metagate_mlu']:.4f}", f"{r['oracle_mlu']:.4f}",
        f"{r['bn_mlu']:.4f}", f"{r['topk_mlu']:.4f}",
        f"{r['sens_mlu']:.4f}", f"{r['gnn_mlu']:.4f}",
        f"{r['metagate_vs_oracle_gap_pct']:+.2f}%",
    ))
add_table(
    ["Topology", "Type", "MetaGate", "Oracle", "BN", "TopK", "Sens", "GNN", "Gap"],
    mlu_rows, font_size=8,
)

doc.add_heading("2.2 MetaGate vs External Baselines", level=2)
doc.add_paragraph(
    "Comparison with ECMP, OSPF, FlexDATE, ERODRL, CFRRL, FlexEntry (from requirements-compliant eval):"
)
baseline_summary = ar.groupby(["dataset", "method"])["mlu"].mean().unstack()
baseline_rows = []
for ds in sorted(NODES_MAP.keys(), key=lambda x: NODES_MAP[x]):
    if ds not in baseline_summary.index:
        continue
    row = baseline_summary.loc[ds]
    mg_mlu = mg_summary[mg_summary["dataset"] == ds]["metagate_mlu"].values
    mg_val = f"{mg_mlu[0]:.4f}" if len(mg_mlu) > 0 else "N/A"
    ecmp_val = f"{row.get('ecmp', 0):.4f}" if "ecmp" in row else "N/A"
    flex_val = f"{row.get('flexdate', 0):.4f}" if "flexdate" in row else "N/A"
    baseline_rows.append((
        TOPO_DISPLAY.get(ds, ds), mg_val, f"{row.get('bottleneck', 0):.4f}",
        f"{row.get('gnn', 0):.4f}", ecmp_val, flex_val,
    ))
add_table(
    ["Topology", "MetaGate", "Bottleneck", "GNN", "ECMP", "FlexDATE"],
    baseline_rows, font_size=9,
)

# ============================================================
# 3. PERFORMANCE RATIO
# ============================================================
doc.add_heading("3. Performance Ratio (PR = Method MLU / Oracle MLU)", level=1)
doc.add_paragraph(
    "PR = 1.000 means the method matches the oracle exactly. Values > 1 indicate suboptimality."
)

mg["pr"] = mg["metagate_mlu"] / mg["oracle_mlu"]
pr_rows = []
for ds in sorted(NODES_MAP.keys(), key=lambda x: NODES_MAP[x]):
    sub = mg[mg["dataset"] == ds]
    if len(sub) == 0:
        continue
    tp = "UNSEEN" if ds in {"germany50", "topologyzoo_vtlwavenet2011"} else "known"
    mean_pr = sub["pr"].mean()
    p95_pr = np.percentile(sub["pr"], 95)
    # Expert PRs
    bn_pr = (sub["bn_mlu"] / sub["oracle_mlu"]).mean()
    gnn_pr = (sub["gnn_mlu"] / sub["oracle_mlu"]).mean()
    pr_rows.append((
        TOPO_DISPLAY.get(ds, ds), tp,
        f"{mean_pr:.6f}", f"{p95_pr:.6f}",
        f"{bn_pr:.6f}", f"{gnn_pr:.6f}",
    ))
add_table(
    ["Topology", "Type", "MetaGate Mean PR", "MetaGate P95 PR", "BN Mean PR", "GNN Mean PR"],
    pr_rows, font_size=9,
)

# ============================================================
# 4. DISTURBANCE
# ============================================================
doc.add_heading("4. Network Disturbance (DB)", level=1)
doc.add_paragraph(
    "Disturbance measures the fraction of OD pairs whose routing changes between consecutive "
    "timesteps. Lower is better (fewer routing disruptions)."
)

experts_ar = ar[ar["method"].isin(["bottleneck", "topk", "sensitivity", "gnn"])]
db_summary = experts_ar.groupby(["dataset", "method"])["disturbance"].agg(["mean", lambda x: np.percentile(x, 95)]).reset_index()
db_summary.columns = ["dataset", "method", "mean_db", "p95_db"]

db_rows = []
for ds in sorted(NODES_MAP.keys(), key=lambda x: NODES_MAP[x]):
    sub = db_summary[db_summary["dataset"] == ds]
    if len(sub) == 0:
        continue
    bn_db = sub[sub["method"] == "bottleneck"]["mean_db"].values
    gnn_db = sub[sub["method"] == "gnn"]["mean_db"].values
    sens_db = sub[sub["method"] == "sensitivity"]["mean_db"].values
    topk_db = sub[sub["method"] == "topk"]["mean_db"].values
    db_rows.append((
        TOPO_DISPLAY.get(ds, ds),
        f"{bn_db[0]:.4f}" if len(bn_db) else "N/A",
        f"{topk_db[0]:.4f}" if len(topk_db) else "N/A",
        f"{sens_db[0]:.4f}" if len(sens_db) else "N/A",
        f"{gnn_db[0]:.4f}" if len(gnn_db) else "N/A",
    ))
add_table(
    ["Topology", "BN DB", "TopK DB", "Sens DB", "GNN DB"],
    db_rows,
)
doc.add_paragraph(
    "Note: MetaGate disturbance is not separately tracked because MetaGate delegates to one "
    "of these 4 experts per timestep. The actual DB depends on which expert is selected and "
    "whether the selected expert changes between consecutive timesteps (selector switches)."
)

# ============================================================
# 5. EXECUTION TIME
# ============================================================
doc.add_heading("5. Execution Time and Decision Latency", level=1)
doc.add_paragraph(
    "Terminology: Decision time = all selector overhead before LP (run 4 experts + extract features "
    "+ MLP forward pass). LP time = LP solver only. Total time = decision + LP (end-to-end per timestep)."
)

time_rows = []
for ds in sorted(NODES_MAP.keys(), key=lambda x: NODES_MAP[x]):
    sub = mg[mg["dataset"] == ds]
    if len(sub) == 0:
        continue
    time_rows.append((
        TOPO_DISPLAY.get(ds, ds), str(NODES_MAP[ds]),
        f"{sub['t_bn_ms'].mean():.1f}", f"{sub['t_topk_ms'].mean():.1f}",
        f"{sub['t_sens_ms'].mean():.1f}", f"{sub['t_gnn_ms'].mean():.1f}",
        f"{sub['t_features_ms'].mean():.2f}", f"{sub['t_mlp_ms'].mean():.2f}",
        f"{sub['t_lp_ms'].mean():.1f}", f"{sub['t_total_ms'].mean():.1f}",
    ))
add_table(
    ["Topology", "N", "BN ms", "TopK ms", "Sens ms", "GNN ms", "Feat ms", "MLP ms", "LP ms", "Total ms"],
    time_rows, font_size=8,
)

p = doc.add_paragraph()
bold_run(p, "MLP gate inference: ")
p.add_run("< 0.3 ms on all topologies (feature extraction + MLP forward pass < 2 ms total). ")
bold_run(p, "Bottleneck: ")
p.add_run("The 4 expert runs dominate decision time (all must run before the gate decides). ")
bold_run(p, "Mean end-to-end: ")
p.add_run("118.8 ms across all topologies.")

# ============================================================
# 6. TRAINING EFFICIENCY
# ============================================================
doc.add_heading("6. Training Efficiency and Stability", level=1)

doc.add_heading("6.1 GNN Expert Training", level=2)
gnn_rows = []
for _, r in gnn_train.iterrows():
    gnn_rows.append((
        str(int(r["epoch"])), f"{r['train_loss']:.3f}", f"{r['val_loss']:.3f}",
        f"{r.get('alpha', 'N/A')}", f"{r['epoch_time_sec']:.1f}s",
    ))
# Show first 5 and last 5
display_rows = gnn_rows[:5] + [("...", "...", "...", "...", "...")] + gnn_rows[-5:]
add_table(
    ["Epoch", "Train Loss", "Val Loss", "Alpha", "Time"],
    display_rows, font_size=9,
)
doc.add_paragraph(
    f"GNN Selector (GraphSAGE): {train_eff.iloc[0]['training_time_sec']:.1f}s total training, "
    f"{int(train_eff.iloc[0]['convergence_epoch'])} epochs, "
    f"best val loss = {train_eff.iloc[0]['best_val_loss']:.3f}. "
    f"Hidden dim = {int(train_eff.iloc[0]['hidden_dim'])}, layers = {int(train_eff.iloc[0]['num_layers'])}."
)

doc.add_heading("6.2 MLP Meta-Gate Training", level=2)
mlp_config = [
    ("Architecture", "3-layer MLP: 49-128-128-64-4 with BatchNorm, Dropout(0.3)"),
    ("Training samples", "2,098 (pooled from 6 known topologies)"),
    ("Validation samples", "450"),
    ("Epochs", "300"),
    ("Batch size", "64"),
    ("Optimizer", "Adam, lr=5e-4"),
    ("Class weights", "Inverse-frequency: w_i = N / (4 * count_i)"),
    ("Train accuracy", "62.3%"),
    ("Val accuracy", "49.8%"),
    ("Training time", "~60 seconds (CPU)"),
]
add_table(["Parameter", "Value"], mlp_config, col_widths=[2.0, 4.5])

doc.add_heading("6.3 Training Stability Assessment", level=2)
doc.add_paragraph(
    "GNN training shows stable convergence: val loss decreases monotonically from 2.769 (epoch 1) "
    "to 2.519 (epoch 30) with no overfitting. MLP gate training converges within 300 epochs; "
    "validation accuracy plateaus at ~50%, indicating the gate is a weak classifier that relies "
    "on calibration for deployment (see Limitations section)."
)

# ============================================================
# 7. ROUTING UPDATE OVERHEAD
# ============================================================
doc.add_heading("7. Routing Update Overhead / Critical Entries", level=1)
doc.add_paragraph(
    "All 4 experts select exactly K_crit = 40 critical OD pairs per timestep. "
    "The LP solver computes new split ratios only for these 40 pairs; all other OD pairs "
    "retain their ECMP routing. This bounds the routing update overhead."
)

doc.add_heading("7.1 Rules Pushed per SDN Cycle", level=2)
rules_rows = []
for _, r in sdn.iterrows():
    rules_rows.append((
        r["topology"], str(int(r["nodes"])),
        f"{r['rules_per_cycle']:.1f}", "40",
        f"{r['decision_ms']:.1f}",
    ))
add_table(
    ["Topology", "Nodes", "Avg Rules/Cycle", "K_crit", "Decision ms"],
    rules_rows,
)
doc.add_paragraph(
    "Rules per cycle ranges from 3.1 (Sprintlink) to 12.1 (Tiscali). "
    "This is far below K_crit=40 because only changed split ratios generate new rules."
)

doc.add_heading("7.2 MetaGate Selector Switches", level=2)
doc.add_paragraph(
    "Selector switches count how many times the MetaGate changes its expert choice between "
    "consecutive timesteps. More switches = more dynamic behavior."
)
switch_rows = []
for ds in sorted(NODES_MAP.keys(), key=lambda x: NODES_MAP[x]):
    sub = mg[mg["dataset"] == ds].sort_values("timestep")
    if len(sub) == 0:
        continue
    switches = (sub["metagate_selector"].values[1:] != sub["metagate_selector"].values[:-1]).sum()
    tp = "UNSEEN" if ds in {"germany50", "topologyzoo_vtlwavenet2011"} else "known"
    sel_counts = sub["metagate_selector"].value_counts()
    dominant = sel_counts.index[0]
    dom_pct = sel_counts.iloc[0] / len(sub) * 100
    switch_rows.append((
        TOPO_DISPLAY.get(ds, ds), tp,
        f"{switches}/{len(sub)-1}", f"{dominant} ({dom_pct:.0f}%)",
        f"{sub['correct'].mean():.1%}",
    ))
add_table(
    ["Topology", "Type", "Switches", "Dominant Expert", "Accuracy"],
    switch_rows,
)

# ============================================================
# 8. FAILURE ROBUSTNESS
# ============================================================
doc.add_heading("8. Failure Robustness", level=1)
doc.add_paragraph(
    "Three failure scenarios tested: single link failure, capacity degradation (50%), "
    "and traffic spike (2x). Results from requirements-compliant evaluation."
)

for ft in ["single_link_failure", "capacity_degradation", "traffic_spike"]:
    ft_display = ft.replace("_", " ").title()
    doc.add_heading(f"8.{['single_link_failure','capacity_degradation','traffic_spike'].index(ft)+1} {ft_display}", level=2)

    fail_sub = fr[(fr["failure_type"] == ft) & (fr["method"].isin(["bottleneck", "topk", "sensitivity", "gnn"]))]
    fail_agg = fail_sub.groupby(["dataset", "method"])["mlu"].mean().unstack()

    fail_rows = []
    for ds in sorted(NODES_MAP.keys(), key=lambda x: NODES_MAP[x]):
        if ds not in fail_agg.index:
            continue
        row = fail_agg.loc[ds]
        fail_rows.append((
            TOPO_DISPLAY.get(ds, ds),
            f"{row.get('bottleneck', 0):.4f}",
            f"{row.get('topk', 0):.4f}",
            f"{row.get('sensitivity', 0):.4f}",
            f"{row.get('gnn', 0):.4f}",
        ))
    if fail_rows:
        add_table(["Topology", "BN MLU", "TopK MLU", "Sens MLU", "GNN MLU"], fail_rows, font_size=9)

# ============================================================
# 9. CDF PLOTS
# ============================================================
doc.add_heading("9. CDF Plots", level=1)
doc.add_paragraph(
    "All CDFs generated from per-timestep raw data. Methods compared: Bottleneck, TopK, "
    "Sensitivity, GNN, ECMP, FlexDATE, and others (10 methods total). "
    "MetaGate CDF overlay is noted where the per-timestep data supports it."
)

plots_dir = PROJECT / "results/requirements_compliant_eval/plots"

doc.add_heading("9.1 CDF of MLU (per topology)", level=2)
for ds in sorted(NODES_MAP.keys(), key=lambda x: NODES_MAP[x]):
    p = plots_dir / f"cdf_mlu_{ds}.png"
    if p.exists():
        doc.add_paragraph(f"{TOPO_DISPLAY.get(ds, ds)}:", style="Normal").bold = True
        add_image_if_exists(p, width_inches=5.0)

doc.add_heading("9.2 CDF of MLU (all topologies combined)", level=2)
add_image_if_exists(plots_dir / "cdf_mlu_all_topologies.png", width_inches=5.5)

doc.add_heading("9.3 CDF of Routing Disturbance (per topology)", level=2)
for ds in sorted(NODES_MAP.keys(), key=lambda x: NODES_MAP[x]):
    p = plots_dir / f"cdf_disturbance_{ds}.png"
    if p.exists():
        doc.add_paragraph(f"{TOPO_DISPLAY.get(ds, ds)}:")
        add_image_if_exists(p, width_inches=5.0)

doc.add_heading("9.4 CDF of Decision/Execution Time (per topology)", level=2)
for ds in sorted(NODES_MAP.keys(), key=lambda x: NODES_MAP[x]):
    p = plots_dir / f"cdf_exec_time_{ds}.png"
    if p.exists():
        doc.add_paragraph(f"{TOPO_DISPLAY.get(ds, ds)}:")
        add_image_if_exists(p, width_inches=5.0)

doc.add_heading("9.5 CDF of MLU under Link Failure", level=2)
add_image_if_exists(plots_dir / "cdf_mlu_failure_single_link_failure.png", width_inches=5.5)

doc.add_heading("9.6 CDF of MLU under Capacity Degradation", level=2)
add_image_if_exists(plots_dir / "cdf_mlu_failure_capacity_degradation.png", width_inches=5.5)

doc.add_heading("9.7 CDF of MLU under Traffic Spike", level=2)
add_image_if_exists(plots_dir / "cdf_mlu_failure_traffic_spike.png", width_inches=5.5)

doc.add_heading("9.8 CDF of Disturbance under Failure", level=2)
for ft in ["single_link_failure", "capacity_degradation", "traffic_spike"]:
    p = plots_dir / f"cdf_disturbance_failure_{ft}.png"
    if p.exists():
        doc.add_paragraph(f"{ft.replace('_', ' ').title()}:")
        add_image_if_exists(p, width_inches=5.0)

doc.add_heading("9.9 Training Convergence", level=2)
add_image_if_exists(plots_dir / "training_convergence.png", width_inches=5.0)

# ============================================================
# 10. SDN DEPLOYMENT
# ============================================================
doc.add_heading("10. SDN Deployment Validation", level=1)

p = doc.add_paragraph()
bold_run(p, "HONESTY NOTE: ")
p.add_run(
    "The SDN results below are from a model-based simulation that mimics the SDN control loop "
    "(read TM -> select expert -> solve LP -> push rules). They use real topology graphs and "
    "traffic matrices, but they are NOT from live Mininet with real packet forwarding. "
    "The simulation computes MLU mathematically from the routing solution rather than measuring "
    "actual packet delays and throughput. See Section 15 for what a live Mininet run would add."
)

doc.add_heading("10.1 SDN Simulation Results", level=2)
sdn_rows = []
for _, r in sdn.iterrows():
    sdn_rows.append((
        r["topology"], str(int(r["nodes"])),
        f"{r['pre_mlu']:.4f}", f"{r['post_mlu']:.4f}",
        f"{r['improvement_pct']:.2f}%",
        f"{r['decision_ms']:.1f}", f"{r['rules_per_cycle']:.1f}",
    ))
add_table(
    ["Topology", "Nodes", "Pre-MLU", "Post-MLU", "Improvement", "Decision ms", "Rules/Cycle"],
    sdn_rows,
)

doc.add_heading("10.2 What These SDN Metrics Mean", level=2)
sdn_metrics = [
    ("Pre-MLU", "ECMP baseline utilization before optimization", "Model-based"),
    ("Post-MLU", "Utilization after LP-optimized rerouting", "Model-based"),
    ("Decision ms", "Time to select expert + solve LP", "Real wall-clock"),
    ("Rules/Cycle", "Number of OpenFlow rule updates per optimization cycle", "Computed from LP solution"),
    ("Throughput", "NOT AVAILABLE - requires live Mininet with iperf", "Needs testbed"),
    ("Packet loss", "NOT AVAILABLE - requires live Mininet", "Needs testbed"),
    ("Jitter", "NOT AVAILABLE - requires live Mininet", "Needs testbed"),
    ("Rule install delay", "NOT AVAILABLE - requires live OVS", "Needs testbed"),
    ("Failure recovery", "NOT AVAILABLE - requires live link-down events", "Needs testbed"),
]
add_table(
    ["Metric", "Description", "Source"],
    sdn_metrics, font_size=9,
)

# ============================================================
# 11. COMPLEXITY ANALYSIS
# ============================================================
doc.add_heading("11. Complexity Analysis", level=1)

doc.add_heading("11.1 Empirical Scaling", level=2)
doc.add_paragraph(
    "Measured per-timestep execution time across 8 topologies (12 to 92 nodes):"
)
mg["nodes"] = mg["dataset"].map(NODES_MAP)
cs = mg.groupby("dataset").agg(
    nodes=("nodes", "first"),
    t_decision=("t_decision_ms", "mean"),
    t_lp=("t_lp_ms", "mean"),
    t_total=("t_total_ms", "mean"),
    t_bn=("t_bn_ms", "mean"),
    t_gnn=("t_gnn_ms", "mean"),
    t_sens=("t_sens_ms", "mean"),
    t_feat=("t_features_ms", "mean"),
    t_mlp=("t_mlp_ms", "mean"),
).sort_values("nodes")

cx_rows = []
for ds, r in cs.iterrows():
    cx_rows.append((
        TOPO_DISPLAY.get(ds, ds), str(int(r["nodes"])),
        f"{r['t_decision']:.1f}", f"{r['t_lp']:.1f}", f"{r['t_total']:.1f}",
        f"{r['t_feat']:.2f}", f"{r['t_mlp']:.2f}",
    ))
add_table(
    ["Topology", "Nodes", "Decision ms", "LP ms", "Total ms", "Feature ms", "MLP ms"],
    cx_rows,
)

doc.add_heading("11.2 Algorithmic Complexity", level=2)
complexity = [
    ("Bottleneck selector", "O(E) per TM", "Scan all edges for max-util link, collect OD pairs"),
    ("TopK selector", "O(D log D)", "Sort all OD demands, take top K"),
    ("Sensitivity selector", "O(K * E)", "Per-flow sensitivity requires edge-level analysis"),
    ("GNN selector", "O(N * H * L)", "GraphSAGE: N nodes, H hidden dim, L layers, forward pass"),
    ("Feature extraction", "O(D + K^2)", "TM stats + pairwise set operations on K=40 selections"),
    ("MLP gate", "O(H^2 * L)", "3-layer MLP: 49->128->128->64->4, ~25K parameters"),
    ("Calibration", "O(10 * 4 * LP)", "One-time: 10 val TMs * 4 experts * LP solve each"),
    ("LP solver", "O(E * D * K)", "LP with E edges, D OD pairs, K critical flows"),
]
add_table(
    ["Component", "Complexity", "Notes"],
    complexity, font_size=9, col_widths=[1.5, 1.2, 4.0],
)

doc.add_heading("11.3 Scaling Observations", level=2)
doc.add_paragraph(
    "Decision time scales roughly quadratically with node count: 4 ms (12 nodes) to 227 ms "
    "(92 nodes). The LP solver scales similarly. The MLP gate itself is constant-time (< 0.3 ms) "
    "regardless of topology size. The scaling bottleneck is running the 4 expert selectors, "
    "particularly Sensitivity analysis which requires O(K*E) edge-level computation."
)
doc.add_paragraph(
    "Total end-to-end time remains under 120 ms for topologies up to 50 nodes, "
    "making the system viable for reactive TE with ~1-second control intervals. "
    "The 92-node VtlWavenet at 452 ms would require ~2-second intervals."
)

# ============================================================
# 12. GERMANY50 DEEP DIVE
# ============================================================
doc.add_heading("12. Germany50 Unseen Topology Deep Dive", level=1)

doc.add_heading("12.1 Before vs After Calibration", level=2)
g50_data = [
    ("BN selections", "44/44 (100%)", "16/44 (36%)"),
    ("GNN selections", "0/44 (0%)", "28/44 (64%)"),
    ("GNN selection rate", "0.0%", "63.6%"),
    ("Selector switches", "0/43", "7/43"),
    ("MetaGate MLU", "19.227", "18.992"),
    ("Oracle MLU", "18.929", "18.929"),
    ("Oracle gap", "+1.57%", "+0.33%"),
    ("Accuracy", "2.3%", "65.9%"),
    ("Performance ratio (PR)", "1.016", "1.003"),
]
add_table(
    ["Metric", "Before Calibration (raw MLP)", "After Calibration (10-val prior)"],
    g50_data, col_widths=[2.0, 2.2, 2.2],
)

doc.add_heading("12.2 Why Calibration Fixed It", level=2)
doc.add_paragraph(
    "The 10 calibration TMs show GNN wins 8/10 times -> prior = [BN=0.07, TopK=0.07, Sens=0.07, GNN=0.79]. "
    "With alpha=5: prior^5 = [BN=0.00002, TopK=0.00002, Sens=0.00002, GNN=0.308]. "
    "The 15,400x GNN/BN ratio overrides the MLP's Bottleneck confidence on most timesteps. "
    "The 7 selector switches across 43 transitions confirm genuine per-timestep dynamic behavior."
)

doc.add_heading("12.3 Oracle Distribution (Test Set)", level=2)
doc.add_paragraph(
    "Germany50 oracle: GNN = 42/44 (95.5%), BN = 1/44 (2.3%), Sens = 1/44 (2.3%). "
    "MetaGate achieves 65.9% accuracy because it still selects BN on 16/44 timesteps where "
    "GNN is optimal. However, the MLU penalty is small: +0.33% oracle gap."
)

# ============================================================
# 13. CERNET CORRECTION
# ============================================================
doc.add_heading("13. CERNET Topology Correction", level=1)
cernet_facts = [
    ("Source", "TopologyZoo (Cernet.graphml), NOT SNDlib"),
    ("Nodes", "41"),
    ("Edges", "59 bidirectional links = 116 directed edges"),
    ("Traffic", "Synthetic MGM (not real SNDlib traffic matrices)"),
    ("Config note", "traffic_mode='real_sndlib' kept to avoid breaking loader; truth documented in comments"),
]
add_table(["Property", "Actual Value"], cernet_facts, col_widths=[1.5, 5.0])

# ============================================================
# 14. LIMITATIONS
# ============================================================
doc.add_heading("14. Limitations (Honest Assessment)", level=1)
limitations = [
    ("Not pure zero-shot",
     "The calibration uses 10 validation TMs from the target topology itself (4 experts x LP each). "
     "This is few-shot adaptation, not pure zero-shot generalization."),
    ("Calibration cost",
     "40 LP solves per topology (10 TMs x 4 experts). One-time cost, not per-timestep, "
     "but non-trivial for very large topologies."),
    ("Alpha hand-tuned",
     "The calibration strength alpha=5 was selected empirically. No principled optimization."),
    ("Raw MLP is weak",
     "Validation accuracy ~50%. Without calibration, the gate collapses to mostly-Bottleneck. "
     "Bayesian calibration compensates, but the MLP itself cannot distinguish when GNN beats BN."),
    ("VtlWavenet never selects GNN",
     "Despite calibration prior GNN=0.57, MLP confidence in BN is too extreme. "
     "Masked by near-optimal performance (+0.06% gap)."),
    ("No live packet measurements",
     "All SDN metrics are model-based. No real throughput, latency, jitter, or packet loss data. "
     "Mininet is not installed in the current environment."),
    ("Accuracy vs MLU disconnect",
     "Some topologies show low accuracy but near-zero MLU gap (Abilene: 40% acc, +0.00% gap). "
     "Expert choices are often interchangeable in MLU terms."),
]
for title, desc in limitations:
    p = doc.add_paragraph()
    bold_run(p, f"{title}. ")
    p.add_run(desc)

# ============================================================
# 15. LIVE MININET: REQUIRED FINAL STEP
# ============================================================
doc.add_heading("15. Live Mininet Testbed: Required Final Step", level=1)

p = doc.add_paragraph()
bold_run(p, "STATUS: ")
p.add_run("Mininet is NOT installed in the current environment. The SDN integration code exists "
          "(sdn/mininet_testbed.py, phase1_reactive/env/mininet_ryu_adapter.py) but has not been "
          "executed with live packet forwarding.")

doc.add_heading("15.1 What Remains To Run", level=2)
steps = [
    "Install Mininet: sudo apt-get install mininet (requires Linux with OVS)",
    "Install Ryu controller: pip install ryu",
    "Run: sudo python -m sdn.mininet_testbed --topology abilene",
    "Repeat for all 8 topologies",
    "Collect iperf throughput, ping latency, tcpdump packet loss measurements",
    "Measure OVS flow-table update timing (ovs-ofctl dump-flows)",
    "Simulate link failures (ip link set down) and measure recovery time",
]
for s in steps:
    doc.add_paragraph(s, style="List Number")

doc.add_heading("15.2 Metrics That Live Mininet Would Produce", level=2)
mininet_metrics = [
    ("Throughput (Mbps)", "iperf between host pairs under optimized routing", "Per OD pair"),
    ("Packet latency (ms)", "ping RTT under optimized vs ECMP routing", "Per OD pair"),
    ("Packet loss (%)", "tcpdump/iperf loss rate under load", "Per OD pair"),
    ("Jitter (ms)", "ping variance under optimized routing", "Per OD pair"),
    ("Rule install delay (ms)", "Time from controller push to OVS flow active", "Per switch"),
    ("Failure recovery (ms)", "Time from link-down to new routes active", "Per failure event"),
    ("End-to-end delay (ms)", "Total path delay including queueing", "Per flow"),
    ("CDF of real packet delay", "Distribution of measured latencies", "Aggregate"),
]
add_table(
    ["Metric", "How Measured", "Granularity"],
    mininet_metrics, font_size=9,
)
doc.add_paragraph(
    "These metrics are the ONLY way to validate that the LP-computed routing solutions "
    "actually improve real packet delivery. Model-based MLU is a necessary but not sufficient "
    "condition for network performance improvement."
)

# ============================================================
# 16. EXACT METHOD DESCRIPTION
# ============================================================
doc.add_heading("16. Exact Thesis Method Description", level=1)
doc.add_paragraph("Copy-paste-ready method description for the thesis:")
doc.add_paragraph("")

method_blocks = [
    "We propose a two-stage expert selection framework for traffic engineering. An MLP-based "
    "meta-gate (3-layer, 128-128-64-4 with BatchNorm, dropout 0.3) selects among four expert "
    "flow-selection strategies: Bottleneck, Top-K by Demand, Sensitivity Analysis, and "
    "GNN-based selection, all operating with K_crit = 40 critical flows for fair comparison.",

    "Stage 1 -- Zero-shot gate training. The MLP gate is trained on pooled oracle labels from "
    "6 known topologies (Abilene, GEANT, CERNET, Ebone, Sprintlink, Tiscali; ~2,098 training "
    "samples). Oracle labels are obtained by running all four experts through LP optimization "
    "and selecting the one achieving minimum MLU per traffic matrix. Inverse-frequency class "
    "weighting addresses the 65% Bottleneck class imbalance. Germany50 and VtlWavenet2011 are "
    "held out entirely -- no gate weight updates use data from these topologies.",

    "Stage 2 -- Few-shot calibration at deployment. Before evaluating on any topology (known "
    "or unseen), a lightweight calibration phase runs 10 validation traffic matrices through "
    "all 4 experts with LP, counting which expert achieves the lowest MLU. This produces a "
    "topology-specific Bayesian prior with Laplace smoothing. At inference time, the MLP softmax "
    "output is fused with this prior: P_final(i) = P_MLP(i) x prior(i)^alpha / Z, where "
    "alpha = 5 controls calibration influence. No gradient updates occur during calibration -- "
    "only 40 LP evaluations (10 TMs x 4 experts).",

    "This is not a pure zero-shot evaluation. The gate weights are zero-shot with respect to "
    "unseen topologies, but the calibration prior requires 10 topology-specific validation "
    "samples. We characterize this as zero-shot learning with few-shot calibration.",
]
for block in method_blocks:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(block)
    run.italic = True
    run.font.size = Pt(10.5)

# ── SAVE ──
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUTPUT))
print(f"Report saved to {OUTPUT}")
print(f"  Pages: ~{len(doc.paragraphs) // 30} (estimated)")
print(f"  Tables: {len(doc.tables)}")
print(f"  Images: {sum(1 for p in doc.paragraphs if p.runs and any(r._element.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing') for r in p.runs))}")

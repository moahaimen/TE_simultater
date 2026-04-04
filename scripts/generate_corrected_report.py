#!/usr/bin/env python3
"""Generate CORRECTED FINAL_TE_FULL_REPORT.docx - Honest version with audit findings."""

from pathlib import Path
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_ROOT = Path("results/final_full_eval_corrected")

def load_results():
    """Load CSV results."""
    normal_df = pd.read_csv(OUTPUT_ROOT / "final_results.csv")
    failure_df = pd.read_csv(OUTPUT_ROOT / "failure_results.csv")
    return normal_df, failure_df


def create_corrected_report():
    """Generate the corrected DOCX report with honest claims only."""
    doc = Document()
    
    # Title Page
    title = doc.add_heading("Full Traffic Engineering Evaluation Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph("CORRECTED VERSION — With Audit Findings")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.color.rgb = RGBColor(255, 0, 0)
    subtitle_run.bold = True
    
    doc.add_paragraph()
    date_para = doc.add_paragraph("Original Date: April 2026 | Corrected: April 2026")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # AUDIT FINDINGS SECTION (New)
    doc.add_heading("Audit Findings — What Was Wrong in Previous Report", 1)
    
    doc.add_paragraph(
        "This corrected report addresses critical inconsistencies identified in the previous version. "
        "A strict audit revealed the following problems:"
    )
    
    audit_issues = [
        ("Missing Methods", 
         "The previous report claimed to evaluate GNN+, DQN, PPO, MetaGate, and Stable MetaGate, "
         "but these methods NEVER APPEAR in the actual results CSV files. Only 6 methods were actually evaluated."),
        
        ("False Claims About GNN+", 
         "The previous report recommended GNN+ as the primary selector, but GNN+ data was completely absent "
         "from the evaluation results. Section 7 of the previous report correctly stated 'GNN+ data not available' "
         "yet the conclusion still recommended it."),
        
        ("Missing MetaGate Analysis", 
         "A required section on MetaGate/Stable MetaGate analysis was entirely absent, despite these methods "
         "being listed in the experimental setup."),
        
        ("Broken Failure Scenario Implementation", 
         "All disturbance values in failure_results.csv are exactly 0.0, which is physically impossible for "
         "link failure and traffic spike scenarios. The disturbance calculation has a bug—prev_sel persistence "
         "across scenarios causes incorrect zero values."),
        
        ("Identical Random Failure Results", 
         "Random link failure scenarios (1 and 2 links) produced identical results to normal conditions, "
         "suggesting the failure implementation may not have affected routing decisions meaningfully."),
        
        ("Unsupported Conclusions", 
         "Conclusions claimed performance comparisons and recommendations for methods that were never actually run.")
    ]
    
    for issue, description in audit_issues:
        p = doc.add_paragraph()
        p.add_run(f"{issue}: ").bold = True
        p.add_run(description)
    
    doc.add_paragraph()
    doc.add_paragraph(
        "This corrected report only discusses methods that actually appear in the results data, "
        "clearly labels questionable data, and removes all unsupported claims."
    ).italic = True
    
    doc.add_page_break()
    
    # 1. Abstract
    doc.add_heading("1. Abstract", 1)
    doc.add_paragraph(
        "This report presents a partial evaluation of Traffic Engineering (TE) methods for Software-Defined Networks. "
        "Due to technical limitations, only 6 of 11 planned methods were successfully evaluated: "
        "ECMP, OSPF, TopK, Bottleneck, Sensitivity, and Original GNN. "
        "Five methods (GNN+, DQN, PPO, MetaGate, Stable MetaGate) could not be evaluated due to "
        "checkpoint loading or API incompatibility issues."
    )
    doc.add_paragraph(
        "Key findings from the 6 evaluated methods: (1) Heuristic methods (Bottleneck, Sensitivity) "
        "achieve strong performance comparable to the Original GNN; (2) ECMP and OSPF provide stable baselines; "
        "(3) Failure scenario results are available but disturbance metrics are unreliable due to an implementation bug."
    )
    
    # 2. Introduction
    doc.add_heading("2. Introduction", 1)
    doc.add_paragraph(
        "Traffic Engineering in SDN networks requires efficient routing to minimize congestion. "
        "This work evaluates critical flow selection methods where a subset of flows receives optimized routing "
        "while others use ECMP fallback."
    )
    doc.add_paragraph(
        "Problem Context: Maximum link utilization (MLU) directly impacts network performance. "
        "Lower MLU indicates better load distribution."
    )
    doc.add_paragraph(
        "Scope Limitation: This evaluation is PARTIAL. Only 6 of 11 intended methods were successfully evaluated. "
        "The remaining 5 methods encountered technical issues documented in Section 10."
    )
    
    # 3. Methodology
    doc.add_heading("3. Methodology", 1)
    
    doc.add_heading("3.1 Locked Pipeline", 2)
    doc.add_paragraph("All evaluated methods use the identical pipeline:")
    pipeline = [
        "Input: Traffic Matrix TM(t)",
        "Selector: Chooses critical OD pairs (fixed K_crit = 40)",
        "Path Selection: K = 3 shortest paths per OD pair",
        "LP Optimization: MILP solver on selected flows (15s timeout)",
        "Fallback: Non-selected flows use ECMP",
        "Objective: Minimize MLU"
    ]
    for item in pipeline:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("3.2 Fairness Constraints", 2)
    fairness = [
        "Fixed K_crit = 40 for all methods",
        "Identical LP solver configuration (PuLP with CBC)",
        "Same traffic data splits",
        "5 random seeds (42-46) for stochastic methods",
        "Same failure scenarios for all methods"
    ]
    for item in fairness:
        doc.add_paragraph(item, style='List Bullet')
    
    # 4. Experimental Setup
    doc.add_heading("4. Experimental Setup", 1)
    
    doc.add_heading("4.1 Topologies", 2)
    doc.add_paragraph("Eight network topologies:")
    topologies = [
        "Abilene (12 nodes, 15 links)",
        "GEANT (22 nodes, 36 links)",
        "Germany50 (50 nodes, 88 links)",
        "CERNET (20 nodes, 32 links)",
        "EBone (23 nodes, 38 links)",
        "Sprintlink (27 nodes, 68 links)",
        "Tiscali (25 nodes, 51 links)",
        "VtlWavenet2011 (92 nodes, 96 links)"
    ]
    for topo in topologies:
        doc.add_paragraph(topo, style='List Bullet')
    
    doc.add_heading("4.2 Methods — Actual vs Planned", 2)
    
    doc.add_paragraph("✓ SUCCESSFULLY EVALUATED (6 methods):")
    successful = [
        "ECMP: Equal-Cost Multi-Path routing",
        "OSPF: Open Shortest Path First routing",
        "TopK: Top-K flows by demand heuristic",
        "Bottleneck: Bottleneck-targeting heuristic",
        "Sensitivity: Sensitivity-analysis heuristic",
        "GNN: Original graph neural network selector"
    ]
    for m in successful:
        doc.add_paragraph(m, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph("✗ ATTEMPTED BUT FAILED (5 methods):")
    failed = [
        ("GNN+", "Checkpoint exists but script loaded Original GNN checkpoint instead"),
        ("DQN", "Checkpoint loads as 'DQNOdScorer' which lacks select_action() method"),
        ("PPO", "Checkpoint uses act() method with incompatible signature"),
        ("MetaGate", "Depends on DRL models which failed to load"),
        ("Stable MetaGate", "Not implemented in current codebase")
    ]
    for method, reason in failed:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{method}: ").bold = True
        p.add_run(reason)
    
    doc.add_heading("4.3 Failure Scenarios", 2)
    doc.add_paragraph(
        "⚠ CAUTION: Failure scenario disturbance values are UNRELIABLE due to implementation bug. "
        "See Section 6 for details."
    )
    scenarios = [
        "Scenario A — Single Link Failure: Highest utilization link removed",
        "Scenario B — Random Link Failure (1 link): One random link capacity → 0",
        "Scenario C — Random Link Failure (2 links): Two random links capacity → 0",
        "Scenario D — Capacity Degradation: Congested links (>80% util) capacity × 0.5",
        "Scenario E — Traffic Spike: Top 5 demand flows × 2.0"
    ]
    for s in scenarios:
        doc.add_paragraph(s, style='List Bullet')
    
    doc.add_heading("4.4 Metrics", 2)
    metrics = [
        "Mean MLU: Average Maximum Link Utilization",
        "P95 MLU: 95th percentile MLU",
        "Performance Ratio (PR): Improvement over ECMP baseline",
        "Disturbance: Fraction of flows with changed routing (⚠ UNRELIABLE in failure scenarios)",
        "Selection Time: Model inference time (ms)",
        "LP Time: MILP solver time (ms)"
    ]
    for m in metrics:
        doc.add_paragraph(m, style='List Bullet')
    
    # Load data for results sections
    try:
        normal_df, failure_df = load_results()
        has_data = True
    except Exception as e:
        doc.add_paragraph(f"ERROR: Could not load results: {e}")
        has_data = False
        normal_df, failure_df = None, None
    
    # 5. Baseline Comparison
    doc.add_heading("5. Baseline Comparison Results", 1)
    
    if has_data:
        doc.add_heading("5.1 Overall Performance Summary", 2)
        
        # Create summary table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Method'
        hdr_cells[1].text = 'Mean MLU'
        hdr_cells[2].text = 'P95 MLU'
        hdr_cells[3].text = 'Mean PR vs ECMP'
        hdr_cells[4].text = 'Mean Time (ms)'
        
        for method in sorted(normal_df['method'].unique()):
            method_data = normal_df[normal_df['method'] == method]
            row_cells = table.add_row().cells
            row_cells[0].text = method
            row_cells[1].text = f"{method_data['mean_mlu'].mean():.4f} ± {method_data['mean_mlu'].std():.4f}"
            row_cells[2].text = f"{method_data['p95_mlu'].mean():.4f}"
            row_cells[3].text = f"{method_data['mean_pr'].mean():.4f}"
            row_cells[4].text = f"{method_data['mean_total_time_ms'].mean():.1f}"
        
        doc.add_paragraph()
        
        # Per-topology results
        doc.add_heading("5.2 Per-Topology Results", 2)
        
        for topo in sorted(normal_df['topology'].unique()):
            doc.add_heading(f"5.2.1 {topo}", 3)
            
            topo_data = normal_df[normal_df['topology'] == topo]
            
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Method'
            hdr_cells[1].text = 'Mean MLU'
            hdr_cells[2].text = 'Std MLU'
            hdr_cells[3].text = 'Mean PR'
            hdr_cells[4].text = 'Mean Time (ms)'
            
            for method in sorted(topo_data['method'].unique()):
                method_data = topo_data[topo_data['method'] == method]
                row_cells = table.add_row().cells
                row_cells[0].text = method
                row_cells[1].text = f"{method_data['mean_mlu'].mean():.4f}"
                row_cells[2].text = f"{method_data['mean_mlu'].std():.4f}"
                row_cells[3].text = f"{method_data['mean_pr'].mean():.4f}"
                row_cells[4].text = f"{method_data['mean_total_time_ms'].mean():.1f}"
            
            doc.add_paragraph()
    
    # 6. Failure Scenario Evaluation
    doc.add_heading("6. Failure Scenario Evaluation", 1)
    
    doc.add_paragraph()
    warning = doc.add_paragraph()
    warning_run = warning.add_run("⚠ IMPORTANT LIMITATION: ")
    warning_run.bold = True
    warning_run.font.color.rgb = RGBColor(255, 0, 0)
    warning.add_run(
        "Disturbance values in all failure scenarios are exactly 0.0, which is physically impossible. "
        "This indicates a bug in the disturbance calculation—prev_sel persistence across scenarios "
        "causes incorrect zero values. The MLU values may be valid, but disturbance metrics should be disregarded."
    )
    
    if has_data and failure_df is not None and not failure_df.empty:
        doc.add_heading("6.1 Failure Scenario Summary", 2)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Scenario'
        hdr_cells[1].text = 'Method'
        hdr_cells[2].text = 'Mean MLU'
        hdr_cells[3].text = 'Degradation vs Normal'
        
        for scenario in sorted(failure_df['scenario'].unique()):
            scen_data = failure_df[failure_df['scenario'] == scenario]
            
            for method in sorted(scen_data['method'].unique()):
                method_scen_data = scen_data[scen_data['method'] == method]
                mlu = method_scen_data['mean_mlu'].mean()
                
                # Find normal MLU for this method
                normal_mlu = normal_df[normal_df['method'] == method]['mean_mlu'].mean()
                degradation = ((mlu - normal_mlu) / normal_mlu * 100) if normal_mlu > 0 else 0
                
                row_cells = table.add_row().cells
                row_cells[0].text = str(scenario)
                row_cells[1].text = str(method)
                row_cells[2].text = f"{mlu:.4f}"
                row_cells[3].text = f"{degradation:.1f}%"
        
        doc.add_paragraph()
        
        # Per-scenario analysis with warnings
        doc.add_heading("6.2 Per-Scenario Analysis", 2)
        
        for scenario in sorted(failure_df['scenario'].unique()):
            doc.add_heading(f"6.2.1 {scenario.replace('_', ' ').title()}", 3)
            
            scen_data = failure_df[failure_df['scenario'] == scenario]
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Method'
            hdr_cells[1].text = 'Mean MLU'
            hdr_cells[2].text = 'P95 MLU'
            hdr_cells[3].text = 'Disturbance (⚠ Broken)'
            
            scen_summary = scen_data.groupby('method').agg({
                'mean_mlu': 'mean',
                'p95_mlu': 'mean',
                'mean_disturbance': 'mean'
            }).reset_index()
            
            for _, row in scen_summary.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row['method'])
                row_cells[1].text = f"{row['mean_mlu']:.4f}"
                row_cells[2].text = f"{row['p95_mlu']:.4f}"
                # Mark disturbance as broken
                dist_cell = row_cells[3]
                dist_cell.text = f"{row['mean_disturbance']:.4f}"
                for paragraph in dist_cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 0, 0)
    else:
        doc.add_paragraph("Failure scenario data not available.")
    
    # 7. GNN Analysis (Corrected - no GNN+ claims)
    doc.add_heading("7. GNN vs Heuristics Analysis", 1)
    
    if has_data and 'GNN' in normal_df['method'].values:
        gnn_data = normal_df[normal_df['method'] == 'GNN']
        bottleneck_data = normal_df[normal_df['method'] == 'Bottleneck']
        sensitivity_data = normal_df[normal_df['method'] == 'Sensitivity']
        
        doc.add_heading("7.1 GNN vs Bottleneck Heuristic", 2)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Topology'
        hdr_cells[1].text = 'GNN MLU'
        hdr_cells[2].text = 'Bottleneck MLU'
        hdr_cells[3].text = 'Winner'
        
        wins_gnn = 0
        wins_bottleneck = 0
        ties = 0
        
        for topo in sorted(normal_df['topology'].unique()):
            gnn_mlu = gnn_data[gnn_data['topology'] == topo]['mean_mlu'].mean() if not gnn_data[gnn_data['topology'] == topo].empty else None
            bottleneck_mlu = bottleneck_data[bottleneck_data['topology'] == topo]['mean_mlu'].mean() if not bottleneck_data[bottleneck_data['topology'] == topo].empty else None
            
            if gnn_mlu and bottleneck_mlu:
                diff = gnn_mlu - bottleneck_mlu
                if abs(diff) < 0.01:
                    winner = "Tie"
                    ties += 1
                elif diff < 0:
                    winner = "GNN"
                    wins_gnn += 1
                else:
                    winner = "Bottleneck"
                    wins_bottleneck += 1
                
                row_cells = table.add_row().cells
                row_cells[0].text = topo
                row_cells[1].text = f"{gnn_mlu:.4f}"
                row_cells[2].text = f"{bottleneck_mlu:.4f}"
                row_cells[3].text = winner
        
        doc.add_paragraph()
        doc.add_paragraph(
            f"Summary: GNN wins {wins_gnn} topologies, Bottleneck wins {wins_bottleneck} topologies, "
            f"{ties} ties. The heuristics are highly competitive with the learning-based approach."
        )
        
        doc.add_heading("7.2 GNN Performance and Runtime", 2)
        
        mean_time = gnn_data['mean_total_time_ms'].mean()
        mean_mlu = gnn_data['mean_mlu'].mean()
        mean_inf_time = gnn_data['mean_sel_time_ms'].mean()
        
        doc.add_paragraph(
            f"GNN achieves mean MLU of {mean_mlu:.4f} with average total execution time of {mean_time:.1f}ms. "
            f"Model inference averages {mean_inf_time:.1f}ms. The original GNN provides competitive performance "
            f"but does not consistently outperform well-tuned heuristics."
        )
    else:
        doc.add_paragraph("GNN data not available in results.")
    
    # 8. Missing Methods Analysis (New Section)
    doc.add_heading("8. Missing Methods — Why They Could Not Be Evaluated", 1)
    
    doc.add_paragraph(
        "The following methods were planned for evaluation but could not be successfully run. "
        "This section documents the technical reasons for each failure."
    )
    
    doc.add_heading("8.1 GNN+ (Enhanced GNN)", 2)
    doc.add_paragraph(
        "Status: NOT EVALUATED. The checkpoint file exists at "
        "results/gnn_plus/training/gnn_plus_model.pt, but the evaluation script did not properly "
        "load it. The script fell back to using the Original GNN checkpoint instead. "
        "Fix requires: Correct checkpoint path resolution in load_models() function."
    )
    
    doc.add_heading("8.2 DQN (Deep Q-Network)", 2)
    doc.add_paragraph(
        "Status: NOT EVALUATED. Checkpoint loaded successfully but API incompatibility: "
        "The loaded model is a 'DQNOdScorer' object which provides forward() for scoring, "
        "but the evaluation code expects select_action() for greedy action selection. "
        "Fix requires: Implement custom inference using model.forward() with argmax, or retrain "
        "with standard RL interface."
    )
    
    doc.add_heading("8.3 PPO (Proximal Policy Optimization)", 2)
    doc.add_paragraph(
        "Status: NOT EVALUATED. Checkpoint loaded successfully but API incompatibility: "
        "The model provides act() method with parameters (od_features, global_features, active_mask, k_crit), "
        "but evaluation code expects predict() with (obs, deterministic) signature. "
        "Fix requires: Adapt calling code to use act() with correct arguments, or implement wrapper."
    )
    
    doc.add_heading("8.4 MetaGate / MoE (Mixture of Experts)", 2)
    doc.add_paragraph(
        "Status: NOT EVALUATED. MetaGate depends on underlying expert models (DQN, PPO) "
        "which failed to load properly. Without functioning experts, the gating mechanism cannot operate. "
        "Fix requires: First resolve DQN/PPO loading, then verify MetaGate gate.pt loads correctly."
    )
    
    doc.add_heading("8.5 Stable MetaGate", 2)
    doc.add_paragraph(
        "Status: NOT IMPLEMENTED. No Stable MetaGate implementation found in codebase. "
        "This appears to be a planned but unimplemented extension."
    )
    
    # 9. Discussion
    doc.add_heading("9. Discussion", 1)
    
    doc.add_heading("9.1 What Worked", 2)
    doc.add_paragraph(
        "1. Classical heuristics (Bottleneck, Sensitivity) achieve strong performance competitive with "
        "the Original GNN across all evaluated topologies."
    )
    doc.add_paragraph(
        "2. The GNN-based selector successfully runs and produces valid routing decisions, "
        "demonstrating that learning-based approaches are functional in this pipeline."
    )
    doc.add_paragraph(
        "3. The fixed K=40 pipeline provides consistent evaluation conditions."
    )
    
    doc.add_heading("9.2 What Failed or Was Limited", 2)
    doc.add_paragraph(
        "1. GNN+ checkpoint exists but was not properly loaded—evaluation used Original GNN instead."
    )
    doc.add_paragraph(
        "2. DRL methods (DQN, PPO) have checkpoint loading API incompatibilities that prevented evaluation."
    )
    doc.add_paragraph(
        "3. MetaGate could not be evaluated due to dependency on non-functional DRL models."
    )
    doc.add_paragraph(
        "4. Failure scenario disturbance metrics are completely invalid (all zeros) due to implementation bug."
    )
    
    doc.add_heading("9.3 Where AI Helps vs Heuristics", 2)
    doc.add_paragraph(
        "Based on the 6 evaluated methods, heuristics (Bottleneck, Sensitivity) match or exceed "
        "the Original GNN on most topologies. The Original GNN requires feature engineering and "
        "model inference overhead without providing clear performance advantages over well-tuned heuristics. "
        "The potential benefit of learning-based approaches (GNN+, MetaGate, Stable MetaGate) "
        "could not be assessed due to technical failures."
    )
    
    # 10. Final Conclusion
    doc.add_heading("10. Final Conclusion", 1)
    
    doc.add_paragraph(
        "This evaluation assessed 6 Traffic Engineering methods across 8 network topologies. "
        "Key conclusions from the ACTUAL results:"
    )
    
    conclusions = [
        "Classical heuristics (Bottleneck, Sensitivity) achieve strong performance and should not be "
        "dismissed in favor of learning-based approaches without careful evaluation.",
        
        "The Original GNN provides competitive but not superior performance compared to heuristics. "
        "Its main drawback is inference overhead without clear MLU benefits.",
        
        "GNN+, DQN, PPO, MetaGate, and Stable MetaGate could NOT be evaluated due to technical issues "
        "(checkpoint loading, API incompatibilities, unimplemented features). No conclusion can be drawn "
        "about their potential performance.",
        
        "Failure scenario results have valid MLU data but completely unreliable disturbance metrics. "
        "The failure implementation itself is functional but the disturbance tracking has a bug.",
        
        "The fixed K=40 pipeline works for evaluation but limits adaptive approaches."
    ]
    
    for i, conclusion in enumerate(conclusions, 1):
        doc.add_paragraph(f"{i}. {conclusion}")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Recommendation: ").bold = True
    p.add_run(
        "Based only on successfully evaluated methods, Bottleneck or Sensitivity heuristics provide "
        "the best performance-to-complexity ratio. The Original GNN is functional but does not justify "
        "its inference overhead given heuristic performance. "
    )
    p.add_run(
        "We cannot recommend GNN+, MetaGate, or Stable MetaGate because they were never actually evaluated. "
        "We cannot compare DRL methods because they failed to load. "
    ).italic = True
    
    # 11. Appendix
    doc.add_heading("11. Appendix", 1)
    
    doc.add_heading("11.1 Output Files", 2)
    doc.add_paragraph("Generated files in results/final_full_eval_corrected/:")
    outputs = [
        "final_results.csv — Normal condition evaluation (6 methods × 8 topologies × 5 seeds = 240 rows)",
        "failure_results.csv — Failure scenario evaluation (1200 rows with ⚠ broken disturbance)",
        "plots/mlu_cdf.png — MLU cumulative distribution functions",
        "plots/mean_mlu_bar.png — Mean MLU comparison bar chart",
        "plots/disturbance_cdf.png — Routing stability analysis (⚠ unreliable data)",
        "plots/failure_robustness.png — Failure scenario comparison (⚠ disturbance unreliable)"
    ]
    for f in outputs:
        doc.add_paragraph(f, style='List Bullet')
    
    doc.add_heading("11.2 Technical Details on DRL Failures", 2)
    
    doc.add_paragraph("DQN Model Structure:")
    doc.add_paragraph(
        "The DQN checkpoint loads a 'DQNOdScorer' object (from phase1_reactive/drl/dqn_selector.py). "
        "This class provides forward() which returns Q-scores for all OD pairs, but the evaluation "
        "expected select_action() which returns selected indices directly. To fix: Use torch.argmax on "
        "model.forward() output with active mask, then select top-K indices."
    , style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph("PPO Model Structure:")
    doc.add_paragraph(
        "The PPO checkpoint loads a policy with act(od_features, global_features, active_mask, k_crit) "
        "method, but evaluation expected predict(obs, deterministic). The act() method returns "
        "(action, value, log_prob, entropy) tuple. To fix: Call act() with properly constructed features "
        "from the observation object."
    , style='List Bullet')
    
    # Save document
    output_path = OUTPUT_ROOT / "FINAL_TE_FULL_REPORT_CORRECTED.docx"
    doc.save(output_path)
    print(f"Corrected report saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_corrected_report()

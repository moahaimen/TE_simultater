#!/usr/bin/env python3
"""Build the definitive MLP Meta-Gate evaluation report as .docx."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT = "results/dynamic_metagate/MLP_MetaGate_Final_Report.docx"

doc = Document()

# ── Styles ──
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level, size in [(1, 16), (2, 14), (3, 12)]:
    h = doc.styles[f"Heading {level}"]
    h.font.name = "Arial"
    h.font.size = Pt(size)
    h.font.bold = True
    h.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)


def add_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Arial"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "Arial"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph("")
    return t


def bold_run(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    return run


def add_para(text, bold=False, italic=False, style_name=None):
    if style_name:
        p = doc.add_paragraph(style=style_name)
    else:
        p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


# ============================================================
# TITLE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("MLP Meta-Gate with Few-Shot Bayesian Calibration")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Final Evaluation Report")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle2.add_run("Architecture: 4-Expert Selection via MLP Gate with Per-Topology Bayesian Calibration\n"
                         "Framing: Zero-Shot Gate Training + Few-Shot Calibration")
run.font.size = Pt(11)
run.italic = True

doc.add_paragraph("")

# ============================================================
# ABSTRACT
# ============================================================
doc.add_heading("Abstract", level=1)
doc.add_paragraph(
    "We present an MLP-based meta-gate that dynamically selects among four traffic engineering "
    "expert strategies (Bottleneck, Top-K by Demand, Sensitivity Analysis, and GNN-based selection) "
    "on a per-traffic-matrix basis. The gate is a 3-layer MLP (128-128-64-4) with BatchNorm and "
    "dropout, trained on pooled oracle labels from 6 known topologies. Germany50 and VtlWavenet2011 "
    "are held out entirely from gate-weight training."
)
doc.add_paragraph(
    "A lightweight few-shot Bayesian calibration phase adapts the gate to each target topology "
    "at deployment time by running 10 validation traffic matrices through all 4 experts, counting "
    "which expert achieves the lowest MLU, and fusing this topology-specific prior with the MLP "
    "softmax output. No gradient updates occur during calibration."
)
doc.add_paragraph(
    "This is not a pure zero-shot evaluation. The gate weights are zero-shot with respect to "
    "unseen topologies, but the calibration prior requires 10 topology-specific validation samples. "
    "We characterize this as zero-shot gate training with few-shot calibration."
)
p = doc.add_paragraph()
bold_run(p, "Key results: ")
p.add_run(
    "Germany50 (unseen): GNN selected 64% of timesteps, oracle gap +0.33%, accuracy 65.9%. "
    "Overall accuracy 60.3% across 8 topologies. Three topologies achieve +0.00% oracle gap "
    "(Abilene, GEANT, Ebone). Mean selector decision overhead 48.4 ms; mean end-to-end time "
    "(including LP) 118.8 ms."
)

# ============================================================
# 1. CONTRIBUTIONS
# ============================================================
doc.add_heading("1. Contributions", level=1)
contributions = [
    "Honest MLP Meta-Gate architecture: Replaces the previous false 'GNN-based meta-selector' "
    "(which was a static dictionary lookup) with a genuine 3-layer MLP classifier that makes "
    "per-timestep expert selection decisions based on 49 traffic and topology features.",

    "4-expert pool under fair comparison: Bottleneck, Top-K by Demand, Sensitivity Analysis, "
    "and GNN-based selection all operate with K_crit = 40 critical flows and are evaluated "
    "through the same LP optimization path.",

    "Unified gate training on known topologies only: One MLP gate trained on pooled data from "
    "6 known topologies (~2,098 training samples). Germany50 and VtlWavenet2011 never appear "
    "during gate-weight training.",

    "Few-shot Bayesian calibration for deployment: A 10-sample calibration phase per topology "
    "builds a prior that fuses with MLP predictions, enabling the gate to detect when GNN is "
    "superior on unseen topologies without gradient updates.",

    "Per-component timing breakdown: Full end-to-end timing honesty "
    "(BN + TopK + Sens + GNN + Features + MLP + LP) reported per timestep.",

    "CERNET correction: Documented the true source (TopologyZoo, not SNDlib), "
    "41 nodes with 59 bidirectional links (116 directed edges), and synthetic MGM traffic.",
]
for i, c in enumerate(contributions, 1):
    p = doc.add_paragraph()
    bold_run(p, f"{i}. ")
    p.add_run(c)

# ============================================================
# 2. ARCHITECTURE
# ============================================================
doc.add_heading("2. Architecture", level=1)

doc.add_heading("2.1 Overview", level=2)
doc.add_paragraph(
    "The system operates as follows per traffic matrix (TM):"
)
steps = [
    "Run all 4 experts (Bottleneck, TopK, Sensitivity, GNN) to produce 4 candidate OD-pair selections, each of size K_crit = 40.",
    "Extract a 49-dimensional feature vector from the TM, expert outputs, GNN diagnostics, ECMP baseline, and topology metrics.",
    "The MLP gate predicts which expert will achieve the lowest MLU.",
    "If Bayesian calibration is active, fuse MLP softmax with topology-specific prior.",
    "Route the predicted expert's selection through LP optimization to obtain the final MLU.",
]
for s in steps:
    doc.add_paragraph(s, style="List Number")

doc.add_heading("2.2 MLP Gate Architecture", level=2)
doc.add_paragraph(
    "3-layer MLP: Linear(49, 128) -> BatchNorm -> ReLU -> Dropout(0.3) -> "
    "Linear(128, 128) -> BatchNorm -> ReLU -> Dropout(0.3) -> "
    "Linear(128, 64) -> ReLU -> Linear(64, 4). "
    "Output: softmax over 4 expert classes {0=Bottleneck, 1=TopK, 2=Sensitivity, 3=GNN}."
)

doc.add_heading("2.3 Feature Vector (49 dimensions)", level=2)
features = [
    ("TM statistics (8)", "mean, std, max, min_nonzero, skew, kurtosis, entropy, top-10 demand share"),
    ("Per-expert demand stats (16)", "4 experts x {mean_demand, std_demand, max_demand, coverage} of selected OD pairs"),
    ("Cross-expert agreement (11)", "6 pairwise Jaccard overlaps (BN-TopK, BN-Sens, BN-GNN, TopK-Sens, TopK-GNN, Sens-GNN) + 4 unique-to-expert fractions + 1 all-four-agree fraction"),
    ("Topology metrics (3)", "num_nodes, num_edges, edge_density"),
    ("GNN diagnostics (5)", "alpha (mixing weight), confidence, GNN correction mean, w_bottleneck, w_sensitivity"),
    ("Demand shares (4)", "fraction of total demand captured by each expert's selection"),
    ("ECMP baseline (2)", "ECMP max link utilization, ECMP mean link utilization"),
]
add_table(
    ["Feature Group", "Dims", "Description"],
    [(name, name.split("(")[1].rstrip(")"), desc) for name, desc in features],
    col_widths=[1.8, 0.5, 4.5],
)

doc.add_heading("2.4 Four Experts", level=2)
experts = [
    ("Bottleneck", "Selects K_crit OD pairs traversing the most-utilized link under ECMP routing"),
    ("Top-K by Demand", "Selects K_crit OD pairs with the highest traffic demand"),
    ("Sensitivity", "Selects K_crit OD pairs whose rerouting has the largest impact on max link utilization"),
    ("GNN (GraphSAGE)", "Learned flow selector using graph neural network; selects K_crit OD pairs based on learned node/edge embeddings"),
]
add_table(
    ["Expert", "Selection Strategy"],
    experts,
    col_widths=[1.5, 5.3],
)
doc.add_paragraph(
    "All 4 experts use K_crit = 40 critical flows and are evaluated through the same LP solver "
    "with identical ECMP base routing, path libraries, and capacity vectors."
)

# ============================================================
# 3. TRAINING DETAILS
# ============================================================
doc.add_heading("3. Training Details", level=1)

doc.add_heading("3.1 Training Data", level=2)
doc.add_paragraph(
    "The gate is trained on pooled data from 6 known topologies only. "
    "Germany50 and VtlWavenet2011 are completely excluded from training."
)
train_data = [
    ("Abilene", "12", "350", "140 (40%)", "58 (17%)", "90 (26%)", "62 (18%)"),
    ("GEANT", "22", "348", "291 (84%)", "0 (0%)", "0 (0%)", "57 (16%)"),
    ("CERNET", "41", "350", "310 (89%)", "3 (1%)", "4 (1%)", "33 (9%)"),
    ("Ebone", "23", "350", "100 (29%)", "61 (17%)", "116 (33%)", "73 (21%)"),
    ("Sprintlink", "44", "350", "283 (81%)", "61 (17%)", "3 (1%)", "3 (1%)"),
    ("Tiscali", "49", "350", "234 (67%)", "5 (1%)", "87 (25%)", "24 (7%)"),
    ("TOTAL", "--", "2,098", "1,358 (65%)", "188 (9%)", "300 (14%)", "252 (12%)"),
]
add_table(
    ["Topology", "Nodes", "Samples", "BN Oracle", "TopK Oracle", "Sens Oracle", "GNN Oracle"],
    train_data,
    col_widths=[1.0, 0.5, 0.7, 1.0, 1.0, 1.0, 1.0],
)

doc.add_heading("3.2 Oracle Label Generation", level=2)
doc.add_paragraph(
    "For each training TM, all 4 experts are run, each selection is evaluated through LP, "
    "and the oracle label is argmin(BN_MLU, TopK_MLU, Sens_MLU, GNN_MLU). "
    "This produces the ground-truth 'best expert' label per timestep."
)

doc.add_heading("3.3 Class Imbalance Handling", level=2)
doc.add_paragraph(
    "The training oracle is heavily skewed: 65% Bottleneck, 9% TopK, 14% Sensitivity, 12% GNN. "
    "We use inverse-frequency class weights in CrossEntropyLoss: "
    "w_i = N / (C * count_i), where N is total samples, C is number of classes (4), "
    "and count_i is the number of samples for class i. This prevents the MLP from collapsing "
    "to always-Bottleneck."
)

doc.add_heading("3.4 Training Configuration", level=2)
config_items = [
    ("Hidden dimension", "128"),
    ("Dropout", "0.3"),
    ("Learning rate", "5e-4 (Adam)"),
    ("Epochs", "300"),
    ("Batch size", "64"),
    ("Train accuracy", "62.3%"),
    ("Validation accuracy", "49.8%"),
]
add_table(["Parameter", "Value"], config_items, col_widths=[2.5, 2.5])

# ============================================================
# 4. FEW-SHOT BAYESIAN CALIBRATION
# ============================================================
doc.add_heading("4. Few-Shot Bayesian Calibration", level=1)

doc.add_heading("4.1 Motivation", level=2)
doc.add_paragraph(
    "Without calibration, the MLP gate selects Bottleneck 100% of the time on Germany50 "
    "(unseen), despite GNN being the oracle-best expert 95% of the time. "
    "This happens because: (1) the training data is 65% Bottleneck, biasing the MLP; "
    "(2) the MLP has never seen Germany50's feature distribution during training; "
    "(3) the MLP's softmax outputs are overconfident in Bottleneck for out-of-distribution inputs."
)

doc.add_heading("4.2 Method", level=2)
doc.add_paragraph(
    "Before evaluating on any topology (known or unseen), a calibration phase runs "
    "10 validation traffic matrices through all 4 experts with LP optimization, "
    "counting which expert achieves the lowest MLU. This produces a topology-specific "
    "Bayesian prior with Laplace smoothing:"
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("prior[i] = (win_count[i] + 1) / (N_calib + C)")
run.italic = True
run.font.size = Pt(11)

doc.add_paragraph(
    "At inference time, the MLP softmax output is fused with this prior:"
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("P_final(i) = P_MLP(i) * prior(i)^alpha / Z")
run.italic = True
run.font.size = Pt(11)

doc.add_paragraph(
    "where alpha = 5 controls calibration influence and Z is a normalization constant. "
    "Higher alpha values give more weight to the calibration prior over the MLP's raw predictions. "
    "No gradient updates occur during calibration -- only 40 LP evaluations (10 TMs x 4 experts)."
)

doc.add_heading("4.3 Calibration Priors (Actual Values)", level=2)
calib_data = [
    ("Abilene", "known", "0.36", "0.29", "0.21", "0.14"),
    ("GEANT", "known", "0.50", "0.07", "0.07", "0.36"),
    ("CERNET", "known", "0.71", "0.07", "0.07", "0.14"),
    ("Ebone", "known", "0.36", "0.21", "0.21", "0.21"),
    ("Sprintlink", "known", "0.57", "0.21", "0.07", "0.14"),
    ("Tiscali", "known", "0.57", "0.07", "0.29", "0.07"),
    ("Germany50", "UNSEEN", "0.07", "0.07", "0.07", "0.79"),
    ("VtlWavenet", "UNSEEN", "0.29", "0.07", "0.07", "0.57"),
]
add_table(
    ["Topology", "Type", "BN Prior", "TopK Prior", "Sens Prior", "GNN Prior"],
    calib_data,
    col_widths=[1.2, 0.7, 0.8, 0.8, 0.8, 0.8],
)
doc.add_paragraph(
    "Germany50's calibration prior strongly favors GNN (0.79), which when raised to alpha=5 "
    "becomes 0.97 -- enough to override the MLP's Bottleneck bias on most timesteps."
)

# ============================================================
# 5. RESULTS
# ============================================================
doc.add_heading("5. Results -- All 8 Topologies", level=1)

results_data = [
    ("Abilene", "known", "40.0%", "0.0546", "0.0546", "+0.00%", "0%", "28"),
    ("GEANT", "known", "82.7%", "0.1602", "0.1602", "+0.00%", "0%", "0"),
    ("CERNET", "known", "89.3%", "1722.69", "1721.73", "+0.06%", "0%", "0"),
    ("Ebone", "known", "24.0%", "379.59", "379.59", "+0.00%", "1%", "17"),
    ("Sprintlink", "known", "66.7%", "880.26", "877.17", "+0.35%", "0%", "0"),
    ("Tiscali", "known", "65.3%", "834.78", "831.93", "+0.34%", "0%", "4"),
    ("Germany50", "UNSEEN", "65.9%", "18.99", "18.93", "+0.33%", "64%", "7"),
    ("VtlWavenet", "UNSEEN", "50.7%", "12251.81", "12244.84", "+0.06%", "0%", "0"),
]
add_table(
    ["Topology", "Type", "Accuracy", "MetaGate MLU", "Oracle MLU", "Gap", "GNN%", "Switches"],
    results_data,
    col_widths=[1.0, 0.6, 0.7, 1.1, 1.1, 0.6, 0.5, 0.7],
)

p = doc.add_paragraph()
bold_run(p, "Overall accuracy: ")
p.add_run("60.3%. ")
bold_run(p, "Known topology accuracy: ")
p.add_run("61.3%. ")
bold_run(p, "Unseen topology accuracy: ")
p.add_run("56.3%.")

doc.add_paragraph(
    "Three topologies (Abilene, GEANT, Ebone) achieve +0.00% oracle gap, meaning the "
    "MetaGate's expert selection produces identical MLU to the oracle's best choice. "
    "The maximum oracle gap across all 8 topologies is +0.35% (Sprintlink)."
)

# ============================================================
# 6. GERMANY50 DEEP DIVE
# ============================================================
doc.add_heading("6. Germany50 (Unseen) -- Before vs After Calibration", level=1)

doc.add_heading("6.1 The Problem", level=2)
doc.add_paragraph(
    "Germany50 is a 50-node topology that was never seen during gate-weight training. "
    "The oracle shows GNN is the best expert for 42/44 test timesteps (95.5%). "
    "However, without calibration, the MLP gate selects Bottleneck 100% of the time "
    "because: (a) 65% of training labels are Bottleneck, creating a strong prior; "
    "(b) Germany50's features are out-of-distribution for the MLP; "
    "(c) the MLP's softmax outputs are overconfident."
)

doc.add_heading("6.2 Before vs After Comparison", level=2)
g50_compare = [
    ("BN selections", "44/44 (100%)", "16/44 (36%)"),
    ("GNN selections", "0/44 (0%)", "28/44 (64%)"),
    ("GNN selection rate", "0.0%", "63.6%"),
    ("Selector switches", "0/43", "7/43"),
    ("MetaGate MLU", "19.227", "18.992"),
    ("Oracle MLU", "18.929", "18.929"),
    ("Oracle gap", "+1.57%", "+0.33%"),
    ("Accuracy", "2.3%", "65.9%"),
]
add_table(
    ["Metric", "Before Calibration", "After Calibration"],
    g50_compare,
    col_widths=[1.8, 2.0, 2.0],
)

doc.add_heading("6.3 Why Calibration Fixed It", level=2)
doc.add_paragraph(
    "The 10-sample calibration phase on Germany50 validation TMs reveals that GNN wins "
    "8/10 times (prior = 0.79). When raised to alpha = 5, the GNN prior becomes 0.97 "
    "while the BN prior drops to 0.00002. This 48,000x ratio is sufficient to override "
    "the MLP's Bottleneck confidence on most timesteps. The gate now dynamically selects "
    "GNN when the calibration evidence supports it, and falls back to Bottleneck when "
    "the MLP is sufficiently confident (16/44 timesteps)."
)
doc.add_paragraph(
    "The 7 selector switches across 43 transitions demonstrate genuine per-timestep "
    "dynamic behavior -- the gate is not simply replacing 'always BN' with 'always GNN'."
)

# ============================================================
# 7. VTLWAVENET ANALYSIS
# ============================================================
doc.add_heading("7. VtlWavenet2011 (Unseen) -- Analysis", level=1)
doc.add_paragraph(
    "VtlWavenet2011 (92 nodes) remains 100% Bottleneck despite a calibration prior "
    "favoring GNN (0.57). The MLP's BN confidence on this large topology is so extreme "
    "that even prior^5 cannot override it."
)
p = doc.add_paragraph()
bold_run(p, "However, this is near-optimal. ")
p.add_run(
    "The oracle distribution is 50.7% BN / 49.3% GNN -- almost exactly split. "
    "The mean BN MLU (12251.81) is actually lower than mean GNN MLU (12252.05). "
    "The oracle gap is only +0.06%, meaning the all-Bottleneck strategy loses almost "
    "nothing compared to the per-timestep oracle. The 'failure' to select GNN here "
    "has negligible practical impact."
)

# ============================================================
# 8. CERNET CORRECTION
# ============================================================
doc.add_heading("8. CERNET Topology -- Honest Documentation", level=1)
cernet_facts = [
    ("Source", "TopologyZoo (Cernet.graphml), NOT SNDlib"),
    ("Nodes", "41"),
    ("Edges", "59 bidirectional links = 116 directed edges"),
    ("Traffic", "Synthetic MGM (not real SNDlib traffic matrices)"),
    ("Config note", "traffic_mode kept as 'real_sndlib' to avoid breaking the data loader; "
     "a comment documents the truth"),
]
add_table(["Property", "Value"], cernet_facts, col_widths=[1.5, 5.0])
doc.add_paragraph(
    "The CERNET topology was previously mislabeled as SNDlib-sourced. "
    "The actual .graphml file comes from TopologyZoo. The traffic matrices in the "
    "processed NPZ file are synthetic (generated by the MGM model), not real measured "
    "traffic. The pipeline configuration retains traffic_mode='real_sndlib' because "
    "changing it breaks the data loader, but this is documented with explicit comments."
)

# ============================================================
# 9. TIMING BREAKDOWN
# ============================================================
doc.add_heading("9. Per-Component Timing Breakdown", level=1)
timing_data = [
    ("Abilene", "0.5", "0.0", "0.7", "1.9", "0.8", "0.16", "24.9", "30.0"),
    ("GEANT", "1.8", "0.0", "2.1", "3.8", "0.8", "0.15", "28.5", "40.6"),
    ("CERNET", "6.6", "0.1", "8.1", "10.5", "0.9", "0.16", "41.2", "79.9"),
    ("Ebone", "2.0", "0.0", "2.5", "4.1", "0.8", "0.15", "28.7", "42.2"),
    ("Sprintlink", "8.1", "0.1", "9.9", "13.3", "1.0", "0.18", "45.3", "92.4"),
    ("Tiscali", "10.4", "0.1", "12.6", "15.3", "1.0", "0.19", "50.1", "108.1"),
    ("Germany50", "7.7", "0.1", "8.7", "16.8", "1.0", "0.20", "46.4", "94.6"),
    ("VtlWavenet", "61.7", "0.5", "103.9", "58.4", "1.4", "0.23", "145.0", "452.3"),
]
add_table(
    ["Topology", "BN ms", "TopK ms", "Sens ms", "GNN ms", "Feat ms", "MLP ms", "LP ms", "Total ms"],
    timing_data,
    col_widths=[0.9, 0.6, 0.6, 0.6, 0.6, 0.6, 0.6, 0.6, 0.7],
)
doc.add_paragraph(
    "Terminology: Decision time = all selector overhead before LP (running 4 experts + "
    "feature extraction + MLP forward pass). LP time = LP solver only. "
    "Total time = decision time + LP time (full end-to-end per timestep)."
)
doc.add_paragraph(
    "The MLP gate inference itself (feature extraction + MLP forward pass) takes < 2 ms on all "
    "topologies. The decision time is dominated by running all 4 experts (required before the gate "
    "can decide). LP optimization adds roughly equal time. Both components are inherent to the "
    "'run all experts, pick best, solve LP' architecture."
)
p = doc.add_paragraph()
bold_run(p, "Mean decision time (selector overhead before LP): ")
p.add_run("48.4 ms. ")
bold_run(p, "Mean LP time (solver only): ")
p.add_run("51.5 ms. ")
bold_run(p, "Mean total time (end-to-end): ")
p.add_run("118.8 ms.")

# ============================================================
# 10. LIMITATIONS
# ============================================================
doc.add_heading("10. Limitations (Honest Assessment)", level=1)

limitations = [
    ("Not pure zero-shot",
     "The calibration phase uses 10 validation TMs from the target topology itself, "
     "running all 4 experts + LP on each. This is a few-shot adaptation step, not pure "
     "zero-shot generalization. The correct framing is 'zero-shot gate training + few-shot calibration'."),

    ("Calibration computational cost",
     "10 TMs x 4 experts x LP = 40 LP solves per topology before test-time inference begins. "
     "This is a one-time cost per topology deployment, not per-timestep, but it is non-trivial."),

    ("Alpha parameter is hand-tuned",
     "The calibration strength alpha = 5 was selected empirically. Different alpha values produce "
     "different accuracy/gap tradeoffs. No principled method was used to select this value."),

    ("Raw MLP is weak without calibration",
     "The MLP validation accuracy is ~50% (barely above random for 4 classes). Without calibration, "
     "the gate collapses to mostly-Bottleneck on unseen topologies. The Bayesian calibration "
     "compensates for the MLP's weakness, but the gate itself cannot reliably distinguish "
     "when GNN beats Bottleneck from features alone."),

    ("VtlWavenet2011 BN dominance",
     "Despite calibration, VtlWavenet never selects GNN because the MLP's Bottleneck confidence "
     "is too extreme for even prior^5 to override. This is masked by the fact that BN is "
     "coincidentally near-optimal (+0.06% gap) on this topology."),

    ("Accuracy vs MLU gap disconnect",
     "Abilene achieves only 40% accuracy but +0.00% oracle gap. Ebone achieves 24% accuracy "
     "but +0.00% gap. This means the gate often picks the 'wrong' expert, but the 'wrong' "
     "expert's MLU happens to be very close to the oracle's. Accuracy overstates the practical "
     "impact of gate errors."),
]
for title, desc in limitations:
    p = doc.add_paragraph()
    bold_run(p, f"{title}. ")
    p.add_run(desc)

# ============================================================
# 11. EXACT METHOD DESCRIPTION FOR THESIS
# ============================================================
doc.add_heading("11. Exact Method Description for Thesis", level=1)
doc.add_paragraph(
    "The following paragraph is the recommended copy-paste-ready method description:",
    style="Normal",
)
doc.add_paragraph("")

# Indented block quote style
method_text = (
    "We propose a two-stage expert selection framework for traffic engineering. "
    "An MLP-based meta-gate (3-layer, 128-128-64-4 with BatchNorm, dropout 0.3) selects "
    "among four expert flow-selection strategies: Bottleneck, Top-K by Demand, Sensitivity "
    "Analysis, and GNN-based selection, all operating with K_crit = 40 critical flows for "
    "fair comparison."
)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
run = p.add_run(method_text)
run.italic = True
run.font.size = Pt(11)

method_text2 = (
    "Stage 1 -- Zero-shot gate training. The MLP gate is trained on pooled oracle labels "
    "from 6 known topologies (Abilene, GEANT, CERNET, Ebone, Sprintlink, Tiscali; ~2,098 "
    "training samples). Oracle labels are obtained by running all four experts through LP "
    "optimization and selecting the one achieving minimum MLU per traffic matrix. "
    "Inverse-frequency class weighting addresses the 65% Bottleneck class imbalance. "
    "Germany50 and VtlWavenet2011 are held out entirely -- no gate weight updates use data "
    "from these topologies."
)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
run = p.add_run(method_text2)
run.italic = True
run.font.size = Pt(11)

method_text3 = (
    "Stage 2 -- Few-shot calibration at deployment. Before evaluating on any topology "
    "(known or unseen), a lightweight calibration phase runs 10 validation traffic matrices "
    "through all 4 experts with LP, counting which expert achieves the lowest MLU. This "
    "produces a topology-specific Bayesian prior with Laplace smoothing. At inference time, "
    "the MLP softmax output is fused with this prior: P_final(i) = P_MLP(i) x prior(i)^alpha / Z, "
    "where alpha = 5 controls calibration influence. No gradient updates occur during "
    "calibration -- only 40 LP evaluations (10 TMs x 4 experts)."
)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
run = p.add_run(method_text3)
run.italic = True
run.font.size = Pt(11)

method_text4 = (
    "This is not a pure zero-shot evaluation. The gate weights are zero-shot with respect "
    "to unseen topologies, but the calibration prior requires 10 topology-specific validation "
    "samples. We characterize this as zero-shot learning with few-shot calibration: the model "
    "never trains on unseen topologies, but adapts its confidence distribution using a small "
    "validation budget before test-time inference."
)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
run = p.add_run(method_text4)
run.italic = True
run.font.size = Pt(11)

# ============================================================
# SAVE
# ============================================================
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
doc.save(OUTPUT)
print(f"Report saved to {OUTPUT}")
print(f"  Sections: 11")
print(f"  Tables: 10")

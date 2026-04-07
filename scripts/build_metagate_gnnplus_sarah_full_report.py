#!/usr/bin/env python3
"""Build a Sarah-style full MetaGate+GNN+ report from the integrated bundle."""

from __future__ import annotations

import json
import math
import os
from datetime import datetime
from pathlib import Path

REPORT_CACHE_ROOT = Path("/tmp") / "metagate_gnnplus_sarah_full_cache"
REPORT_CACHE_ROOT.mkdir(parents=True, exist_ok=True)
os.environ.setdefault("MPLCONFIGDIR", str(REPORT_CACHE_ROOT / "mplconfig"))
os.environ.setdefault("XDG_CACHE_HOME", str(REPORT_CACHE_ROOT / "xdg_cache"))

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parent.parent
os.chdir(PROJECT_ROOT)

OUTPUT_DIR = Path(
    os.environ.get(
        "METAGATE_GNNPLUS_REPORT_INPUT_DIR",
        str(PROJECT_ROOT / "results" / "dynamic_metagate_gnnplus_sarah_full"),
    )
).resolve()
PLOTS_DIR = OUTPUT_DIR / "plots"
REPORT_DOC = OUTPUT_DIR / os.environ.get(
    "METAGATE_GNNPLUS_REPORT_NAME",
    "AI_Driven_TE_MLP_MetaGate_GNNPlus_Full_Report.docx",
)
AUDIT_MD = OUTPUT_DIR / os.environ.get(
    "METAGATE_GNNPLUS_REPORT_AUDIT",
    "report_audit_sarah_full.md",
)

BASELINE_NORMAL_CSV = PROJECT_ROOT / "results" / "final_full_eval_corrected" / "final_results.csv"
GNNPLUS_STAGE1_SUMMARY = PROJECT_ROOT / "results" / "gnn_plus" / "stage1_regularization" / "training_d02" / "summary.json"
GNNPLUS_SUP_LOG = PROJECT_ROOT / "results" / "gnn_plus" / "stage1_regularization" / "training_d02" / "supervised_log.csv"
GNNPLUS_RL_LOG = PROJECT_ROOT / "results" / "gnn_plus" / "stage1_regularization" / "training_d02" / "reinforce_log.csv"

SUMMARY_CSV = OUTPUT_DIR / "metagate_summary.csv"
RESULTS_CSV = OUTPUT_DIR / "metagate_results.csv"
DECISIONS_CSV = OUTPUT_DIR / "metagate_decisions.csv"
TIMING_CSV = OUTPUT_DIR / "metagate_timing.csv"
SDN_METRICS_CSV = OUTPUT_DIR / "metagate_sdn_metrics.csv"
TRAINING_SUMMARY_JSON = OUTPUT_DIR / "training_summary.json"
TRAIN_DIST_CSV = OUTPUT_DIR / "train_oracle_distribution.csv"
CALIBRATION_CSV = OUTPUT_DIR / "calibration_priors.csv"
FAILURE_RESULTS_CSV = OUTPUT_DIR / "metagate_failure_results.csv"
FAILURE_SUMMARY_CSV = OUTPUT_DIR / "metagate_failure_summary.csv"
FAILURE_SDN_CSV = OUTPUT_DIR / "metagate_failure_sdn_metrics.csv"
ZERO_SHOT_UNSEEN_CSV = OUTPUT_DIR / "zero_shot_unseen_summary.csv"

TOPOLOGY_ORDER = [
    "abilene",
    "cernet",
    "geant",
    "ebone",
    "sprintlink",
    "tiscali",
    "germany50",
    "vtlwavenet2011",
]

TOPOLOGY_INFO = {
    "abilene": {"display": "Abilene", "type": "known", "nodes": 12},
    "cernet": {"display": "CERNET", "type": "known", "nodes": 41},
    "geant": {"display": "GEANT", "type": "known", "nodes": 22},
    "ebone": {"display": "Ebone", "type": "known", "nodes": 23},
    "sprintlink": {"display": "Sprintlink", "type": "known", "nodes": 44},
    "tiscali": {"display": "Tiscali", "type": "known", "nodes": 49},
    "germany50": {"display": "Germany50", "type": "UNSEEN", "nodes": 50},
    "vtlwavenet2011": {"display": "VtlWavenet2011", "type": "UNSEEN", "nodes": 92},
}

TOPOLOGY_ALIASES = {
    "abilene": "abilene",
    "abilene_backbone": "abilene",
    "cernet": "cernet",
    "geant": "geant",
    "rocketfuel_ebone": "ebone",
    "ebone": "ebone",
    "rocketfuel_sprintlink": "sprintlink",
    "sprintlink": "sprintlink",
    "rocketfuel_tiscali": "tiscali",
    "tiscali": "tiscali",
    "germany50": "germany50",
    "topologyzoo_vtlwavenet2011": "vtlwavenet2011",
    "vtlwavenet2011": "vtlwavenet2011",
}

METHOD_COLORS = {
    "metagate": "#1f77b4",
    "bottleneck": "#2ca02c",
    "topk": "#9467bd",
    "sensitivity": "#8c564b",
    "gnnplus": "#ff7f0e",
    "ecmp": "#7f7f7f",
    "ospf": "#17becf",
    "gnn": "#d62728",
}

FAILURE_ORDER = [
    "single_link_failure",
    "random_link_failure_1",
    "random_link_failure_2",
    "capacity_degradation_50",
    "traffic_spike_2x",
]

FAILURE_LABELS = {
    "single_link_failure": "Single Link Failure",
    "random_link_failure_1": "Random Link Failure (1)",
    "random_link_failure_2": "Random Link Failure (2)",
    "capacity_degradation_50": "Capacity Degradation (50%)",
    "traffic_spike_2x": "Traffic Spike (2x)",
}

EXPERT_ORDER = ["bottleneck", "topk", "sensitivity", "gnnplus"]
EXPERT_LABELS = {
    "bottleneck": "Bottleneck",
    "topk": "TopK",
    "sensitivity": "Sensitivity",
    "gnnplus": "GNN+",
}


def canon_topology(name: str) -> str:
    return TOPOLOGY_ALIASES.get(str(name).strip().lower(), str(name).strip().lower())


def topo_display(topo: str) -> str:
    return TOPOLOGY_INFO[topo]["display"]


def topo_with_size(topo: str) -> str:
    info = TOPOLOGY_INFO[topo]
    return f"{info['display']} ({info['nodes']}n)"


def topo_type(topo: str) -> str:
    return TOPOLOGY_INFO[topo]["type"]


def fmt(value, digits: int = 4) -> str:
    try:
        v = float(value)
    except Exception:
        return str(value)
    if not math.isfinite(v):
        return "N/A"
    if abs(v) >= 1e6 or (0 < abs(v) < 1e-4):
        return f"{v:.2e}"
    return f"{v:.{digits}f}"


def fmt_pct(value, digits: int = 2) -> str:
    try:
        v = float(value)
    except Exception:
        return str(value)
    if not math.isfinite(v):
        return "N/A"
    return f"{v:.{digits}f}%"


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    for level in range(1, 4):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x1A, 0x47, 0x80)


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], font_size: int = 8) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(header))
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.name = "Arial"
        set_cell_shading(cell, "D9EAF7")
    for row_vals in rows:
        row = table.add_row()
        for i, value in enumerate(row_vals):
            cell = row.cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(value))
            run.font.size = Pt(font_size)
            run.font.name = "Arial"
    doc.add_paragraph("")


def add_image(doc: Document, path: Path, caption: str, width: float = 6.5) -> None:
    if not path.exists():
        doc.add_paragraph(f"[Missing image: {path}]")
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_title_page(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI-Driven Traffic Engineering\nwith MLP Meta-Gate Expert Selection")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x80)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "Comprehensive Evaluation Report\n"
        "Method: MLP Meta-Gate with Few-Shot Bayesian Calibration\n"
        "4 Experts: Bottleneck, TopK, Sensitivity, GNN+\n"
        "Framing: Zero-Shot Gate Training + Few-Shot Calibration\n"
        "This is NOT a pure zero-shot evaluation"
    )
    run.font.size = Pt(11)

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run(
        "Covers: MLU, Performance Ratio, Disturbance, Execution Time,\n"
        "Training Efficiency, Failure Robustness, SDN Deployment,\n"
        "Complexity Analysis, CDF Plots, and Live Mininet Status"
    )
    run.font.size = Pt(10)
    run.italic = True

    gen = doc.add_paragraph()
    gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    items = [
        "1. Method Summary",
        "2. MLU Results (Maximum Link Utilization)",
        "3. Performance Ratio (PR)",
        "4. Network Disturbance (DB)",
        "5. Execution Time and Decision Latency",
        "6. Training Efficiency and Stability",
        "7. Routing Update Overhead",
        "8. Failure Robustness",
        "9. CDF Plots",
        "10. SDN Deployment Validation",
        "11. Complexity Analysis",
        "12. Germany50 Unseen Topology Deep Dive",
        "13. CERNET Correction",
        "14. Limitations (Honest Assessment)",
        "15. Live Mininet Testbed: Required Final Step",
        "16. Exact Thesis Method Description",
    ]
    for item in items:
        doc.add_paragraph(item)
    doc.add_page_break()


def load_inputs():
    required = [
        SUMMARY_CSV,
        RESULTS_CSV,
        TIMING_CSV,
        SDN_METRICS_CSV,
        TRAINING_SUMMARY_JSON,
        TRAIN_DIST_CSV,
        CALIBRATION_CSV,
        FAILURE_RESULTS_CSV,
        FAILURE_SUMMARY_CSV,
        FAILURE_SDN_CSV,
    ]
    missing = [str(p) for p in required if not p.exists()]
    if missing:
        raise FileNotFoundError("Missing required input files:\n" + "\n".join(missing))

    summary = pd.read_csv(SUMMARY_CSV)
    results = pd.read_csv(RESULTS_CSV)
    timing = pd.read_csv(TIMING_CSV)
    sdn_metrics = pd.read_csv(SDN_METRICS_CSV)
    decisions = pd.read_csv(DECISIONS_CSV) if DECISIONS_CSV.exists() else pd.DataFrame()
    failure_results = pd.read_csv(FAILURE_RESULTS_CSV)
    failure_summary = pd.read_csv(FAILURE_SUMMARY_CSV)
    failure_sdn = pd.read_csv(FAILURE_SDN_CSV)
    train_dist = pd.read_csv(TRAIN_DIST_CSV)
    calibration = pd.read_csv(CALIBRATION_CSV)
    training_summary = json.loads(TRAINING_SUMMARY_JSON.read_text(encoding="utf-8"))
    zero_shot = pd.read_csv(ZERO_SHOT_UNSEEN_CSV) if ZERO_SHOT_UNSEEN_CSV.exists() else pd.DataFrame()
    stage1_summary = json.loads(GNNPLUS_STAGE1_SUMMARY.read_text(encoding="utf-8"))
    sup_log = pd.read_csv(GNNPLUS_SUP_LOG)
    rl_log = pd.read_csv(GNNPLUS_RL_LOG)
    baseline = pd.read_csv(BASELINE_NORMAL_CSV) if BASELINE_NORMAL_CSV.exists() else pd.DataFrame()

    for frame in [summary, results, timing, sdn_metrics, decisions, failure_results, failure_summary, failure_sdn, train_dist, calibration, zero_shot]:
        if "dataset" in frame.columns:
            frame["topology"] = frame["dataset"].map(canon_topology)
        elif "topology" in frame.columns:
            frame["topology"] = frame["topology"].map(canon_topology)
        elif "Topology" in frame.columns:
            frame["topology"] = frame["Topology"].map(canon_topology)

    if not baseline.empty:
        baseline["topology"] = baseline["topology"].map(canon_topology)
        baseline["method_norm"] = baseline["method"].astype(str).str.lower()

    return {
        "summary": summary,
        "results": results,
        "timing": timing,
        "sdn_metrics": sdn_metrics,
        "decisions": decisions,
        "failure_results": failure_results,
        "failure_summary": failure_summary,
        "failure_sdn": failure_sdn,
        "train_dist": train_dist,
        "calibration": calibration,
        "training_summary": training_summary,
        "zero_shot": zero_shot,
        "stage1_summary": stage1_summary,
        "sup_log": sup_log,
        "rl_log": rl_log,
        "baseline": baseline,
    }


def value_counts_pct(series: pd.Series, key: str) -> float:
    total = max(len(series), 1)
    return 100.0 * int(series.value_counts().get(key, 0)) / total


def selector_switches(series: pd.Series) -> int:
    arr = series.astype(str).to_numpy()
    if len(arr) <= 1:
        return 0
    return int((arr[1:] != arr[:-1]).sum())


def sorted_cdf(arr: np.ndarray) -> tuple[np.ndarray, np.ndarray]:
    vals = np.asarray(arr, dtype=float)
    vals = vals[np.isfinite(vals)]
    if vals.size == 0:
        return np.array([]), np.array([])
    vals.sort()
    return vals, np.arange(1, len(vals) + 1) / len(vals)


def plot_selector_distribution(results: pd.DataFrame) -> Path:
    path = PLOTS_DIR / "selector_distribution.png"
    counts = (
        results.groupby(["topology", "metagate_selector"])
        .size()
        .unstack(fill_value=0)
        .reindex(index=TOPOLOGY_ORDER, fill_value=0)
        .reindex(columns=EXPERT_ORDER, fill_value=0)
    )
    fig, ax = plt.subplots(figsize=(11, 5))
    bottom = np.zeros(len(counts))
    x = np.arange(len(counts))
    for expert in EXPERT_ORDER:
        vals = counts[expert].to_numpy()
        ax.bar(x, vals, bottom=bottom, color=METHOD_COLORS[expert], label=EXPERT_LABELS[expert])
        bottom += vals
    ax.set_xticks(x)
    ax.set_xticklabels([topo_display(t) for t in counts.index], rotation=20, ha="right")
    ax.set_ylabel("Selected Timesteps")
    ax.set_title("MetaGate Expert Selection Distribution")
    ax.legend(ncols=4, fontsize=8)
    ax.grid(True, axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def plot_accuracy_gap(summary: pd.DataFrame) -> Path:
    path = PLOTS_DIR / "accuracy_gap.png"
    ordered = summary.set_index("topology").reindex(TOPOLOGY_ORDER).reset_index()
    gaps = ((ordered["metagate_mlu"] - ordered["oracle_mlu"]) / ordered["oracle_mlu"] * 100.0).to_numpy()
    acc = (ordered["accuracy"] * 100.0).to_numpy()
    x = np.arange(len(ordered))
    fig, ax1 = plt.subplots(figsize=(11, 4.8))
    ax1.bar(x, gaps, color="#4c78a8", alpha=0.85, label="Oracle Gap (%)")
    ax1.set_xticks(x)
    ax1.set_xticklabels([topo_display(t) for t in ordered["topology"]], rotation=20, ha="right")
    ax1.set_ylabel("Oracle Gap (%)")
    ax1.axhline(0.0, color="black", linewidth=0.8)
    ax1.grid(True, axis="y", alpha=0.25)
    ax2 = ax1.twinx()
    ax2.plot(x, acc, color="#d62728", marker="o", linewidth=2.0, label="Accuracy (%)")
    ax2.set_ylabel("Accuracy (%)")
    ax2.set_ylim(0, 100)
    handles1, labels1 = ax1.get_legend_handles_labels()
    handles2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(handles1 + handles2, labels1 + labels2, loc="upper left")
    ax1.set_title("MetaGate Accuracy and Oracle Gap by Topology")
    fig.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def plot_mlu_cdf_grid(results: pd.DataFrame) -> Path:
    path = PLOTS_DIR / "cdf_mlu_per_topology.png"
    fig, axes = plt.subplots(2, 4, figsize=(14, 7), sharey=True)
    methods = [
        ("MetaGate", "metagate_mlu", METHOD_COLORS["metagate"]),
        ("Bottleneck", "bn_mlu", METHOD_COLORS["bottleneck"]),
        ("TopK", "topk_mlu", METHOD_COLORS["topk"]),
        ("Sensitivity", "sens_mlu", METHOD_COLORS["sensitivity"]),
        ("GNN+", "gnnplus_mlu", METHOD_COLORS["gnnplus"]),
    ]
    for ax, topo in zip(axes.flatten(), TOPOLOGY_ORDER):
        sub = results[results["topology"] == topo]
        for label, col, color in methods:
            x, y = sorted_cdf(sub[col].to_numpy())
            if x.size:
                ax.plot(x, y, linewidth=1.5, label=label, color=color)
        ax.set_title(topo_display(topo), fontsize=10)
        ax.set_xlabel("MLU")
        ax.grid(True, alpha=0.2)
    axes[0, 0].set_ylabel("CDF")
    axes[1, 0].set_ylabel("CDF")
    handles, labels = axes[0, 0].get_legend_handles_labels()
    fig.legend(handles, labels, loc="upper center", ncols=5, fontsize=8)
    fig.suptitle("CDF of MLU by Topology", fontsize=14, fontweight="bold", y=0.98)
    fig.tight_layout(rect=[0, 0, 1, 0.94])
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def plot_mlu_cdf_global(results: pd.DataFrame) -> Path:
    path = PLOTS_DIR / "cdf_mlu_global.png"
    fig, ax = plt.subplots(figsize=(8.5, 5))
    methods = [
        ("MetaGate", "metagate_mlu", METHOD_COLORS["metagate"]),
        ("Bottleneck", "bn_mlu", METHOD_COLORS["bottleneck"]),
        ("TopK", "topk_mlu", METHOD_COLORS["topk"]),
        ("Sensitivity", "sens_mlu", METHOD_COLORS["sensitivity"]),
        ("GNN+", "gnnplus_mlu", METHOD_COLORS["gnnplus"]),
    ]
    for label, col, color in methods:
        x, y = sorted_cdf(results[col].to_numpy())
        if x.size:
            ax.plot(x, y, linewidth=2.0, label=label, color=color)
    ax.set_xlabel("MLU")
    ax.set_ylabel("CDF")
    ax.set_title("CDF of MLU Across All Topologies")
    ax.grid(True, alpha=0.25)
    ax.legend()
    fig.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def plot_disturbance_cdf_grid(results: pd.DataFrame) -> Path:
    path = PLOTS_DIR / "cdf_disturbance_per_topology.png"
    fig, axes = plt.subplots(2, 4, figsize=(14, 7), sharey=True)
    for ax, topo in zip(axes.flatten(), TOPOLOGY_ORDER):
        sub = results[results["topology"] == topo]
        x, y = sorted_cdf(sub["disturbance"].to_numpy())
        if x.size:
            ax.plot(x, y, linewidth=1.8, color=METHOD_COLORS["metagate"])
        ax.set_title(topo_display(topo), fontsize=10)
        ax.set_xlabel("Disturbance")
        ax.grid(True, alpha=0.2)
    axes[0, 0].set_ylabel("CDF")
    axes[1, 0].set_ylabel("CDF")
    fig.suptitle("CDF of Routing Disturbance by Topology (MetaGate)", fontsize=14, fontweight="bold", y=0.98)
    fig.tight_layout(rect=[0, 0, 1, 0.95])
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def plot_time_cdf_grid(results: pd.DataFrame) -> Path:
    path = PLOTS_DIR / "cdf_total_time_per_topology.png"
    fig, axes = plt.subplots(2, 4, figsize=(14, 7), sharey=True)
    for ax, topo in zip(axes.flatten(), TOPOLOGY_ORDER):
        sub = results[results["topology"] == topo]
        x, y = sorted_cdf(sub["t_total_ms"].to_numpy())
        if x.size:
            ax.plot(x, y, linewidth=1.8, color=METHOD_COLORS["metagate"])
        ax.set_title(topo_display(topo), fontsize=10)
        ax.set_xlabel("Total Time (ms)")
        ax.grid(True, alpha=0.2)
    axes[0, 0].set_ylabel("CDF")
    axes[1, 0].set_ylabel("CDF")
    fig.suptitle("CDF of Decision / Execution Time by Topology (MetaGate)", fontsize=14, fontweight="bold", y=0.98)
    fig.tight_layout(rect=[0, 0, 1, 0.95])
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def plot_failure_mlu_cdf(failure_results: pd.DataFrame, scenario: str, output_name: str, title: str) -> Path:
    path = PLOTS_DIR / output_name
    sub = failure_results[failure_results["scenario"] == scenario]
    fig, ax = plt.subplots(figsize=(8.5, 5))
    methods = [
        ("MetaGate", "metagate_mlu", METHOD_COLORS["metagate"]),
        ("Bottleneck", "bn_mlu", METHOD_COLORS["bottleneck"]),
        ("TopK", "topk_mlu", METHOD_COLORS["topk"]),
        ("Sensitivity", "sens_mlu", METHOD_COLORS["sensitivity"]),
        ("GNN+", "gnnplus_mlu", METHOD_COLORS["gnnplus"]),
    ]
    for label, col, color in methods:
        x, y = sorted_cdf(sub[col].to_numpy())
        if x.size:
            ax.plot(x, y, linewidth=2.0, label=label, color=color)
    ax.set_xlabel("Post-Failure MLU")
    ax.set_ylabel("CDF")
    ax.set_title(title)
    ax.grid(True, alpha=0.25)
    ax.legend()
    fig.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def plot_failure_disturbance_cdf(failure_results: pd.DataFrame) -> Path:
    path = PLOTS_DIR / "cdf_failure_disturbance.png"
    fig, ax = plt.subplots(figsize=(8.5, 5))
    for scenario in ["single_link_failure", "capacity_degradation_50", "traffic_spike_2x"]:
        sub = failure_results[failure_results["scenario"] == scenario]
        x, y = sorted_cdf(sub["disturbance"].to_numpy())
        if x.size:
            ax.plot(x, y, linewidth=2.0, label=FAILURE_LABELS[scenario])
    ax.set_xlabel("Disturbance")
    ax.set_ylabel("CDF")
    ax.set_title("CDF of Disturbance Under Failure")
    ax.grid(True, alpha=0.25)
    ax.legend()
    fig.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def plot_training_convergence(sup_log: pd.DataFrame, rl_log: pd.DataFrame) -> Path:
    path = PLOTS_DIR / "training_convergence.png"
    fig, axes = plt.subplots(1, 2, figsize=(11, 4.5))
    axes[0].plot(sup_log["epoch"], sup_log["train_loss"], label="Train Loss", color="#1f77b4")
    axes[0].plot(sup_log["epoch"], sup_log["val_loss"], label="Val Loss", color="#d62728")
    axes[0].set_title("GNN+ Supervised Training")
    axes[0].set_xlabel("Epoch")
    axes[0].set_ylabel("Loss")
    axes[0].grid(True, alpha=0.25)
    axes[0].legend()

    axes[1].plot(rl_log["epoch"], rl_log["train_mlu"], label="Train MLU", color="#2ca02c")
    axes[1].plot(rl_log["epoch"], rl_log["val_mlu"], label="Val MLU", color="#ff7f0e")
    axes[1].set_title("GNN+ RL Fine-Tuning")
    axes[1].set_xlabel("Epoch")
    axes[1].set_ylabel("MLU")
    axes[1].grid(True, alpha=0.25)
    axes[1].legend()

    fig.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def aggregate_external_baselines(baseline: pd.DataFrame) -> pd.DataFrame:
    if baseline.empty:
        return pd.DataFrame()
    return (
        baseline.groupby(["topology", "method_norm"], as_index=False)
        .agg(
            mean_mlu=("mean_mlu", "mean"),
            mean_pr=("mean_pr", "mean"),
            mean_disturbance=("mean_disturbance", "mean"),
            mean_total_time_ms=("mean_total_time_ms", "mean"),
        )
    )


def build_report():
    data = load_inputs()
    PLOTS_DIR.mkdir(parents=True, exist_ok=True)

    summary = data["summary"].set_index("topology").reindex(TOPOLOGY_ORDER).reset_index()
    results = data["results"]
    timing = data["timing"].set_index("topology").reindex(TOPOLOGY_ORDER).reset_index()
    sdn_metrics = data["sdn_metrics"].set_index("topology").reindex(TOPOLOGY_ORDER).reset_index()
    failure_results = data["failure_results"]
    failure_summary = data["failure_summary"]
    failure_sdn = data["failure_sdn"]
    train_dist = data["train_dist"].set_index("topology").reindex(TOPOLOGY_ORDER).reset_index()
    calibration = data["calibration"].set_index("topology").reindex(TOPOLOGY_ORDER).reset_index()
    training_summary = data["training_summary"]
    zero_shot = data["zero_shot"]
    stage1_summary = data["stage1_summary"]
    sup_log = data["sup_log"]
    rl_log = data["rl_log"]
    baseline = aggregate_external_baselines(data["baseline"])

    summary["gap_pct"] = (summary["metagate_mlu"] - summary["oracle_mlu"]) / summary["oracle_mlu"] * 100.0
    summary["meta_pr"] = summary["metagate_mlu"] / summary["oracle_mlu"]
    summary["bn_pr"] = summary["bn_mlu"] / summary["oracle_mlu"]
    summary["topk_pr"] = summary["topk_mlu"] / summary["oracle_mlu"]
    summary["sens_pr"] = summary["sens_mlu"] / summary["oracle_mlu"]
    summary["gnnplus_pr"] = summary["gnnplus_mlu"] / summary["oracle_mlu"]

    selector_plot = plot_selector_distribution(results)
    gap_plot = plot_accuracy_gap(summary)
    mlu_cdf_grid = plot_mlu_cdf_grid(results)
    mlu_cdf_global = plot_mlu_cdf_global(results)
    dist_cdf_grid = plot_disturbance_cdf_grid(results)
    time_cdf_grid = plot_time_cdf_grid(results)
    failure_link_cdf = plot_failure_mlu_cdf(failure_results, "single_link_failure", "cdf_mlu_single_link_failure.png", "CDF of MLU under Single Link Failure")
    failure_cap_cdf = plot_failure_mlu_cdf(failure_results, "capacity_degradation_50", "cdf_mlu_capacity_degradation.png", "CDF of MLU under Capacity Degradation (50%)")
    failure_spike_cdf = plot_failure_mlu_cdf(failure_results, "traffic_spike_2x", "cdf_mlu_traffic_spike.png", "CDF of MLU under Traffic Spike (2x)")
    failure_dist_cdf = plot_failure_disturbance_cdf(failure_results)
    train_plot = plot_training_convergence(sup_log, rl_log)

    overall_acc = float(training_summary["overall_test_accuracy"]) * 100.0
    known_acc = float(training_summary["known_test_accuracy"]) * 100.0
    unseen_acc = float(training_summary["unseen_test_accuracy"]) * 100.0

    doc = Document()
    style_doc(doc)
    add_title_page(doc)
    add_toc(doc)

    # 1
    doc.add_heading("1. Method Summary", level=1)
    doc.add_paragraph(
        "An MLP-based meta-gate (3-layer, 128-128-64-4, BatchNorm, dropout 0.3) dynamically selects among four "
        "traffic-engineering experts per traffic matrix: Bottleneck, TopK by Demand, Sensitivity Analysis, and GNN+-based selection. "
        "The gate is trained on pooled oracle labels from 6 known topologies only."
    )
    doc.add_paragraph(
        "At deployment, a lightweight few-shot calibration phase runs 10 validation traffic matrices through all four "
        "experts, counts which expert wins, and fuses the resulting topology-specific Bayesian prior with the MLP softmax. "
        "This is zero-shot gate training with few-shot calibration, not pure zero-shot inference."
    )
    add_bullet(doc, f"Overall accuracy across 8 topologies: {overall_acc:.1f}%")
    add_bullet(doc, f"Known-topology accuracy: {known_acc:.1f}%")
    add_bullet(doc, f"Unseen-topology accuracy: {unseen_acc:.1f}%")
    add_bullet(doc, "This report includes integrated MetaGate+GNN+ MLU, disturbance, timing, failure, and SDN-style analytical metrics from the same rerun bundle.")
    add_bullet(doc, "The SDN metrics are model-based analytical control-loop metrics, not live Mininet packet measurements.")
    doc.add_page_break()

    # 2
    doc.add_heading("2. MLU Results (Maximum Link Utilization)", level=1)
    doc.add_heading("2.1 MetaGate vs Oracle vs Individual Experts", level=2)
    rows = []
    for topo in TOPOLOGY_ORDER:
        row = summary[summary["topology"] == topo].iloc[0]
        rows.append([
            topo_with_size(topo),
            topo_type(topo),
            fmt(row["metagate_mlu"]),
            fmt(row["oracle_mlu"]),
            fmt(row["bn_mlu"]),
            fmt(row["topk_mlu"]),
            fmt(row["sens_mlu"]),
            fmt(row["gnnplus_mlu"]),
            f"{row['gap_pct']:+.2f}%",
        ])
    add_table(doc, ["Topology", "Type", "MetaGate", "Oracle", "BN", "TopK", "Sens", "GNN+", "Gap"], rows, font_size=7)
    add_image(doc, gap_plot, "Figure 2. MetaGate accuracy and oracle gap by topology.", width=6.7)

    doc.add_heading("2.2 MetaGate vs External Baselines", level=2)
    doc.add_paragraph(
        "The table below adds standalone external baselines from the corrected baseline bundle. These baselines are "
        "not from the same integrated run, so they should be read as external comparison context rather than strict same-bundle significance evidence."
    )
    ext_rows = []
    for topo in TOPOLOGY_ORDER:
        mg_row = summary[summary["topology"] == topo].iloc[0]
        def base(method_name: str) -> str:
            row = baseline[(baseline["topology"] == topo) & (baseline["method_norm"] == method_name.lower())]
            return fmt(row.iloc[0]["mean_mlu"]) if not row.empty else "N/A"
        ext_rows.append([
            topo_with_size(topo),
            fmt(mg_row["metagate_mlu"]),
            base("bottleneck"),
            base("topk"),
            base("sensitivity"),
            base("gnn"),
            base("ecmp"),
            base("ospf"),
            "Unavailable",
            "Unavailable",
            "Unavailable",
        ])
    add_table(doc, ["Topology", "MetaGate", "BN", "TopK", "Sens", "GNN", "ECMP", "OSPF", "FlexDATE", "ERODRL", "FlexEntry"], ext_rows, font_size=6)
    doc.add_page_break()

    # 3
    doc.add_heading("3. Performance Ratio (PR)", level=1)
    doc.add_paragraph("PR = Method MLU / Oracle MLU. PR = 1.000 means exact oracle matching. Values above 1 indicate suboptimality.")
    pr_rows = []
    for topo in TOPOLOGY_ORDER:
        row = summary[summary["topology"] == topo].iloc[0]
        pr_rows.append([
            topo_with_size(topo),
            topo_type(topo),
            fmt(row["meta_pr"], 6),
            fmt(row["bn_pr"], 6),
            fmt(row["topk_pr"], 6),
            fmt(row["sens_pr"], 6),
            fmt(row["gnnplus_pr"], 6),
        ])
    add_table(doc, ["Topology", "Type", "MetaGate PR", "BN PR", "TopK PR", "Sens PR", "GNN+ PR"], pr_rows, font_size=8)
    doc.add_page_break()

    # 4
    doc.add_heading("4. Network Disturbance (DB)", level=1)
    doc.add_paragraph(
        "Unlike the older MetaGate report, this rerun tracks actual realized MetaGate routing disturbance directly in the integrated pipeline."
    )
    db_rows = []
    for topo in TOPOLOGY_ORDER:
        row = sdn_metrics[sdn_metrics["topology"] == topo].iloc[0]
        sub = results[results["topology"] == topo]
        route_change_freq = float((sub["disturbance"] > 1e-9).mean() * 100.0) if not sub.empty else 0.0
        db_rows.append([
            topo_with_size(topo),
            fmt(row["mean_disturbance"]),
            f"{route_change_freq:.1f}%",
            fmt(sub["disturbance"].quantile(0.95) if not sub.empty else np.nan),
        ])
    add_table(doc, ["Topology", "Mean MetaGate DB", "Route Change Frequency", "P95 DB"], db_rows, font_size=8)
    doc.add_page_break()

    # 5
    doc.add_heading("5. Execution Time and Decision Latency", level=1)
    doc.add_paragraph(
        "Decision time = all selector overhead before LP (run 4 experts + extract features + MLP forward pass). "
        "LP time = the selected LP solve only. Total time = decision time + selected LP."
    )
    timing_rows = []
    for topo in TOPOLOGY_ORDER:
        row = timing[timing["topology"] == topo].iloc[0]
        timing_rows.append([
            topo_with_size(topo),
            fmt(row["t_bn_ms"], 1),
            fmt(row["t_topk_ms"], 1),
            fmt(row["t_sens_ms"], 1),
            fmt(row["t_gnnplus_ms"], 1),
            fmt(row["t_features_ms"], 2),
            fmt(row["t_mlp_ms"], 3),
            fmt(row["t_lp_ms"], 1),
            fmt(row["t_total_ms"], 1),
        ])
    add_table(doc, ["Topology", "BN ms", "TopK ms", "Sens ms", "GNN+ ms", "Feat ms", "MLP ms", "LP ms", "Total ms"], timing_rows, font_size=7)
    add_image(doc, time_cdf_grid, "Figure 5. CDF of decision / execution time by topology.", width=6.8)
    doc.add_page_break()

    # 6
    doc.add_heading("6. Training Efficiency and Stability", level=1)
    doc.add_heading("6.1 GNN+ Expert Training", level=2)
    gnn_rows = [
        ["Supervised stage time (s)", fmt(stage1_summary["sup_time"], 2)],
        ["Supervised best epoch", str(stage1_summary["sup_best_epoch"])],
        ["RL fine-tuning time (s)", fmt(stage1_summary["rl_time"], 2)],
        ["RL best epoch", str(stage1_summary["rl_best_epoch"])],
        ["RL best val MLU", fmt(stage1_summary["rl_best_val_mlu"])],
        ["Total training time (s)", fmt(stage1_summary["total_time"], 2)],
    ]
    add_table(doc, ["Metric", "Value"], gnn_rows, font_size=8)
    doc.add_heading("6.2 MLP Meta-Gate Training", level=2)
    meta_rows = [
        ["Train samples", str(training_summary["pooled_train_samples"])],
        ["Validation samples", str(training_summary["pooled_val_samples"])],
        ["Train accuracy", fmt_pct(float(training_summary["train_accuracy"]) * 100.0, 1)],
        ["Validation accuracy", fmt_pct(float(training_summary["val_accuracy"]) * 100.0, 1)],
        ["Epochs / batch size", f"{training_summary['metagate_config']['num_epochs']} / {training_summary['metagate_config']['batch_size']}"],
        ["Soft labels enabled", str(training_summary["metagate_config"].get("use_soft_labels", False))],
        ["Regret loss weight", str(training_summary["metagate_config"].get("regret_loss_weight", "N/A"))],
        ["Feature clip", str(training_summary["metagate_config"].get("feature_clip", "N/A"))],
    ]
    add_table(doc, ["Metric", "Value"], meta_rows, font_size=8)
    doc.add_heading("6.3 Training Stability Assessment", level=2)
    doc.add_paragraph(
        "The GNN+ supervised stage shows a monotonic reduction in validation loss through the best epoch, and the "
        "RL fine-tuning stage is short and intentionally conservative. The MetaGate training bundle stores final train "
        "and validation accuracies, but it does not persist a per-epoch CSV log in this branch."
    )
    add_image(doc, train_plot, "Figure 6. GNN+ supervised and RL fine-tuning convergence.", width=6.7)
    doc.add_page_break()

    # 7
    doc.add_heading("7. Routing Update Overhead", level=1)
    doc.add_paragraph(
        "All four experts operate with K_crit = 40 candidate critical flows. The realized routing-update overhead, "
        "however, depends on how many flow-group entries actually change after LP rerouting."
    )
    overhead_rows = []
    for topo in TOPOLOGY_ORDER:
        row = sdn_metrics[sdn_metrics["topology"] == topo].iloc[0]
        sub = results[results["topology"] == topo]
        overhead_rows.append([
            topo_with_size(topo),
            fmt(row["flow_table_updates"], 2),
            fmt(row["rule_install_delay_ms"], 3),
            f"{selector_switches(sub['metagate_selector'])}/{max(len(sub)-1,1)}",
        ])
    add_table(doc, ["Topology", "Rules Pushed / Cycle", "Rule Install Delay (ms)", "Selector Switches"], overhead_rows, font_size=8)
    add_image(doc, selector_plot, "Figure 7. MetaGate selected-expert distribution by topology.", width=6.6)
    doc.add_page_break()

    # 8
    doc.add_heading("8. Failure Robustness", level=1)
    doc.add_paragraph(
        "Five failure scenarios are reported here from the integrated MetaGate+GNN+ rerun: Single Link Failure, "
        "Random Link Failure (1), Random Link Failure (2), Capacity Degradation (50%), and Traffic Spike (2x)."
    )
    scenario_rows = []
    for scenario in FAILURE_ORDER:
        sub = failure_results[failure_results["scenario"] == scenario]
        summ = failure_summary[failure_summary["scenario"] == scenario]
        if sub.empty or summ.empty:
            continue
        counts = sub["metagate_selector"].value_counts()
        scenario_rows.append([
            FAILURE_LABELS[scenario],
            fmt_pct(float(summ["accuracy"].mean() * 100.0), 1),
            f"{((summ['metagate_mlu'].mean() - summ['oracle_mlu'].mean()) / summ['oracle_mlu'].mean()) * 100.0:+.2f}%",
            fmt(summ["failure_recovery_ms"].mean(), 1),
            f"{100.0 * counts.get('bottleneck', 0) / len(sub):.0f}%",
            f"{100.0 * counts.get('topk', 0) / len(sub):.0f}%",
            f"{100.0 * counts.get('sensitivity', 0) / len(sub):.0f}%",
            f"{100.0 * counts.get('gnnplus', 0) / len(sub):.0f}%",
        ])
    add_table(doc, ["Scenario", "Accuracy", "Oracle Gap", "Recovery (ms)", "BN%", "TopK%", "Sens%", "GNN+%"], scenario_rows, font_size=8)

    for idx, scenario in enumerate(FAILURE_ORDER, start=1):
        doc.add_heading(f"8.{idx} {FAILURE_LABELS[scenario]}", level=2)
        sub = failure_results[failure_results["scenario"] == scenario]
        rows = []
        for topo in TOPOLOGY_ORDER:
            topo_sub = sub[sub["topology"] == topo]
            if topo_sub.empty:
                continue
            counts = topo_sub["metagate_selector"].value_counts()
            dominant = counts.idxmax()
            rows.append([
                topo_with_size(topo),
                EXPERT_LABELS[dominant],
                f"{100.0 * counts.get('bottleneck', 0) / len(topo_sub):.0f}%",
                f"{100.0 * counts.get('topk', 0) / len(topo_sub):.0f}%",
                f"{100.0 * counts.get('sensitivity', 0) / len(topo_sub):.0f}%",
                f"{100.0 * counts.get('gnnplus', 0) / len(topo_sub):.0f}%",
                fmt_pct(float(topo_sub['correct'].mean() * 100.0), 1),
                f"{((topo_sub['metagate_mlu'].mean() - topo_sub['oracle_mlu'].mean()) / topo_sub['oracle_mlu'].mean()) * 100.0:+.2f}%",
                fmt(topo_sub["failure_recovery_ms"].mean(), 1),
            ])
        add_table(doc, ["Topology", "Dominant", "BN%", "TopK%", "Sens%", "GNN+%", "Accuracy", "Gap", "Recovery ms"], rows, font_size=7)
    doc.add_page_break()

    # 9
    doc.add_heading("9. CDF Plots", level=1)
    doc.add_paragraph("All CDFs in this section are generated from per-timestep raw data in the integrated MetaGate+GNN+ bundle.")
    doc.add_heading("9.1 CDF of MLU (per topology)", level=2)
    add_image(doc, mlu_cdf_grid, "Figure 8. CDF of MLU by topology for MetaGate and its four experts.", width=6.8)
    doc.add_heading("9.2 CDF of MLU (all topologies combined)", level=2)
    add_image(doc, mlu_cdf_global, "Figure 9. CDF of MLU across all topologies.", width=6.3)
    doc.add_heading("9.3 CDF of Routing Disturbance (per topology)", level=2)
    add_image(doc, dist_cdf_grid, "Figure 10. CDF of MetaGate routing disturbance by topology.", width=6.8)
    doc.add_heading("9.4 CDF of Decision/Execution Time (per topology)", level=2)
    add_image(doc, time_cdf_grid, "Figure 11. CDF of MetaGate execution time by topology.", width=6.8)
    doc.add_heading("9.5 CDF of MLU under Link Failure", level=2)
    add_image(doc, failure_link_cdf, "Figure 12. CDF of MLU under single-link failure.", width=6.3)
    doc.add_heading("9.6 CDF of MLU under Capacity Degradation", level=2)
    add_image(doc, failure_cap_cdf, "Figure 13. CDF of MLU under capacity degradation (50%).", width=6.3)
    doc.add_heading("9.7 CDF of MLU under Traffic Spike", level=2)
    add_image(doc, failure_spike_cdf, "Figure 14. CDF of MLU under traffic spike (2x).", width=6.3)
    doc.add_heading("9.8 CDF of Disturbance under Failure", level=2)
    add_image(doc, failure_dist_cdf, "Figure 15. CDF of disturbance under failure scenarios.", width=6.3)
    doc.add_heading("9.9 Training Convergence", level=2)
    add_image(doc, train_plot, "Figure 16. GNN+ training convergence used by the integrated MetaGate+GNN+ system.", width=6.3)
    doc.add_page_break()

    # 10
    doc.add_heading("10. SDN Deployment Validation", level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        "HONESTY NOTE: The SDN results below are from a model-based simulation that mimics the SDN control loop "
        "(read TM -> select expert -> solve LP -> push rules). They use real topology graphs and traffic matrices, "
        "but they are NOT from live Mininet with real packet forwarding."
    )
    r.bold = True
    r.font.color.rgb = RGBColor(0xAA, 0x00, 0x00)
    doc.add_heading("10.1 SDN Simulation Results", level=2)
    sdn_rows = []
    for topo in TOPOLOGY_ORDER:
        row = sdn_metrics[sdn_metrics["topology"] == topo].iloc[0]
        sdn_rows.append([
            topo_with_size(topo),
            fmt(row["mean_mlu"]),
            fmt(row["throughput"]),
            fmt(row["mean_latency_au"]),
            fmt(row["p95_latency_au"]),
            fmt(row["packet_loss"]),
            fmt(row["jitter_au"]),
            fmt(row["decision_time_ms"], 1),
            fmt(row["flow_table_updates"], 2),
            fmt(row["rule_install_delay_ms"], 3),
        ])
    add_table(doc, ["Topology", "MLU", "Throughput", "Mean Delay", "P95 Delay", "Packet Loss", "Jitter", "Decision ms", "Rules/Cycle", "Rule Delay ms"], sdn_rows, font_size=7)
    doc.add_heading("10.2 What These SDN Metrics Mean", level=2)
    metric_rows = [
        ["Throughput", "Fraction of routed demand over total demand"],
        ["Mean / P95 Delay", "Telemetry-derived end-to-end delay under current routing"],
        ["Packet Loss", "Fraction of demand not routed successfully"],
        ["Jitter", "Delay variation relative to the previous timestep"],
        ["Rules/Cycle", "Average changed OpenFlow group entries pushed per control cycle"],
        ["Rule Delay", "Measured time to build and diff the rule updates in the adapter"],
    ]
    add_table(doc, ["Metric", "Interpretation"], metric_rows, font_size=8)
    doc.add_page_break()

    # 11
    doc.add_heading("11. Complexity Analysis", level=1)
    doc.add_heading("11.1 Empirical Scaling", level=2)
    scale_rows = []
    for topo in TOPOLOGY_ORDER:
        row = timing[timing["topology"] == topo].iloc[0]
        scale_rows.append([
            topo_with_size(topo),
            str(TOPOLOGY_INFO[topo]["nodes"]),
            fmt(row["t_decision_ms"], 1),
            fmt(row["t_lp_ms"], 1),
            fmt(row["t_total_ms"], 1),
        ])
    add_table(doc, ["Topology", "Nodes", "Decision ms", "LP ms", "Total ms"], scale_rows, font_size=8)
    doc.add_heading("11.2 Algorithmic Complexity", level=2)
    complexity_rows = [
        ["Bottleneck", "Heuristic critical-flow selection + LP"],
        ["TopK", "Demand ranking + LP"],
        ["Sensitivity", "Impact scoring + LP"],
        ["GNN+", "Graph inference + LP"],
        ["MetaGate", "Run 4 experts + 49-d feature extraction + MLP + one selected LP"],
    ]
    add_table(doc, ["Component", "Controller Work"], complexity_rows, font_size=8)
    doc.add_heading("11.3 Scaling Observations", level=2)
    doc.add_paragraph(
        f"Decision time grows from {timing['t_decision_ms'].min():.1f} ms on the smallest topology to "
        f"{timing['t_decision_ms'].max():.1f} ms on the largest. The MLP itself remains tiny; the dominant "
        "cost is running all four experts before the gate decides, followed by the selected LP solve."
    )
    doc.add_page_break()

    # 12
    doc.add_heading("12. Germany50 Unseen Topology Deep Dive", level=1)
    germany = summary[summary["topology"] == "germany50"].iloc[0]
    germany_sub = results[results["topology"] == "germany50"]
    germany_counts = germany_sub["metagate_selector"].value_counts()
    germany_prior = calibration[calibration["topology"] == "germany50"]
    deep_rows = [
        ["Accuracy", fmt_pct(germany["accuracy"] * 100.0, 1)],
        ["MetaGate MLU", fmt(germany["metagate_mlu"])],
        ["Oracle MLU", fmt(germany["oracle_mlu"])],
        ["Oracle gap", f"{germany['gap_pct']:+.2f}%"],
        ["Predicted Bottleneck selections", str(int(germany_counts.get("bottleneck", 0)))],
        ["Predicted GNN+ selections", str(int(germany_counts.get("gnnplus", 0)))],
    ]
    if not germany_prior.empty:
        grow = germany_prior.iloc[0]
        deep_rows.extend(
            [
                ["Calibration prior BN", fmt(grow["bottleneck_prior"], 2)],
                ["Calibration prior TopK", fmt(grow["topk_prior"], 2)],
                ["Calibration prior Sens", fmt(grow["sensitivity_prior"], 2)],
                ["Calibration prior GNN+", fmt(grow["gnnplus_prior"], 2)],
            ]
        )
    if not zero_shot.empty:
        zrow = zero_shot[zero_shot["topology"] == "germany50"]
        if not zrow.empty:
            z = zrow.iloc[0]
            deep_rows.extend(
                [
                    ["Zero-shot Germany50 accuracy", fmt_pct(float(z["accuracy"]) * 100.0, 1)],
                    ["Zero-shot Germany50 gap", f"{float(z['oracle_gap_pct']):+.2f}%"],
                ]
            )
    add_table(doc, ["Metric", "Value"], deep_rows, font_size=8)
    doc.add_paragraph(
        "Germany50 remains the clearest unseen-topology case where calibration matters. The table above combines the "
        "actual calibrated selector behavior and, when available, the zero-shot baseline for the same MetaGate+GNN+ system."
    )
    doc.add_page_break()

    # 13
    doc.add_heading("13. CERNET Correction", level=1)
    cernet_rows = [
        ["Source", "TopologyZoo graph, not SNDlib"],
        ["Nodes", "41"],
        ["Edges", "59 bidirectional links = 116 directed edges"],
        ["Traffic", "Synthetic MGM traffic, not measured SNDlib matrices"],
    ]
    add_table(doc, ["Property", "Value"], cernet_rows, font_size=8)
    doc.add_page_break()

    # 14
    doc.add_heading("14. Limitations (Honest Assessment)", level=1)
    limitations = [
        "This is not pure zero-shot inference. The report uses zero-shot gate training with few-shot Bayesian calibration.",
        "Paper baselines such as FlexDATE, FlexEntry, and ERODRL are not runnable in the current repository and are therefore shown only as unavailable entries.",
        "The SDN metrics are model-based analytical control-loop outputs, not live packet measurements.",
        "The MLP classifier remains imperfect; some topologies still show low exact expert accuracy even when the oracle gap is small.",
        "The MetaGate architecture incurs the cost of running all four experts before making a decision.",
    ]
    for item in limitations:
        add_bullet(doc, item)
    doc.add_page_break()

    # 15
    doc.add_heading("15. Live Mininet Testbed: Required Final Step", level=1)
    doc.add_paragraph(
        "STATUS: Mininet is not part of this evaluation bundle. The present SDN sections validate the control loop analytically, "
        "but a real Mininet deployment would still be required to measure actual packet delay, throughput, OVS flow-install timing, "
        "and failure recovery in a running emulated network."
    )
    doc.add_page_break()

    # 16
    doc.add_heading("16. Exact Thesis Method Description", level=1)
    doc.add_paragraph(
        "We propose a two-stage expert selection framework for traffic engineering. An MLP-based meta-gate "
        "(3-layer, 128-128-64-4 with BatchNorm and dropout 0.3) selects among four expert flow-selection "
        "strategies: Bottleneck, Top-K by Demand, Sensitivity Analysis, and GNN+-based selection, all operating "
        "with K_crit = 40 critical flows for fair comparison."
    )
    doc.add_paragraph(
        "Stage 1 -- Zero-shot gate training. The MLP gate is trained on pooled oracle labels from six known "
        "topologies only. Oracle labels are obtained by running all four experts through LP optimization and "
        "selecting the expert with minimum MLU per traffic matrix. Germany50 and VtlWavenet2011 are held out "
        "from gate-weight training."
    )
    doc.add_paragraph(
        "Stage 2 -- Few-shot Bayesian calibration at deployment. Before evaluation on each topology, a lightweight "
        "calibration phase runs 10 validation traffic matrices through all four experts to estimate a topology-specific "
        "prior. At inference time, the MLP softmax is fused with this prior. No gradient updates occur during calibration."
    )

    REPORT_DOC.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_DOC)

    audit_lines = [
        "# MetaGate+GNN+ Sarah-Style Full Report Audit",
        "",
        f"- Report: `{REPORT_DOC.relative_to(PROJECT_ROOT)}`",
        f"- Source bundle: `{OUTPUT_DIR.relative_to(PROJECT_ROOT)}`",
        "- Integrated same-bundle metrics included: MLU, PR, disturbance, execution time, routing updates, failure robustness, SDN metrics.",
        "- External standalone baselines appear only in the comparison table and are clearly marked as separate-bundle context.",
        f"- Zero-shot comparison available in report: {'yes' if not zero_shot.empty else 'no'}",
    ]
    AUDIT_MD.write_text("\n".join(audit_lines), encoding="utf-8")
    print(f"Report saved to {REPORT_DOC}")


if __name__ == "__main__":
    build_report()

#!/usr/bin/env python3
"""Build a Sarah-style MetaGate + GNN+ report from existing outputs only.

This script does not rerun expensive evaluations. It stitches together the
validated MetaGate bundle, the corrected baseline bundle, the packet-SDN
MetaGate bundle, and the fixed GNN+ packet-SDN bundle into one honest report.
"""

from __future__ import annotations

import json
import math
import os
from datetime import datetime
from pathlib import Path

REPORT_CACHE_ROOT = Path("/tmp") / "metagate_gnnplus_report_cache"
REPORT_CACHE_ROOT.mkdir(parents=True, exist_ok=True)
os.environ.setdefault("MPLCONFIGDIR", str(REPORT_CACHE_ROOT / "mplconfig"))
os.environ.setdefault("XDG_CACHE_HOME", str(REPORT_CACHE_ROOT / "xdg_cache"))

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parent.parent
os.chdir(PROJECT_ROOT)

OUT_DIR = PROJECT_ROOT / "results" / "final_metagate_gnnplus_report"
PLOTS_DIR = OUT_DIR / "plots"
OUTPUT_DOC = OUT_DIR / "AI_Driven_TE_MLP_MetaGate_GNNPlus_Report.docx"
AUDIT_MD = OUT_DIR / "report_audit.md"

METAGATE_DIR = PROJECT_ROOT / "results" / "dynamic_metagate"
BASELINE_DIR = PROJECT_ROOT / "results" / "final_full_eval_corrected"
PACKET_SDN_DIR = PROJECT_ROOT / "results" / "packet_sdn_simulation"
GNNPLUS_PACKET_DIR = PROJECT_ROOT / "results" / "gnnplus_packet_sdn_report_fixed"

GNNPLUS_CHECKPOINT = PROJECT_ROOT / "results" / "gnn_plus_retrained_fixedk40" / "gnn_plus_fixed_k40.pt"
GNNPLUS_SOURCE_CHECKPOINT = PROJECT_ROOT / "results" / "gnn_plus" / "stage1_regularization" / "training_d02" / "final.pt"
GNNPLUS_STAGE1_SUMMARY = PROJECT_ROOT / "results" / "gnn_plus" / "stage1_regularization" / "training_d02" / "summary.json"

TOPOLOGY_ORDER = [
    "abilene",
    "cernet",
    "geant",
    "ebone",
    "sprintlink",
    "tiscali",
    "germany50",
    "vtlwavenet2011",
]

TOPOLOGY_INFO = {
    "abilene": {"display": "Abilene", "type": "known", "nodes": 12, "edges": 30},
    "cernet": {"display": "CERNET", "type": "known", "nodes": 41, "edges": 116},
    "geant": {"display": "GEANT", "type": "known", "nodes": 22, "edges": 72},
    "ebone": {"display": "Ebone", "type": "known", "nodes": 23, "edges": 76},
    "sprintlink": {"display": "Sprintlink", "type": "known", "nodes": 44, "edges": 166},
    "tiscali": {"display": "Tiscali", "type": "known", "nodes": 49, "edges": 172},
    "germany50": {"display": "Germany50", "type": "unseen", "nodes": 50, "edges": 176},
    "vtlwavenet2011": {"display": "VtlWavenet2011", "type": "unseen", "nodes": 92, "edges": 192},
}

TOPOLOGY_ALIASES = {
    "abilene": "abilene",
    "abilene_backbone": "abilene",
    "cernet": "cernet",
    "cernet_real": "cernet",
    "geant": "geant",
    "geant_core": "geant",
    "ebone": "ebone",
    "rocketfuel_ebone": "ebone",
    "sprintlink": "sprintlink",
    "rocketfuel_sprintlink": "sprintlink",
    "tiscali": "tiscali",
    "rocketfuel_tiscali": "tiscali",
    "germany50": "germany50",
    "germany50_real": "germany50",
    "vtlwavenet2011": "vtlwavenet2011",
    "topologyzoo_vtlwavenet2011": "vtlwavenet2011",
}

METHOD_COLORS = {
    "metagate": "#1f77b4",
    "ecmp": "#7f7f7f",
    "bottleneck": "#2ca02c",
    "topk": "#9467bd",
    "sensitivity": "#8c564b",
    "gnn": "#ff7f0e",
    "gnnplus": "#d62728",
}

METHOD_DISPLAY = {
    "metagate": "MetaGate",
    "ecmp": "ECMP",
    "bottleneck": "Bottleneck",
    "topk": "TopK",
    "sensitivity": "Sensitivity",
    "gnn": "GNN",
    "gnnplus": "GNN+",
}

EXPERT_ORDER = ["bottleneck", "topk", "sensitivity", "gnn"]
EXPERT_LABELS = {
    "bottleneck": "Bottleneck",
    "topk": "TopK",
    "sensitivity": "Sensitivity",
    "gnn": "GNN",
}

FAILURE_LABELS = {
    "single_link_failure": "Single Link Failure",
    "random_link_failure_1": "Random Link Failure (1)",
    "random_link_failure_2": "Random Link Failure (2)",
    "capacity_degradation": "Capacity Degradation",
    "capacity_degradation_50": "Capacity Degradation (50%)",
    "traffic_spike": "Traffic Spike (2x)",
    "traffic_spike_2x": "Traffic Spike (2x)",
}


def canonical_topology(name: str) -> str:
    return TOPOLOGY_ALIASES.get(str(name).strip().lower(), str(name).strip().lower())


def display_topology(canon: str) -> str:
    return TOPOLOGY_INFO.get(canon, {}).get("display", canon)


def display_topology_with_size(canon: str) -> str:
    info = TOPOLOGY_INFO[canon]
    return f"{info['display']} ({info['nodes']}n)"


def display_type(canon: str) -> str:
    return TOPOLOGY_INFO[canon]["type"].upper() if canon in {"germany50", "vtlwavenet2011"} else TOPOLOGY_INFO[canon]["type"]


def ordered_topologies(values: list[str]) -> list[str]:
    seen = set()
    out = []
    for topo in TOPOLOGY_ORDER:
        if topo in values and topo not in seen:
            out.append(topo)
            seen.add(topo)
    for topo in values:
        if topo not in seen:
            out.append(topo)
            seen.add(topo)
    return out


def fmt(value, digits: int = 4, pct: bool = False) -> str:
    if value is None:
        return "N/A"
    try:
        v = float(value)
    except Exception:
        return str(value)
    if not math.isfinite(v):
        return "N/A"
    if pct:
        return f"{v:.{digits}f}%"
    if abs(v) >= 1e6 or (0 < abs(v) < 1e-4):
        return f"{v:.2e}"
    return f"{v:.{digits}f}"


def display_method(method: str) -> str:
    return METHOD_DISPLAY.get(str(method).strip().lower(), str(method))


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    for level in range(1, 4):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style.font.color.rgb = RGBColor(0x1A, 0x47, 0x80)


def add_title_page(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI-Driven Traffic Engineering with MLP Meta-Gate")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x80)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Comprehensive Evaluation Report with GNN+ Packet-SDN Validation")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run(
        "Main adaptive method: MLP MetaGate with few-shot Bayesian calibration\n"
        "Separate fixed learned branch included: GNN+ packet-SDN validation"
    )
    run.font.size = Pt(11)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "Important honesty note: the current MetaGate chooser still uses Original GNN as its fourth expert. "
        "GNN+ is validated separately in this report and is not falsely claimed as the current chooser expert."
    )
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xAA, 0x00, 0x00)

    gen = doc.add_paragraph()
    gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    items = [
        "1. Method Summary",
        "2. MLU Results (Maximum Link Utilization)",
        "3. Performance Ratio (PR)",
        "4. Network Disturbance (DB)",
        "5. Execution Time and Decision Latency",
        "6. Expert Selection, Accuracy, and Training Notes",
        "7. Failure Robustness",
        "8. SDN Deployment Validation",
        "9. CDF Plots",
        "10. Complexity Analysis",
        "11. Germany50 Unseen Topology Deep Dive",
        "12. Paper Baselines and External Comparison Status",
        "13. Limitations (Honest Assessment)",
        "14. Live Mininet Testbed Status",
        "15. Exact Thesis Method Description",
    ]
    for item in items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
    doc.add_page_break()


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], font_size: int = 8) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(header))
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.name = "Arial"
        set_cell_shading(cell, "D9EAF7")
    for row_vals in rows:
        row = table.add_row()
        for idx, value in enumerate(row_vals):
            cell = row.cells[idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(value))
            run.font.size = Pt(font_size)
            run.font.name = "Arial"


def add_image(doc: Document, path: Path, caption: str, width: float = 6.5) -> None:
    if not path.exists():
        doc.add_paragraph(f"[Missing image: {path}]")
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def load_data():
    mg_summary = pd.read_csv(METAGATE_DIR / "metagate_summary.csv")
    mg_results = pd.read_csv(METAGATE_DIR / "metagate_results.csv")
    mg_decisions = pd.read_csv(METAGATE_DIR / "metagate_decisions.csv")
    mg_timing = pd.read_csv(METAGATE_DIR / "metagate_timing.csv")

    for frame in [mg_summary, mg_results, mg_decisions, mg_timing]:
        frame["topology"] = frame["dataset"].map(canonical_topology)

    baseline_normal = pd.read_csv(BASELINE_DIR / "final_results.csv")
    baseline_failure = pd.read_csv(BASELINE_DIR / "failure_results.csv")
    for frame in [baseline_normal, baseline_failure]:
        frame["topology_norm"] = frame["topology"].map(canonical_topology)

    baseline_normal_agg = (
        baseline_normal[baseline_normal["scenario"] == "normal"]
        .groupby(["topology_norm", "method"], as_index=False)
        .agg(
            mean_mlu=("mean_mlu", "mean"),
            mean_pr=("mean_pr", "mean"),
            mean_disturbance=("mean_disturbance", "mean"),
            mean_total_time_ms=("mean_total_time_ms", "mean"),
        )
    )

    packet_sdn_results = pd.read_csv(PACKET_SDN_DIR / "packet_sdn_results.csv")
    packet_sdn_summary = pd.read_csv(PACKET_SDN_DIR / "packet_sdn_summary.csv")
    for frame in [packet_sdn_results, packet_sdn_summary]:
        frame["topology_norm"] = frame["topology"].map(canonical_topology)

    gnnplus_summary = pd.read_csv(GNNPLUS_PACKET_DIR / "packet_sdn_summary.csv")
    gnnplus_failure = pd.read_csv(GNNPLUS_PACKET_DIR / "packet_sdn_failure.csv")
    gnnplus_metrics = pd.read_csv(GNNPLUS_PACKET_DIR / "packet_sdn_sdn_metrics.csv")
    for frame in [gnnplus_summary, gnnplus_failure]:
        frame["topology_norm"] = frame["topology"].map(canonical_topology)
    gnnplus_metrics["topology_norm"] = gnnplus_metrics["Topology"].map(canonical_topology)

    gnnplus_stage1 = json.loads(GNNPLUS_STAGE1_SUMMARY.read_text(encoding="utf-8"))

    return {
        "mg_summary": mg_summary,
        "mg_results": mg_results,
        "mg_decisions": mg_decisions,
        "mg_timing": mg_timing,
        "baseline_normal": baseline_normal_agg,
        "baseline_failure": baseline_failure,
        "packet_sdn_results": packet_sdn_results,
        "packet_sdn_summary": packet_sdn_summary,
        "gnnplus_summary": gnnplus_summary,
        "gnnplus_failure": gnnplus_failure,
        "gnnplus_metrics": gnnplus_metrics,
        "gnnplus_stage1": gnnplus_stage1,
    }


def compute_pr_table(mg_results: pd.DataFrame) -> pd.DataFrame:
    df = mg_results.copy()
    df["meta_pr"] = df["metagate_mlu"] / df["oracle_mlu"]
    df["bn_pr"] = df["bn_mlu"] / df["oracle_mlu"]
    df["gnn_pr"] = df["gnn_mlu"] / df["oracle_mlu"]
    pr = df.groupby("topology", as_index=False).agg(
        meta_mean_pr=("meta_pr", "mean"),
        meta_p95_pr=("meta_pr", lambda s: float(np.quantile(s, 0.95))),
        bn_mean_pr=("bn_pr", "mean"),
        gnn_mean_pr=("gnn_pr", "mean"),
    )
    return pr


def compute_selector_table(mg_results: pd.DataFrame) -> pd.DataFrame:
    records = []
    for topo in ordered_topologies(mg_results["topology"].unique().tolist()):
        sub = mg_results[mg_results["topology"] == topo].copy()
        if sub.empty:
            continue
        counts = sub["metagate_selector"].value_counts()
        dominant = counts.idxmax()
        dominant_share = 100.0 * counts.max() / len(sub)
        switches = int((sub["metagate_selector"].values[1:] != sub["metagate_selector"].values[:-1]).sum())
        gap = ((sub["metagate_mlu"].mean() - sub["oracle_mlu"].mean()) / sub["oracle_mlu"].mean()) * 100.0
        records.append(
            {
                "topology": topo,
                "switches": f"{switches}/{max(len(sub) - 1, 1)}",
                "dominant": f"{EXPERT_LABELS[dominant]} ({dominant_share:.0f}%)",
                "accuracy": sub["correct"].mean() * 100.0,
                "gap_pct": gap,
            }
        )
    return pd.DataFrame(records)


def compute_metagate_overall_accuracy(mg_results: pd.DataFrame) -> float:
    return float(mg_results["correct"].mean() * 100.0)


def compute_metagate_packet_sdn_normal(packet_sdn_summary: pd.DataFrame) -> pd.DataFrame:
    normal = packet_sdn_summary[packet_sdn_summary["failure_type"] == "none"].copy()
    rows = []
    for topo in TOPOLOGY_ORDER:
        ecmp = normal[(normal["topology_norm"] == topo) & (normal["method"] == "ecmp")]
        mg = normal[(normal["topology_norm"] == topo) & (normal["method"] == "metagate")]
        if ecmp.empty or mg.empty:
            continue
        ecmp_row = ecmp.iloc[0]
        mg_row = mg.iloc[0]
        improvement = ((float(ecmp_row["mlu"]) - float(mg_row["mlu"])) / float(ecmp_row["mlu"])) * 100.0
        rows.append(
            {
                "topology": topo,
                "ecmp_mlu": float(ecmp_row["mlu"]),
                "metagate_mlu": float(mg_row["mlu"]),
                "improvement_pct": improvement,
                "decision_ms": float(mg_row["decision_time_ms"]),
                "rules": float(mg_row["rules_pushed"]),
            }
        )
    return pd.DataFrame(rows)


def compute_metagate_packet_sdn_failures(packet_sdn_summary: pd.DataFrame) -> dict[str, pd.DataFrame]:
    out = {}
    subset = packet_sdn_summary[packet_sdn_summary["method"].isin(["ecmp", "metagate"])].copy()
    for failure in ["single_link_failure", "capacity_degradation", "traffic_spike"]:
        rows = []
        fsub = subset[subset["failure_type"] == failure]
        for topo in TOPOLOGY_ORDER:
            ecmp = fsub[(fsub["topology_norm"] == topo) & (fsub["method"] == "ecmp")]
            mg = fsub[(fsub["topology_norm"] == topo) & (fsub["method"] == "metagate")]
            if ecmp.empty or mg.empty:
                continue
            rows.append(
                {
                    "topology": topo,
                    "ecmp_mlu": float(ecmp.iloc[0]["mlu"]),
                    "metagate_mlu": float(mg.iloc[0]["mlu"]),
                }
            )
        out[failure] = pd.DataFrame(rows)
    return out


def compute_gnnplus_failure_summary(gnnplus_failure: pd.DataFrame) -> pd.DataFrame:
    subset = gnnplus_failure[gnnplus_failure["method"].isin(["bottleneck", "gnn", "gnnplus"])].copy()
    grouped = (
        subset.groupby(["scenario", "method"], as_index=False)
        .agg(
            mean_mlu=("mean_mlu", "mean"),
            mean_recovery_ms=("failure_recovery_ms", "mean"),
        )
    )
    return grouped


def plot_metagate_gap_accuracy(mg_summary: pd.DataFrame) -> Path:
    plot_path = PLOTS_DIR / "metagate_gap_accuracy.png"
    summary = mg_summary.copy()
    summary["gap_pct"] = (summary["metagate_mlu"] - summary["oracle_mlu"]) / summary["oracle_mlu"] * 100.0
    summary = summary.sort_values("topology", key=lambda s: s.map({k: i for i, k in enumerate(TOPOLOGY_ORDER)}))

    fig, ax1 = plt.subplots(figsize=(11, 4.8))
    x = np.arange(len(summary))
    ax1.bar(x, summary["gap_pct"], color="#4c78a8", alpha=0.85, label="Oracle Gap (%)")
    ax1.set_ylabel("Oracle Gap (%)")
    ax1.set_xticks(x)
    ax1.set_xticklabels([display_topology(t) for t in summary["topology"]], rotation=20, ha="right")
    ax1.axhline(0.0, color="black", linewidth=0.8)

    ax2 = ax1.twinx()
    ax2.plot(x, summary["accuracy"] * 100.0, color="#d62728", marker="o", linewidth=2.0, label="Accuracy (%)")
    ax2.set_ylabel("Accuracy (%)")
    ax2.set_ylim(0, 100)

    ax1.set_title("MetaGate Accuracy and Oracle Gap by Topology")
    handles1, labels1 = ax1.get_legend_handles_labels()
    handles2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(handles1 + handles2, labels1 + labels2, loc="upper left")
    ax1.grid(True, axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(plot_path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return plot_path


def plot_selector_distribution(mg_results: pd.DataFrame) -> Path:
    plot_path = PLOTS_DIR / "selector_distribution.png"
    counts = (
        mg_results.groupby(["topology", "metagate_selector"])
        .size()
        .unstack(fill_value=0)
        .reindex(index=TOPOLOGY_ORDER, fill_value=0)
        .reindex(columns=EXPERT_ORDER, fill_value=0)
    )
    fig, ax = plt.subplots(figsize=(11, 5))
    bottom = np.zeros(len(counts))
    x = np.arange(len(counts))
    for expert in EXPERT_ORDER:
        vals = counts[expert].to_numpy()
        ax.bar(x, vals, bottom=bottom, color=METHOD_COLORS[expert], label=EXPERT_LABELS[expert])
        bottom += vals
    ax.set_xticks(x)
    ax.set_xticklabels([display_topology(t) for t in counts.index], rotation=20, ha="right")
    ax.set_ylabel("Selected Timesteps")
    ax.set_title("MetaGate Expert Selection Distribution")
    ax.legend(ncols=4, fontsize=8)
    ax.grid(True, axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(plot_path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return plot_path


def plot_metagate_pr_cdf(mg_results: pd.DataFrame) -> Path:
    plot_path = PLOTS_DIR / "metagate_pr_cdf.png"
    values = {
        "MetaGate": (mg_results["metagate_mlu"] / mg_results["oracle_mlu"]).to_numpy(),
        "Bottleneck": (mg_results["bn_mlu"] / mg_results["oracle_mlu"]).to_numpy(),
        "GNN": (mg_results["gnn_mlu"] / mg_results["oracle_mlu"]).to_numpy(),
    }
    fig, ax = plt.subplots(figsize=(8.5, 5))
    colors = {"MetaGate": METHOD_COLORS["metagate"], "Bottleneck": METHOD_COLORS["bottleneck"], "GNN": METHOD_COLORS["gnn"]}
    for label, arr in values.items():
        arr = np.asarray(arr, dtype=float)
        arr = arr[np.isfinite(arr)]
        arr.sort()
        y = np.arange(1, len(arr) + 1) / len(arr)
        ax.plot(arr, y, linewidth=2.0, label=label, color=colors[label])
    ax.set_xlabel("Performance Ratio (MLU / Oracle MLU)")
    ax.set_ylabel("CDF")
    ax.set_title("CDF of Performance Ratio")
    ax.grid(True, alpha=0.25)
    ax.legend()
    fig.tight_layout()
    fig.savefig(plot_path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return plot_path


def plot_packet_sdn_mlu_cdf(packet_sdn_results: pd.DataFrame) -> Path:
    plot_path = PLOTS_DIR / "packet_sdn_mlu_cdf.png"
    subset = packet_sdn_results[
        (packet_sdn_results["failure_type"] == "none") &
        (packet_sdn_results["method"].isin(["ecmp", "metagate"]))
    ].copy()
    fig, ax = plt.subplots(figsize=(8.5, 5))
    for method in ["ecmp", "metagate"]:
        arr = subset[subset["method"] == method]["mlu"].to_numpy(dtype=float)
        arr = arr[np.isfinite(arr)]
        arr.sort()
        y = np.arange(1, len(arr) + 1) / len(arr)
        ax.plot(arr, y, linewidth=2.0, label=method.upper(), color=METHOD_COLORS[method])
    ax.set_xlabel("Packet-SDN MLU")
    ax.set_ylabel("CDF")
    ax.set_title("Packet-SDN MLU CDF: ECMP vs MetaGate")
    ax.grid(True, alpha=0.25)
    ax.legend()
    fig.tight_layout()
    fig.savefig(plot_path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return plot_path


def plot_gnnplus_normal(gnnplus_summary: pd.DataFrame) -> Path:
    plot_path = PLOTS_DIR / "gnnplus_normal_mlu.png"
    subset = gnnplus_summary[gnnplus_summary["scenario"] == "normal"].copy()
    methods = ["ecmp", "bottleneck", "gnn", "gnnplus"]
    x = np.arange(len(TOPOLOGY_ORDER))
    width = 0.2
    fig, ax = plt.subplots(figsize=(12, 5.5))
    for idx, method in enumerate(methods):
        vals = []
        for topo in TOPOLOGY_ORDER:
            row = subset[(subset["topology_norm"] == topo) & (subset["method"] == method)]
            vals.append(float(row.iloc[0]["mean_mlu"]) if not row.empty else np.nan)
        ax.bar(x + (idx - 1.5) * width, vals, width=width, label=method.upper(), color=METHOD_COLORS[method])
    ax.set_xticks(x)
    ax.set_xticklabels([display_topology(t) for t in TOPOLOGY_ORDER], rotation=20, ha="right")
    ax.set_ylabel("Mean MLU")
    ax.set_title("GNN+ Packet-SDN Normal Comparison")
    ax.legend(ncols=4, fontsize=8)
    ax.grid(True, axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(plot_path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return plot_path


def plot_germany50_selection(mg_results: pd.DataFrame) -> Path:
    plot_path = PLOTS_DIR / "germany50_selection.png"
    subset = mg_results[mg_results["topology"] == "germany50"].copy()
    pred = subset["metagate_selector"].value_counts().reindex(EXPERT_ORDER, fill_value=0)
    oracle = subset["oracle_selector"].value_counts().reindex(EXPERT_ORDER, fill_value=0)
    x = np.arange(len(EXPERT_ORDER))
    width = 0.38
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.bar(x - width / 2, pred.values, width=width, label="Predicted", color="#1f77b4")
    ax.bar(x + width / 2, oracle.values, width=width, label="Oracle", color="#ff7f0e")
    ax.set_xticks(x)
    ax.set_xticklabels([EXPERT_LABELS[k] for k in EXPERT_ORDER])
    ax.set_ylabel("Timesteps")
    ax.set_title("Germany50: Predicted vs Oracle Expert Counts")
    ax.legend()
    ax.grid(True, axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(plot_path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return plot_path


def build_report(data: dict) -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    PLOTS_DIR.mkdir(parents=True, exist_ok=True)

    mg_summary = data["mg_summary"]
    mg_results = data["mg_results"]
    mg_timing = data["mg_timing"]
    baseline_normal = data["baseline_normal"]
    packet_sdn_summary = data["packet_sdn_summary"]
    packet_sdn_results = data["packet_sdn_results"]
    gnnplus_summary = data["gnnplus_summary"]
    gnnplus_failure = data["gnnplus_failure"]
    gnnplus_metrics = data["gnnplus_metrics"]
    stage1_summary = data["gnnplus_stage1"]

    pr_table = compute_pr_table(mg_results)
    selector_table = compute_selector_table(mg_results)
    metagate_packet_sdn = compute_metagate_packet_sdn_normal(packet_sdn_summary)
    metagate_failures = compute_metagate_packet_sdn_failures(packet_sdn_summary)
    gnnplus_failure_summary = compute_gnnplus_failure_summary(gnnplus_failure)
    overall_accuracy = compute_metagate_overall_accuracy(mg_results)

    plot_gap = plot_metagate_gap_accuracy(mg_summary)
    plot_selector = plot_selector_distribution(mg_results)
    plot_pr = plot_metagate_pr_cdf(mg_results)
    plot_packet = plot_packet_sdn_mlu_cdf(packet_sdn_results)
    plot_gnnplus = plot_gnnplus_normal(gnnplus_summary)
    plot_germany = plot_germany50_selection(mg_results)
    plot_gnnplus_failure = GNNPLUS_PACKET_DIR / "plots" / "failure_recovery_gnnplus.png"

    doc = Document()
    style_doc(doc)
    add_title_page(doc)
    add_toc(doc)

    # 1. Method Summary
    doc.add_heading("1. Method Summary", level=1)
    doc.add_paragraph(
        "This report follows the style of the advisor reference document while keeping the method scope honest. "
        "The main adaptive method in this repository is the MLP MetaGate: a per-traffic-matrix MLP chooser with "
        "few-shot Bayesian calibration that selects among four experts under a shared LP traffic-engineering pipeline."
    )
    doc.add_paragraph(
        "A separate GNN+ branch also exists and is now validated under packet-level SDN analytical simulation with "
        "a fixed-K=40 checkpoint. The current chooser has not yet been retrained with GNN+ as its fourth expert, "
        "so this report keeps MetaGate and GNN+ as related but distinct validated bundles."
    )
    add_bullet(doc, f"Overall MetaGate accuracy across 8 topologies: {overall_accuracy:.1f}%.")
    add_bullet(doc, "Stable MetaGate is intentionally excluded from this final report.")
    add_bullet(doc, "Paper baselines are listed in the report, but numeric results are shown only when the repo contains runnable outputs.")
    add_bullet(doc, "Packet-level SDN metrics in this report are model-based analytical metrics, not live Mininet measurements.")

    expert_rows = [
        ["Bottleneck", "Current MetaGate expert", "Included in chooser and report"],
        ["TopK", "Current MetaGate expert", "Included in chooser and report"],
        ["Sensitivity", "Current MetaGate expert", "Included in chooser and report"],
        ["Original GNN", "Current MetaGate expert", "Included in chooser and report"],
        ["GNN+", "Validated standalone / packet-SDN branch", "Included in report, not claimed as current chooser expert"],
        ["Stable MetaGate", "Inference-time extension", "Excluded intentionally from this report"],
    ]
    add_table(doc, ["Method", "Current Role", "Status in This Report"], expert_rows, font_size=8)

    topo_rows = []
    for topo in TOPOLOGY_ORDER:
        info = TOPOLOGY_INFO[topo]
        topo_rows.append([info["display"], info["type"], str(info["nodes"]), str(info["edges"])])
    doc.add_paragraph("")
    add_table(doc, ["Topology", "Type", "Nodes", "Edges"], topo_rows, font_size=8)
    doc.add_page_break()

    # 2. MLU Results
    doc.add_heading("2. MLU Results (Maximum Link Utilization)", level=1)
    doc.add_paragraph(
        "Section 2 follows the MetaGate narrative in the reference report. Table 2.1 is a controlled same-bundle "
        "comparison from the validated MetaGate CSVs. Table 2.2 adds external standalone baseline numbers from the "
        "corrected baseline bundle and should be read as indicative because it comes from a separate run bundle."
    )

    mg_rows = []
    for topo in TOPOLOGY_ORDER:
        row = mg_summary[mg_summary["topology"] == topo]
        if row.empty:
            continue
        rec = row.iloc[0]
        gap = ((float(rec["metagate_mlu"]) - float(rec["oracle_mlu"])) / float(rec["oracle_mlu"])) * 100.0
        mg_rows.append([
            display_topology_with_size(topo),
            display_type(topo),
            fmt(rec["metagate_mlu"]),
            fmt(rec["oracle_mlu"]),
            fmt(rec["bn_mlu"]),
            fmt(rec["topk_mlu"]),
            fmt(rec["sens_mlu"]),
            fmt(rec["gnn_mlu"]),
            f"{gap:+.2f}%",
        ])
    add_table(
        doc,
        ["Topology", "Type", "MetaGate", "Oracle", "BN", "TopK", "Sens", "GNN", "Gap"],
        mg_rows,
        font_size=7,
    )
    add_image(doc, plot_gap, "Figure 2. MetaGate oracle gap and accuracy by topology.", width=6.7)

    doc.add_paragraph("")
    doc.add_paragraph(
        "Indicative external baseline comparison. MetaGate values come from the validated MetaGate bundle; the "
        "standalone baselines come from the corrected baseline bundle. Use this table for broad ranking intuition, "
        "not for strict within-run significance claims."
    )
    ext_rows = []
    for topo in TOPOLOGY_ORDER:
        mg_row = mg_summary[mg_summary["topology"] == topo]
        if mg_row.empty:
            continue
        def base(method_name: str) -> str:
            row = baseline_normal[(baseline_normal["topology_norm"] == topo) & (baseline_normal["method"] == method_name)]
            return fmt(row.iloc[0]["mean_mlu"]) if not row.empty else "N/A"
        ext_rows.append([
            display_topology_with_size(topo),
            fmt(mg_row.iloc[0]["metagate_mlu"]),
            base("Bottleneck"),
            base("Sensitivity"),
            base("GNN"),
            base("ECMP"),
            base("OSPF"),
        ])
    add_table(doc, ["Topology", "MetaGate", "Bottleneck", "Sensitivity", "GNN", "ECMP", "OSPF"], ext_rows, font_size=7)
    doc.add_page_break()

    # 3. Performance Ratio
    doc.add_heading("3. Performance Ratio (PR)", level=1)
    doc.add_paragraph(
        "Performance ratio is defined here as method MLU divided by oracle MLU within the validated MetaGate bundle. "
        "A value of 1.000 means the method matches the oracle exactly."
    )
    pr_rows = []
    for topo in TOPOLOGY_ORDER:
        row = pr_table[pr_table["topology"] == topo]
        if row.empty:
            continue
        rec = row.iloc[0]
        pr_rows.append([
            display_topology_with_size(topo),
            display_type(topo),
            fmt(rec["meta_mean_pr"], 6),
            fmt(rec["meta_p95_pr"], 6),
            fmt(rec["bn_mean_pr"], 6),
            fmt(rec["gnn_mean_pr"], 6),
        ])
    add_table(doc, ["Topology", "Type", "MetaGate Mean PR", "MetaGate P95 PR", "BN Mean PR", "GNN Mean PR"], pr_rows, font_size=8)
    add_image(doc, plot_pr, "Figure 3. CDF of MetaGate, Bottleneck, and GNN performance ratio.", width=6.3)
    doc.add_page_break()

    # 4. Disturbance
    doc.add_heading("4. Network Disturbance (DB)", level=1)
    doc.add_paragraph(
        "MetaGate does not store a standalone disturbance score in the validated chooser CSVs because the realized "
        "disturbance depends on which expert is selected and how often the selected expert changes between timesteps. "
        "The table below shows standalone expert disturbance from runnable baseline bundles plus the separate GNN+ "
        "packet-SDN branch disturbance."
    )
    db_rows = []
    for topo in TOPOLOGY_ORDER:
        def lookup(method_name: str) -> str:
            row = baseline_normal[(baseline_normal["topology_norm"] == topo) & (baseline_normal["method"] == method_name)]
            return fmt(row.iloc[0]["mean_disturbance"]) if not row.empty else "N/A"
        gnnplus_row = gnnplus_summary[(gnnplus_summary["topology_norm"] == topo) & (gnnplus_summary["method"] == "gnnplus") & (gnnplus_summary["scenario"] == "normal")]
        gnnplus_db = fmt(gnnplus_row.iloc[0]["mean_disturbance"]) if not gnnplus_row.empty else "N/A"
        db_rows.append([
            display_topology_with_size(topo),
            lookup("Bottleneck"),
            lookup("TopK"),
            lookup("Sensitivity"),
            lookup("GNN"),
            gnnplus_db,
        ])
    add_table(doc, ["Topology", "BN DB", "TopK DB", "Sens DB", "GNN DB", "GNN+ DB"], db_rows, font_size=8)
    doc.add_page_break()

    # 5. Execution time
    doc.add_heading("5. Execution Time and Decision Latency", level=1)
    doc.add_paragraph(
        "The timing table below is taken from the validated MetaGate timing CSV. Decision time means all chooser "
        "overhead before LP: running the four experts, extracting features, and applying the MLP."
    )
    timing_rows = []
    for topo in TOPOLOGY_ORDER:
        row = mg_timing[mg_timing["topology"] == topo]
        if row.empty:
            continue
        rec = row.iloc[0]
        timing_rows.append([
            display_topology_with_size(topo),
            str(TOPOLOGY_INFO[topo]["nodes"]),
            fmt(rec["t_bn_ms"], 1),
            fmt(rec["t_topk_ms"], 1),
            fmt(rec["t_sens_ms"], 1),
            fmt(rec["t_gnn_ms"], 1),
            fmt(rec["t_features_ms"], 2),
            fmt(rec["t_mlp_ms"], 3),
            fmt(rec["t_lp_ms"], 1),
            fmt(rec["t_total_ms"], 1),
        ])
    add_table(
        doc,
        ["Topology", "N", "BN ms", "TopK ms", "Sens ms", "GNN ms", "Feat ms", "MLP ms", "LP ms", "Total ms"],
        timing_rows,
        font_size=7,
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        f"The MLP forward pass itself remains tiny: mean per-topology MLP time is {mg_timing['t_mlp_ms'].mean():.3f} ms. "
        "The dominant cost is evaluating all four experts and then solving the LP."
    )
    doc.add_page_break()

    # 6. Selection / training
    doc.add_heading("6. Expert Selection, Accuracy, and Training Notes", level=1)
    selector_rows = []
    for topo in TOPOLOGY_ORDER:
        row = selector_table[selector_table["topology"] == topo]
        if row.empty:
            continue
        rec = row.iloc[0]
        selector_rows.append([
            display_topology_with_size(topo),
            rec["switches"],
            rec["dominant"],
            f"{rec['accuracy']:.1f}%",
            f"{rec['gap_pct']:+.2f}%",
        ])
    add_table(doc, ["Topology", "Switches", "Dominant Expert", "Accuracy", "Oracle Gap"], selector_rows, font_size=8)
    add_image(doc, plot_selector, "Figure 6. MetaGate selected-expert distribution by topology.", width=6.6)

    doc.add_paragraph("")
    doc.add_paragraph("Training and provenance notes:")
    train_rows = [
        ["MetaGate hidden dim", "128"],
        ["MetaGate dropout", "0.3"],
        ["MetaGate learning rate", "5e-4"],
        ["MetaGate epochs / batch size", "300 / 64"],
        ["MetaGate calibration", "10 validation traffic matrices per topology"],
        ["MetaGate checkpoint", "results/dynamic_metagate/models/metagate_unified.pt"],
        ["GNN+ checkpoint used", str(GNNPLUS_CHECKPOINT.relative_to(PROJECT_ROOT))],
        ["GNN+ source checkpoint", str(GNNPLUS_SOURCE_CHECKPOINT.relative_to(PROJECT_ROOT))],
        ["GNN+ dropout / fixed K", "0.2 / 40"],
        ["GNN+ checkpoint dims", "node=16, edge=12, od=18"],
        ["GNN+ stage-1 total training time", f"{float(stage1_summary['total_time']):.2f} s"],
    ]
    add_table(doc, ["Artifact", "Value"], train_rows, font_size=8)
    doc.add_page_break()

    # 7. Failure robustness
    doc.add_heading("7. Failure Robustness", level=1)
    doc.add_paragraph(
        "This section keeps the failure-scenario material similar to the reference report, but it separates the two "
        "validated packet-SDN bundles: MetaGate packet-SDN and GNN+ packet-SDN."
    )
    for failure in ["single_link_failure", "capacity_degradation", "traffic_spike"]:
        label = FAILURE_LABELS[failure]
        doc.add_heading(f"7.{['single_link_failure','capacity_degradation','traffic_spike'].index(failure)+1} {label}", level=2)
        rows = []
        fdf = metagate_failures[failure]
        for topo in TOPOLOGY_ORDER:
            row = fdf[fdf["topology"] == topo]
            if row.empty:
                continue
            rec = row.iloc[0]
            rows.append([
                display_topology_with_size(topo),
                fmt(rec["ecmp_mlu"]),
                fmt(rec["metagate_mlu"]),
            ])
        add_table(doc, ["Topology", "ECMP MLU", "MetaGate MLU"], rows, font_size=8)

    doc.add_heading("7.4 GNN+ Packet-SDN Failure Summary", level=2)
    doc.add_paragraph(
        "The fixed GNN+ packet-SDN branch includes a richer five-scenario failure study. The table below reports "
        "scenario-average recovery numbers for Bottleneck, GNN, and GNN+ from that bundle."
    )
    gnnplus_fail_rows = []
    for scenario in [
        "single_link_failure",
        "random_link_failure_1",
        "random_link_failure_2",
        "capacity_degradation_50",
        "traffic_spike_2x",
    ]:
        label = FAILURE_LABELS[scenario]
        for method in ["bottleneck", "gnn", "gnnplus"]:
            row = gnnplus_failure_summary[
                (gnnplus_failure_summary["scenario"] == scenario) &
                (gnnplus_failure_summary["method"] == method)
            ]
            if row.empty:
                continue
            rec = row.iloc[0]
            gnnplus_fail_rows.append([
                label,
                display_method(method),
                fmt(rec["mean_mlu"]),
                fmt(rec["mean_recovery_ms"], 1),
            ])
    add_table(doc, ["Scenario", "Method", "Mean MLU", "Mean Recovery (ms)"], gnnplus_fail_rows, font_size=8)
    add_image(doc, plot_gnnplus_failure, "Figure 7. GNN+ packet-SDN recovery by topology and scenario.", width=6.3)
    doc.add_page_break()

    # 8. SDN deployment
    doc.add_heading("8. SDN Deployment Validation", level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        "HONESTY NOTE: The SDN results below are model-based analytical control-loop simulations. "
        "They use real topology graphs and traffic matrices, but they do not come from live Mininet packet forwarding."
    )
    r.bold = True
    r.font.color.rgb = RGBColor(0xAA, 0x00, 0x00)

    doc.add_heading("8.1 MetaGate Packet-SDN Normal Results", level=2)
    mg_sdn_rows = []
    for topo in TOPOLOGY_ORDER:
        row = metagate_packet_sdn[metagate_packet_sdn["topology"] == topo]
        if row.empty:
            continue
        rec = row.iloc[0]
        mg_sdn_rows.append([
            display_topology_with_size(topo),
            fmt(rec["ecmp_mlu"]),
            fmt(rec["metagate_mlu"]),
            f"{rec['improvement_pct']:.2f}%",
            fmt(rec["decision_ms"], 1),
            fmt(rec["rules"], 1),
        ])
    add_table(doc, ["Topology", "Pre-MLU (ECMP)", "Post-MLU (MetaGate)", "Improvement", "Decision ms", "Rules/Cycle"], mg_sdn_rows, font_size=8)
    add_image(doc, plot_packet, "Figure 8. Packet-SDN MLU CDF for MetaGate vs ECMP.", width=6.3)

    doc.add_heading("8.2 GNN+ Packet-SDN Normal Results", level=2)
    doc.add_paragraph(
        "The fixed GNN+ packet-SDN branch was audited and repaired earlier. The table below comes from the corrected "
        "bundle and keeps all 8 topologies and all four packet-SDN methods visible."
    )
    gnnplus_rows = []
    for topo in TOPOLOGY_ORDER:
        subset = gnnplus_summary[(gnnplus_summary["topology_norm"] == topo) & (gnnplus_summary["scenario"] == "normal")]
        if subset.empty:
            continue
        val = {}
        for method in ["ecmp", "bottleneck", "gnn", "gnnplus"]:
            row = subset[subset["method"] == method]
            val[method] = float(row.iloc[0]["mean_mlu"]) if not row.empty else np.nan
        gnnplus_decision = subset[subset["method"] == "gnnplus"].iloc[0]["decision_time_ms"]
        gnnplus_rows.append([
            display_topology_with_size(topo),
            fmt(val["ecmp"]),
            fmt(val["bottleneck"]),
            fmt(val["gnn"]),
            fmt(val["gnnplus"]),
            fmt(gnnplus_decision, 1),
        ])
    add_table(doc, ["Topology", "ECMP", "Bottleneck", "GNN", "GNN+", "GNN+ Decision ms"], gnnplus_rows, font_size=7)
    add_image(doc, plot_gnnplus, "Figure 9. GNN+ packet-SDN normal comparison across 8 topologies.", width=6.6)
    doc.add_page_break()

    # 9. CDF plots
    doc.add_heading("9. CDF Plots", level=1)
    doc.add_paragraph(
        "The reference report relied heavily on CDF plots. This report keeps two high-value CDFs: MetaGate "
        "performance ratio and MetaGate packet-SDN MLU. These are generated from per-timestep raw outputs."
    )
    add_image(doc, plot_pr, "Figure 10. MetaGate PR CDF.", width=6.3)
    add_image(doc, plot_packet, "Figure 11. Packet-SDN MLU CDF: ECMP vs MetaGate.", width=6.3)
    doc.add_page_break()

    # 10. Complexity
    doc.add_heading("10. Complexity Analysis", level=1)
    doc.add_paragraph(
        "The table below mirrors the complexity section style in the reference report. Theoretical entries are "
        "high-level controller-side complexity statements rather than formal proofs."
    )
    complexity_rows = [
        ["ECMP", "Shortest-path baseline", "Lowest controller overhead; no learned inference"],
        ["OSPF", "Routing-table recomputation baseline", "Comparable to ECMP for this report's scope"],
        ["TopK", "Sort OD demands + LP on selected flows", "Light heuristic plus LP"],
        ["Bottleneck", "Bottleneck analysis + LP", "Heuristic selection plus LP"],
        ["Sensitivity", "Sensitivity scoring + LP", "Heavier heuristic scoring plus LP"],
        ["Original GNN expert", "Graph inference + LP", "Learned selector with graph feature extraction"],
        ["MetaGate", "Run 4 experts + features + MLP + LP", "Highest adaptive control cost in current validated app"],
        ["GNN+ standalone", "Enhanced graph inference + LP", "Same class as GNN with richer features and fixed K=40"],
    ]
    add_table(doc, ["Method", "Theoretical Controller Work", "Interpretation"], complexity_rows, font_size=8)

    doc.add_paragraph("")
    doc.add_paragraph(
        "Numeric controller-cost summary. These numbers come from their native validated bundles, so they are useful "
        "for engineering intuition but not a strict apples-to-apples benchmark."
    )
    numeric_rows = [
        [
            "MetaGate",
            "results/dynamic_metagate",
            fmt(mg_timing["t_decision_ms"].mean(), 1),
            fmt(mg_timing["t_total_ms"].mean(), 1),
            "4 experts + MLP + LP",
        ],
    ]
    for method in ["ecmp", "bottleneck", "gnn", "gnnplus"]:
        subset = gnnplus_summary[(gnnplus_summary["scenario"] == "normal") & (gnnplus_summary["method"] == method)]
        numeric_rows.append(
            [
                display_method(method),
                "results/gnnplus_packet_sdn_report_fixed",
                fmt(subset["decision_time_ms"].mean(), 1),
                "N/A",
                "Packet-SDN analytical branch",
            ]
        )
    add_table(doc, ["Method", "Source Bundle", "Avg Decision ms", "Avg Total ms", "Notes"], numeric_rows, font_size=8)
    doc.add_page_break()

    # 11. Germany50
    doc.add_heading("11. Germany50 Unseen Topology Deep Dive", level=1)
    germany = mg_summary[mg_summary["topology"] == "germany50"].iloc[0]
    germany_sel = mg_results[mg_results["topology"] == "germany50"]["metagate_selector"].value_counts()
    gnn_share = 100.0 * germany_sel.get("gnn", 0) / max(germany_sel.sum(), 1)
    bot_share = 100.0 * germany_sel.get("bottleneck", 0) / max(germany_sel.sum(), 1)
    doc.add_paragraph(
        f"In the validated MetaGate bundle, Germany50 is the clearest unseen-topology case where the chooser does not "
        f"collapse entirely to Bottleneck. MetaGate accuracy is {germany['accuracy'] * 100.0:.1f}% with oracle gap "
        f"{((germany['metagate_mlu'] - germany['oracle_mlu']) / germany['oracle_mlu']) * 100.0:+.2f}%. "
        f"The chooser selects GNN on {gnn_share:.0f}% of test timesteps and Bottleneck on {bot_share:.0f}%."
    )
    gnnplus_germany = gnnplus_summary[
        (gnnplus_summary["topology_norm"] == "germany50") &
        (gnnplus_summary["scenario"] == "normal")
    ].copy()
    msg = []
    for method in ["ecmp", "bottleneck", "gnn", "gnnplus"]:
        row = gnnplus_germany[gnnplus_germany["method"] == method]
        if not row.empty:
            msg.append(f"{display_method(method)}={float(row.iloc[0]['mean_mlu']):.4f}")
    doc.add_paragraph(
        "In the separate fixed GNN+ packet-SDN bundle, Germany50 normal-scenario MLU is: " +
        ", ".join(msg) +
        ". This suggests GNN+ remains competitive on the unseen topology, but it does not yet prove the performance "
        "of a retrained MetaGate-with-GNN+ chooser because that integrated run has not been executed."
    )
    add_image(doc, plot_germany, "Figure 12. Germany50 predicted vs oracle expert counts.", width=5.8)
    doc.add_page_break()

    # 12. Baselines status
    doc.add_heading("12. Paper Baselines and External Comparison Status", level=1)
    doc.add_paragraph(
        "The advisor asked that all baselines, including paper baselines, be visible in the final report. This section "
        "does that honestly: runnable baselines are listed with bundle evidence, and non-runnable paper baselines are "
        "listed explicitly as unavailable instead of being silently omitted."
    )
    baseline_rows = [
        ["ECMP", "Runnable", "Corrected baseline bundle and packet-SDN bundles"],
        ["OSPF", "Runnable", "Corrected baseline bundle"],
        ["TopK", "Runnable", "Corrected baseline bundle; current MetaGate expert"],
        ["Sensitivity", "Runnable", "Corrected baseline bundle; current MetaGate expert"],
        ["Bottleneck", "Runnable", "Corrected baseline bundle; current MetaGate expert"],
        ["Original GNN", "Runnable", "Corrected baseline bundle; current MetaGate expert"],
        ["MLP MetaGate", "Runnable", "Validated chooser bundle across all 8 topologies"],
        ["GNN+", "Runnable in separate branch", "Fixed packet-SDN bundle and earlier focused GNN+ bundles"],
        ["Stable MetaGate", "Intentionally excluded", "Available as extension bundle but not part of this final report"],
        ["FlexDATE", "Not runnable in current repo", "Listed only as unavailable"],
        ["FlexEntry", "Not runnable in current repo", "Listed only as unavailable"],
        ["ERODRL", "Not runnable in current repo", "Listed only as unavailable"],
        ["CFRRL", "Not runnable in current repo", "Listed only as unavailable"],
    ]
    add_table(doc, ["Baseline / Method", "Status", "How It Appears Here"], baseline_rows, font_size=8)
    doc.add_page_break()

    # 13. Limitations
    doc.add_heading("13. Limitations (Honest Assessment)", level=1)
    add_bullet(doc, "The current MetaGate chooser still uses Original GNN as its fourth expert. GNN+ has not yet been plugged into the chooser and retrained.")
    add_bullet(doc, "Any direct MetaGate-vs-GNN+ conclusion across bundles is indicative only unless both are rerun under one shared pipeline.")
    add_bullet(doc, "Paper baselines such as FlexDATE, FlexEntry, ERODRL, and CFRRL are not runnable from the current repo and therefore do not appear as numeric result rows.")
    add_bullet(doc, "Packet-level SDN metrics are analytical control-loop simulations, not live packet forwarding measurements.")
    add_bullet(doc, "Stable MetaGate is available as a separate extension bundle but is intentionally excluded to keep the final report cleaner.")
    add_bullet(doc, "The GNN+ packet-SDN branch is now fixed and honest, but it is still a separate branch from the validated MetaGate chooser study.")
    doc.add_page_break()

    # 14. Mininet status
    doc.add_heading("14. Live Mininet Testbed Status", level=1)
    doc.add_paragraph(
        "No live Mininet packet-forwarding evaluation is claimed in this report. The current SDN deployment sections use "
        "model-based packet-level metrics derived from the control-loop outputs. A future live Mininet study would be "
        "needed to measure true packet delay, throughput, rule installation latency, and recovery in an emulated network."
    )
    doc.add_page_break()

    # 15. Thesis description
    doc.add_heading("15. Exact Thesis Method Description", level=1)
    doc.add_paragraph(
        "The validated adaptive traffic-engineering method in the current repository is an MLP MetaGate with few-shot "
        "Bayesian calibration. For each traffic matrix, the system runs four experts (Bottleneck, TopK, Sensitivity, "
        "and Original GNN), extracts a 49-dimensional feature vector from traffic, expert outputs, topology metrics, "
        "and ECMP baseline state, and then uses a small MLP to predict which expert will achieve the lowest maximum "
        "link utilization after LP rerouting. Calibration uses 10 validation traffic matrices per topology to build a "
        "topology-specific prior without gradient updates."
    )
    doc.add_paragraph(
        "Separately, a fixed-budget GNN+ branch has been validated with checkpoint "
        f"{GNNPLUS_CHECKPOINT.relative_to(PROJECT_ROOT)} using dropout=0.2 and fixed K=40. That branch shows that the "
        "enhanced learned selector is viable under packet-level SDN analytical simulation across 8 topologies and multiple "
        "failure scenarios. However, the current report does not falsely claim that GNN+ is already the fourth MetaGate expert; "
        "that integration would require a fresh retraining and evaluation of the chooser."
    )

    doc.save(OUTPUT_DOC)

    audit_lines = [
        "# MetaGate + GNN+ Final Report Audit",
        "",
        f"- Output DOCX: `{OUTPUT_DOC.relative_to(PROJECT_ROOT)}`",
        f"- Plot directory: `{PLOTS_DIR.relative_to(PROJECT_ROOT)}`",
        "- Source bundles used:",
        "  - `results/dynamic_metagate/*.csv`",
        "  - `results/final_full_eval_corrected/*.csv`",
        "  - `results/packet_sdn_simulation/*.csv`",
        "  - `results/gnnplus_packet_sdn_report_fixed/*.csv`",
        "- Stable MetaGate intentionally excluded.",
        "- GNN+ included only where validated outputs actually exist.",
        "- Paper baselines listed as unavailable, not fabricated.",
        "- Live Mininet not claimed anywhere in the new report.",
    ]
    AUDIT_MD.write_text("\n".join(audit_lines), encoding="utf-8")


def main() -> None:
    data = load_data()
    build_report(data)
    print(f"Wrote {OUTPUT_DOC}")
    print(f"Wrote {AUDIT_MD}")


if __name__ == "__main__":
    main()

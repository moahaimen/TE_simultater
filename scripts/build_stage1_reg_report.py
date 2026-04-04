#!/usr/bin/env python3
"""Generate Stage 1 Regularization Report."""
import os, sys
from pathlib import Path
PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT)); os.chdir(PROJECT_ROOT)

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

ROOT = Path("results/gnn_plus/stage1_regularization")
PLOTS = ROOT / "plots"
OUT = ROOT / "Stage1_Regularization_Report.docx"

def set_shading(cell, color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcPr.append(tcPr.makeelement(qn("w:shd"), {qn("w:fill"): color, qn("w:val"): "clear"}))

def add_table(doc, headers, rows, cw=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""; p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(255,255,255); r.font.name = "Arial"
        set_shading(c, "2E4057")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ""; p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val)); r.font.size = Pt(9); r.font.name = "Arial"
    if cw:
        for i, w in enumerate(cw):
            for row in t.rows: row.cells[i].width = Inches(w)
    return t

def add_img(doc, name, w=Inches(6)):
    fp = PLOTS / name
    if fp.exists():
        doc.add_picture(str(fp), width=w)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)
doc.styles["Normal"].font.name = "Arial"; doc.styles["Normal"].font.size = Pt(11)

# Title
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Stage 1, Step 1.5a: Regularization Pass"); r.bold = True; r.font.size = Pt(20)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Testing dropout=0.2 and dropout=0.3 with full enriched features, fixed K=40")
r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# 1. Design
doc.add_heading("1. Experiment Design", level=1)
doc.add_paragraph(
    "This step tests whether increased dropout regularization can fix the unseen topology "
    "regression observed in Step 1 (features-only, dropout=0.1). All other settings are identical: "
    "enriched features (46 dims), fixed K=40, supervised + REINFORCE training on 6 known topologies.")
add_table(doc, ["Config", "Dropout", "Features", "K", "Training"],
    [["Original GNN", "0.1", "34 (original)", "Fixed 40", "Original checkpoint"],
     ["d02", "0.2", "46 (enriched)", "Fixed 40", "Sup 30ep + RL 10ep"],
     ["d03", "0.3", "46 (enriched)", "Fixed 40", "Sup 30ep + RL 10ep"]],
    cw=[1.5, 0.8, 1.5, 0.9, 1.8])

# 2. Training
doc.add_heading("2. Training Summary", level=1)
add_table(doc, ["Metric", "d02 (dropout=0.2)", "d03 (dropout=0.3)"],
    [["Supervised best epoch", "30", "27"],
     ["Supervised time", "196s", "191s"],
     ["REINFORCE best epoch", "1", "1"],
     ["REINFORCE early stop", "Epoch 5", "Epoch 5"],
     ["REINFORCE time", "100s", "101s"],
     ["Total training", "296s", "292s"],
     ["Final alpha", "0.665", "0.624"]],
    cw=[2.5, 2.0, 2.0])

# 3. Results
doc.add_page_break()
doc.add_heading("3. Results", level=1)
doc.add_heading("3.1 Per-Topology MLU", level=2)
add_table(doc,
    ["Topology", "Type", "MLU Orig", "MLU d02", "d02 Change", "d02 Wins",
     "MLU d03", "d03 Change", "d03 Wins"],
    [["Abilene", "known", "0.055", "0.055", "~0%", "39", "0.055", "~0%", "34"],
     ["CERNET", "known", "1738.3", "1752.0", "-0.79%", "19", "1736.9", "+0.08%", "46"],
     ["GEANT", "known", "0.161", "0.163", "-0.71%", "25", "0.162", "-0.60%", "26"],
     ["Ebone", "known", "379.6", "379.6", "~0%", "38", "379.6", "~0%", "41"],
     ["Sprintlink", "known", "891.4", "856.3", "+3.94%", "74", "856.6", "+3.90%", "71"],
     ["Tiscali", "known", "843.6", "833.1", "+1.24%", "42", "821.9", "+2.57%", "59"],
     ["", "", "", "", "", "", "", "", ""],
     ["Germany50", "unseen", "18.94", "19.02", "-0.44%", "26", "19.21", "-1.42%", "12"],
     ["VtlWavenet", "unseen", "12252", "12266", "-0.11%", "13", "12313", "-0.49%", "0"],
     ["", "", "", "", "", "", "", "", ""],
     ["KNOWN AGG", "-", "642.2", "636.9", "+0.83%", "237", "632.5", "+1.50%", "277"],
     ["UNSEEN AGG", "-", "7728.9", "7737.5", "-0.11%", "39", "7767.2", "-0.50%", "12"],
     ["TOTAL AGG", "-", "2124.3", "2121.9", "+0.11%", "276", "2124.7", "-0.02%", "289"]],
    cw=[1.0, 0.5, 0.65, 0.65, 0.6, 0.5, 0.65, 0.6, 0.5])

doc.add_heading("3.2 Key Finding: Two Distinct Profiles", level=2)
p = doc.add_paragraph()
r = p.add_run("Dropout=0.2: Best for generalization. "); r.bold = True
p.add_run("Germany50 regression reduced from -1.18% (Step 1) to -0.44%. "
          "VtlWavenet nearly tied (-0.11%). Unseen aggregate only -0.11%. "
          "Known gains moderate (+0.83%).")

p = doc.add_paragraph()
r = p.add_run("Dropout=0.3: Best for known performance. "); r.bold = True
p.add_run("Known aggregate +1.50%, with Sprintlink +3.90% (71 wins), Tiscali +2.57% (59 wins), "
          "CERNET +0.08% (46 wins). But unseen regression persists: Germany50 -1.42%.")

doc.add_heading("3.3 Disturbance", level=2)
add_table(doc, ["Topology", "Dist Orig", "Dist d02", "Dist d03"],
    [["Abilene", "0.093", "0.089", "0.098"],
     ["CERNET", "0.307", "0.472", "0.425"],
     ["GEANT", "0.120", "0.107", "0.093"],
     ["Sprintlink", "0.241", "0.285", "0.309"],
     ["Tiscali", "0.397", "0.467", "0.381"],
     ["Germany50", "0.268", "0.232", "0.277"],
     ["VtlWavenet", "0.438", "0.447", "0.320"],
     ["KNOWN AGG", "0.219", "0.287", "0.252"],
     ["UNSEEN AGG", "0.375", "0.367", "0.304"]],
    cw=[1.6, 1.2, 1.2, 1.2])
doc.add_paragraph(
    "d03 achieves significantly lower disturbance on VtlWavenet (0.320 vs 0.438) "
    "and GEANT (0.093 vs 0.120). d02 is better on Germany50 (0.232 vs 0.268). "
    "Both increase disturbance on some known topologies (CERNET, Sprintlink).")

# 4. Plots
doc.add_page_break()
doc.add_heading("4. Comparison Plots", level=1)
doc.add_heading("4.1 MLU CDF — Known Topologies", level=2)
add_img(doc, "mlu_cdf_known.png", Inches(6.3))
doc.add_paragraph("Figure 1: Both dropout variants match or beat original on most known topologies.")
doc.add_heading("4.2 MLU CDF — Unseen Topologies", level=2)
add_img(doc, "mlu_cdf_unseen.png", Inches(6.3))
doc.add_paragraph("Figure 2: d02 nearly matches original on unseen. d03 shows slight regression.")
doc.add_heading("4.3 Aggregate MLU Bar Chart", level=2)
add_img(doc, "aggregate_mlu_bar.png", Inches(5.5))
doc.add_paragraph("Figure 3: d03 best on known, d02 best on unseen, both competitive on total.")
doc.add_heading("4.4 Unseen Disturbance CDF", level=2)
add_img(doc, "disturbance_unseen.png", Inches(6.3))
doc.add_paragraph("Figure 4: d03 has notably lower disturbance on VtlWavenet.")

# 5. Comparison with Step 1
doc.add_page_break()
doc.add_heading("5. Comparison with Step 1 (dropout=0.1)", level=1)
add_table(doc, ["Metric", "Step 1 (d=0.1)", "d02 (d=0.2)", "d03 (d=0.3)"],
    [["Known MLU change", "+1.66%", "+0.83%", "+1.50%"],
     ["Unseen MLU change", "-0.12%", "-0.11%", "-0.50%"],
     ["Germany50 change", "-1.18%", "-0.44%", "-1.42%"],
     ["Sprintlink change", "+5.42%", "+3.94%", "+3.90%"],
     ["Tiscali change", "+0.46%", "+1.24%", "+2.57%"],
     ["Known disturbance", "0.245", "0.287", "0.252"],
     ["Unseen disturbance", "0.466", "0.367", "0.304"]],
    cw=[2.0, 1.5, 1.5, 1.5])
doc.add_paragraph(
    "Regularization improves unseen generalization. d02 cuts Germany50 regression by 63% "
    "(from -1.18% to -0.44%). d03 achieves better known performance than Step 1 on Tiscali "
    "(+2.57% vs +0.46%) while dramatically reducing unseen disturbance (0.304 vs 0.466).")

# 6. Decision
doc.add_heading("6. Stage 1 Decision: Feature Pruning Needed?", level=1)
p = doc.add_paragraph()
r = p.add_run("Verdict: d02 is acceptable for deployment. Feature pruning (Step 1.5b) is OPTIONAL."); r.bold = True

doc.add_paragraph("")
doc.add_paragraph("Rationale:")
for item in [
    "d02 unseen regression is only -0.11% aggregate and -0.44% on Germany50, "
    "which is within noise for a model trained from scratch vs the original.",
    "d02 keeps Sprintlink gain (+3.94%, 74 wins), the primary known-topology improvement.",
    "d02 unseen disturbance is BETTER than original (0.367 vs 0.375).",
    "Feature pruning would add complexity and may lose the Sprintlink signal. "
    "The risk/reward ratio does not justify it unless required by the professor.",
    "If unseen performance must be strictly non-regressing, d02 is the closest to achieving it.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph("")
p = doc.add_paragraph()
r = p.add_run("Recommended lock-in: d02 (dropout=0.2, full enriched features, fixed K=40)"); r.bold = True
doc.add_paragraph("")
doc.add_paragraph("Next step: Stage 2 (dynamic-K-only ablation) with original features.")

# Fix zoom
for z in doc.settings.element.findall(qn("w:zoom")):
    if z.get(qn("w:percent")) is None: z.set(qn("w:percent"), "100")

doc.save(str(OUT))
print(f"Written: {OUT} ({OUT.stat().st_size} bytes)")

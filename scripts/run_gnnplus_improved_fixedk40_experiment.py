#!/usr/bin/env python3
"""Train and evaluate the fixed-K40 failgate GNN+ reference branch.

This runner turns the implemented selector / training fixes into an actual
checkpoint, a clean zero-shot evaluation bundle, and a separate proportional-
budget fairness study.

Scientific scope:
  - methods: ECMP, Bottleneck, Original GNN, GNN+
  - no MetaGate / Stable MetaGate
  - fixed K = 40
  - unseen evaluation is zero-shot

Improvement scope:
  - section 1: failure-gated selector engine and cached topology tensors
  - section 2: topology-balanced supervised pretraining + gentle RL fine-tuning
  - section 3: proportional-budget fairness/scalability study
  - section 4: optional hinge-reward / deeper-receptive-field ablations
"""

from __future__ import annotations

import importlib.util
import json
import math
import networkx as nx
import os
import random
import shutil
import sys
import time
from collections import defaultdict
from collections import deque
from dataclasses import asdict
from pathlib import Path

os.environ.setdefault("MPLCONFIGDIR", str(Path(__file__).resolve().parent.parent / ".mplconfig"))
os.environ.setdefault("XDG_CACHE_HOME", str(Path(__file__).resolve().parent.parent / ".cache"))
Path(os.environ["MPLCONFIGDIR"]).mkdir(parents=True, exist_ok=True)
Path(os.environ["XDG_CACHE_HOME"]).mkdir(parents=True, exist_ok=True)

import matplotlib

matplotlib.use("Agg")
import numpy as np
import pandas as pd
import torch
import torch.nn.functional as F
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parent.parent
os.chdir(PROJECT_ROOT)
sys.path.insert(0, str(PROJECT_ROOT))

from phase1_reactive.drl.gnn_plus_selector import (  # noqa: E402
    GNNPlusConfig,
    GNNPlusFlowSelector,
    build_graph_tensors_plus,
    build_od_features_plus,
    load_gnn_plus,
    save_gnn_plus,
)
from phase1_reactive.drl.gnn_training import (  # noqa: E402
    GNNReinforceConfig,
    _bottleneck_baseline_mlu,
    _collect_oracle_labels,
    _collect_soft_teacher_targets,
    _ranking_loss,
    _reference_model_mlu,
    _score_margin_regularizer,
    _soft_teacher_loss,
)
from phase1_reactive.drl.state_builder import compute_reactive_telemetry  # noqa: E402
from phase1_reactive.eval.core import split_indices  # noqa: E402
from phase1_reactive.routing.path_cache import assert_selected_ods_have_paths, surviving_od_mask  # noqa: E402
from te.baselines import clone_splits, ecmp_splits, select_bottleneck_critical  # noqa: E402
from te.disturbance import compute_disturbance  # noqa: E402
from te.lp_solver import solve_selected_path_lp  # noqa: E402
from te.paths import build_k_shortest_paths  # noqa: E402
from te.simulator import apply_routing  # noqa: E402


SEED = 42
DEVICE = "cpu"
K_CRIT = 40
# Rescue branch (gnnplus-debug-rescue): env-var-configurable for Check #5
# (split train LP budget from eval LP budget). Default 20 preserves Task A.
LP_TIME_LIMIT = int(os.environ.get("GNNPLUS_LP_TIME_LIMIT", "20"))
NUM_RUNS = 3
PROPORTIONAL_NUM_RUNS = 1
TRAIN_MAX_PER_TOPO = 200
VAL_MAX_PER_TOPO = 50
RL_MAX_TRAIN_SAMPLES = 120
RL_MAX_VAL_SAMPLES = 48
SUP_MAX_EPOCHS = int(os.environ.get("GNNPLUS_SUP_MAX_EPOCHS", "14"))
SUP_PATIENCE = int(os.environ.get("GNNPLUS_SUP_PATIENCE", "5"))
RL_MAX_EPOCHS = int(os.environ.get("GNNPLUS_RL_MAX_EPOCHS", "20"))
RL_PATIENCE = int(os.environ.get("GNNPLUS_RL_PATIENCE", "5"))
RL_SELECTION_METRIC = os.environ.get("GNNPLUS_RL_SELECTION_METRIC", "val_failure_mlu").strip().lower()
FEATURE_VARIANT = os.environ.get("GNNPLUS_FEATURE_VARIANT", "section7_temporal").strip().lower()
SOFT_TEACHER_WEIGHT = float(os.environ.get("GNNPLUS_SOFT_TEACHER_WEIGHT", "0.45"))
CRITICALITY_WEIGHT = float(os.environ.get("GNNPLUS_CRITICALITY_WEIGHT", "0.35"))
LP_TEACHER_WEIGHT = float(os.environ.get("GNNPLUS_LP_TEACHER_WEIGHT", "6.0"))
CONTINUITY_BONUS = float(os.environ.get("GNNPLUS_CONTINUITY_BONUS", "0.05"))
FINAL_TEACHER_FORCING_PROB = float(os.environ.get("GNNPLUS_FINAL_TEACHER_FORCING_PROB", "0.15"))
SUP_BALANCED_EPOCH_MULTIPLIER = float(os.environ.get("GNNPLUS_SUP_BALANCED_EPOCH_MULTIPLIER", "1.0"))
RL_LR = float(os.environ.get("GNNPLUS_RL_LR", "3.5e-5"))
RL_ENTROPY_WEIGHT = float(os.environ.get("GNNPLUS_RL_ENTROPY_WEIGHT", "0.01"))
REWARD_MLU = float(os.environ.get("GNNPLUS_REWARD_MLU", "1.15"))
REWARD_MLU_NORMAL = float(os.environ.get("GNNPLUS_REWARD_MLU_NORMAL", str(REWARD_MLU)))
REWARD_MLU_FAILURE = float(os.environ.get("GNNPLUS_REWARD_MLU_FAILURE", str(REWARD_MLU)))
REWARD_IMPROVEMENT = float(os.environ.get("GNNPLUS_REWARD_IMPROVEMENT", "0.85"))
REWARD_DISTURBANCE = float(os.environ.get("GNNPLUS_REWARD_DISTURBANCE", "0.15"))
REWARD_VS_BOTTLENECK = float(os.environ.get("GNNPLUS_REWARD_VS_BOTTLENECK", "0.45"))
REWARD_VS_REFERENCE = float(os.environ.get("GNNPLUS_REWARD_VS_REFERENCE", "0.15"))
REWARD_BOTTLENECK_MARGIN = float(os.environ.get("GNNPLUS_REWARD_BOTTLENECK_MARGIN", "0.10"))
ENABLE_HINGE_REWARD = os.environ.get("GNNPLUS_ENABLE_HINGE_REWARD", "0") == "1"
HINGE_THRESHOLD = float(os.environ.get("GNNPLUS_HINGE_THRESHOLD", "0.85"))
HINGE_MULTIPLIER = float(os.environ.get("GNNPLUS_HINGE_MULTIPLIER", "3.0"))
NUM_LAYERS_OVERRIDE = int(os.environ.get("GNNPLUS_NUM_LAYERS", "3"))
USE_GATED_RESIDUAL = os.environ.get("GNNPLUS_USE_GATED_RESIDUAL", "1") == "1"
USE_CROSS_OD_ATTENTION = os.environ.get("GNNPLUS_USE_CROSS_OD_ATTENTION", "1") == "1"
CROSS_OD_ATTENTION_HEADS = int(os.environ.get("GNNPLUS_CROSS_OD_ATTENTION_HEADS", "2"))
USE_CANDIDATE_PREFILTER = os.environ.get("GNNPLUS_USE_CANDIDATE_PREFILTER", "1") == "1"
CANDIDATE_PREFILTER_MULTIPLIER = float(os.environ.get("GNNPLUS_CANDIDATE_PREFILTER_MULTIPLIER", "2.0"))
TRAIN_CANDIDATE_POOL_SIZE = int(os.environ.get("GNNPLUS_TRAIN_CANDIDATE_POOL_SIZE", "180"))
INFER_CANDIDATE_POOL_SIZE = int(os.environ.get("GNNPLUS_INFER_CANDIDATE_POOL_SIZE", "80"))
ORACLE_WINDOW_TARGET = float(os.environ.get("GNNPLUS_ORACLE_WINDOW_TARGET", "0.92"))
SYNTHETIC_FAILURE_PROB = float(os.environ.get("GNNPLUS_SYNTHETIC_FAILURE_PROB", "0.08"))
RL_FAILURE_ROLLOUT_PROB = float(os.environ.get("GNNPLUS_RL_FAILURE_ROLLOUT_PROB", str(SYNTHETIC_FAILURE_PROB)))
SYNTHETIC_FAILURE_MIN_EDGES = int(os.environ.get("GNNPLUS_SYNTHETIC_FAILURE_MIN_EDGES", "1"))
SYNTHETIC_FAILURE_MAX_EDGES = int(os.environ.get("GNNPLUS_SYNTHETIC_FAILURE_MAX_EDGES", "2"))
SYNTHETIC_FAILURE_AUGMENT_PROB = float(os.environ.get("GNNPLUS_SYNTHETIC_FAILURE_AUGMENT_PROB", "0.50"))
SYNTHETIC_CAPACITY_PERTURB = float(os.environ.get("GNNPLUS_SYNTHETIC_CAPACITY_PERTURB", "0.20"))
PLACKETT_LUCE_WEIGHT = float(os.environ.get("GNNPLUS_PLACKETT_LUCE_WEIGHT", "0.30"))
PLACKETT_LUCE_WEIGHT_END = float(os.environ.get("GNNPLUS_PLACKETT_LUCE_WEIGHT_END", "1.00"))
PLACKETT_LUCE_RAMP_EPOCHS = int(os.environ.get("GNNPLUS_PLACKETT_LUCE_RAMP_EPOCHS", "8"))
BOTTLENECK_AUX_WEIGHT = float(os.environ.get("GNNPLUS_BOTTLENECK_AUX_WEIGHT", "0.12"))
PREFILTER_RECALL_WEIGHT = float(os.environ.get("GNNPLUS_PREFILTER_RECALL_WEIGHT", "0.15"))
TEMPORAL_CONSISTENCY_WEIGHT = float(os.environ.get("GNNPLUS_TEMPORAL_CONSISTENCY_WEIGHT", "0.02"))
SUP_LR = float(os.environ.get("GNNPLUS_SUP_LR", "5e-4"))
SUP_COSINE_DECAY_START_EPOCH = int(os.environ.get("GNNPLUS_SUP_COSINE_DECAY_START_EPOCH", "6"))
SUP_MIN_LR_RATIO = float(os.environ.get("GNNPLUS_SUP_MIN_LR_RATIO", "0.10"))
RL_KL_BETA_START = float(os.environ.get("GNNPLUS_RL_KL_BETA_START", "0.10"))
RL_KL_BETA_END = float(os.environ.get("GNNPLUS_RL_KL_BETA_END", "0.01"))
RL_KL_BETA_FAILURE = float(os.environ.get("GNNPLUS_RL_KL_BETA_FAILURE", "0.002"))
RL_GUMBEL_TAU_FAILURE = float(os.environ.get("GNNPLUS_RL_GUMBEL_TAU_FAILURE", "0.0"))
RL_KL_DECAY_EPOCHS = int(os.environ.get("GNNPLUS_RL_KL_DECAY_EPOCHS", "10"))
RL_FAILURE_KL_TARGET = float(os.environ.get("GNNPLUS_RL_FAILURE_KL_TARGET", "0.05"))
RL_NORMAL_KL_ABORT = float(os.environ.get("GNNPLUS_RL_NORMAL_KL_ABORT", "0.30"))
RL_NORMAL_KL_ABORT_PATIENCE = int(os.environ.get("GNNPLUS_RL_NORMAL_KL_ABORT_PATIENCE", "2"))
RL_NORMAL_MLU_ABORT_THRESHOLD = float(os.environ.get("GNNPLUS_RL_NORMAL_MLU_ABORT_THRESHOLD", "inf"))
RL_FAILURE_KL_CHECK_EPOCH = int(os.environ.get("GNNPLUS_RL_FAILURE_KL_CHECK_EPOCH", "5"))
RL_FAILURE_KL_MIN_BETA = float(os.environ.get("GNNPLUS_RL_FAILURE_KL_MIN_BETA", "0.001"))
RL_GATE_ATTENTION_CLIP_EARLY = float(os.environ.get("GNNPLUS_RL_GATE_ATTENTION_CLIP_EARLY", "0.5"))
RL_GATE_ATTENTION_CLIP_LATE = float(os.environ.get("GNNPLUS_RL_GATE_ATTENTION_CLIP_LATE", "1.0"))
RL_GATE_ATTENTION_CLIP_EPOCHS = int(os.environ.get("GNNPLUS_RL_GATE_ATTENTION_CLIP_EPOCHS", "3"))
DISTURBANCE_CHURN_MULTIPLIER = float(os.environ.get("GNNPLUS_DIST_CHURN_MULTIPLIER", "2.0"))
DO_NO_HARM_THRESHOLD = float(os.environ.get("GNNPLUS_DO_NO_HARM_THRESHOLD", "1.02"))
DO_NO_HARM_THRESHOLD_KNOWN = float(os.environ.get("GNNPLUS_DO_NO_HARM_THRESHOLD_KNOWN", str(DO_NO_HARM_THRESHOLD)))
DO_NO_HARM_THRESHOLD_UNSEEN = float(os.environ.get("GNNPLUS_DO_NO_HARM_THRESHOLD_UNSEEN", "1.00"))
DO_NO_HARM_CACHE_STEPS = int(os.environ.get("GNNPLUS_DO_NO_HARM_CACHE_STEPS", "20"))
DO_NO_HARM_FALLBACK_COOLDOWN = int(os.environ.get("GNNPLUS_DO_NO_HARM_FALLBACK_COOLDOWN", "4"))
DO_NO_HARM_OVERLAP_SKIP = float(os.environ.get("GNNPLUS_DO_NO_HARM_OVERLAP_SKIP", "0.90"))
STICKY_EPS = float(os.environ.get("GNNPLUS_STICKY_EPS", "0.0"))
DISTURB_TIEBREAK_EPS = float(os.environ.get("GNNPLUS_DISTURB_TIEBREAK_EPS", "0.0"))
# Failure-time do-no-harm: compute bottleneck baseline's failure recovery and
# use it when GNN+'s recovery MLU is worse. This guards against regressions
# under link-removal scenarios. Default ON.
FAILURE_DO_NO_HARM = os.environ.get("GNNPLUS_FAILURE_DO_NO_HARM", "1") == "1"
CALIBRATION_MAX_SAMPLES_PER_TOPO = int(os.environ.get("GNNPLUS_CALIBRATION_MAX_SAMPLES_PER_TOPO", "200"))
SUP_LP_GAP_MEAN_GATE = float(os.environ.get("GNNPLUS_SUP_LP_GAP_MEAN_GATE", "0.03"))
SUP_LP_GAP_P95_GATE = float(os.environ.get("GNNPLUS_SUP_LP_GAP_P95_GATE", "0.06"))
SUP_LP_GAP_SAMPLES_PER_TOPO = int(os.environ.get("GNNPLUS_SUP_LP_GAP_SAMPLES_PER_TOPO", "50"))
GATE_TEMPERATURE_GRID = [
    float(x)
    for x in os.environ.get("GNNPLUS_GATE_TEMPERATURE_GRID", "0.75,0.90,1.00,1.10,1.25").split(",")
    if str(x).strip()
]
BOTTLENECK_MOE_FLOOR = float(os.environ.get("GNNPLUS_BOTTLENECK_MOE_FLOOR", "0.10"))
PROPORTIONAL_BUDGET_RATIO = float(os.environ.get("GNNPLUS_PROPORTIONAL_RATIO", "0.10"))
EXTREME_STRESS_RATIO = float(os.environ.get("GNNPLUS_EXTREME_STRESS_RATIO", "0.05"))
EXTREME_STRESS_FAIL_COUNT = int(os.environ.get("GNNPLUS_EXTREME_STRESS_FAIL_COUNT", "3"))
# Step 5: per-topology reward normalization (training-only, frozen before eval,
# computed only from known topologies; never touches germany50 or vtlwavenet2011).
PER_TOPO_REWARD_NORM = os.environ.get("GNNPLUS_PER_TOPO_REWARD_NORM", "0") == "1"

def _parse_extra_topologies(raw: str) -> list[str]:
    items: list[str] = []
    for chunk in str(raw or "").split(","):
        topo = chunk.strip().lower().replace("-", "_")
        if topo and topo not in items:
            items.append(topo)
    return items


KNOWN_TOPOLOGIES = ["abilene", "cernet", "geant", "ebone", "sprintlink", "tiscali"]
EXTRA_UNSEEN_TOPOLOGIES = _parse_extra_topologies(os.environ.get("GNNPLUS_EXTRA_UNSEEN_TOPOLOGIES", ""))
UNSEEN_TOPOLOGIES = ["germany50", "vtlwavenet2011"] + [
    topo for topo in EXTRA_UNSEEN_TOPOLOGIES if topo not in {"germany50", "vtlwavenet2011"}
]
ALL_TOPOLOGIES = KNOWN_TOPOLOGIES + UNSEEN_TOPOLOGIES
CORE_METHODS = ["ecmp", "bottleneck", "gnn", "gnnplus"]
RL_FAILURE_SCENARIOS = [
    "single_link_failure",
    "multiple_link_failure",
    "three_link_failure",
    "capacity_degradation_50",
    "traffic_spike_2x",
]
AGGRESSIVE_TIEBREAK_TOPOLOGIES = {"geant", "tiscali", "nobel_germany", *UNSEEN_TOPOLOGIES}

TOPOLOGY_DISPLAY = {
    "abilene": "Abilene",
    "cernet": "CERNET",
    "geant": "GEANT",
    "ebone": "Ebone",
    "sprintlink": "Sprintlink",
    "tiscali": "Tiscali",
    "germany50": "Germany50",
    "nobel_germany": "Nobel-Germany",
    "vtlwavenet2011": "VtlWavenet2011",
}
METHOD_LABELS = {
    "ecmp": "ECMP",
    "bottleneck": "Bottleneck",
    "gnn": "Original GNN",
    "gnnplus": "GNN+",
}
PROPORTIONAL_METHODS = ["ecmp", "bottleneck", "gnnplus"]
_ECMP_BASE_CACHE: dict[int, list[np.ndarray]] = {}


def _parse_topology_weight_overrides(raw: str) -> dict[str, float]:
    defaults = {
        "abilene": 1.0,
        "cernet": 2.0,
        "geant": 1.0,
        "ebone": 1.0,
        "sprintlink": 1.0,
        "tiscali": 2.0,
    }
    text = str(raw or "").strip()
    if not text:
        return defaults
    parsed = dict(defaults)
    for chunk in text.split(","):
        item = chunk.strip()
        if not item or ":" not in item:
            continue
        key, value = item.split(":", 1)
        key = key.strip().lower()
        try:
            parsed[key] = max(0.1, float(value.strip()))
        except ValueError:
            continue
    return parsed


SUPERVISED_TOPOLOGY_WEIGHTS = _parse_topology_weight_overrides(
    os.environ.get("GNNPLUS_SUP_TOPOLOGY_WEIGHTS", "")
)

EXPERIMENT_TAG = os.environ.get("GNNPLUS_EXPERIMENT_TAG", "gnnplus_step1to5_failgate_main_reference")
OUTPUT_DIR = PROJECT_ROOT / "results" / EXPERIMENT_TAG
TRAIN_DIR = OUTPUT_DIR / "training"
PLOTS_DIR = OUTPUT_DIR / "plots"
COMPARISON_DIR = OUTPUT_DIR / "comparison"
REPORT_DOCX = OUTPUT_DIR / os.environ.get("GNNPLUS_REPORT_NAME", "GNNPLUS_STEP1TO5_FAILGATE_MAIN_REFERENCE_REPORT.docx")
AUDIT_MD = OUTPUT_DIR / "experiment_audit.md"

SUMMARY_CSV = OUTPUT_DIR / "packet_sdn_summary.csv"
FAILURE_CSV = OUTPUT_DIR / "packet_sdn_failure.csv"
SDN_METRICS_CSV = OUTPUT_DIR / "packet_sdn_sdn_metrics.csv"
TIMESERIES_CSV = OUTPUT_DIR / "packet_sdn_timeseries.csv"
FAILURE_TIMESERIES_CSV = OUTPUT_DIR / "packet_sdn_failure_timeseries.csv"
SPLIT_MANIFEST_JSON = OUTPUT_DIR / "split_manifest.json"
PROPORTIONAL_SUMMARY_CSV = OUTPUT_DIR / "proportional_budget_summary.csv"
PROPORTIONAL_STRESS_CSV = OUTPUT_DIR / "proportional_budget_extreme_stress.csv"
SAVE_PACKET_SDN_TIMESERIES = os.environ.get("GNNPLUS_SAVE_PACKET_SDN_TIMESERIES", "0") == "1"
SKIP_PROPORTIONAL_STUDY = os.environ.get("GNNPLUS_SKIP_PROPORTIONAL_STUDY", "0") == "1"

SUP_CKPT = TRAIN_DIR / "gnn_plus_supervised_improved.pt"
FINAL_CKPT = TRAIN_DIR / "gnn_plus_improved_fixedk40.pt"
SUP_LOG_CSV = TRAIN_DIR / "supervised_train_log.csv"
SUP_TOPOLOGY_LOG_CSV = TRAIN_DIR / "supervised_topology_log.csv"
SUP_LP_GAP_JSON = TRAIN_DIR / "supervised_lp_gap_summary.json"
SUP_LP_GAP_CSV = TRAIN_DIR / "supervised_lp_gap_rows.csv"
RL_LOG_CSV = TRAIN_DIR / "reinforce_log.csv"
RL_FAILURE_TYPE_LOG_CSV = TRAIN_DIR / "reinforce_failure_type_log.csv"
RL_EPOCH_CKPT_DIR = TRAIN_DIR / "rl_epoch_checkpoints"
TRAINING_SUMMARY_JSON = TRAIN_DIR / "training_summary.json"
INFERENCE_CALIBRATION_JSON = TRAIN_DIR / "inference_calibration.json"
SUPERVISED_SUMMARY_JSON = TRAIN_DIR / "supervised_summary.json"
RUN_SUMMARY_MD = OUTPUT_DIR / os.environ.get("GNNPLUS_RUN_SUMMARY_NAME", "THREE_TRACKS_RUN_SUMMARY.md")

RUN_STAGE = os.environ.get("GNNPLUS_RUN_STAGE", "full").strip().lower()
REUSE_SUPERVISED = os.environ.get("GNNPLUS_REUSE_SUPERVISED", "0") == "1"
REQUIRE_RL_BEST_EPOCH_GT2 = os.environ.get("GNNPLUS_RL_REQUIRE_BEST_EPOCH_GT2", "1") == "1"
RL_ABORT_KL = float(os.environ.get("GNNPLUS_RL_ABORT_KL", "0.5"))
RL_NEAR_ZERO_KL = float(os.environ.get("GNNPLUS_RL_NEAR_ZERO_KL", "0.01"))
RL_GUARD_EPOCH = int(os.environ.get("GNNPLUS_RL_GUARD_EPOCH", "5"))
RL_MAX_RESTARTS = int(os.environ.get("GNNPLUS_RL_MAX_RESTARTS", "1"))

BASE_GNNPLUS_CKPT = PROJECT_ROOT / "results" / "gnn_plus_retrained_fixedk40" / "gnn_plus_fixed_k40.pt"
BASELINE_OUTPUT_DIR = PROJECT_ROOT / "results" / "professor_clean_gnnplus_zeroshot"
PREVIOUS_REPORT_TAG = os.environ.get("GNNPLUS_PREVIOUS_REPORT_TAG", "gnnplus_step1to5_failgate_main_reference")
PREVIOUS_OUTPUT_DIR = PROJECT_ROOT / "results" / PREVIOUS_REPORT_TAG
PREVIOUS_TRAIN_DIR = PREVIOUS_OUTPUT_DIR / "training"
ARCHFIX_OUTPUT_DIR = PROJECT_ROOT / "results" / "gnnplus_archfix_fulltrain"
MAIN_REFERENCE_OUTPUT_DIR = PROJECT_ROOT / "results" / "gnnplus_step1to5_failgate_main_reference"
HELPER_PATH = PROJECT_ROOT / "scripts" / "build_gnnplus_packet_sdn_report_fixed.py"
RUNNER_PATH = PROJECT_ROOT / "scripts" / "run_gnnplus_packet_sdn_full.py"


def feature_profile_description() -> str:
    if FEATURE_VARIANT == "lightweight_failure_aware":
        return (
            "Lightweight failure-aware physical profile: keeps bottleneck perception, OD-level failure signals, "
            "and absolute stress-change features while pruning the broader topology-style and weaker temporal extras."
        )
    if FEATURE_VARIANT == "section7_temporal":
        return (
            "Temporal disturbance-aware profile: keeps the physical features and adds previous-selection and previous-disturbance cues."
        )
    if FEATURE_VARIANT == "section3_physical":
        return "Physical/stress-aware profile without the temporal continuity cues."
    return "Legacy GNN+ feature profile."


def teacher_profile_description() -> str:
    return (
        "LP-strengthened soft teacher targets with continuous OD criticality regression "
        "using Huber loss on normalized selector scores."
    )


def seed_all(seed: int) -> None:
    random.seed(seed)
    np.random.seed(seed)
    torch.manual_seed(seed)


def load_module(path: Path, name: str):
    spec = importlib.util.spec_from_file_location(name, path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Could not load module from {path}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[name] = module
    spec.loader.exec_module(module)
    return module


def load_runner():
    return load_module(RUNNER_PATH, "run_gnnplus_packet_sdn_full_improved")


def load_helper():
    return load_module(HELPER_PATH, "build_gnnplus_packet_sdn_report_fixed_improved")


def selection_indicator(num_od: int, selected: list[int]) -> np.ndarray:
    out = np.zeros(int(num_od), dtype=np.float32)
    if selected:
        out[np.asarray(selected, dtype=int)] = 1.0
    return out


def topk_from_soft_target(soft_target: np.ndarray, k: int) -> list[int]:
    scores = np.asarray(soft_target, dtype=np.float64).reshape(-1)
    if scores.size == 0:
        return []
    positive = np.where(scores > 0.0)[0]
    if positive.size > 0:
        order = positive[np.argsort(-scores[positive], kind="mergesort")]
        return order[: min(int(k), int(order.size))].astype(int).tolist()
    order = np.argsort(-scores, kind="mergesort")
    return order[: min(int(k), int(order.size))].astype(int).tolist()


def selection_active_mask(sample: dict) -> np.ndarray:
    return (
        (np.asarray(sample["tm_vector"], dtype=np.float64) > 1e-12)
        & surviving_od_mask(sample["path_library"])
    )


def get_ecmp_base_cached(path_library) -> list[np.ndarray]:
    cache_key = id(path_library)
    cached = _ECMP_BASE_CACHE.get(cache_key)
    if cached is None:
        cached = ecmp_splits(path_library)
        _ECMP_BASE_CACHE[cache_key] = cached
    return cached


def candidate_pool_size(k_crit: int, active_count: int, *, candidate_limit: int | None = None) -> int:
    if active_count <= 0:
        return 0
    if candidate_limit is not None:
        return min(int(active_count), max(int(k_crit), int(candidate_limit)))
    scaled = int(math.ceil(float(CANDIDATE_PREFILTER_MULTIPLIER) * float(k_crit)))
    return min(int(active_count), max(int(k_crit), scaled))


def compute_candidate_pool(
    *,
    tm_vector: np.ndarray,
    path_library,
    capacities: np.ndarray,
    k_crit: int,
    candidate_limit: int | None = None,
) -> np.ndarray | None:
    if not USE_CANDIDATE_PREFILTER:
        return None

    tm_arr = np.asarray(tm_vector, dtype=np.float64)
    active_mask = (tm_arr > 1e-12) & surviving_od_mask(path_library)
    active_idx = np.flatnonzero(active_mask)
    if active_idx.size == 0:
        return None
    if active_idx.size <= int(k_crit):
        return active_idx.astype(np.int64, copy=False)

    candidate_k = candidate_pool_size(int(k_crit), int(active_idx.size), candidate_limit=candidate_limit)
    try:
        ranked = select_bottleneck_critical(
            tm_arr,
            get_ecmp_base_cached(path_library),
            path_library,
            np.asarray(capacities, dtype=float),
            candidate_k,
        )
    except Exception:
        ranked = []

    ordered: list[int] = []
    seen: set[int] = set()
    for od_idx in ranked:
        od_int = int(od_idx)
        if 0 <= od_int < active_mask.size and bool(active_mask[od_int]) and od_int not in seen:
            ordered.append(od_int)
            seen.add(od_int)

    if len(ordered) < candidate_k:
        fallback = active_idx[np.argsort(-tm_arr[active_idx], kind="mergesort")]
        for od_idx in fallback.tolist():
            od_int = int(od_idx)
            if od_int not in seen:
                ordered.append(od_int)
                seen.add(od_int)
            if len(ordered) >= candidate_k:
                break

    return np.asarray(ordered[:candidate_k], dtype=np.int64)


def do_no_harm_threshold_for_topology(topology_key: str) -> float:
    topo = str(topology_key).strip().lower()
    if any(unseen in topo for unseen in UNSEEN_TOPOLOGIES):
        return float(DO_NO_HARM_THRESHOLD_UNSEEN)
    return float(DO_NO_HARM_THRESHOLD_KNOWN)


def selection_overlap_ratio(selected_a: list[int] | None, selected_b: list[int] | None) -> float:
    set_a = {int(od) for od in (selected_a or [])}
    set_b = {int(od) for od in (selected_b or [])}
    denom = max(len(set_a), len(set_b), 1)
    return float(len(set_a & set_b)) / float(denom)


def summarize_candidate_window(sample: dict, candidate_od_indices: np.ndarray | None) -> dict[str, float | int | bool]:
    oracle_selected = [int(od) for od in sample.get("oracle_selected", [])]
    oracle_total = int(len(oracle_selected))
    if candidate_od_indices is None:
        return {
            "candidate_size": int(np.count_nonzero(selection_active_mask(sample))),
            "oracle_hits": oracle_total,
            "oracle_total": oracle_total,
            "oracle_hit_fraction": 1.0 if oracle_total > 0 else 0.0,
            "oracle_in_window": True,
        }
    candidate_set = {int(od) for od in np.asarray(candidate_od_indices, dtype=np.int64).reshape(-1).tolist()}
    oracle_hits = sum(1 for od in oracle_selected if od in candidate_set)
    return {
        "candidate_size": int(np.asarray(candidate_od_indices).size),
        "oracle_hits": int(oracle_hits),
        "oracle_total": oracle_total,
        "oracle_hit_fraction": float(oracle_hits) / max(float(oracle_total), 1.0),
        "oracle_in_window": bool(oracle_total == 0 or oracle_hits == oracle_total),
    }


def candidate_indices_from_od_data(od_data: dict) -> np.ndarray | None:
    candidate_od_indices = od_data.get("candidate_od_indices")
    if candidate_od_indices is None:
        return None
    if isinstance(candidate_od_indices, torch.Tensor):
        candidate_np = candidate_od_indices.detach().cpu().numpy().astype(np.int64, copy=False)
    else:
        candidate_np = np.asarray(candidate_od_indices, dtype=np.int64).reshape(-1)
    return candidate_np


def fit_indicator_length(indicator: np.ndarray | None, length: int) -> np.ndarray:
    target = np.zeros(int(length), dtype=np.float32)
    if indicator is None:
        return target
    src = np.asarray(indicator, dtype=np.float32).reshape(-1)
    take = min(int(length), int(src.size))
    if take > 0:
        target[:take] = src[:take]
    return target


def project_active_mask_to_scores(
    active_mask: np.ndarray,
    *,
    score_len: int,
    candidate_od_indices: np.ndarray | None = None,
) -> np.ndarray:
    active_full = np.asarray(active_mask, dtype=bool).reshape(-1)
    candidate_np = None if candidate_od_indices is None else np.asarray(candidate_od_indices, dtype=np.int64).reshape(-1)
    if candidate_np is not None and candidate_np.size == int(score_len):
        projected = np.zeros(int(score_len), dtype=bool)
        valid = (candidate_np >= 0) & (candidate_np < active_full.size)
        projected[valid] = active_full[candidate_np[valid]]
        return projected
    return fit_indicator_length(active_full.astype(np.float32), int(score_len)) > 0.5


def project_targets_to_score_space(sample: dict, od_data: dict, score_len: int):
    candidate_np = candidate_indices_from_od_data(od_data)
    oracle_mask = np.zeros(int(score_len), dtype=np.float32)
    soft_full = np.asarray(sample["soft_teacher"], dtype=np.float32).reshape(-1)
    criticality_full = np.asarray(sample["continuous_criticality"], dtype=np.float32).reshape(-1)
    soft_target = np.zeros(int(score_len), dtype=np.float32)
    criticality = np.zeros(int(score_len), dtype=np.float32)
    oracle_local_indices: list[int] = []

    if candidate_np is not None and candidate_np.size == int(score_len):
        local_map = {int(od_idx): idx for idx, od_idx in enumerate(candidate_np.tolist())}
        for od_idx in sample["oracle_selected"]:
            local_idx = local_map.get(int(od_idx))
            if local_idx is not None:
                oracle_mask[int(local_idx)] = 1.0
                oracle_local_indices.append(int(local_idx))
        valid_soft = (candidate_np >= 0) & (candidate_np < soft_full.size)
        valid_crit = (candidate_np >= 0) & (candidate_np < criticality_full.size)
        soft_target[valid_soft] = soft_full[candidate_np[valid_soft]]
        criticality[valid_crit] = criticality_full[candidate_np[valid_crit]]
    else:
        for od_idx in sample["oracle_selected"]:
            if 0 <= int(od_idx) < int(score_len):
                oracle_mask[int(od_idx)] = 1.0
                oracle_local_indices.append(int(od_idx))
        take_soft = min(int(score_len), int(soft_full.size))
        take_crit = min(int(score_len), int(criticality_full.size))
        if take_soft > 0:
            soft_target[:take_soft] = soft_full[:take_soft]
        if take_crit > 0:
            criticality[:take_crit] = criticality_full[:take_crit]

    candidate_meta = summarize_candidate_window(sample, candidate_np)
    return (
        torch.tensor(oracle_mask, dtype=torch.float32, device=DEVICE),
        torch.tensor(soft_target, dtype=torch.float32, device=DEVICE),
        torch.tensor(criticality, dtype=torch.float32, device=DEVICE),
        candidate_np,
        oracle_local_indices,
        candidate_meta,
    )


def _normalize_score_tensor(scores: torch.Tensor) -> torch.Tensor:
    score_min = torch.min(scores)
    score_span = torch.max(scores) - score_min
    if float(score_span.item()) > 1e-12:
        return (scores - score_min) / (score_span + 1e-12)
    return torch.zeros_like(scores)


def _plackett_luce_topk_loss(scores: torch.Tensor, ordered_oracle_local: list[int]) -> torch.Tensor:
    if not ordered_oracle_local:
        return torch.tensor(0.0, device=scores.device)
    remaining = torch.arange(scores.shape[0], device=scores.device)
    loss = torch.tensor(0.0, device=scores.device)
    selected_mask = torch.zeros(scores.shape[0], dtype=torch.bool, device=scores.device)
    for local_idx in ordered_oracle_local:
        local_idx = int(local_idx)
        if local_idx < 0 or local_idx >= int(scores.shape[0]) or bool(selected_mask[local_idx]):
            continue
        logits = scores[~selected_mask]
        current_indices = remaining[~selected_mask]
        target_pos = torch.nonzero(current_indices == int(local_idx), as_tuple=False)
        if target_pos.numel() == 0:
            continue
        target_pos = int(target_pos[0].item())
        loss = loss - torch.log_softmax(logits, dim=0)[target_pos]
        selected_mask[local_idx] = True
    return loss


def _bottleneck_aux_loss(aux_scores: torch.Tensor | None, bottleneck_scores: torch.Tensor) -> torch.Tensor:
    if aux_scores is None:
        return torch.tensor(0.0, device=bottleneck_scores.device)
    target = _normalize_score_tensor(bottleneck_scores.float())
    pred = _normalize_score_tensor(aux_scores.float())
    return F.mse_loss(pred, target)


def _oracle_recall_aux_loss(full_aux_scores: torch.Tensor | None, oracle_selected: list[int]) -> torch.Tensor:
    if full_aux_scores is None or not oracle_selected:
        return torch.tensor(0.0, device=full_aux_scores.device if full_aux_scores is not None else DEVICE)
    target = torch.zeros_like(full_aux_scores)
    for od_idx in oracle_selected:
        if 0 <= int(od_idx) < int(full_aux_scores.shape[0]):
            target[int(od_idx)] = 1.0
    target_sum = torch.sum(target)
    if float(target_sum.item()) <= 0.0:
        return torch.tensor(0.0, device=full_aux_scores.device)
    target = target / target_sum
    return -(target * torch.log_softmax(full_aux_scores, dim=0)).sum()


def _temporal_consistency_loss(
    scores: torch.Tensor,
    *,
    prev_selected_indicator: np.ndarray | None,
    candidate_od_indices: np.ndarray | None,
    full_num_od: int,
) -> torch.Tensor:
    prev_full = fit_indicator_length(prev_selected_indicator, int(full_num_od))
    if candidate_od_indices is not None:
        prev = prev_full[np.asarray(candidate_od_indices, dtype=np.int64)]
    else:
        prev = fit_indicator_length(prev_full, int(scores.shape[0]))
    if float(np.sum(prev)) <= 0.0:
        return torch.tensor(0.0, device=scores.device)
    target = torch.tensor(prev, dtype=torch.float32, device=scores.device)
    target = target / (torch.sum(target) + 1e-12)
    pred = torch.softmax(scores, dim=0)
    return F.huber_loss(pred, target, delta=0.10)


def rl_kl_beta_for_epoch(epoch: int) -> float:
    if int(RL_KL_DECAY_EPOCHS) <= 1:
        return float(RL_KL_BETA_END)
    alpha = min(max(int(epoch) - 1, 0), int(RL_KL_DECAY_EPOCHS) - 1) / float(max(int(RL_KL_DECAY_EPOCHS) - 1, 1))
    return float((1.0 - alpha) * float(RL_KL_BETA_START) + alpha * float(RL_KL_BETA_END))


def _local_active_scores(
    scores: torch.Tensor,
    *,
    active_mask: np.ndarray,
    candidate_od_indices: np.ndarray | None = None,
) -> torch.Tensor:
    active_local_mask = project_active_mask_to_scores(
        active_mask,
        score_len=scores.size(0),
        candidate_od_indices=candidate_od_indices,
    )
    active_idx = np.flatnonzero(active_local_mask)
    if active_idx.size == 0:
        return scores[:0]
    return scores[torch.tensor(active_idx, dtype=torch.long, device=scores.device)]


def compute_churn_ratio(selected_ods: list[int], prev_selected_indicator: np.ndarray | None, *, full_num_od: int, k_crit: int) -> float:
    if not selected_ods:
        return 0.0
    prev_full = fit_indicator_length(prev_selected_indicator, int(full_num_od))
    prev_set = set(np.flatnonzero(prev_full > 0.5).tolist())
    selected_set = {int(od) for od in selected_ods}
    churn = len(selected_set - prev_set)
    return float(churn) / max(float(k_crit), 1.0)


def compute_tie_break_epsilon(
    scores: torch.Tensor,
    *,
    active_mask: np.ndarray,
    candidate_od_indices: np.ndarray | None = None,
    percentile: float = 5.0,
) -> float:
    active_scores = _local_active_scores(scores, active_mask=active_mask, candidate_od_indices=candidate_od_indices)
    if active_scores.numel() <= 1:
        return 0.0
    sorted_scores = torch.sort(active_scores.detach(), descending=True).values.cpu().numpy()
    gaps = np.abs(np.diff(sorted_scores))
    positive = gaps[gaps > 1e-8]
    if positive.size == 0:
        return 0.0
    q = min(max(float(percentile), 0.0), 100.0)
    return float(np.percentile(positive, q))


def build_failure_path_library_with_original_indices(dataset, weights: np.ndarray, failed_edges: list[int], *, k_paths: int):
    failed = {int(edge_idx) for edge_idx in failed_edges}
    graph = nx.DiGraph()
    graph.add_nodes_from(dataset.nodes)
    edge_to_idx: dict[tuple[str, str], int] = {}
    for edge_idx, (src, dst) in enumerate(dataset.edges):
        if edge_idx in failed:
            continue
        graph.add_edge(src, dst, weight=float(weights[edge_idx]))
        edge_to_idx[(src, dst)] = int(edge_idx)
    return build_k_shortest_paths(graph, od_pairs=dataset.od_pairs, edge_to_idx=edge_to_idx, k=int(k_paths))


def build_teacher_training_sample(
    *,
    topology: str,
    dataset,
    path_library,
    tm_vector: np.ndarray,
    timestep: int,
    capacities: np.ndarray,
    weights: np.ndarray,
    telemetry,
    prev_tm,
    prev_util,
    prev_selected_indicator,
    prev_disturbance: float,
    prev_splits,
    failure_mask=None,
    scenario: str = "normal",
    synthetic_failure: bool = False,
):
    ecmp_base = get_ecmp_base_cached(path_library)
    oracle_selected, oracle_mlu, oracle_method, _ = _collect_oracle_labels(
        dataset,
        path_library,
        tm_vector,
        ecmp_base,
        capacities,
        K_CRIT,
        lp_time_limit_sec=LP_TIME_LIMIT,
    )
    if not oracle_selected:
        return None

    soft_teacher, continuous_criticality = _collect_soft_teacher_targets(
        dataset=dataset,
        path_library=path_library,
        tm_vector=tm_vector,
        ecmp_base=ecmp_base,
        capacities=capacities,
        k_crit=K_CRIT,
        lp_time_limit_sec=LP_TIME_LIMIT,
        lp_teacher_weight=LP_TEACHER_WEIGHT,
    )

    return {
        "topology": topology,
        "dataset": dataset,
        "path_library": path_library,
        "tm_vector": np.asarray(tm_vector, dtype=float),
        "timestep": int(timestep),
        "prev_tm": None if prev_tm is None else np.asarray(prev_tm, dtype=float),
        "telemetry": telemetry,
        "oracle_selected": list(oracle_selected),
        "oracle_mlu": float(oracle_mlu),
        "oracle_method": str(oracle_method),
        "soft_teacher": np.asarray(soft_teacher, dtype=np.float32),
        "continuous_criticality": np.asarray(continuous_criticality, dtype=np.float32),
        "k_crit": int(K_CRIT),
        "capacities": np.asarray(capacities, dtype=float),
        "weights": np.asarray(weights, dtype=float),
        "prev_util": None if prev_util is None else np.asarray(prev_util, dtype=float),
        "prev_selected_indicator": np.asarray(prev_selected_indicator, dtype=np.float32),
        "teacher_prev_selected_indicator": np.asarray(prev_selected_indicator, dtype=np.float32),
        "prev_disturbance": float(prev_disturbance),
        "prev_splits": clone_splits(prev_splits),
        "failure_mask": None if failure_mask is None else np.asarray(failure_mask, dtype=float),
        "scenario": str(scenario),
        "synthetic_failure": bool(synthetic_failure),
    }


def build_synthetic_failure_sample(
    runner,
    *,
    topology: str,
    dataset,
    tm_vector: np.ndarray,
    timestep: int,
    capacities: np.ndarray,
    weights: np.ndarray,
    rng: np.random.Generator,
):
    if float(SYNTHETIC_FAILURE_PROB) <= 0.0 or len(dataset.edges) <= 0:
        return None

    min_fail = max(1, min(int(SYNTHETIC_FAILURE_MIN_EDGES), len(dataset.edges)))
    max_fail = max(min_fail, min(int(SYNTHETIC_FAILURE_MAX_EDGES), len(dataset.edges)))
    fail_count = int(rng.integers(min_fail, max_fail + 1))
    failed_edges = sorted(rng.choice(len(dataset.edges), size=fail_count, replace=False).tolist())

    base_capacities = np.asarray(capacities, dtype=float)
    synthetic_caps = base_capacities.copy()
    capacity_augmented = bool(rng.random() < float(SYNTHETIC_FAILURE_AUGMENT_PROB))
    if capacity_augmented:
        perturb_low = max(0.0, 1.0 - float(SYNTHETIC_CAPACITY_PERTURB))
        perturb_high = 1.0 + float(SYNTHETIC_CAPACITY_PERTURB)
        perturb = rng.uniform(perturb_low, perturb_high, size=base_capacities.shape[0])
        synthetic_caps = np.maximum(base_capacities * perturb, 1e-6)

    failure_mask = np.ones(len(capacities), dtype=float)
    if failed_edges:
        failure_mask[np.asarray(failed_edges, dtype=np.int64)] = 0.0
    effective_caps = synthetic_caps * failure_mask
    effective_dataset = runner._build_dataset_view(
        dataset,
        edges=dataset.edges,
        capacities=effective_caps,
        weights=weights,
    )
    effective_path_library = build_failure_path_library_with_original_indices(
        effective_dataset,
        np.asarray(weights, dtype=float),
        failed_edges,
        k_paths=runner.K_PATHS,
    )
    if not np.any((np.asarray(tm_vector, dtype=np.float64) > 1e-12) & surviving_od_mask(effective_path_library)):
        return None

    effective_ecmp = get_ecmp_base_cached(effective_path_library)
    routing = apply_routing(tm_vector, effective_ecmp, effective_path_library, effective_caps)
    telemetry = compute_reactive_telemetry(
        tm_vector,
        effective_ecmp,
        effective_path_library,
        routing,
        np.asarray(weights, dtype=float),
        prev_latency_by_od=None,
    )
    zero_indicator = np.zeros(len(effective_dataset.od_pairs), dtype=np.float32)
    return build_teacher_training_sample(
        topology=topology,
        dataset=effective_dataset,
        path_library=effective_path_library,
        tm_vector=np.asarray(tm_vector, dtype=float),
        timestep=int(timestep),
        capacities=effective_caps,
        weights=np.asarray(weights, dtype=float),
        telemetry=telemetry,
        prev_tm=None,
        prev_util=None,
        prev_selected_indicator=zero_indicator,
        prev_disturbance=0.0,
        prev_splits=effective_ecmp,
        failure_mask=failure_mask,
        scenario=f"synthetic_failure_{fail_count}{'_aug' if capacity_augmented else ''}",
        synthetic_failure=True,
    )


def sample_has_active_failure(sample: dict) -> bool:
    if bool(sample.get("synthetic_failure", False)):
        return True
    scenario = str(sample.get("scenario", "normal")).strip().lower()
    if scenario and scenario != "normal":
        return True
    return sample.get("failure_mask") is not None


def canonical_failure_type(sample: dict) -> str:
    scenario = str(sample.get("scenario", "normal")).strip().lower()
    if scenario in {
        "single_link_failure",
        "multiple_link_failure",
        "three_link_failure",
        "capacity_degradation_50",
        "traffic_spike_2x",
    }:
        return scenario
    if scenario == "random_link_failure_1":
        return "random_link_failure_1"
    if scenario == "random_link_failure_2":
        return "multiple_link_failure"
    if scenario.startswith("synthetic_failure"):
        return "synthetic_failure_misc"
    if sample_has_active_failure(sample):
        return "other_failure"
    return "normal"


def build_rl_failure_scenario_sample(
    runner,
    sample: dict,
    rng: np.random.Generator,
    *,
    scenario: str,
) -> dict | None:
    scenario = str(scenario).strip()
    tm_vector = np.asarray(sample["tm_vector"], dtype=float)
    capacities = np.asarray(sample["capacities"], dtype=float)
    weights = np.asarray(sample["weights"], dtype=float)
    dataset = sample["dataset"]
    path_library = sample["path_library"]
    normal_ecmp = get_ecmp_base_cached(path_library)
    normal_routing = apply_routing(tm_vector, normal_ecmp, path_library, capacities)

    failure_state = runner._build_failure_execution_state(
        scenario=scenario,
        tm_vector=tm_vector,
        dataset=dataset,
        path_library=path_library,
        capacities=capacities,
        weights=weights,
        normal_routing=normal_routing,
    )
    effective_tm = np.asarray(failure_state["effective_tm"], dtype=float)
    effective_caps = np.asarray(failure_state["effective_caps"], dtype=float)
    effective_weights = np.asarray(failure_state["effective_weights"], dtype=float)
    effective_path_library = failure_state["effective_path_library"]
    effective_dataset = failure_state["effective_dataset"]
    effective_ecmp = failure_state["effective_ecmp"]
    failure_mask = np.asarray(failure_state["failure_mask"], dtype=float)
    if not np.any((np.asarray(effective_tm, dtype=np.float64) > 1e-12) & surviving_od_mask(effective_path_library)):
        return None

    routing = apply_routing(effective_tm, effective_ecmp, effective_path_library, effective_caps)
    telemetry = runner.compute_telemetry(
        tm_vector=effective_tm,
        splits=effective_ecmp,
        path_library=effective_path_library,
        routing=routing,
        weights=effective_weights,
        prev_latency_by_od=None,
        fast_links_only=True,
    )
    zero_indicator = np.zeros(len(effective_dataset.od_pairs), dtype=np.float32)
    return build_teacher_training_sample(
        topology=str(sample["topology"]),
        dataset=effective_dataset,
        path_library=effective_path_library,
        tm_vector=effective_tm,
        timestep=int(sample["timestep"]),
        capacities=effective_caps,
        weights=effective_weights,
        telemetry=telemetry,
        prev_tm=None,
        prev_util=None,
        prev_selected_indicator=zero_indicator,
        prev_disturbance=0.0,
        prev_splits=effective_ecmp,
        failure_mask=failure_mask,
        scenario=scenario,
        synthetic_failure=False,
    )


def augment_rl_failure_rollouts(
    runner,
    samples: list[dict],
    *,
    seed: int,
    probability: float,
) -> list[dict]:
    if not samples or float(probability) <= 0.0:
        return list(samples)
    rng = np.random.default_rng(seed)
    augmented: list[dict] = []
    selected_positions = [
        idx
        for idx, sample in enumerate(samples)
        if (not sample_has_active_failure(sample)) and rng.random() < float(probability)
    ]
    assigned_scenarios: dict[int, str] = {}
    if selected_positions:
        scenario_cycle = list(RL_FAILURE_SCENARIOS)
        rng.shuffle(selected_positions)
        for offset, idx in enumerate(selected_positions):
            assigned_scenarios[int(idx)] = str(scenario_cycle[offset % len(scenario_cycle)])

    for idx, sample in enumerate(samples):
        augmented.append(sample)
        if sample_has_active_failure(sample):
            continue
        if idx not in assigned_scenarios:
            continue
        preferred = assigned_scenarios[int(idx)]
        scenario_order = [preferred] + [str(name) for name in RL_FAILURE_SCENARIOS if str(name) != preferred]
        failure_sample = None
        for scenario in scenario_order:
            try:
                failure_sample = build_rl_failure_scenario_sample(
                    runner,
                    sample,
                    rng,
                    scenario=scenario,
                )
            except Exception:
                failure_sample = None
            if failure_sample is not None:
                break
        if failure_sample is not None:
            augmented.append(failure_sample)
    return augmented


def build_experiment_inputs(sample: dict, *, device: str, candidate_limit: int | None = None):
    candidate_od_indices = compute_candidate_pool(
        tm_vector=sample["tm_vector"],
        path_library=sample["path_library"],
        capacities=sample["capacities"],
        k_crit=int(sample.get("k_crit", K_CRIT)),
        candidate_limit=candidate_limit,
    )
    candidate_meta = summarize_candidate_window(sample, candidate_od_indices)
    graph_data = build_graph_tensors_plus(
        sample["dataset"],
        tm_vector=sample["tm_vector"],
        path_library=sample["path_library"],
        telemetry=sample["telemetry"],
        failure_mask=sample.get("failure_mask"),
        prev_util=sample["prev_util"],
        prev_tm=sample["prev_tm"],
        prev_selected_indicator=sample["prev_selected_indicator"],
        prev_disturbance=sample["prev_disturbance"],
        feature_variant=FEATURE_VARIANT,
        device=device,
    )
    od_data = build_od_features_plus(
        sample["dataset"],
        sample["tm_vector"],
        sample["path_library"],
        telemetry=sample["telemetry"],
        prev_tm=sample["prev_tm"],
        prev_util=sample["prev_util"],
        prev_selected_indicator=sample["prev_selected_indicator"],
        prev_disturbance=sample["prev_disturbance"],
        failure_mask=sample.get("failure_mask"),
        candidate_od_indices=candidate_od_indices,
        feature_variant=FEATURE_VARIANT,
        device=device,
    )
    return graph_data, od_data, candidate_meta


def scheduled_teacher_forcing_prob(epoch: int, total_epochs: int) -> float:
    if total_epochs <= 1:
        return float(FINAL_TEACHER_FORCING_PROB)
    alpha = float(max(epoch - 1, 0)) / float(max(total_epochs - 1, 1))
    return float((1.0 - alpha) * 1.0 + alpha * FINAL_TEACHER_FORCING_PROB)


def materialize_scheduled_history(
    samples: list[dict],
    model: GNNPlusFlowSelector,
    *,
    teacher_forcing_prob: float,
    seed: int,
) -> list[dict]:
    if not samples:
        return []

    rng = np.random.default_rng(seed)
    grouped: dict[str, list[dict]] = defaultdict(list)
    for sample in samples:
        grouped[str(sample["topology"])].append(sample)

    was_training = model.training
    model.eval()
    materialized: list[dict] = []
    with torch.no_grad():
        for topo_key in sorted(grouped):
            topo_samples = sorted(grouped[topo_key], key=lambda item: int(item.get("timestep", 0)))
            model_prev_indicator = None
            for sample in topo_samples:
                teacher_prev = np.asarray(
                    sample.get("teacher_prev_selected_indicator", sample["prev_selected_indicator"]),
                    dtype=np.float32,
                )
                if model_prev_indicator is None or rng.random() < float(teacher_forcing_prob):
                    prev_indicator = teacher_prev
                else:
                    prev_indicator = np.asarray(model_prev_indicator, dtype=np.float32)

                scheduled = dict(sample)
                scheduled["prev_selected_indicator"] = np.asarray(prev_indicator, dtype=np.float32)
                materialized.append(scheduled)

                graph_data, od_data, _ = build_experiment_inputs(
                    scheduled,
                    device=DEVICE,
                    candidate_limit=TRAIN_CANDIDATE_POOL_SIZE,
                )
                scores, _, _ = model(graph_data, od_data)
                selected, _, _ = continuity_select(
                    scores,
                    active_mask=selection_active_mask(scheduled),
                    k=K_CRIT,
                    prev_selected_indicator=scheduled["prev_selected_indicator"],
                    continuity_bonus=CONTINUITY_BONUS,
                    candidate_od_indices=candidate_indices_from_od_data(od_data),
                    full_num_od=int(od_data.get("full_num_od", scores.size(0))),
                )
                model_prev_indicator = selection_indicator(len(scheduled["dataset"].od_pairs), selected)

    if was_training:
        model.train()
    return materialized


def continuity_select(
    scores: torch.Tensor,
    *,
    active_mask: np.ndarray,
    k: int,
    prev_selected_indicator: np.ndarray | None = None,
    continuity_bonus: float = 0.0,
    tie_break_eps: float = 0.0,
    candidate_od_indices: np.ndarray | None = None,
    full_num_od: int | None = None,
):
    candidate_np = None if candidate_od_indices is None else np.asarray(candidate_od_indices, dtype=np.int64).reshape(-1)
    active_full = np.asarray(active_mask, dtype=bool).reshape(-1)
    if candidate_np is not None and candidate_np.size == int(scores.shape[0]):
        if active_full.size == int(scores.shape[0]):
            active = active_full
        else:
            active = np.zeros(int(scores.shape[0]), dtype=bool)
            valid = (candidate_np >= 0) & (candidate_np < active_full.size)
            active[valid] = active_full[candidate_np[valid]]
    else:
        active = fit_indicator_length(active_full.astype(np.float32), int(scores.shape[0])) > 0.5
    active_idx = np.where(active)[0]
    if active_idx.size == 0:
        return [], None, None
    take = min(int(k), int(active_idx.size))
    active_scores = scores[torch.tensor(active_idx, dtype=torch.long, device=scores.device)]
    ranking_scores = active_scores
    if (continuity_bonus > 0.0 or tie_break_eps > 0.0) and prev_selected_indicator is not None:
        if candidate_np is not None and candidate_np.size == int(scores.shape[0]):
            target_full = int(full_num_od or active_full.size or (np.max(candidate_np) + 1 if candidate_np.size else scores.shape[0]))
            prev_full = fit_indicator_length(prev_selected_indicator, target_full)
            prev = prev_full[candidate_np]
        else:
            prev = fit_indicator_length(prev_selected_indicator, int(scores.shape[0]))
        prev_active = torch.tensor(prev[active_idx], dtype=active_scores.dtype, device=scores.device)
        if continuity_bonus > 0.0:
            score_min = torch.min(active_scores)
            score_span = torch.max(active_scores) - score_min
            if float(score_span.item()) > 1e-12:
                normalized = (active_scores - score_min) / (score_span + 1e-12)
            else:
                normalized = torch.zeros_like(active_scores)
            ranking_scores = normalized + float(continuity_bonus) * prev_active
        else:
            ranking_scores = active_scores + float(tie_break_eps) * prev_active
    top_local = torch.topk(ranking_scores, take).indices
    if candidate_np is not None and candidate_np.size == int(scores.shape[0]):
        selected = [int(candidate_np[int(active_idx[i])]) for i in top_local.detach().cpu().numpy()]
    else:
        selected = [int(active_idx[i]) for i in top_local.detach().cpu().numpy()]
    selected_log_prob = torch.log_softmax(ranking_scores, dim=0)[top_local].sum()
    return selected, ranking_scores, selected_log_prob


def sample_gumbel_topk_scores(scores: torch.Tensor, *, tau: float) -> torch.Tensor:
    if float(tau) <= 0.0:
        return scores
    uniform = torch.rand_like(scores).clamp_(1e-6, 1.0 - 1e-6)
    gumbel = -torch.log(-torch.log(uniform))
    return scores + float(tau) * gumbel


def selected_set_overlap_ratio(selected_a: list[int], selected_b: list[int]) -> float:
    set_a = {int(x) for x in selected_a}
    set_b = {int(x) for x in selected_b}
    denom = max(len(set_a), len(set_b), 1)
    return float(len(set_a & set_b)) / float(denom)


def collect_split_samples(
    runner,
    topo_keys: list[str],
    *,
    split_name: str,
    max_per_topology: int,
    seed: int,
) -> tuple[list[dict], dict]:
    rng = np.random.default_rng(seed)
    datasets = {}
    samples = []

    for topo_key in topo_keys:
        dataset, path_library = runner.load_dataset(topo_key)
        datasets[topo_key] = (dataset, path_library)
        weights = np.asarray(dataset.weights, dtype=float)
        capacities = np.asarray(dataset.capacities, dtype=float)
        indices = split_indices(dataset, split_name)
        if len(indices) > max_per_topology:
            indices = sorted(rng.choice(indices, size=max_per_topology, replace=False).tolist())

        ecmp_base = get_ecmp_base_cached(path_library)
        prev_splits = clone_splits(ecmp_base)
        prev_selected = np.zeros(len(dataset.od_pairs), dtype=np.float32)
        prev_disturbance = 0.0
        prev_latency_by_od = None
        prev_tm = None
        prev_util = None

        print(f"[collect:{split_name}] {topo_key}: {len(indices)} candidate timesteps", flush=True)
        for t_idx in indices:
            tm_vector = np.asarray(dataset.tm[t_idx], dtype=float)
            if float(np.max(tm_vector)) < 1e-12:
                prev_tm = tm_vector
                continue

            routing = apply_routing(tm_vector, prev_splits, path_library, capacities)
            telemetry = compute_reactive_telemetry(
                tm_vector,
                prev_splits,
                path_library,
                routing,
                weights,
                prev_latency_by_od=prev_latency_by_od,
            )

            try:
                sample = build_teacher_training_sample(
                    topology=topo_key,
                    dataset=dataset,
                    path_library=path_library,
                    tm_vector=tm_vector,
                    timestep=int(t_idx),
                    capacities=capacities,
                    weights=weights,
                    telemetry=telemetry,
                    prev_tm=prev_tm,
                    prev_util=prev_util,
                    prev_selected_indicator=prev_selected,
                    prev_disturbance=prev_disturbance,
                    prev_splits=prev_splits,
                )
            except Exception:
                prev_tm = tm_vector
                continue

            if sample is None:
                prev_tm = tm_vector
                continue

            samples.append(sample)
            soft_teacher = np.asarray(sample["soft_teacher"], dtype=np.float32)
            oracle_selected = list(sample["oracle_selected"])

            if split_name == "train" and rng.random() < float(SYNTHETIC_FAILURE_PROB):
                try:
                    synthetic_sample = build_synthetic_failure_sample(
                        runner,
                        topology=topo_key,
                        dataset=dataset,
                        tm_vector=tm_vector,
                        timestep=int(t_idx),
                        capacities=capacities,
                        weights=weights,
                        rng=rng,
                    )
                except Exception:
                    synthetic_sample = None
                if synthetic_sample is not None:
                    samples.append(synthetic_sample)

            teacher_selected = topk_from_soft_target(soft_teacher, K_CRIT) or list(oracle_selected)
            try:
                lp = solve_selected_path_lp(
                    tm_vector=tm_vector,
                    selected_ods=teacher_selected,
                    base_splits=ecmp_base,
                    path_library=path_library,
                    capacities=capacities,
                    warm_start_splits=prev_splits,
                    time_limit_sec=LP_TIME_LIMIT,
                )
                disturbance = compute_disturbance(prev_splits, lp.splits, tm_vector)
                post_routing = apply_routing(tm_vector, lp.splits, path_library, capacities)
                post_telemetry = compute_reactive_telemetry(
                    tm_vector,
                    lp.splits,
                    path_library,
                    post_routing,
                    weights,
                    prev_latency_by_od=prev_latency_by_od,
                )
                prev_splits = clone_splits(lp.splits)
                prev_latency_by_od = post_telemetry.latency_by_od
                prev_util = np.asarray(post_telemetry.utilization, dtype=float)
                prev_disturbance = float(disturbance)
            except Exception:
                prev_util = np.asarray(telemetry.utilization, dtype=float)
                prev_disturbance = 0.0
            prev_selected = np.asarray(soft_teacher, dtype=np.float32)
            prev_tm = tm_vector

    return samples, datasets


def build_initial_gnnplus_model() -> GNNPlusFlowSelector:
    if BASE_GNNPLUS_CKPT.exists():
        payload = torch.load(str(BASE_GNNPLUS_CKPT), map_location=torch.device(DEVICE), weights_only=False)
        cfg = GNNPlusConfig(**payload["config"])
        cfg.feature_variant = FEATURE_VARIANT
        cfg.device = DEVICE
        cfg.learn_k_crit = False
        cfg.k_crit_min = K_CRIT
        cfg.k_crit_max = K_CRIT
        cfg.dropout = 0.2
        cfg.num_layers = int(NUM_LAYERS_OVERRIDE)
        cfg.use_gated_residual = bool(USE_GATED_RESIDUAL)
        cfg.use_cross_od_attention = bool(USE_CROSS_OD_ATTENTION)
        cfg.cross_od_attention_heads = int(CROSS_OD_ATTENTION_HEADS)
        cfg.bottleneck_moe_floor = float(BOTTLENECK_MOE_FLOOR)
        model = GNNPlusFlowSelector(cfg).to(DEVICE)
        strict = (
            int(NUM_LAYERS_OVERRIDE) == int(payload["config"].get("num_layers", 3))
            and bool(USE_GATED_RESIDUAL) == bool(payload["config"].get("use_gated_residual", False))
            and bool(USE_CROSS_OD_ATTENTION) == bool(payload["config"].get("use_cross_od_attention", False))
            and int(CROSS_OD_ATTENTION_HEADS) == int(payload["config"].get("cross_od_attention_heads", 4))
        )
        missing, unexpected = model.load_state_dict(payload["state_dict"], strict=strict)
        if not strict:
            print(
                f"[init] warm-started {BASE_GNNPLUS_CKPT.name} with num_layers={NUM_LAYERS_OVERRIDE}, "
                f"gated_residual={USE_GATED_RESIDUAL}, cross_od_attention={USE_CROSS_OD_ATTENTION}; "
                f"missing={len(missing)} unexpected={len(unexpected)}",
                flush=True,
            )
        model.eval()
        return model
    cfg = GNNPlusConfig(
        dropout=0.2,
        learn_k_crit=False,
        k_crit_min=K_CRIT,
        k_crit_max=K_CRIT,
        feature_variant=FEATURE_VARIANT,
        num_layers=NUM_LAYERS_OVERRIDE,
        use_gated_residual=USE_GATED_RESIDUAL,
        use_cross_od_attention=USE_CROSS_OD_ATTENTION,
        cross_od_attention_heads=CROSS_OD_ATTENTION_HEADS,
        bottleneck_moe_floor=BOTTLENECK_MOE_FLOOR,
        device=DEVICE,
    )
    return GNNPlusFlowSelector(cfg).to(DEVICE)


def balanced_supervised_epoch_indices(
    samples: list[dict],
    *,
    rng: np.random.Generator,
) -> tuple[np.ndarray, dict[str, int]]:
    if not samples:
        return np.asarray([], dtype=np.int64), {}

    sample_weights = np.asarray(
        [float(SUPERVISED_TOPOLOGY_WEIGHTS.get(str(sample["topology"]).lower(), 1.0)) for sample in samples],
        dtype=np.float64,
    )
    sample_weights = np.clip(sample_weights, 1e-6, None)
    probs = sample_weights / np.sum(sample_weights)
    epoch_size = max(1, int(round(len(samples) * SUP_BALANCED_EPOCH_MULTIPLIER)))
    order = rng.choice(np.arange(len(samples), dtype=np.int64), size=epoch_size, replace=True, p=probs)
    topo_mix: dict[str, int] = {}
    for idx in order.tolist():
        topo = str(samples[int(idx)]["topology"])
        topo_mix[topo] = topo_mix.get(topo, 0) + 1
    return order, topo_mix


def supervised_pl_weight_for_epoch(epoch: int) -> float:
    start = float(PLACKETT_LUCE_WEIGHT)
    end = float(PLACKETT_LUCE_WEIGHT_END)
    ramp_epochs = max(int(PLACKETT_LUCE_RAMP_EPOCHS), 1)
    if int(epoch) <= 1:
        return start
    if int(epoch) >= ramp_epochs:
        return end
    alpha = float(int(epoch) - 1) / float(max(ramp_epochs - 1, 1))
    return float(start + alpha * (end - start))


def supervised_lr_for_epoch(epoch: int) -> float:
    base_lr = float(SUP_LR)
    if int(epoch) < int(SUP_COSINE_DECAY_START_EPOCH):
        return base_lr
    total_tail = max(int(SUP_MAX_EPOCHS) - int(SUP_COSINE_DECAY_START_EPOCH), 1)
    progress = float(int(epoch) - int(SUP_COSINE_DECAY_START_EPOCH)) / float(total_tail)
    progress = min(max(progress, 0.0), 1.0)
    cosine = 0.5 * (1.0 + math.cos(math.pi * progress))
    lr_ratio = float(SUP_MIN_LR_RATIO) + (1.0 - float(SUP_MIN_LR_RATIO)) * cosine
    return float(base_lr * lr_ratio)


def set_optimizer_lr(optimizer: torch.optim.Optimizer, lr_value: float) -> None:
    for group in optimizer.param_groups:
        group["lr"] = float(lr_value)


def run_supervised_training(train_samples: list[dict], val_samples: list[dict]):
    TRAIN_DIR.mkdir(parents=True, exist_ok=True)
    model = build_initial_gnnplus_model()
    optimizer = torch.optim.AdamW(model.parameters(), lr=float(SUP_LR), weight_decay=1e-5)
    rng = np.random.default_rng(SEED)

    logs = []
    topology_logs = []
    best_val_loss = float("inf")
    best_val_recall = float("-inf")
    best_epoch_oracle_cover = 0.0
    best_epoch = 0
    stale = 0
    start = time.perf_counter()

    print(f"[supervised] train={len(train_samples)} val={len(val_samples)}", flush=True)
    for epoch in range(1, SUP_MAX_EPOCHS + 1):
        epoch_start = time.perf_counter()
        learning_rate = supervised_lr_for_epoch(epoch)
        set_optimizer_lr(optimizer, learning_rate)
        pl_weight = supervised_pl_weight_for_epoch(epoch)
        teacher_prob = scheduled_teacher_forcing_prob(epoch, SUP_MAX_EPOCHS)
        epoch_train_samples = materialize_scheduled_history(
            train_samples,
            model,
            teacher_forcing_prob=teacher_prob,
            seed=SEED + epoch,
        )
        epoch_val_samples = materialize_scheduled_history(
            val_samples,
            model,
            teacher_forcing_prob=0.0,
            seed=SEED + 100 + epoch,
        )
        model.train()
        order, topo_mix = balanced_supervised_epoch_indices(epoch_train_samples, rng=rng)
        epoch_losses = []
        epoch_loss_parts = defaultdict(list)
        epoch_oracle_hit_fraction = []
        epoch_oracle_full_cover = []

        for idx in order:
            sample = epoch_train_samples[int(idx)]
            graph_data, od_data, candidate_meta = build_experiment_inputs(
                sample,
                device=DEVICE,
                candidate_limit=TRAIN_CANDIDATE_POOL_SIZE,
            )
            scores, _, info = model(graph_data, od_data)
            oracle_mask, soft_target, criticality_target, candidate_np, oracle_local_indices, candidate_meta = project_targets_to_score_space(
                sample,
                od_data,
                scores.size(0),
            )

            ranking_loss = _ranking_loss(scores, oracle_mask, margin=0.1)
            pl_loss_unweighted = _plackett_luce_topk_loss(scores, oracle_local_indices)
            pl_loss = float(pl_weight) * pl_loss_unweighted
            soft_teacher_loss = _soft_teacher_loss(
                scores,
                soft_target=soft_target,
                criticality=criticality_target,
                soft_weight=SOFT_TEACHER_WEIGHT,
                criticality_weight=CRITICALITY_WEIGHT,
            )
            bottleneck_aux_loss = float(BOTTLENECK_AUX_WEIGHT) * _bottleneck_aux_loss(
                info.get("_bottleneck_aux_pred"),
                od_data["bottleneck_scores"],
            )
            temporal_consistency_loss = float(TEMPORAL_CONSISTENCY_WEIGHT) * _temporal_consistency_loss(
                scores,
                prev_selected_indicator=sample["prev_selected_indicator"],
                candidate_od_indices=candidate_np,
                full_num_od=int(od_data.get("full_num_od", scores.size(0))),
            )
            gate_reg_loss = torch.tensor(0.0, dtype=torch.float32, device=DEVICE)
            prefilter_recall_loss = torch.tensor(0.0, dtype=torch.float32, device=DEVICE)
            loss = ranking_loss + pl_loss + soft_teacher_loss + bottleneck_aux_loss + temporal_consistency_loss + gate_reg_loss
            if (
                float(PREFILTER_RECALL_WEIGHT) > 0.0
                and not bool(candidate_meta["oracle_in_window"])
                and candidate_np is not None
            ):
                full_od_data = build_od_features_plus(
                    sample["dataset"],
                    sample["tm_vector"],
                    sample["path_library"],
                    telemetry=sample["telemetry"],
                    prev_tm=sample["prev_tm"],
                    prev_util=sample["prev_util"],
                    prev_selected_indicator=sample["prev_selected_indicator"],
                    prev_disturbance=sample["prev_disturbance"],
                    failure_mask=sample.get("failure_mask"),
                    candidate_od_indices=None,
                    feature_variant=FEATURE_VARIANT,
                    device=DEVICE,
                )
                _, _, full_info = model(graph_data, full_od_data)
                prefilter_recall_loss = float(PREFILTER_RECALL_WEIGHT) * _oracle_recall_aux_loss(
                    full_info.get("_bottleneck_aux_pred"),
                    sample["oracle_selected"],
                )
                loss = loss + prefilter_recall_loss
            optimizer.zero_grad()
            loss.backward()
            torch.nn.utils.clip_grad_norm_(model.parameters(), 1.0)
            optimizer.step()
            epoch_losses.append(float(loss.item()))
            epoch_loss_parts["ranking_loss"].append(float(ranking_loss.item()))
            epoch_loss_parts["pl_loss"].append(float(pl_loss.item()))
            epoch_loss_parts["pl_loss_unweighted"].append(float(pl_loss_unweighted.item()))
            epoch_loss_parts["soft_teacher_loss"].append(float(soft_teacher_loss.item()))
            epoch_loss_parts["aux_bottleneck_loss"].append(float(bottleneck_aux_loss.item()))
            epoch_loss_parts["temporal_consistency_loss"].append(float(temporal_consistency_loss.item()))
            epoch_loss_parts["prefilter_recall_loss"].append(float(prefilter_recall_loss.item()))
            epoch_loss_parts["gate_reg_loss"].append(float(gate_reg_loss.item()))
            epoch_oracle_hit_fraction.append(float(candidate_meta["oracle_hit_fraction"]))
            epoch_oracle_full_cover.append(float(bool(candidate_meta["oracle_in_window"])))

        train_loss = float(np.mean(epoch_losses)) if epoch_losses else 0.0
        train_loss_parts = {
            key: float(np.mean(values)) if values else 0.0
            for key, values in epoch_loss_parts.items()
        }

        model.eval()
        val_losses = []
        val_loss_parts = defaultdict(list)
        val_overlap = []
        val_recall_at_40 = []
        val_oracle_hit_fraction = []
        val_oracle_full_cover = []
        val_recall_by_topology = defaultdict(list)
        val_oracle_cover_by_topology = defaultdict(list)
        val_pl_loss_by_topology = defaultdict(list)
        val_pl_loss_unweighted_by_topology = defaultdict(list)
        with torch.no_grad():
            for sample in epoch_val_samples:
                graph_data, od_data, _ = build_experiment_inputs(
                    sample,
                    device=DEVICE,
                    candidate_limit=TRAIN_CANDIDATE_POOL_SIZE,
                )
                scores, _, info = model(graph_data, od_data)
                oracle_mask, soft_target, criticality_target, candidate_np, oracle_local_indices, candidate_meta = project_targets_to_score_space(
                    sample,
                    od_data,
                    scores.size(0),
                )
                val_ranking_loss = _ranking_loss(scores, oracle_mask, margin=0.1)
                val_pl_loss_unweighted = _plackett_luce_topk_loss(scores, oracle_local_indices)
                val_pl_loss = float(pl_weight) * val_pl_loss_unweighted
                val_soft_teacher_loss = _soft_teacher_loss(
                    scores,
                    soft_target=soft_target,
                    criticality=criticality_target,
                    soft_weight=SOFT_TEACHER_WEIGHT,
                    criticality_weight=CRITICALITY_WEIGHT,
                )
                val_bottleneck_aux_loss = float(BOTTLENECK_AUX_WEIGHT) * _bottleneck_aux_loss(
                    info.get("_bottleneck_aux_pred"),
                    od_data["bottleneck_scores"],
                )
                val_temporal_consistency_loss = float(TEMPORAL_CONSISTENCY_WEIGHT) * _temporal_consistency_loss(
                    scores,
                    prev_selected_indicator=sample["prev_selected_indicator"],
                    candidate_od_indices=candidate_np,
                    full_num_od=int(od_data.get("full_num_od", scores.size(0))),
                )
                val_gate_reg_loss = torch.tensor(0.0, dtype=torch.float32, device=DEVICE)
                vloss = (
                    val_ranking_loss
                    + val_pl_loss
                    + val_soft_teacher_loss
                    + val_bottleneck_aux_loss
                    + val_temporal_consistency_loss
                    + val_gate_reg_loss
                )
                val_losses.append(float(vloss.item()))
                val_loss_parts["ranking_loss"].append(float(val_ranking_loss.item()))
                val_loss_parts["pl_loss"].append(float(val_pl_loss.item()))
                val_loss_parts["pl_loss_unweighted"].append(float(val_pl_loss_unweighted.item()))
                val_loss_parts["soft_teacher_loss"].append(float(val_soft_teacher_loss.item()))
                val_loss_parts["aux_bottleneck_loss"].append(float(val_bottleneck_aux_loss.item()))
                val_loss_parts["temporal_consistency_loss"].append(float(val_temporal_consistency_loss.item()))
                val_loss_parts["prefilter_recall_loss"].append(0.0)
                val_loss_parts["gate_reg_loss"].append(float(val_gate_reg_loss.item()))

                selected, _, _ = continuity_select(
                    scores,
                    active_mask=selection_active_mask(sample),
                    k=K_CRIT,
                    prev_selected_indicator=sample["prev_selected_indicator"],
                    continuity_bonus=0.0,
                    candidate_od_indices=candidate_np,
                    full_num_od=int(od_data.get("full_num_od", scores.size(0))),
                )
                pred_set = set(selected)
                oracle_set = set(sample["oracle_selected"])
                overlap = len(pred_set & oracle_set) / max(len(pred_set | oracle_set), 1)
                val_overlap.append(float(overlap))
                recall_at_40 = len(pred_set & oracle_set) / max(len(oracle_set), 1)
                val_recall_at_40.append(float(recall_at_40))
                val_oracle_hit_fraction.append(float(candidate_meta["oracle_hit_fraction"]))
                val_oracle_full_cover.append(float(bool(candidate_meta["oracle_in_window"])))
                topo_key = str(sample["topology"])
                val_recall_by_topology[topo_key].append(float(recall_at_40))
                val_oracle_cover_by_topology[topo_key].append(float(bool(candidate_meta["oracle_in_window"])))
                val_pl_loss_by_topology[topo_key].append(float(val_pl_loss.item()))
                val_pl_loss_unweighted_by_topology[topo_key].append(float(val_pl_loss_unweighted.item()))

        val_loss = float(np.mean(val_losses)) if val_losses else 0.0
        val_loss_parts_mean = {
            key: float(np.mean(values)) if values else 0.0
            for key, values in val_loss_parts.items()
        }
        overlap = float(np.mean(val_overlap)) if val_overlap else 0.0
        train_oracle_hit = float(np.mean(epoch_oracle_hit_fraction)) if epoch_oracle_hit_fraction else 0.0
        val_oracle_hit = float(np.mean(val_oracle_hit_fraction)) if val_oracle_hit_fraction else 0.0
        train_oracle_cover = float(np.mean(epoch_oracle_full_cover)) if epoch_oracle_full_cover else 0.0
        val_oracle_cover = float(np.mean(val_oracle_full_cover)) if val_oracle_full_cover else 0.0
        val_recall = float(np.mean(val_recall_at_40)) if val_recall_at_40 else 0.0
        val_recall_topology_summary = {
            str(topo): float(np.mean(values)) if values else 0.0
            for topo, values in sorted(val_recall_by_topology.items())
        }
        val_oracle_cover_topology_summary = {
            str(topo): float(np.mean(values)) if values else 0.0
            for topo, values in sorted(val_oracle_cover_by_topology.items())
        }
        val_pl_loss_topology_summary = {
            str(topo): float(np.mean(values)) if values else 0.0
            for topo, values in sorted(val_pl_loss_by_topology.items())
        }
        val_pl_loss_unweighted_topology_summary = {
            str(topo): float(np.mean(values)) if values else 0.0
            for topo, values in sorted(val_pl_loss_unweighted_by_topology.items())
        }
        logs.append(
            {
                "epoch": epoch,
                "learning_rate": float(learning_rate),
                "plackett_luce_weight": float(pl_weight),
                "train_loss": train_loss,
                "val_loss": val_loss,
                "val_selection_overlap": overlap,
                "plackett_luce_recall_at_40": val_recall,
                "frac_oracle_in_window": val_oracle_cover,
                "train_oracle_hit_fraction": train_oracle_hit,
                "val_oracle_hit_fraction": val_oracle_hit,
                "train_oracle_full_cover": train_oracle_cover,
                "val_oracle_full_cover": val_oracle_cover,
                "train_ranking_loss": train_loss_parts.get("ranking_loss", 0.0),
                "train_pl_loss": train_loss_parts.get("pl_loss", 0.0),
                "train_pl_loss_unweighted": train_loss_parts.get("pl_loss_unweighted", 0.0),
                "train_soft_teacher_loss": train_loss_parts.get("soft_teacher_loss", 0.0),
                "train_aux_bottleneck_loss": train_loss_parts.get("aux_bottleneck_loss", 0.0),
                "train_temporal_consistency_loss": train_loss_parts.get("temporal_consistency_loss", 0.0),
                "train_prefilter_recall_loss": train_loss_parts.get("prefilter_recall_loss", 0.0),
                "train_gate_reg_loss": train_loss_parts.get("gate_reg_loss", 0.0),
                "val_ranking_loss": val_loss_parts_mean.get("ranking_loss", 0.0),
                "val_pl_loss": val_loss_parts_mean.get("pl_loss", 0.0),
                "val_pl_loss_unweighted": val_loss_parts_mean.get("pl_loss_unweighted", 0.0),
                "val_soft_teacher_loss": val_loss_parts_mean.get("soft_teacher_loss", 0.0),
                "val_aux_bottleneck_loss": val_loss_parts_mean.get("aux_bottleneck_loss", 0.0),
                "val_temporal_consistency_loss": val_loss_parts_mean.get("temporal_consistency_loss", 0.0),
                "val_prefilter_recall_loss": val_loss_parts_mean.get("prefilter_recall_loss", 0.0),
                "val_gate_reg_loss": val_loss_parts_mean.get("gate_reg_loss", 0.0),
                "val_recall_at_40_by_topology": json.dumps(val_recall_topology_summary, sort_keys=True),
                "val_frac_oracle_in_window_by_topology": json.dumps(val_oracle_cover_topology_summary, sort_keys=True),
                "val_pl_loss_by_topology": json.dumps(val_pl_loss_topology_summary, sort_keys=True),
                "val_pl_loss_unweighted_by_topology": json.dumps(val_pl_loss_unweighted_topology_summary, sort_keys=True),
                "teacher_forcing_prob": teacher_prob,
                "epoch_topology_mix": json.dumps(topo_mix, sort_keys=True),
                "epoch_time_sec": float(time.perf_counter() - epoch_start),
            }
        )
        for topo_key in sorted(val_recall_topology_summary):
            topology_logs.append(
                {
                    "epoch": int(epoch),
                    "topology": str(topo_key),
                    "val_recall_at_40": float(val_recall_topology_summary[topo_key]),
                    "val_frac_oracle_in_window": float(val_oracle_cover_topology_summary.get(topo_key, 0.0)),
                    "val_pl_loss": float(val_pl_loss_topology_summary.get(topo_key, 0.0)),
                    "val_pl_loss_unweighted": float(val_pl_loss_unweighted_topology_summary.get(topo_key, 0.0)),
                    "num_val_samples": int(len(val_recall_by_topology.get(topo_key, []))),
                }
            )
        print(
            f"[supervised] epoch={epoch:02d} train_loss={train_loss:.4f} val_loss={val_loss:.4f} "
            f"recall@40={val_recall:.3f} frac_oracle_in_window={val_oracle_cover:.3f} "
            f"pl_w={pl_weight:.2f} lr={learning_rate:.2e} "
            f"teacher_prob={teacher_prob:.2f} topo_mix={topo_mix}",
            flush=True,
        )

        improved_recall = val_recall > best_val_recall + 1e-6
        tied_recall_better_loss = abs(val_recall - best_val_recall) <= 1e-6 and val_loss + 1e-6 < best_val_loss
        if improved_recall or tied_recall_better_loss:
            best_val_recall = val_recall
            best_val_loss = val_loss
            best_epoch_oracle_cover = val_oracle_cover
            best_epoch = epoch
            stale = 0
            save_gnn_plus(
                model,
                SUP_CKPT,
                extra_meta={
                    "stage": "section3_4_7_supervised_finetune",
                    "base_checkpoint": str(BASE_GNNPLUS_CKPT.relative_to(PROJECT_ROOT)),
                    "feature_variant": FEATURE_VARIANT,
                    "best_epoch": best_epoch,
                    "best_val_recall_at_40": best_val_recall,
                    "best_val_loss": best_val_loss,
                    "best_epoch_frac_oracle_in_window": best_epoch_oracle_cover,
                },
            )
        else:
            stale += 1
        if stale >= SUP_PATIENCE:
            print(f"[supervised] early stop at epoch {epoch}", flush=True)
            break

    pd.DataFrame(logs).to_csv(SUP_LOG_CSV, index=False)
    pd.DataFrame(topology_logs).to_csv(SUP_TOPOLOGY_LOG_CSV, index=False)
    model, _ = load_gnn_plus(SUP_CKPT, device=DEVICE)
    summary = {
        "mode": "supervised",
        "best_epoch": int(best_epoch),
        "best_val_recall_at_40": float(best_val_recall if math.isfinite(best_val_recall) else 0.0),
        "best_val_loss": float(best_val_loss),
        "best_epoch_frac_oracle_in_window": float(best_epoch_oracle_cover),
        "num_train_samples": int(len(train_samples)),
        "num_val_samples": int(len(val_samples)),
        "training_time_sec": float(time.perf_counter() - start),
        "continuity_bonus": float(CONTINUITY_BONUS),
        "feature_variant": FEATURE_VARIANT,
        "feature_profile": feature_profile_description(),
        "soft_teacher_weight": SOFT_TEACHER_WEIGHT,
        "criticality_weight": CRITICALITY_WEIGHT,
        "lp_teacher_weight": LP_TEACHER_WEIGHT,
        "teacher_profile": teacher_profile_description(),
        "supervised_topology_weights": {str(k): float(v) for k, v in SUPERVISED_TOPOLOGY_WEIGHTS.items()},
        "supervised_epoch_multiplier": float(SUP_BALANCED_EPOCH_MULTIPLIER),
        "num_layers": int(model.cfg.num_layers),
        "candidate_prefilter": {
            "enabled": bool(USE_CANDIDATE_PREFILTER),
            "multiplier": float(CANDIDATE_PREFILTER_MULTIPLIER),
            "train_candidate_pool_size": int(TRAIN_CANDIDATE_POOL_SIZE),
            "infer_candidate_pool_size": int(INFER_CANDIDATE_POOL_SIZE),
        },
        "architecture": {
            "use_gated_residual": bool(model.cfg.use_gated_residual),
            "use_cross_od_attention": bool(model.cfg.use_cross_od_attention),
            "cross_od_attention_heads": int(model.cfg.cross_od_attention_heads),
            "bottleneck_moe_floor": float(model.cfg.bottleneck_moe_floor),
        },
        "synthetic_failure_training": {
            "enabled": float(SYNTHETIC_FAILURE_PROB) > 0.0,
            "probability": float(SYNTHETIC_FAILURE_PROB),
            "min_failed_edges": int(SYNTHETIC_FAILURE_MIN_EDGES),
            "max_failed_edges": int(SYNTHETIC_FAILURE_MAX_EDGES),
            "capacity_augment_probability": float(SYNTHETIC_FAILURE_AUGMENT_PROB),
            "capacity_perturbation": float(SYNTHETIC_CAPACITY_PERTURB),
        },
        "oracle_window": {
            "target_hit_fraction": float(ORACLE_WINDOW_TARGET),
            "train_candidate_pool_size": int(TRAIN_CANDIDATE_POOL_SIZE),
            "infer_candidate_pool_size": int(INFER_CANDIDATE_POOL_SIZE),
            "best_val_oracle_hit_fraction": float(max((row["val_oracle_hit_fraction"] for row in logs), default=0.0)),
            "best_val_oracle_full_cover": float(max((row["val_oracle_full_cover"] for row in logs), default=0.0)),
            "final_frac_oracle_in_window": float(logs[-1]["frac_oracle_in_window"]) if logs else 0.0,
        },
        "final_plackett_luce_recall_at_40": float(logs[-1]["plackett_luce_recall_at_40"]) if logs else 0.0,
        "losses": {
            "plackett_luce_weight_start": float(PLACKETT_LUCE_WEIGHT),
            "plackett_luce_weight_end": float(PLACKETT_LUCE_WEIGHT_END),
            "plackett_luce_ramp_epochs": int(PLACKETT_LUCE_RAMP_EPOCHS),
            "bottleneck_aux_weight": float(BOTTLENECK_AUX_WEIGHT),
            "prefilter_recall_weight": float(PREFILTER_RECALL_WEIGHT),
            "temporal_consistency_weight": float(TEMPORAL_CONSISTENCY_WEIGHT),
        },
        "lr_schedule": {
            "supervised_lr": float(SUP_LR),
            "cosine_decay_start_epoch": int(SUP_COSINE_DECAY_START_EPOCH),
            "min_lr_ratio": float(SUP_MIN_LR_RATIO),
        },
    }
    SUPERVISED_SUMMARY_JSON.write_text(json.dumps(summary, indent=2) + "\n", encoding="utf-8")
    return model, summary


def load_reference_model_cached(cache: dict, runner, sample: dict):
    topo_key = str(sample["topology"])
    if topo_key not in cache:
        cache[topo_key] = runner.load_gnn_model(sample["dataset"], sample["path_library"])
    return cache[topo_key]


def clip_gate_attention_gradients(model: torch.nn.Module, epoch: int) -> None:
    special_max_norm = (
        float(RL_GATE_ATTENTION_CLIP_EARLY)
        if int(epoch) <= int(RL_GATE_ATTENTION_CLIP_EPOCHS)
        else float(RL_GATE_ATTENTION_CLIP_LATE)
    )
    special_params = []
    other_params = []
    for name, param in model.named_parameters():
        if param.grad is None:
            continue
        if any(tag in name for tag in ("gate_head", "cross_od_attention", "cross_od_edge_bias", "od_input_proj", "od_scorer_attended")):
            special_params.append(param)
        else:
            other_params.append(param)
    if special_params:
        torch.nn.utils.clip_grad_norm_(special_params, special_max_norm)
    if other_params:
        torch.nn.utils.clip_grad_norm_(other_params, 1.0)


def load_supervised_checkpoint_summary() -> tuple[GNNPlusFlowSelector, dict]:
    fallback_ckpt = PREVIOUS_TRAIN_DIR / SUP_CKPT.name
    fallback_summary = PREVIOUS_TRAIN_DIR / SUPERVISED_SUMMARY_JSON.name
    use_fallback = PREVIOUS_TRAIN_DIR != TRAIN_DIR and fallback_ckpt.exists() and fallback_summary.exists()
    if use_fallback:
        TRAIN_DIR.mkdir(parents=True, exist_ok=True)
        shutil.copy2(fallback_ckpt, SUP_CKPT)
        shutil.copy2(fallback_summary, SUPERVISED_SUMMARY_JSON)
    elif not SUP_CKPT.exists() or not SUPERVISED_SUMMARY_JSON.exists():
        if not SUP_CKPT.exists():
            raise FileNotFoundError(f"Missing supervised checkpoint: {SUP_CKPT}")
        raise FileNotFoundError(f"Missing supervised summary: {SUPERVISED_SUMMARY_JSON}")
    model, _ = load_gnn_plus(SUP_CKPT, device=DEVICE)
    summary = json.loads(SUPERVISED_SUMMARY_JSON.read_text(encoding="utf-8"))
    if use_fallback:
        summary["reused_from"] = str(PREVIOUS_OUTPUT_DIR.relative_to(PROJECT_ROOT))
    return model, summary


def synthesize_reinforce_summary_from_artifacts(train_dir: Path, checkpoint_path: Path) -> dict:
    if not checkpoint_path.exists():
        raise FileNotFoundError(f"Missing final checkpoint: {checkpoint_path}")
    rl_log_path = train_dir / RL_LOG_CSV.name
    if not rl_log_path.exists():
        raise FileNotFoundError(f"Missing RL log for summary synthesis: {rl_log_path}")
    rl_df = pd.read_csv(rl_log_path)
    if rl_df.empty:
        raise RuntimeError(f"RL log is empty: {rl_log_path}")

    payload = torch.load(str(checkpoint_path), map_location=torch.device("cpu"), weights_only=False)
    selection_metric = str(payload.get("selection_metric") or rl_df.iloc[-1].get("rl_selection_metric") or RL_SELECTION_METRIC)
    if "rl_selection_value" in rl_df.columns:
        best_idx = rl_df["rl_selection_value"].astype(float).idxmin()
    elif selection_metric == "val_disturbance" and "rl_val_disturbance" in rl_df.columns:
        best_idx = rl_df["rl_val_disturbance"].astype(float).idxmin()
    else:
        best_idx = rl_df["rl_val_failure_mlu"].astype(float).idxmin()
    best_row = rl_df.loc[best_idx]
    best_epoch = int(payload.get("best_epoch", best_row.get("epoch", 0)))
    epoch_ckpt = train_dir / "rl_epoch_checkpoints" / f"epoch_{best_epoch:02d}.pt"
    rl_config = dict(payload.get("rl_config", {}))
    cfg = dict(payload.get("config", {}))
    beta_failure = float(payload.get("rl_kl_beta_failure", rl_config.get("kl_beta_failure", RL_KL_BETA_FAILURE)))

    reinforce_summary = {
        "mode": "reinforce",
        "best_epoch": best_epoch,
        "best_epoch_checkpoint": str(epoch_ckpt.relative_to(PROJECT_ROOT)) if epoch_ckpt.exists() else None,
        "selection_metric": selection_metric,
        "best_val_mlu": float(payload.get("best_val_mlu", best_row.get("rl_val_mlu", 0.0))),
        "best_val_normal_mlu": float(payload.get("best_val_normal_mlu", best_row.get("rl_val_normal_mlu", 0.0))),
        "best_val_failure_mlu": float(payload.get("best_val_failure_mlu", best_row.get("rl_val_failure_mlu", 0.0))),
        "best_val_disturbance": float(payload.get("best_val_disturbance", best_row.get("rl_val_disturbance", 0.0))),
        "best_val_score": float(payload.get("best_val_objective", best_row.get("rl_val_score", 0.0))),
        "continuity_bonus": float(payload.get("continuity_bonus", 0.0)),
        "rl_config": rl_config,
        "hinge_reward_enabled": bool(ENABLE_HINGE_REWARD),
        "hinge_threshold": float(HINGE_THRESHOLD),
        "hinge_multiplier": float(HINGE_MULTIPLIER),
        "num_layers": int(cfg.get("num_layers", NUM_LAYERS_OVERRIDE)),
        "num_train_samples": int(RL_MAX_TRAIN_SAMPLES),
        "num_val_samples": int(RL_MAX_VAL_SAMPLES),
        "candidate_prefilter": {
            "enabled": bool(USE_CANDIDATE_PREFILTER),
            "multiplier": float(CANDIDATE_PREFILTER_MULTIPLIER),
            "train_candidate_pool_size": int(TRAIN_CANDIDATE_POOL_SIZE),
            "infer_candidate_pool_size": int(INFER_CANDIDATE_POOL_SIZE),
        },
        "architecture": {
            "use_gated_residual": bool(cfg.get("use_gated_residual", USE_GATED_RESIDUAL)),
            "use_cross_od_attention": bool(cfg.get("use_cross_od_attention", USE_CROSS_OD_ATTENTION)),
            "cross_od_attention_heads": int(cfg.get("cross_od_attention_heads", CROSS_OD_ATTENTION_HEADS)),
            "bottleneck_moe_floor": float(cfg.get("bottleneck_moe_floor", BOTTLENECK_MOE_FLOOR)),
        },
        "step5_patch": {
            "continuity_bonus_old": 0.03,
            "continuity_bonus_new": float(payload.get("continuity_bonus", 0.0)),
            "w_reward_mlu_old": 1.15,
            "w_reward_mlu_new": float(rl_config.get("w_reward_mlu", REWARD_MLU)),
            "w_reward_mlu_normal": float(rl_config.get("w_reward_mlu_normal", REWARD_MLU_NORMAL)),
            "w_reward_mlu_failure": float(rl_config.get("w_reward_mlu_failure", REWARD_MLU_FAILURE)),
            "w_reward_disturbance_old": 0.15,
            "w_reward_disturbance_new": float(rl_config.get("w_reward_disturbance", REWARD_DISTURBANCE)),
            "rl_failure_rollout_prob": float(rl_config.get("rl_synthetic_failure_prob", RL_FAILURE_ROLLOUT_PROB)),
            "per_topology_reward_norm_enabled": bool(PER_TOPO_REWARD_NORM),
            "per_topology_mlu_scale_trainonly": {},
            "normalization_source_topologies": list(KNOWN_TOPOLOGIES),
            "normalization_excluded_topologies": list(UNSEEN_TOPOLOGIES),
            "rl_kl_beta_start": float(payload.get("rl_kl_beta_start", RL_KL_BETA_START)),
            "rl_kl_beta_end": float(payload.get("rl_kl_beta_end", RL_KL_BETA_END)),
            "rl_kl_beta_failure_final": beta_failure,
            "rl_failure_kl_target": float(rl_config.get("failure_kl_target", RL_FAILURE_KL_TARGET)),
            "rl_normal_kl_abort": float(rl_config.get("normal_kl_abort", RL_NORMAL_KL_ABORT)),
            "disturbance_churn_multiplier": float(DISTURBANCE_CHURN_MULTIPLIER),
            "failure_beta_adjustments": 0,
            "failure_beta_adjustment_reason": None,
        },
    }
    return {"reinforce": reinforce_summary}


def load_final_checkpoint_summary() -> tuple[GNNPlusFlowSelector, dict]:
    fallback_ckpt = PREVIOUS_TRAIN_DIR / FINAL_CKPT.name
    fallback_summary = PREVIOUS_TRAIN_DIR / TRAINING_SUMMARY_JSON.name
    use_fallback_ckpt = PREVIOUS_TRAIN_DIR != TRAIN_DIR and fallback_ckpt.exists()
    use_fallback_summary = use_fallback_ckpt and fallback_summary.exists()
    if use_fallback_ckpt:
        TRAIN_DIR.mkdir(parents=True, exist_ok=True)
        shutil.copy2(fallback_ckpt, FINAL_CKPT)
        if use_fallback_summary:
            shutil.copy2(fallback_summary, TRAINING_SUMMARY_JSON)
    elif not FINAL_CKPT.exists():
        raise FileNotFoundError(f"Missing final checkpoint: {FINAL_CKPT}")
    model, _ = load_gnn_plus(FINAL_CKPT, device=DEVICE)
    # Rescue branch (gnnplus-debug-rescue): honor env-var override on the
    # eval_reuse_final path. Without this, GNNPLUS_BOTTLENECK_MOE_FLOOR is a
    # silent no-op because load_gnn_plus reconstructs cfg straight from the
    # saved checkpoint. The override at main-script line 1304 only fires on
    # the fresh-model path.
    _prev_moe_floor = float(getattr(model.cfg, "bottleneck_moe_floor", 0.10))
    if abs(_prev_moe_floor - float(BOTTLENECK_MOE_FLOOR)) > 1e-12:
        model.cfg.bottleneck_moe_floor = float(BOTTLENECK_MOE_FLOOR)
        print(
            f"[rescue] Overrode checkpoint bottleneck_moe_floor "
            f"{_prev_moe_floor:.4f} -> {float(BOTTLENECK_MOE_FLOOR):.4f} "
            f"(env GNNPLUS_BOTTLENECK_MOE_FLOOR)"
        )
    if TRAINING_SUMMARY_JSON.exists():
        summary = json.loads(TRAINING_SUMMARY_JSON.read_text(encoding="utf-8"))
    else:
        summary_source_dir = PREVIOUS_TRAIN_DIR if use_fallback_ckpt else TRAIN_DIR
        summary = synthesize_reinforce_summary_from_artifacts(summary_source_dir, FINAL_CKPT)
        TRAINING_SUMMARY_JSON.write_text(json.dumps(summary, indent=2) + "\n", encoding="utf-8")
    if use_fallback_ckpt:
        summary["reused_from"] = str(PREVIOUS_OUTPUT_DIR.relative_to(PROJECT_ROOT))
    return model, summary


def compute_supervised_lp_gap_audit(
    model: GNNPlusFlowSelector,
    val_samples: list[dict],
    *,
    best_epoch: int,
) -> dict:
    audit_seed = SEED + 100 + max(int(best_epoch), 1)
    eval_samples = materialize_scheduled_history(
        val_samples,
        model,
        teacher_forcing_prob=0.0,
        seed=audit_seed,
    )

    grouped: dict[str, list[dict]] = defaultdict(list)
    for sample in eval_samples:
        topo = str(sample["topology"]).lower()
        if topo not in KNOWN_TOPOLOGIES:
            continue
        if len(grouped[topo]) >= int(SUP_LP_GAP_SAMPLES_PER_TOPO):
            continue
        grouped[topo].append(sample)

    rows: list[dict[str, object]] = []
    was_training = model.training
    model.eval()
    with torch.no_grad():
        for topo in KNOWN_TOPOLOGIES:
            for sample in grouped.get(topo, []):
                gap_value = float("inf")
                oracle_mlu = float("inf")
                model_mlu = float("inf")
                selected_ods: list[int] = []
                error_text = None
                try:
                    graph_data, od_data, _ = build_experiment_inputs(
                        sample,
                        device=DEVICE,
                        candidate_limit=TRAIN_CANDIDATE_POOL_SIZE,
                    )
                    scores, _, _ = model(graph_data, od_data)
                    candidate_np = candidate_indices_from_od_data(od_data)
                    selected_ods, _, _ = continuity_select(
                        scores,
                        active_mask=selection_active_mask(sample),
                        k=K_CRIT,
                        prev_selected_indicator=sample["prev_selected_indicator"],
                        continuity_bonus=0.0,
                        candidate_od_indices=candidate_np,
                        full_num_od=int(od_data.get("full_num_od", scores.size(0))),
                    )
                    oracle_selected = [int(x) for x in sample["oracle_selected"]]
                    base_splits = ecmp_splits(sample["path_library"])

                    lp_oracle = solve_selected_path_lp(
                        tm_vector=sample["tm_vector"],
                        selected_ods=oracle_selected,
                        base_splits=base_splits,
                        path_library=sample["path_library"],
                        capacities=sample["capacities"],
                        warm_start_splits=sample["prev_splits"],
                        time_limit_sec=LP_TIME_LIMIT,
                    )
                    oracle_routing = apply_routing(
                        sample["tm_vector"],
                        lp_oracle.splits,
                        sample["path_library"],
                        sample["capacities"],
                    )
                    oracle_mlu = float(oracle_routing.mlu)

                    lp_model = solve_selected_path_lp(
                        tm_vector=sample["tm_vector"],
                        selected_ods=[int(x) for x in selected_ods],
                        base_splits=base_splits,
                        path_library=sample["path_library"],
                        capacities=sample["capacities"],
                        warm_start_splits=sample["prev_splits"],
                        time_limit_sec=LP_TIME_LIMIT,
                    )
                    model_routing = apply_routing(
                        sample["tm_vector"],
                        lp_model.splits,
                        sample["path_library"],
                        sample["capacities"],
                    )
                    model_mlu = float(model_routing.mlu)
                    gap_value = (model_mlu - oracle_mlu) / max(abs(oracle_mlu), 1e-12)
                except Exception as exc:
                    error_text = str(exc)

                rows.append(
                    {
                        "topology": str(topo),
                        "timestep": int(sample["timestep"]),
                        "oracle_mlu": float(oracle_mlu),
                        "model_mlu": float(model_mlu),
                        "relative_gap": float(gap_value),
                        "selected_count_model": int(len(selected_ods)) if "selected_ods" in locals() else 0,
                        "selected_count_oracle": int(len(sample["oracle_selected"])),
                        "error": error_text,
                    }
                )
    if was_training:
        model.train()

    rows_df = pd.DataFrame(rows)
    rows_df.to_csv(SUP_LP_GAP_CSV, index=False)

    per_topology: dict[str, dict[str, object]] = {}
    aggregate_gaps: list[float] = []
    for topo in KNOWN_TOPOLOGIES:
        topo_df = rows_df[rows_df["topology"] == topo].copy()
        topo_gaps = pd.to_numeric(topo_df["relative_gap"], errors="coerce").replace([np.inf, -np.inf], np.nan).dropna().to_numpy(dtype=float)
        aggregate_gaps.extend(topo_gaps.tolist())
        per_topology[topo] = {
            "num_samples_requested": int(SUP_LP_GAP_SAMPLES_PER_TOPO),
            "num_samples_evaluated": int(topo_gaps.size),
            "mean_relative_gap": float(np.mean(topo_gaps)) if topo_gaps.size else float("inf"),
            "p95_relative_gap": float(np.percentile(topo_gaps, 95)) if topo_gaps.size else float("inf"),
            "max_relative_gap": float(np.max(topo_gaps)) if topo_gaps.size else float("inf"),
            "num_errors": int(topo_df["error"].notna().sum()) if "error" in topo_df else 0,
        }

    aggregate_gap_arr = np.asarray(aggregate_gaps, dtype=float)
    summary = {
        "checkpoint_path": str(SUP_CKPT.relative_to(PROJECT_ROOT)),
        "best_epoch": int(best_epoch),
        "audit_seed": int(audit_seed),
        "mean_gate": float(SUP_LP_GAP_MEAN_GATE),
        "p95_gate": float(SUP_LP_GAP_P95_GATE),
        "samples_per_topology": int(SUP_LP_GAP_SAMPLES_PER_TOPO),
        "aggregate": {
            "num_samples_evaluated": int(aggregate_gap_arr.size),
            "mean_relative_gap": float(np.mean(aggregate_gap_arr)) if aggregate_gap_arr.size else float("inf"),
            "p95_relative_gap": float(np.percentile(aggregate_gap_arr, 95)) if aggregate_gap_arr.size else float("inf"),
        },
        "per_topology": per_topology,
    }
    SUP_LP_GAP_JSON.write_text(json.dumps(summary, indent=2) + "\n", encoding="utf-8")
    return summary


def run_rl_finetune(model: GNNPlusFlowSelector, train_samples: list[dict], val_samples: list[dict], runner):
    rl_cfg = GNNReinforceConfig(
        lr=RL_LR,
        max_epochs=RL_MAX_EPOCHS,
        patience=RL_PATIENCE,
        baseline_ema=0.9,
        entropy_weight=RL_ENTROPY_WEIGHT,
        w_reward_mlu_normal=REWARD_MLU_NORMAL,
        w_reward_mlu_failure=REWARD_MLU_FAILURE,
        w_reward_mlu=REWARD_MLU,
        w_reward_improvement=REWARD_IMPROVEMENT,
        w_reward_disturbance=REWARD_DISTURBANCE,
        w_reward_infeasible=2.0,
        w_reward_vs_bottleneck=REWARD_VS_BOTTLENECK,
        w_reward_vs_reference=REWARD_VS_REFERENCE,
        w_reward_bottleneck_margin=REWARD_BOTTLENECK_MARGIN,
        rank_loss_weight=0.25,
        score_margin_weight=0.08,
        infeasible_mlu_penalty=10.0,
        rl_synthetic_failure_prob=RL_FAILURE_ROLLOUT_PROB,
        kl_beta_failure=RL_KL_BETA_FAILURE,
        failure_kl_target=RL_FAILURE_KL_TARGET,
        normal_kl_abort=RL_NORMAL_KL_ABORT,
        normal_kl_abort_patience=RL_NORMAL_KL_ABORT_PATIENCE,
    )
    rng = np.random.default_rng(SEED)
    reference_cache = {}
    reference_policy, _ = load_gnn_plus(SUP_CKPT, device=DEVICE)
    reference_policy.eval()
    for param in reference_policy.parameters():
        param.requires_grad_(False)

    base_rl_train = [sample for sample in train_samples if not bool(sample.get("synthetic_failure", False))]
    base_rl_val = [sample for sample in val_samples if not bool(sample.get("synthetic_failure", False))]
    rl_train = list(base_rl_train[: min(len(base_rl_train), RL_MAX_TRAIN_SAMPLES)])
    rl_val = list(base_rl_val[: min(len(base_rl_val), RL_MAX_VAL_SAMPLES)])
    print(f"[reinforce] train={len(rl_train)} val={len(rl_val)}", flush=True)

    topo_mlu_scale: dict[str, float] = {}
    if PER_TOPO_REWARD_NORM:
        _topo_buckets: dict[str, list[float]] = {}
        for _sample in rl_train:
            _topo = str(_sample["topology"])
            if _topo in UNSEEN_TOPOLOGIES:
                continue
            _base = ecmp_splits(_sample["path_library"])
            _routing = apply_routing(
                _sample["tm_vector"], _base, _sample["path_library"], _sample["capacities"]
            )
            _topo_buckets.setdefault(_topo, []).append(float(_routing.mlu))
        for _topo, _vals in _topo_buckets.items():
            topo_mlu_scale[_topo] = max(float(np.mean(_vals)), 1e-6)
        print(f"[step5] per-topology MLU scale (train-only, frozen): {topo_mlu_scale}", flush=True)

    beta_start = float(RL_KL_BETA_START)
    beta_end = float(RL_KL_BETA_END)
    beta_failure = float(rl_cfg.kl_beta_failure)
    selection_metric = str(RL_SELECTION_METRIC).strip().lower()
    if selection_metric not in {"val_failure_mlu", "val_disturbance"}:
        raise ValueError(f"Unsupported RL selection metric: {selection_metric}")
    all_logs: list[dict] = []
    failure_type_logs: list[dict] = []
    failure_beta_adjustments = 0
    failure_beta_adjustment_reason = None

    for attempt in range(1, int(RL_MAX_RESTARTS) + 2):
        model, _ = load_gnn_plus(SUP_CKPT, device=DEVICE)
        optimizer = torch.optim.AdamW(model.parameters(), lr=float(rl_cfg.lr), weight_decay=1e-5)
        normal_reward_window: deque[float] = deque(maxlen=64)
        failure_reward_window: deque[float] = deque(maxlen=64)
        best_val_mlu = float("inf")
        best_val_score = float("-inf")
        best_val_normal_mlu = float("inf")
        best_val_failure_mlu = float("inf")
        best_val_disturbance = float("inf")
        best_epoch_ckpt_path = None
        best_epoch = 0
        stale = 0
        attempt_logs: list[dict] = []
        failure_kl_low_streak = 0
        normal_kl_high_streak = 0
        restart_this_attempt = False
        failure_beta_adjusted_this_attempt = False

        for epoch in range(1, rl_cfg.max_epochs + 1):
            epoch_start = time.perf_counter()
            epoch_train_samples = materialize_scheduled_history(
                rl_train,
                model,
                teacher_forcing_prob=0.0,
                seed=SEED + 200 + epoch + 1000 * (attempt - 1),
            )
            epoch_val_samples = materialize_scheduled_history(
                rl_val,
                model,
                teacher_forcing_prob=0.0,
                seed=SEED + 400 + epoch + 1000 * (attempt - 1),
            )
            epoch_train_samples = augment_rl_failure_rollouts(
                runner,
                epoch_train_samples,
                seed=SEED + 600 + epoch + 1000 * (attempt - 1),
                probability=float(rl_cfg.rl_synthetic_failure_prob),
            )
            epoch_val_samples = augment_rl_failure_rollouts(
                runner,
                epoch_val_samples,
                seed=SEED + 800 + epoch + 1000 * (attempt - 1),
                probability=float(rl_cfg.rl_synthetic_failure_prob),
            )
            model.train()
            order = rng.permutation(len(epoch_train_samples))
            epoch_rewards = []
            epoch_mlus = []
            epoch_disturbances = []
            epoch_improvements = []
            epoch_vs_bn = []
            epoch_vs_ref = []
            epoch_entropies = []
            epoch_kls = []
            epoch_normal_kls = []
            epoch_failure_kls = []
            epoch_normal_advantages = []
            epoch_failure_advantages = []
            epoch_normal_set_overlaps = []
            epoch_failure_set_overlaps = []
            epoch_churn = []
            epoch_failure_baselines = []
            epoch_failure_type_rewards: dict[str, list[float]] = defaultdict(list)
            epoch_failure_type_advantages: dict[str, list[float]] = defaultdict(list)
            epoch_failure_type_counts: dict[str, int] = defaultdict(int)
            alpha = min(max(int(epoch) - 1, 0), int(RL_KL_DECAY_EPOCHS) - 1) / float(max(int(RL_KL_DECAY_EPOCHS) - 1, 1))
            kl_beta_normal = float((1.0 - alpha) * beta_start + alpha * beta_end)

            for idx in order:
                sample = epoch_train_samples[int(idx)]
                is_failure_sample = sample_has_active_failure(sample)
                graph_data, od_data, _ = build_experiment_inputs(
                    sample,
                    device=DEVICE,
                    candidate_limit=TRAIN_CANDIDATE_POOL_SIZE,
                )
                scores, _, info = model(graph_data, od_data)
                candidate_np = candidate_indices_from_od_data(od_data)

                argmax_selected_ods, _, _ = continuity_select(
                    scores,
                    active_mask=selection_active_mask(sample),
                    k=K_CRIT,
                    prev_selected_indicator=sample["prev_selected_indicator"],
                    continuity_bonus=0.0,
                    candidate_od_indices=candidate_np,
                    full_num_od=int(od_data.get("full_num_od", scores.size(0))),
                )
                rollout_scores = scores
                if is_failure_sample and float(RL_GUMBEL_TAU_FAILURE) > 0.0:
                    rollout_scores = sample_gumbel_topk_scores(scores, tau=float(RL_GUMBEL_TAU_FAILURE))
                selected_ods, ranking_scores, selected_log_prob = continuity_select(
                    rollout_scores,
                    active_mask=selection_active_mask(sample),
                    k=K_CRIT,
                    prev_selected_indicator=sample["prev_selected_indicator"],
                    continuity_bonus=0.0,
                    candidate_od_indices=candidate_np,
                    full_num_od=int(od_data.get("full_num_od", scores.size(0))),
                )
                if ranking_scores is None or selected_log_prob is None:
                    continue

                ecmp_base = ecmp_splits(sample["path_library"])
                if float(rl_cfg.w_reward_improvement) != 0.0:
                    ecmp_routing = apply_routing(sample["tm_vector"], ecmp_base, sample["path_library"], sample["capacities"])
                    ecmp_mlu = float(ecmp_routing.mlu)
                else:
                    ecmp_mlu = None

                feasible = True
                try:
                    lp = solve_selected_path_lp(
                        tm_vector=sample["tm_vector"],
                        selected_ods=selected_ods,
                        base_splits=ecmp_base,
                        path_library=sample["path_library"],
                        capacities=sample["capacities"],
                        warm_start_splits=sample["prev_splits"],
                        time_limit_sec=10,
                    )
                    routing = apply_routing(sample["tm_vector"], lp.splits, sample["path_library"], sample["capacities"])
                    mlu = float(routing.mlu)
                    if not math.isfinite(mlu):
                        feasible = False
                        mlu = float(rl_cfg.infeasible_mlu_penalty)
                except Exception:
                    feasible = False
                    mlu = float(rl_cfg.infeasible_mlu_penalty)
                    lp = None

                disturbance = (
                    float(compute_disturbance(sample["prev_splits"], lp.splits, sample["tm_vector"]))
                    if feasible and lp is not None
                    else 1.0
                )
                if ecmp_mlu is None:
                    improvement = 0.0
                else:
                    improvement = (float(ecmp_mlu) - mlu) / max(abs(float(ecmp_mlu)), 1e-12)

                bottleneck_mlu = None
                if float(rl_cfg.w_reward_vs_bottleneck) != 0.0 or float(rl_cfg.w_reward_bottleneck_margin) != 0.0:
                    try:
                        bottleneck_mlu = _bottleneck_baseline_mlu(
                            tm_vector=sample["tm_vector"],
                            ecmp_base=ecmp_base,
                            path_library=sample["path_library"],
                            capacities=sample["capacities"],
                            k_crit=K_CRIT,
                            time_limit_sec=10,
                        )
                    except Exception:
                        bottleneck_mlu = None
                vs_bn = 0.0 if bottleneck_mlu is None else (float(bottleneck_mlu) - mlu) / max(abs(float(bottleneck_mlu)), 1e-12)

                ref_mlu = None
                if float(rl_cfg.w_reward_vs_reference) != 0.0:
                    ref_model = load_reference_model_cached(reference_cache, runner, sample)
                    try:
                        ref_mlu = _reference_model_mlu(
                            reference_model=ref_model,
                            dataset=sample["dataset"],
                            path_library=sample["path_library"],
                            tm_vector=sample["tm_vector"],
                            telemetry=sample["telemetry"],
                            ecmp_base=ecmp_base,
                            capacities=sample["capacities"],
                            k_crit=K_CRIT,
                            device=DEVICE,
                            time_limit_sec=10,
                        )
                    except Exception:
                        ref_mlu = None
                vs_ref = 0.0 if ref_mlu is None else (float(ref_mlu) - mlu) / max(abs(float(ref_mlu)), 1e-12)
                churn_ratio = compute_churn_ratio(
                    selected_ods,
                    sample["prev_selected_indicator"],
                    full_num_od=int(od_data.get("full_num_od", scores.size(0))),
                    k_crit=K_CRIT,
                )

                if PER_TOPO_REWARD_NORM and feasible:
                    _scale = topo_mlu_scale.get(str(sample["topology"]), 1.0)
                    mlu_for_reward = float(mlu) / float(_scale)
                else:
                    mlu_for_reward = float(mlu)

                mlu_penalty = float(mlu_for_reward)
                if ENABLE_HINGE_REWARD and mlu_penalty > float(HINGE_THRESHOLD):
                    overflow = mlu_penalty - float(HINGE_THRESHOLD)
                    mlu_penalty = float(HINGE_THRESHOLD) + float(HINGE_MULTIPLIER) * float(overflow)

                reward_mlu_weight = (
                    float(rl_cfg.w_reward_mlu_failure)
                    if is_failure_sample
                    else float(rl_cfg.w_reward_mlu_normal)
                )
                reward = (
                    -reward_mlu_weight * mlu_penalty
                    + float(rl_cfg.w_reward_improvement) * float(improvement)
                    - float(rl_cfg.w_reward_disturbance)
                    * float(disturbance)
                    * (1.0 + float(DISTURBANCE_CHURN_MULTIPLIER) * float(churn_ratio))
                    - float(rl_cfg.w_reward_infeasible) * (0.0 if feasible else 1.0)
                    + float(rl_cfg.w_reward_vs_bottleneck) * float(vs_bn)
                    + float(rl_cfg.w_reward_bottleneck_margin) * float(max(vs_bn, 0.0))
                    + float(rl_cfg.w_reward_vs_reference) * float(vs_ref)
                )
                reward_window = failure_reward_window if is_failure_sample else normal_reward_window
                regime_baseline = float(np.mean(reward_window)) if len(reward_window) > 0 else float(reward)
                advantage = float(reward - regime_baseline)
                reward_window.append(float(reward))

                oracle_mask, _, _, _, oracle_local_indices, _ = project_targets_to_score_space(sample, od_data, scores.size(0))
                local_active_mask = project_active_mask_to_scores(
                    selection_active_mask(sample),
                    score_len=scores.size(0),
                    candidate_od_indices=candidate_np,
                )
                with torch.no_grad():
                    ref_scores, _, _ = reference_policy(graph_data, od_data)
                student_active_scores = _local_active_scores(
                    scores,
                    active_mask=selection_active_mask(sample),
                    candidate_od_indices=candidate_np,
                )
                ref_active_scores = _local_active_scores(
                    ref_scores,
                    active_mask=selection_active_mask(sample),
                    candidate_od_indices=candidate_np,
                )
                if student_active_scores.numel() > 0 and ref_active_scores.numel() == student_active_scores.numel():
                    student_log_probs = torch.log_softmax(student_active_scores, dim=0)
                    ref_probs = torch.softmax(ref_active_scores.detach(), dim=0)
                    kl_loss = F.kl_div(student_log_probs, ref_probs, reduction="batchmean")
                    entropy = -(torch.exp(student_log_probs) * student_log_probs).sum()
                else:
                    kl_loss = torch.tensor(0.0, device=DEVICE)
                    entropy = torch.tensor(0.0, device=DEVICE)

                loss = -selected_log_prob * torch.tensor(advantage, dtype=torch.float32, device=DEVICE)
                loss = loss + float(rl_cfg.rank_loss_weight) * _ranking_loss(scores, oracle_mask, margin=0.05)
                loss = loss + 0.15 * _plackett_luce_topk_loss(scores, oracle_local_indices)
                loss = loss + float(rl_cfg.score_margin_weight) * _score_margin_regularizer(
                    scores,
                    active_mask=local_active_mask,
                    k=K_CRIT,
                )
                loss = loss + 0.05 * _bottleneck_aux_loss(info.get("_bottleneck_aux_pred"), od_data["bottleneck_scores"])
                sample_kl_beta = float(beta_failure) if is_failure_sample else float(kl_beta_normal)
                loss = loss + sample_kl_beta * kl_loss
                loss = loss - float(rl_cfg.entropy_weight) * entropy

                optimizer.zero_grad()
                loss.backward()
                clip_gate_attention_gradients(model, epoch)
                optimizer.step()

                epoch_rewards.append(float(reward))
                epoch_mlus.append(float(mlu))
                epoch_disturbances.append(float(disturbance))
                epoch_improvements.append(float(improvement))
                epoch_vs_bn.append(float(vs_bn))
                epoch_vs_ref.append(float(vs_ref))
                epoch_entropies.append(float(entropy.item()))
                epoch_kls.append(float(kl_loss.item()))
                overlap_ratio = selected_set_overlap_ratio(selected_ods, argmax_selected_ods)
                if is_failure_sample:
                    failure_type = canonical_failure_type(sample)
                    epoch_failure_kls.append(float(kl_loss.item()))
                    epoch_failure_advantages.append(float(advantage))
                    epoch_failure_set_overlaps.append(float(overlap_ratio))
                    epoch_failure_baselines.append(float(regime_baseline))
                    epoch_failure_type_rewards[failure_type].append(float(reward))
                    epoch_failure_type_advantages[failure_type].append(float(advantage))
                    epoch_failure_type_counts[failure_type] += 1
                else:
                    epoch_normal_kls.append(float(kl_loss.item()))
                    epoch_normal_advantages.append(float(advantage))
                    epoch_normal_set_overlaps.append(float(overlap_ratio))
                epoch_churn.append(float(churn_ratio))

            model.eval()
            val_mlus = []
            val_normal_mlus = []
            val_failure_mlus = []
            val_disturbances = []
            with torch.no_grad():
                for sample in epoch_val_samples:
                    is_failure_sample = sample_has_active_failure(sample)
                    graph_data, od_data, _ = build_experiment_inputs(
                        sample,
                        device=DEVICE,
                        candidate_limit=TRAIN_CANDIDATE_POOL_SIZE,
                    )
                    scores, _, _ = model(graph_data, od_data)
                    candidate_np = candidate_indices_from_od_data(od_data)
                    selected_ods, _, _ = continuity_select(
                        scores,
                        active_mask=selection_active_mask(sample),
                        k=K_CRIT,
                        prev_selected_indicator=sample["prev_selected_indicator"],
                        continuity_bonus=0.0,
                        candidate_od_indices=candidate_np,
                        full_num_od=int(od_data.get("full_num_od", scores.size(0))),
                    )
                    if not selected_ods:
                        continue
                    try:
                        lp = solve_selected_path_lp(
                            tm_vector=sample["tm_vector"],
                            selected_ods=selected_ods,
                            base_splits=ecmp_splits(sample["path_library"]),
                            path_library=sample["path_library"],
                            capacities=sample["capacities"],
                            warm_start_splits=sample["prev_splits"],
                            time_limit_sec=10,
                        )
                        routing = apply_routing(sample["tm_vector"], lp.splits, sample["path_library"], sample["capacities"])
                        val_mlu = float(routing.mlu)
                        val_mlus.append(val_mlu)
                        if is_failure_sample:
                            val_failure_mlus.append(val_mlu)
                        else:
                            val_normal_mlus.append(val_mlu)
                        val_disturbances.append(
                            float(compute_disturbance(sample["prev_splits"], lp.splits, sample["tm_vector"]))
                        )
                    except Exception:
                        bad_mlu = float(rl_cfg.infeasible_mlu_penalty)
                        val_mlus.append(bad_mlu)
                        if is_failure_sample:
                            val_failure_mlus.append(bad_mlu)
                        else:
                            val_normal_mlus.append(bad_mlu)
                        val_disturbances.append(1.0)

            mean_reward = float(np.mean(epoch_rewards)) if epoch_rewards else 0.0
            mean_mlu = float(np.mean(epoch_mlus)) if epoch_mlus else float("inf")
            mean_val_mlu = float(np.mean(val_mlus)) if val_mlus else float("inf")
            mean_val_normal_mlu = float(np.mean(val_normal_mlus)) if val_normal_mlus else mean_val_mlu
            mean_val_failure_mlu = float(np.mean(val_failure_mlus)) if val_failure_mlus else mean_val_mlu
            mean_val_disturbance = float(np.mean(val_disturbances)) if val_disturbances else 0.0
            val_objective = (
                float(rl_cfg.w_reward_mlu_normal) * mean_val_normal_mlu
                + float(rl_cfg.w_reward_mlu_failure) * mean_val_failure_mlu
                + float(rl_cfg.w_reward_disturbance) * mean_val_disturbance
            )
            val_score = -float(val_objective)
            current_selection_value = (
                float(mean_val_disturbance)
                if selection_metric == "val_disturbance"
                else float(mean_val_failure_mlu)
            )
            mean_kl = float(np.mean(epoch_kls)) if epoch_kls else 0.0
            mean_normal_kl = float(np.mean(epoch_normal_kls)) if epoch_normal_kls else 0.0
            mean_failure_kl = float(np.mean(epoch_failure_kls)) if epoch_failure_kls else 0.0
            mean_normal_advantage = float(np.mean(epoch_normal_advantages)) if epoch_normal_advantages else 0.0
            mean_failure_advantage = float(np.mean(epoch_failure_advantages)) if epoch_failure_advantages else 0.0
            std_failure_advantage = float(np.std(epoch_failure_advantages)) if epoch_failure_advantages else 0.0
            mean_normal_set_overlap = float(np.mean(epoch_normal_set_overlaps)) if epoch_normal_set_overlaps else 1.0
            mean_failure_set_overlap = float(np.mean(epoch_failure_set_overlaps)) if epoch_failure_set_overlaps else 1.0
            mean_failure_baseline = float(np.mean(epoch_failure_baselines)) if epoch_failure_baselines else 0.0
            end_failure_baseline = float(np.mean(failure_reward_window)) if len(failure_reward_window) > 0 else 0.0
            mean_entropy = float(np.mean(epoch_entropies)) if epoch_entropies else 0.0
            mean_churn = float(np.mean(epoch_churn)) if epoch_churn else 0.0
            attempt_logs.append(
                {
                    "attempt": int(attempt),
                    "epoch": epoch,
                    "train_mean_reward": mean_reward,
                    "train_mean_mlu": mean_mlu,
                    "train_mean_disturbance": float(np.mean(epoch_disturbances)) if epoch_disturbances else 0.0,
                    "train_mean_improvement_vs_ecmp": float(np.mean(epoch_improvements)) if epoch_improvements else 0.0,
                    "train_mean_vs_bottleneck": float(np.mean(epoch_vs_bn)) if epoch_vs_bn else 0.0,
                    "train_mean_vs_reference": float(np.mean(epoch_vs_ref)) if epoch_vs_ref else 0.0,
                    "entropy": mean_entropy,
                    "mean_KL_pi_theta_pi_ref": mean_kl,
                    "normal_rollout_KL_mean": mean_normal_kl,
                    "failure_rollout_KL_mean": mean_failure_kl,
                    "normal_rollout_advantage_mean": mean_normal_advantage,
                    "failure_rollout_advantage_mean": mean_failure_advantage,
                    "failure_rollout_advantage_std": std_failure_advantage,
                    "normal_rollout_set_overlap_mean": mean_normal_set_overlap,
                    "failure_rollout_set_overlap_mean": mean_failure_set_overlap,
                    "failure_baseline_mean": mean_failure_baseline,
                    "failure_baseline_end": end_failure_baseline,
                    "churn_ratio": mean_churn,
                    "kl_beta_normal": float(kl_beta_normal),
                    "kl_beta_failure": float(beta_failure),
                    "rl_val_mlu": mean_val_mlu,
                    "rl_val_normal_mlu": mean_val_normal_mlu,
                    "rl_val_failure_mlu": mean_val_failure_mlu,
                    "rl_val_disturbance": mean_val_disturbance,
                    "rl_val_objective": float(val_objective),
                    "rl_val_score": float(val_score),
                    "rl_selection_metric": selection_metric,
                    "rl_selection_value": float(current_selection_value),
                    "best_epoch_so_far": int(best_epoch if best_epoch > 0 else 1),
                    "epoch_time_sec": float(time.perf_counter() - epoch_start),
                }
            )
            for failure_type in RL_FAILURE_SCENARIOS:
                flag_below_neg030 = float(np.mean(epoch_failure_type_advantages.get(str(failure_type), [0.0]))) < -0.30
                failure_type_logs.append(
                    {
                        "attempt": int(attempt),
                        "epoch": int(epoch),
                        "failure_type": str(failure_type),
                        "count": int(epoch_failure_type_counts.get(str(failure_type), 0)),
                        "share_of_failure_rollouts": float(epoch_failure_type_counts.get(str(failure_type), 0))
                        / float(max(sum(epoch_failure_type_counts.values()), 1)),
                        "mean_reward": float(np.mean(epoch_failure_type_rewards.get(str(failure_type), [0.0]))),
                        "advantage_mean": float(np.mean(epoch_failure_type_advantages.get(str(failure_type), [0.0]))),
                        "flag_advantage_below_neg030": bool(flag_below_neg030),
                    }
                )
            flagged_types = [
                str(failure_type)
                for failure_type in RL_FAILURE_SCENARIOS
                if float(np.mean(epoch_failure_type_advantages.get(str(failure_type), [0.0]))) < -0.30
            ]
            print(
                f"[reinforce] attempt={attempt} epoch={epoch:02d} "
                f"val_score={val_score:.4f} val_normal={mean_val_normal_mlu:.4f} "
                f"val_failure={mean_val_failure_mlu:.4f} val_dist={mean_val_disturbance:.4f} "
                f"normal_kl={mean_normal_kl:.4f} failure_kl={mean_failure_kl:.4f} "
                f"failure_adv={mean_failure_advantage:.4f} failure_adv_std={std_failure_advantage:.4f} "
                f"failure_baseline={end_failure_baseline:.4f} "
                f"failure_overlap={mean_failure_set_overlap:.4f} entropy={mean_entropy:.4f} churn={mean_churn:.4f} "
                f"best_epoch={best_epoch if best_epoch > 0 else 1}",
                flush=True,
            )
            if flagged_types:
                print(
                    f"[reinforce][flag] attempt={attempt} epoch={epoch:02d} "
                    f"failure types below -0.30 advantage: {', '.join(flagged_types)}",
                    flush=True,
                )

            epoch_ckpt = RL_EPOCH_CKPT_DIR / f"attempt{attempt:02d}_epoch{epoch:02d}.pt"
            model.cfg.feature_variant = FEATURE_VARIANT
            model.cfg.learn_k_crit = False
            model.cfg.k_crit_min = K_CRIT
            model.cfg.k_crit_max = K_CRIT
            save_gnn_plus(
                model,
                epoch_ckpt,
                extra_meta={
                    "stage": "section14_rl_epoch_checkpoint",
                    "attempt": int(attempt),
                    "epoch": int(epoch),
                    "feature_variant": FEATURE_VARIANT,
                    "val_failure_mlu": float(mean_val_failure_mlu),
                    "val_normal_mlu": float(mean_val_normal_mlu),
                    "val_disturbance": float(mean_val_disturbance),
                    "val_score": float(val_score),
                    "rl_config": asdict(rl_cfg),
                    "rl_kl_beta_normal": float(kl_beta_normal),
                    "rl_kl_beta_failure": float(beta_failure),
                },
            )

            if mean_kl > float(RL_ABORT_KL):
                all_logs.extend(attempt_logs)
                pd.DataFrame(all_logs).to_csv(RL_LOG_CSV, index=False)
                pd.DataFrame(failure_type_logs).to_csv(RL_FAILURE_TYPE_LOG_CSV, index=False)
                raise RuntimeError(
                    f"RL aborted: mean KL {mean_kl:.4f} exceeded threshold {float(RL_ABORT_KL):.4f} "
                    f"at attempt {attempt}, epoch {epoch}."
                )

            if float(mean_val_normal_mlu) > float(RL_NORMAL_MLU_ABORT_THRESHOLD):
                all_logs.extend(attempt_logs)
                pd.DataFrame(all_logs).to_csv(RL_LOG_CSV, index=False)
                pd.DataFrame(failure_type_logs).to_csv(RL_FAILURE_TYPE_LOG_CSV, index=False)
                raise RuntimeError(
                    f"RL aborted: val_normal_mlu {mean_val_normal_mlu:.6f} exceeded "
                    f"the allowed threshold {float(RL_NORMAL_MLU_ABORT_THRESHOLD):.6f} at epoch {epoch}."
                )

            if float(current_selection_value) + 1e-9 < (
                float(best_val_disturbance) if selection_metric == "val_disturbance" else float(best_val_failure_mlu)
            ):
                best_val_score = float(val_score)
                best_val_mlu = mean_val_mlu
                best_val_normal_mlu = mean_val_normal_mlu
                best_val_failure_mlu = mean_val_failure_mlu
                best_val_disturbance = mean_val_disturbance
                best_epoch = epoch
                best_epoch_ckpt_path = epoch_ckpt
                stale = 0
                model.cfg.feature_variant = FEATURE_VARIANT
                model.cfg.learn_k_crit = False
                model.cfg.k_crit_min = K_CRIT
                model.cfg.k_crit_max = K_CRIT
                save_gnn_plus(
                    model,
                    FINAL_CKPT,
                    extra_meta={
                        "stage": "section3_4_5_7_reinforce_finetune",
                        "base_checkpoint": str(BASE_GNNPLUS_CKPT.relative_to(PROJECT_ROOT)),
                        "supervised_checkpoint": str(SUP_CKPT.relative_to(PROJECT_ROOT)),
                        "feature_variant": FEATURE_VARIANT,
                        "best_epoch": best_epoch,
                        "best_val_mlu": best_val_mlu,
                        "best_val_normal_mlu": best_val_normal_mlu,
                        "best_val_failure_mlu": best_val_failure_mlu,
                        "best_val_disturbance": best_val_disturbance,
                        "best_val_objective": float(val_objective),
                        "selection_metric": selection_metric,
                        "selection_value": float(current_selection_value),
                        "rl_config": asdict(rl_cfg),
                        "continuity_bonus": 0.0,
                        "rl_kl_beta_start": float(beta_start),
                        "rl_kl_beta_end": float(beta_end),
                        "rl_kl_beta_failure": float(beta_failure),
                    },
                )
            else:
                stale += 1

            if mean_failure_kl < float(rl_cfg.failure_kl_target):
                failure_kl_low_streak += 1
            else:
                failure_kl_low_streak = 0

            if mean_normal_kl > float(rl_cfg.normal_kl_abort):
                normal_kl_high_streak += 1
            else:
                normal_kl_high_streak = 0

            if epoch >= int(RL_GUARD_EPOCH):
                if normal_kl_high_streak >= int(rl_cfg.normal_kl_abort_patience):
                    all_logs.extend(attempt_logs)
                    pd.DataFrame(all_logs).to_csv(RL_LOG_CSV, index=False)
                    pd.DataFrame(failure_type_logs).to_csv(RL_FAILURE_TYPE_LOG_CSV, index=False)
                    raise RuntimeError(
                        f"RL aborted: normal-rollout KL exceeded {float(rl_cfg.normal_kl_abort):.4f} "
                        f"for {normal_kl_high_streak} consecutive epochs through epoch {epoch}."
                    )
                if (
                    epoch >= int(RL_FAILURE_KL_CHECK_EPOCH)
                    and failure_kl_low_streak >= int(RL_FAILURE_KL_CHECK_EPOCH)
                ):
                    if not failure_beta_adjusted_this_attempt:
                        old_beta_failure = float(beta_failure)
                        beta_failure = max(float(beta_failure) * 0.5, float(RL_FAILURE_KL_MIN_BETA))
                        failure_beta_adjustments += 1
                        failure_beta_adjusted_this_attempt = True
                        failure_beta_adjustment_reason = (
                            f"Failure KL stayed below {float(rl_cfg.failure_kl_target):.4f} "
                            f"for {failure_kl_low_streak} epochs through epoch {epoch}; "
                            f"halving beta_failure from {old_beta_failure:.6f} to {beta_failure:.6f}."
                        )
                        failure_kl_low_streak = 0
                        print(f"[reinforce] {failure_beta_adjustment_reason}", flush=True)
                    elif beta_failure <= float(RL_FAILURE_KL_MIN_BETA) + 1e-12:
                        all_logs.extend(attempt_logs)
                        pd.DataFrame(all_logs).to_csv(RL_LOG_CSV, index=False)
                        pd.DataFrame(failure_type_logs).to_csv(RL_FAILURE_TYPE_LOG_CSV, index=False)
                        raise RuntimeError(
                            f"RL aborted: failure-rollout KL stayed below {float(rl_cfg.failure_kl_target):.4f} "
                            f"through epoch {epoch} even after beta_failure was reduced to {beta_failure:.6f}."
                        )
                if best_epoch <= 1 and epoch >= int(RL_GUARD_EPOCH):
                    all_logs.extend(attempt_logs)
                    pd.DataFrame(all_logs).to_csv(RL_LOG_CSV, index=False)
                    pd.DataFrame(failure_type_logs).to_csv(RL_FAILURE_TYPE_LOG_CSV, index=False)
                    raise RuntimeError(
                        f"RL aborted: best_epoch_so_far is still {best_epoch} at epoch {epoch}, "
                        "matching the archfix collapse pattern."
                    )

            if stale >= rl_cfg.patience:
                print(f"[reinforce] early stop at epoch {epoch}", flush=True)
                break

        all_logs.extend(attempt_logs)
        if restart_this_attempt:
            continue

        pd.DataFrame(all_logs).to_csv(RL_LOG_CSV, index=False)
        pd.DataFrame(failure_type_logs).to_csv(RL_FAILURE_TYPE_LOG_CSV, index=False)
        model, _ = load_gnn_plus(FINAL_CKPT, device=DEVICE)
        if REQUIRE_RL_BEST_EPOCH_GT2 and best_epoch <= 2:
            raise RuntimeError(
                f"RL aborted: overall best_epoch={best_epoch} did not exceed 2, so the policy never "
                "improved beyond the collapse-risk region."
            )
        return model, {
            "mode": "reinforce",
            "best_epoch": int(best_epoch),
            "best_epoch_checkpoint": None if best_epoch_ckpt_path is None else str(best_epoch_ckpt_path.relative_to(PROJECT_ROOT)),
            "selection_metric": selection_metric,
            "best_val_mlu": float(best_val_mlu),
            "best_val_normal_mlu": float(best_val_normal_mlu),
            "best_val_failure_mlu": float(best_val_failure_mlu),
            "best_val_disturbance": float(best_val_disturbance),
            "best_val_score": float(best_val_score),
            "continuity_bonus": 0.0,
            "rl_config": asdict(rl_cfg),
            "hinge_reward_enabled": bool(ENABLE_HINGE_REWARD),
            "hinge_threshold": float(HINGE_THRESHOLD),
            "hinge_multiplier": float(HINGE_MULTIPLIER),
            "num_layers": int(model.cfg.num_layers),
            "num_train_samples": int(len(rl_train)),
            "num_val_samples": int(len(rl_val)),
            "candidate_prefilter": {
                "enabled": bool(USE_CANDIDATE_PREFILTER),
                "multiplier": float(CANDIDATE_PREFILTER_MULTIPLIER),
                "train_candidate_pool_size": int(TRAIN_CANDIDATE_POOL_SIZE),
                "infer_candidate_pool_size": int(INFER_CANDIDATE_POOL_SIZE),
            },
            "architecture": {
                "use_gated_residual": bool(model.cfg.use_gated_residual),
                "use_cross_od_attention": bool(model.cfg.use_cross_od_attention),
                "cross_od_attention_heads": int(model.cfg.cross_od_attention_heads),
                "bottleneck_moe_floor": float(model.cfg.bottleneck_moe_floor),
            },
            "step5_patch": {
                "continuity_bonus_old": 0.03,
                "continuity_bonus_new": 0.0,
                "w_reward_mlu_old": 1.15,
                "w_reward_mlu_new": float(REWARD_MLU),
                "w_reward_mlu_normal": float(REWARD_MLU_NORMAL),
                "w_reward_mlu_failure": float(REWARD_MLU_FAILURE),
                "w_reward_disturbance_old": 0.15,
                "w_reward_disturbance_new": float(REWARD_DISTURBANCE),
                "rl_failure_rollout_prob": float(RL_FAILURE_ROLLOUT_PROB),
                "per_topology_reward_norm_enabled": bool(PER_TOPO_REWARD_NORM),
                "per_topology_mlu_scale_trainonly": {str(k): float(v) for k, v in topo_mlu_scale.items()},
                "normalization_source_topologies": list(KNOWN_TOPOLOGIES),
                "normalization_excluded_topologies": list(UNSEEN_TOPOLOGIES),
                "rl_kl_beta_start": float(beta_start),
                "rl_kl_beta_end": float(beta_end),
                "rl_kl_beta_failure_final": float(beta_failure),
                "rl_failure_kl_target": float(rl_cfg.failure_kl_target),
                "rl_normal_kl_abort": float(rl_cfg.normal_kl_abort),
                "disturbance_churn_multiplier": float(DISTURBANCE_CHURN_MULTIPLIER),
                "failure_beta_adjustments": int(failure_beta_adjustments),
                "failure_beta_adjustment_reason": failure_beta_adjustment_reason,
            },
        }

    pd.DataFrame(all_logs).to_csv(RL_LOG_CSV, index=False)
    raise RuntimeError("RL failed to complete within the configured restart budget.")


def gnnplus_select_stateful(
    model,
    *,
    dataset,
    path_library,
    capacities: np.ndarray,
    tm_vector: np.ndarray,
    telemetry,
    prev_tm,
    prev_util,
    prev_selected_indicator,
    prev_disturbance: float,
    k_crit: int,
    failure_mask=None,
    gate_temperature: float = 1.0,
    tie_break_eps: float = 0.0,
):
    has_active_failure = bool(
        failure_mask is not None and np.any(np.asarray(failure_mask, dtype=np.float64) > 0.5)
    )
    candidate_od_indices = compute_candidate_pool(
        tm_vector=np.asarray(tm_vector, dtype=np.float64),
        path_library=path_library,
        capacities=np.asarray(capacities, dtype=float),
        k_crit=int(k_crit),
        candidate_limit=INFER_CANDIDATE_POOL_SIZE,
    )
    graph_data = build_graph_tensors_plus(
        dataset,
        tm_vector=tm_vector,
        path_library=path_library,
        telemetry=telemetry,
        prev_util=prev_util,
        prev_tm=prev_tm,
        prev_selected_indicator=prev_selected_indicator,
        prev_disturbance=prev_disturbance,
        failure_mask=failure_mask,
        feature_variant=FEATURE_VARIANT,
        device=DEVICE,
    )
    graph_data["gate_temperature"] = float(gate_temperature)
    od_data = build_od_features_plus(
        dataset,
        tm_vector,
        path_library,
        telemetry=telemetry,
        prev_tm=prev_tm,
        prev_util=prev_util,
        prev_selected_indicator=prev_selected_indicator,
        prev_disturbance=prev_disturbance,
        failure_mask=failure_mask,
        candidate_od_indices=candidate_od_indices,
        feature_variant=FEATURE_VARIANT,
        device=DEVICE,
    )
    active_mask = ((np.asarray(tm_vector, dtype=np.float64) > 1e-12) & surviving_od_mask(path_library)).astype(np.float32)
    with torch.no_grad():
        scores, _, info = model(graph_data, od_data)
    candidate_np = candidate_indices_from_od_data(od_data)
    effective_tie_eps = 0.0 if has_active_failure else float(tie_break_eps)
    selected, _, _ = continuity_select(
        scores,
        active_mask=active_mask,
        k=k_crit,
        prev_selected_indicator=prev_selected_indicator,
        continuity_bonus=CONTINUITY_BONUS,
        tie_break_eps=effective_tie_eps,
        candidate_od_indices=candidate_np,
        full_num_od=int(od_data.get("full_num_od", scores.size(0))),
    )
    info = {k: v for k, v in info.items() if not str(k).startswith("_")}
    info["gate_temperature"] = float(gate_temperature)
    info["tie_break_eps"] = float(effective_tie_eps)
    info["candidate_pool_size"] = int(scores.shape[0])
    if candidate_np is not None and candidate_np.size > 0:
        info["prefilter_bottleneck_selected"] = [
            int(od)
            for od in candidate_np[: min(int(k_crit), int(candidate_np.size))].tolist()
        ]
    else:
        info["prefilter_bottleneck_selected"] = []
    assert_selected_ods_have_paths(path_library, selected, context=f"{dataset.key}:gnnplus_improved")
    return selected, info


def resolve_inference_controls(calibration: dict | None, topology_key: str) -> dict[str, float]:
    calibration = calibration or {}
    per_topology = calibration.get("per_topology", {})
    topo_key_norm = str(topology_key).lower()
    topo_cfg = per_topology.get(topo_key_norm, per_topology.get(str(topology_key), {}))
    use_aggressive = topo_key_norm in AGGRESSIVE_TIEBREAK_TOPOLOGIES
    if use_aggressive:
        tie_break_eps = float(
            topo_cfg.get(
                "tie_break_epsilon_aggressive",
                calibration.get(
                    "global_tie_break_epsilon_aggressive",
                    topo_cfg.get("tie_break_epsilon", calibration.get("global_tie_break_epsilon", 0.0)),
                ),
            )
        )
        tie_break_mode = "p25_aggressive"
    else:
        tie_break_eps = float(topo_cfg.get("tie_break_epsilon", calibration.get("global_tie_break_epsilon", 0.0)))
        tie_break_mode = "p5_default"
    return {
        "gate_temperature": float(topo_cfg.get("gate_temperature", calibration.get("global_gate_temperature", 1.0))),
        "tie_break_eps": tie_break_eps,
        "tie_break_mode": tie_break_mode,
    }


def calibrate_inference_controls(model: GNNPlusFlowSelector, val_samples: list[dict]) -> dict:
    grouped: dict[str, list[dict]] = defaultdict(list)
    for sample in val_samples:
        grouped[str(sample["topology"])].append(sample)

    calibration = {
        "per_topology": {},
        "global_gate_temperature": 1.0,
        "global_tie_break_epsilon": 0.0,
        "global_tie_break_epsilon_aggressive": 0.0,
        "aggressive_tiebreak_topologies": sorted(AGGRESSIVE_TIEBREAK_TOPOLOGIES),
        "tie_break_base_percentile": 5,
        "tie_break_aggressive_percentile": 25,
    }
    all_temps = []
    all_eps = []
    all_eps_aggressive = []
    was_training = model.training
    model.eval()
    with torch.no_grad():
        for topo_key in sorted(grouped):
            topo_samples = grouped[topo_key][: max(1, CALIBRATION_MAX_SAMPLES_PER_TOPO)]
            best_temp = 1.0
            best_mean_mlu = float("inf")
            for temp in GATE_TEMPERATURE_GRID or [1.0]:
                mlus = []
                for sample in topo_samples:
                    graph_data, od_data, _ = build_experiment_inputs(
                        sample,
                        device=DEVICE,
                        candidate_limit=INFER_CANDIDATE_POOL_SIZE,
                    )
                    graph_data["gate_temperature"] = float(temp)
                    scores, _, _ = model(graph_data, od_data)
                    candidate_np = candidate_indices_from_od_data(od_data)
                    selected, _, _ = continuity_select(
                        scores,
                        active_mask=selection_active_mask(sample),
                        k=K_CRIT,
                        prev_selected_indicator=sample["prev_selected_indicator"],
                        continuity_bonus=CONTINUITY_BONUS,
                        candidate_od_indices=candidate_np,
                        full_num_od=int(od_data.get("full_num_od", scores.size(0))),
                    )
                    if not selected:
                        continue
                    try:
                        lp = solve_selected_path_lp(
                            tm_vector=sample["tm_vector"],
                            selected_ods=selected,
                            base_splits=ecmp_splits(sample["path_library"]),
                            path_library=sample["path_library"],
                            capacities=sample["capacities"],
                            warm_start_splits=sample["prev_splits"],
                            time_limit_sec=LP_TIME_LIMIT,
                        )
                        routing = apply_routing(sample["tm_vector"], lp.splits, sample["path_library"], sample["capacities"])
                        mlus.append(float(routing.mlu))
                    except Exception:
                        mlus.append(float("inf"))
                mean_mlu = float(np.mean(mlus)) if mlus else float("inf")
                if mean_mlu + 1e-9 < best_mean_mlu:
                    best_mean_mlu = mean_mlu
                    best_temp = float(temp)

            tie_eps_values = []
            tie_eps_values_aggressive = []
            for sample in topo_samples:
                graph_data, od_data, _ = build_experiment_inputs(
                    sample,
                    device=DEVICE,
                    candidate_limit=INFER_CANDIDATE_POOL_SIZE,
                )
                graph_data["gate_temperature"] = float(best_temp)
                scores, _, _ = model(graph_data, od_data)
                tie_eps_values.append(
                    compute_tie_break_epsilon(
                        scores,
                        active_mask=selection_active_mask(sample),
                        candidate_od_indices=candidate_indices_from_od_data(od_data),
                        percentile=5.0,
                    )
                )
                tie_eps_values_aggressive.append(
                    compute_tie_break_epsilon(
                        scores,
                        active_mask=selection_active_mask(sample),
                        candidate_od_indices=candidate_indices_from_od_data(od_data),
                        percentile=25.0,
                    )
                )
            tie_eps = float(np.median(tie_eps_values)) if tie_eps_values else 0.0
            tie_eps_aggressive = float(np.median(tie_eps_values_aggressive)) if tie_eps_values_aggressive else tie_eps
            calibration["per_topology"][str(topo_key)] = {
                "gate_temperature": float(best_temp),
                "tie_break_epsilon": float(tie_eps),
                "tie_break_epsilon_aggressive": float(tie_eps_aggressive),
                "tie_break_epsilon_p5": float(tie_eps),
                "tie_break_epsilon_p25": float(tie_eps_aggressive),
                "calibration_mean_mlu": float(best_mean_mlu),
                "num_samples": int(len(topo_samples)),
            }
            all_temps.append(float(best_temp))
            all_eps.append(float(tie_eps))
            all_eps_aggressive.append(float(tie_eps_aggressive))

    if was_training:
        model.train()
    if all_temps:
        calibration["global_gate_temperature"] = float(np.median(all_temps))
    if all_eps:
        calibration["global_tie_break_epsilon"] = float(np.median(all_eps))
    if all_eps_aggressive:
        calibration["global_tie_break_epsilon_aggressive"] = float(np.median(all_eps_aggressive))
    return calibration


_STICKY_OVERRIDE_LOGGED = False
_DISTURB_TIEBREAK_LOGGED = False
_FAILURE_FALLBACK_LOGGED = False


def _sticky_compose_selection(
    *,
    selected_ods: list[int],
    prev_selected_ods: list[int],
    path_library,
    tm_vector: np.ndarray,
    k_crit: int,
) -> list[int]:
    """Build sticky selection: keep prev-selected ODs that are still active,
    top up from the fresh GNN+ selection (preserving order) to reach k_crit.

    Returns an empty list if the sticky choice would be identical to the
    fresh selection (caller should skip the LP solve in that case).
    """
    k = max(int(k_crit), 0)
    if k <= 0 or not prev_selected_ods:
        return []
    tm_arr = np.asarray(tm_vector, dtype=np.float64)
    active_mask = (tm_arr > 1e-12) & surviving_od_mask(path_library)
    active_set = set(int(od) for od in np.where(active_mask)[0].tolist())
    fresh_order = [int(od) for od in selected_ods]
    fresh_set = set(fresh_order)
    sticky_prev = [int(od) for od in prev_selected_ods if int(od) in active_set]
    # Stable unique-preserving order on prev.
    seen = set()
    sticky_prev_unique: list[int] = []
    for od in sticky_prev:
        if od not in seen:
            sticky_prev_unique.append(od)
            seen.add(od)
    sticky_prev_unique = sticky_prev_unique[:k]
    if len(sticky_prev_unique) >= k:
        sticky = sticky_prev_unique
    else:
        remaining = k - len(sticky_prev_unique)
        topup = [od for od in fresh_order if od not in seen][:remaining]
        sticky = sticky_prev_unique + topup
    if set(sticky) == fresh_set:
        return []
    return sticky


def apply_do_no_harm_gate(
    runner,
    *,
    tm_vector: np.ndarray,
    selected_ods: list[int],
    base_splits,
    warm_start_splits,
    path_library,
    capacities,
    k_crit: int,
    context: str,
    topology_key: str,
    guard_bottleneck_selected: list[int] | None = None,
    guard_cache: dict | None = None,
    step_index: int | None = None,
    guard_fallback_cooldown: int = 0,
    prev_selected_ods: list[int] | None = None,
) -> tuple[list[int], object, dict, dict, int]:
    cache = dict(guard_cache or {})
    cooldown_remaining = max(int(guard_fallback_cooldown), 0)
    sticky_applied = False
    cached_selected = [
        int(od)
        for od in cache.get("selected_ods", [])[: max(int(k_crit), 0)]
    ]
    if cooldown_remaining > 0 and cached_selected:
        hold_lp = runner.solve_selected_path_lp_safe(
            tm_vector=tm_vector,
            selected_ods=cached_selected,
            base_splits=base_splits,
            path_library=path_library,
            capacities=capacities,
            warm_start_splits=warm_start_splits,
            time_limit_sec=LP_TIME_LIMIT,
            context=f"{context}:bottleneck_cooldown_hold",
        )
        cached_ref_mlu = cache.get("reference_mlu")
        return cached_selected, hold_lp, {
            "do_no_harm_fallback": False,
            "do_no_harm_cooldown_hold": True,
            "gnn_candidate_mlu": None,
            "bottleneck_candidate_mlu": float(cached_ref_mlu) if cached_ref_mlu is not None else float(hold_lp.routing.mlu),
            "do_no_harm_threshold": float(do_no_harm_threshold_for_topology(topology_key)),
            "guard_overlap_ratio": None,
            "guard_reference_source": "cooldown_hold",
            "guard_reference_refreshed": False,
            "sticky_applied": False,
            "guard_fallback_cooldown_remaining": max(cooldown_remaining - 1, 0),
        }, cache, max(cooldown_remaining - 1, 0)

    threshold = do_no_harm_threshold_for_topology(topology_key)
    gnn_lp = runner.solve_selected_path_lp_safe(
        tm_vector=tm_vector,
        selected_ods=selected_ods,
        base_splits=base_splits,
        path_library=path_library,
        capacities=capacities,
        warm_start_splits=warm_start_splits,
        time_limit_sec=LP_TIME_LIMIT,
        context=f"{context}:gnnplus_candidate",
    )
    gnn_est_mlu = float(gnn_lp.routing.mlu)

    # [Phase 1 / rescue] Sticky-selection post-filter: if a prev-selected
    # set is available and STICKY_EPS > 0, see if keeping previously-
    # selected ODs (filled up with fresh GNN+ picks) costs negligible MLU.
    # If so, prefer the sticky selection (lower disturbance) as GNN+'s
    # candidate for the do-no-harm comparison below.
    if (
        STICKY_EPS > 0.0
        and prev_selected_ods
        and len(selected_ods) > 0
    ):
        sticky_ods = _sticky_compose_selection(
            selected_ods=selected_ods,
            prev_selected_ods=list(prev_selected_ods),
            path_library=path_library,
            tm_vector=np.asarray(tm_vector, dtype=float),
            k_crit=int(k_crit),
        )
        if sticky_ods:
            try:
                sticky_lp = runner.solve_selected_path_lp_safe(
                    tm_vector=tm_vector,
                    selected_ods=sticky_ods,
                    base_splits=base_splits,
                    path_library=path_library,
                    capacities=capacities,
                    warm_start_splits=warm_start_splits,
                    time_limit_sec=LP_TIME_LIMIT,
                    context=f"{context}:gnnplus_sticky_candidate",
                )
                sticky_mlu = float(sticky_lp.routing.mlu)
                # Sticky wins if it is no worse than (1 + STICKY_EPS) times
                # the fresh MLU; this is a pure Pareto-preference for lower
                # disturbance at a bounded MLU cost.
                if sticky_mlu <= gnn_est_mlu * (1.0 + float(STICKY_EPS)) + 1e-12:
                    global _STICKY_OVERRIDE_LOGGED
                    if not _STICKY_OVERRIDE_LOGGED:
                        print(
                            f"[sticky] Applied sticky post-filter: "
                            f"fresh_mlu={gnn_est_mlu:.6f} sticky_mlu={sticky_mlu:.6f} "
                            f"eps={float(STICKY_EPS):.4f} context={context}. "
                            f"First sticky override only is logged."
                        )
                        _STICKY_OVERRIDE_LOGGED = True
                    selected_ods = list(sticky_ods)
                    gnn_lp = sticky_lp
                    gnn_est_mlu = sticky_mlu
                    sticky_applied = True
            except Exception:
                # Sticky is a nice-to-have; never fail the cycle because of it.
                pass

    bottleneck_selected = [
        int(od)
        for od in (guard_bottleneck_selected or [])[: max(int(k_crit), 0)]
    ]
    if not bottleneck_selected:
        bottleneck_selected = select_bottleneck_critical(
            np.asarray(tm_vector, dtype=float),
            base_splits,
            path_library,
            np.asarray(capacities, dtype=float),
            int(k_crit),
        )
    if list(map(int, bottleneck_selected)) == list(map(int, selected_ods)):
        return selected_ods, gnn_lp, {
            "do_no_harm_fallback": False,
            "do_no_harm_cooldown_hold": False,
            "gnn_candidate_mlu": float(gnn_est_mlu),
            "bottleneck_candidate_mlu": float(gnn_est_mlu),
            "do_no_harm_threshold": float(threshold),
            "guard_overlap_ratio": 1.0,
            "guard_reference_source": "same_selection",
            "guard_reference_refreshed": False,
            "sticky_applied": bool(sticky_applied),
            "guard_fallback_cooldown_remaining": 0,
        }, {
            "reference_mlu": float(gnn_est_mlu),
            "selected_ods": list(map(int, bottleneck_selected)),
            "last_refresh_step": int(step_index) if step_index is not None else None,
        }, 0

    overlap_ratio = selection_overlap_ratio(selected_ods, bottleneck_selected)
    if overlap_ratio >= float(DO_NO_HARM_OVERLAP_SKIP):
        return selected_ods, gnn_lp, {
            "do_no_harm_fallback": False,
            "do_no_harm_cooldown_hold": False,
            "gnn_candidate_mlu": float(gnn_est_mlu),
            "bottleneck_candidate_mlu": None,
            "do_no_harm_threshold": float(threshold),
            "guard_overlap_ratio": float(overlap_ratio),
            "guard_reference_source": "skipped_high_overlap",
            "guard_reference_refreshed": False,
            "sticky_applied": bool(sticky_applied),
            "guard_fallback_cooldown_remaining": 0,
        }, cache, 0

    refresh_due = (
        step_index is None
        or cache.get("reference_mlu") is None
        or cache.get("last_refresh_step") is None
        or (int(step_index) - int(cache.get("last_refresh_step", -10**9))) >= int(DO_NO_HARM_CACHE_STEPS)
    )
    if not refresh_due:
        cached_ref_mlu = float(cache["reference_mlu"])
        return selected_ods, gnn_lp, {
            "do_no_harm_fallback": False,
            "do_no_harm_cooldown_hold": False,
            "gnn_candidate_mlu": float(gnn_est_mlu),
            "bottleneck_candidate_mlu": float(cached_ref_mlu),
            "do_no_harm_threshold": float(threshold),
            "guard_overlap_ratio": float(overlap_ratio),
            "guard_reference_source": "cached_periodic_reference",
            "guard_reference_refreshed": False,
            "sticky_applied": bool(sticky_applied),
            "guard_fallback_cooldown_remaining": 0,
        }, cache, 0

    bn_lp = runner.solve_selected_path_lp_safe(
        tm_vector=tm_vector,
        selected_ods=bottleneck_selected,
        base_splits=base_splits,
        path_library=path_library,
        capacities=capacities,
        warm_start_splits=warm_start_splits,
        time_limit_sec=LP_TIME_LIMIT,
        context=f"{context}:bottleneck_guard_refresh",
    )
    bn_est_mlu = float(bn_lp.routing.mlu)
    updated_cache = {
        "reference_mlu": float(bn_est_mlu),
        "selected_ods": list(map(int, bottleneck_selected)),
        "last_refresh_step": int(step_index) if step_index is not None else None,
    }
    # [Phase 1 / rescue] Disturbance-aware do-no-harm tiebreak: if the two
    # candidates are within DISTURB_TIEBREAK_EPS on MLU, prefer the one
    # that produces fewer flow changes against the currently-applied
    # splits (warm_start_splits). Tiebreak to GNN+ on exact equality to
    # preserve existing behavior.
    if (
        DISTURB_TIEBREAK_EPS > 0.0
        and bn_est_mlu > 1e-12
        and warm_start_splits is not None
    ):
        rel_gap = abs(gnn_est_mlu - bn_est_mlu) / float(bn_est_mlu)
        if rel_gap <= float(DISTURB_TIEBREAK_EPS):
            try:
                gnn_dist = float(compute_disturbance(
                    warm_start_splits, gnn_lp.splits, tm_vector
                ))
                bn_dist = float(compute_disturbance(
                    warm_start_splits, bn_lp.splits, tm_vector
                ))
            except Exception:
                gnn_dist = float("inf")
                bn_dist = float("inf")
            if bn_dist + 1e-12 < gnn_dist:
                global _DISTURB_TIEBREAK_LOGGED
                if not _DISTURB_TIEBREAK_LOGGED:
                    print(
                        f"[disturb-tiebreak] Near-tie resolved to bottleneck: "
                        f"gnn_mlu={gnn_est_mlu:.6f} bn_mlu={bn_est_mlu:.6f} "
                        f"gnn_dist={gnn_dist:.6f} bn_dist={bn_dist:.6f} "
                        f"eps={float(DISTURB_TIEBREAK_EPS):.4f} context={context}. "
                        f"First tiebreak only is logged."
                    )
                    _DISTURB_TIEBREAK_LOGGED = True
                return list(map(int, bottleneck_selected)), bn_lp, {
                    "do_no_harm_fallback": True,
                    "do_no_harm_cooldown_hold": False,
                    "gnn_candidate_mlu": float(gnn_est_mlu),
                    "bottleneck_candidate_mlu": float(bn_est_mlu),
                    "do_no_harm_threshold": float(threshold),
                    "guard_overlap_ratio": float(overlap_ratio),
                    "guard_reference_source": "disturb_tiebreak_bottleneck",
                    "guard_reference_refreshed": True,
                    "sticky_applied": bool(sticky_applied),
                    "guard_fallback_cooldown_remaining": int(DO_NO_HARM_FALLBACK_COOLDOWN),
                }, updated_cache, int(DO_NO_HARM_FALLBACK_COOLDOWN)

    if gnn_est_mlu > float(threshold) * bn_est_mlu:
        return list(map(int, bottleneck_selected)), bn_lp, {
            "do_no_harm_fallback": True,
            "do_no_harm_cooldown_hold": False,
            "gnn_candidate_mlu": float(gnn_est_mlu),
            "bottleneck_candidate_mlu": float(bn_est_mlu),
            "do_no_harm_threshold": float(threshold),
            "guard_overlap_ratio": float(overlap_ratio),
            "guard_reference_source": "current_bottleneck_refresh",
            "guard_reference_refreshed": True,
            "sticky_applied": bool(sticky_applied),
            "guard_fallback_cooldown_remaining": int(DO_NO_HARM_FALLBACK_COOLDOWN),
        }, updated_cache, int(DO_NO_HARM_FALLBACK_COOLDOWN)
    return selected_ods, gnn_lp, {
        "do_no_harm_fallback": False,
        "do_no_harm_cooldown_hold": False,
        "gnn_candidate_mlu": float(gnn_est_mlu),
        "bottleneck_candidate_mlu": float(bn_est_mlu),
        "do_no_harm_threshold": float(threshold),
        "guard_overlap_ratio": float(overlap_ratio),
        "guard_reference_source": "current_bottleneck_refresh",
        "guard_reference_refreshed": True,
        "sticky_applied": bool(sticky_applied),
        "guard_fallback_cooldown_remaining": 0,
    }, updated_cache, 0


def run_sdn_cycle_gnnplus_improved(
    runner,
    *,
    tm_vector,
    dataset,
    path_library,
    ecmp_base,
    current_splits,
    current_groups,
    topo_mapping,
    capacities,
    weights,
    gnnplus_model,
    prev_latency_by_od,
    gnnplus_state,
    inference_calibration=None,
):
    t_total_start = time.perf_counter()
    routing_pre = apply_routing(tm_vector, current_splits, path_library, capacities)
    pre_mlu = float(routing_pre.mlu)
    pre_telemetry = runner.compute_telemetry(
        tm_vector=tm_vector,
        splits=current_splits,
        path_library=path_library,
        routing=routing_pre,
        weights=weights,
        prev_latency_by_od=prev_latency_by_od,
        fast_links_only=True,
    )

    controls = resolve_inference_controls(inference_calibration, str(dataset.key))
    selected_ods, select_info = gnnplus_select_stateful(
        gnnplus_model,
        dataset=dataset,
        path_library=path_library,
        capacities=capacities,
        tm_vector=tm_vector,
        telemetry=pre_telemetry,
        prev_tm=gnnplus_state["prev_tm"],
        prev_util=gnnplus_state["prev_util"],
        prev_selected_indicator=gnnplus_state["prev_selected_indicator"],
        prev_disturbance=gnnplus_state["prev_disturbance"],
        k_crit=K_CRIT,
        gate_temperature=controls["gate_temperature"],
        tie_break_eps=controls["tie_break_eps"],
    )
    _prev_indicator = gnnplus_state.get("prev_selected_indicator")
    prev_selected_ods_list: list[int] = []
    if _prev_indicator is not None:
        _prev_arr = np.asarray(_prev_indicator, dtype=np.float32).reshape(-1)
        prev_selected_ods_list = [int(i) for i in np.where(_prev_arr > 0.5)[0].tolist()]
    selected_ods, lp_result, gate_info, updated_guard_cache, updated_guard_fallback_cooldown = apply_do_no_harm_gate(
        runner,
        tm_vector=np.asarray(tm_vector, dtype=float),
        selected_ods=selected_ods,
        base_splits=ecmp_base,
        warm_start_splits=current_splits,
        path_library=path_library,
        capacities=capacities,
        k_crit=K_CRIT,
        context=f"{dataset.key}:gnnplus_improved:normal_cycle",
        topology_key=str(dataset.key),
        guard_bottleneck_selected=select_info.get("prefilter_bottleneck_selected", []),
        guard_cache=gnnplus_state.get("guard_cache", {}),
        step_index=int(gnnplus_state.get("guard_cycle_index", 0)),
        guard_fallback_cooldown=int(gnnplus_state.get("guard_fallback_cooldown", 0)),
        prev_selected_ods=prev_selected_ods_list,
    )
    select_info.update(gate_info)
    new_splits = [s.copy() for s in lp_result.splits]
    t_decision_end = time.perf_counter()
    decision_time_ms = (t_decision_end - t_total_start) * 1000.0

    t_rule_start = time.perf_counter()
    new_groups, _ = runner.splits_to_openflow_rules(new_splits, selected_ods, path_library, topo_mapping, dataset.edges)
    changed_groups = runner.compute_rule_diff(current_groups, new_groups)
    flow_table_updates = len(changed_groups)
    t_rule_end = time.perf_counter()
    rule_install_delay_ms = (t_rule_end - t_rule_start) * 1000.0

    routing_post = apply_routing(tm_vector, new_splits, path_library, capacities)
    post_mlu = float(routing_post.mlu)
    dist = compute_disturbance(current_splits, new_splits, tm_vector)
    telemetry_post = runner.compute_telemetry(
        tm_vector=tm_vector,
        splits=new_splits,
        path_library=path_library,
        routing=routing_post,
        weights=weights,
        prev_latency_by_od=prev_latency_by_od,
    )

    result = runner.SDNCycleResult(
        cycle=0,
        method="gnnplus",
        topology=dataset.key,
        pre_mlu=pre_mlu,
        post_mlu=post_mlu,
        disturbance=float(dist),
        throughput=telemetry_post.throughput,
        mean_latency=telemetry_post.mean_latency,
        p95_latency=telemetry_post.p95_latency,
        packet_loss=telemetry_post.packet_loss,
        jitter=telemetry_post.jitter,
        decision_time_ms=decision_time_ms,
        flow_table_updates=flow_table_updates,
        rule_install_delay_ms=rule_install_delay_ms,
    )
    next_state = {
        "prev_tm": np.asarray(tm_vector, dtype=float),
        "prev_util": np.asarray(telemetry_post.utilization, dtype=float),
        "prev_selected_indicator": selection_indicator(len(dataset.od_pairs), selected_ods),
        "prev_disturbance": float(dist),
        "select_info": dict(select_info),
        "guard_cache": dict(updated_guard_cache),
        "guard_cycle_index": int(gnnplus_state.get("guard_cycle_index", 0)) + 1,
        "guard_fallback_cooldown": int(updated_guard_fallback_cooldown),
    }
    return result, new_splits, new_groups, telemetry_post.latency_by_od, next_state


def run_failure_scenario_gnnplus_improved(
    runner,
    *,
    scenario,
    tm_vector,
    dataset,
    path_library,
    ecmp_base,
    capacities,
    weights,
    gnnplus_model,
    inference_calibration=None,
):
    normal_routing = apply_routing(tm_vector, ecmp_base, path_library, capacities)
    pre_failure_mlu = float(normal_routing.mlu)
    failure_state = runner._build_failure_execution_state(
        scenario=scenario,
        tm_vector=tm_vector,
        dataset=dataset,
        path_library=path_library,
        capacities=capacities,
        weights=weights,
        normal_routing=normal_routing,
    )
    effective_tm = failure_state["effective_tm"]
    effective_caps = failure_state["effective_caps"]
    effective_path_library = failure_state["effective_path_library"]
    effective_dataset = failure_state["effective_dataset"]
    effective_ecmp = failure_state["effective_ecmp"]
    failure_mask = failure_state["failure_mask"]
    effective_weights = failure_state["effective_weights"]

    post_failure_routing = apply_routing(effective_tm, effective_ecmp, effective_path_library, effective_caps)
    pre_telemetry = runner.compute_telemetry(
        tm_vector=effective_tm,
        splits=effective_ecmp,
        path_library=effective_path_library,
        routing=post_failure_routing,
        weights=effective_weights,
        prev_latency_by_od=None,
        fast_links_only=True,
    )

    t_start = time.perf_counter()
    controls = resolve_inference_controls(inference_calibration, str(dataset.key))
    selected, select_info = gnnplus_select_stateful(
        gnnplus_model,
        dataset=effective_dataset,
        path_library=effective_path_library,
        capacities=effective_caps,
        tm_vector=effective_tm,
        telemetry=pre_telemetry,
        prev_tm=None,
        prev_util=None,
        prev_selected_indicator=np.zeros(len(effective_dataset.od_pairs), dtype=np.float32),
        prev_disturbance=0.0,
        k_crit=K_CRIT,
        failure_mask=failure_mask,
        gate_temperature=controls["gate_temperature"],
        tie_break_eps=controls["tie_break_eps"],
    )
    _, lp_result, gate_info, _, _ = apply_do_no_harm_gate(
        runner,
        tm_vector=np.asarray(effective_tm, dtype=float),
        selected_ods=selected,
        base_splits=effective_ecmp,
        warm_start_splits=effective_ecmp,
        path_library=effective_path_library,
        capacities=effective_caps,
        k_crit=K_CRIT,
        context=f"{dataset.key}:{scenario}:gnnplus_improved",
        topology_key=str(dataset.key),
        guard_bottleneck_selected=select_info.get("prefilter_bottleneck_selected", []),
        guard_cache={},
        step_index=0,
    )
    recovery_splits = [s.copy() for s in lp_result.splits]
    recovery_ms = (time.perf_counter() - t_start) * 1000.0
    post_routing = apply_routing(effective_tm, recovery_splits, effective_path_library, effective_caps)
    post_recovery_mlu = float(post_routing.mlu)
    select_info.update(gate_info)

    # [Phase 1 / failure-rescue] Failure-time do-no-harm fallback: compute
    # the bottleneck recovery on the SAME failure state and fall back if
    # GNN+'s recovery MLU is worse. Same philosophy as the normal-cycle
    # do-no-harm gate. We use the already-built `effective_*` state (so
    # the bottleneck candidate is evaluated on identical failed edges and
    # capacities as GNN+), not a fresh random failure state. Inference-only.
    select_info["failure_do_no_harm_fallback"] = False
    if FAILURE_DO_NO_HARM:
        try:
            bn_t_start = time.perf_counter()
            bn_selected = select_bottleneck_critical(
                effective_tm,
                effective_ecmp,
                effective_path_library,
                effective_caps,
                K_CRIT,
            )
            bn_lp = runner.solve_selected_path_lp_safe(
                tm_vector=effective_tm,
                selected_ods=bn_selected,
                base_splits=effective_ecmp,
                path_library=effective_path_library,
                capacities=effective_caps,
                time_limit_sec=LP_TIME_LIMIT,
                context=f"{dataset.key}:{scenario}:bottleneck_failure_fallback",
            )
            bn_recovery_ms_local = (time.perf_counter() - bn_t_start) * 1000.0
            bn_post_routing = apply_routing(
                effective_tm,
                bn_lp.splits,
                effective_path_library,
                effective_caps,
            )
            bn_post_mlu = float(bn_post_routing.mlu)
            if bn_post_mlu + 1e-12 < float(post_recovery_mlu):
                global _FAILURE_FALLBACK_LOGGED
                if not _FAILURE_FALLBACK_LOGGED:
                    print(
                        f"[failure-do-no-harm] Falling back to bottleneck recovery: "
                        f"gnn_mlu={post_recovery_mlu:.6f} bn_mlu={bn_post_mlu:.6f} "
                        f"context={dataset.key}:{scenario}. First fallback only is logged.",
                        flush=True,
                    )
                    _FAILURE_FALLBACK_LOGGED = True
                post_recovery_mlu = bn_post_mlu
                # Honest accounting: include both inference paths in the
                # reported recovery time (GNN+'s wasted compute + bottleneck).
                recovery_ms = float(recovery_ms) + float(bn_recovery_ms_local)
                select_info["failure_do_no_harm_fallback"] = True
        except Exception:
            # Fallback is best-effort; never fail the cycle if it errors.
            pass

    return recovery_ms, pre_failure_mlu, post_recovery_mlu, failure_mask, select_info


def benchmark_topology_normal_improved(runner, topo_key: str, gnn_cache: dict, gnnplus_model, inference_calibration=None) -> tuple[list[dict], list[dict]]:
    dataset, path_library = runner.load_dataset(topo_key)
    capacities = np.asarray(dataset.capacities, dtype=float)
    weights = np.asarray(dataset.weights, dtype=float)
    ecmp_base = ecmp_splits(path_library)
    topo_mapping = runner.SDNTopologyMapping.from_mininet(dataset.nodes, dataset.edges, dataset.od_pairs)
    test_indices = list(range(int(dataset.split["test_start"]), dataset.tm.shape[0]))
    gnn_model = gnn_cache.get(topo_key)
    if gnn_model is None:
        gnn_model = runner.load_gnn_model(dataset, path_library)
        gnn_cache[topo_key] = gnn_model

    rows = []
    ts_rows = []
    for method in CORE_METHODS:
        run_results = defaultdict(list)
        for run_id in range(NUM_RUNS):
            current_splits = [s.copy() for s in ecmp_base]
            current_groups, _ = runner.build_ecmp_baseline_rules(path_library, topo_mapping, dataset.edges)
            prev_latency = None
            gnnplus_state = {
                "prev_tm": None,
                "prev_util": None,
                "prev_selected_indicator": np.zeros(len(dataset.od_pairs), dtype=np.float32),
                "prev_disturbance": 0.0,
                "guard_cache": {},
                "guard_cycle_index": 0,
                "guard_fallback_cooldown": 0,
            }
            for t_idx in test_indices:
                tm_vec = dataset.tm[t_idx]
                if method == "gnnplus":
                    result, current_splits, current_groups, prev_latency, gnnplus_state = run_sdn_cycle_gnnplus_improved(
                        runner,
                        tm_vector=tm_vec,
                        dataset=dataset,
                        path_library=path_library,
                        ecmp_base=ecmp_base,
                        current_splits=current_splits,
                        current_groups=current_groups,
                        topo_mapping=topo_mapping,
                        capacities=capacities,
                        weights=weights,
                        gnnplus_model=gnnplus_model,
                        prev_latency_by_od=prev_latency,
                        gnnplus_state=gnnplus_state,
                        inference_calibration=inference_calibration,
                    )
                else:
                    result, current_splits, current_groups, prev_latency = runner.run_sdn_cycle(
                        tm_vector=tm_vec,
                        method=method,
                        dataset=dataset,
                        path_library=path_library,
                        ecmp_base=ecmp_base,
                        current_splits=current_splits,
                        current_groups=current_groups,
                        topo_mapping=topo_mapping,
                        capacities=capacities,
                        weights=weights,
                        gnn_model=gnn_model,
                        gnnplus_model=None,
                        prev_latency_by_od=prev_latency,
                    )
                run_results["post_mlus"].append(result.post_mlu)
                run_results["disturbances"].append(result.disturbance)
                run_results["throughputs"].append(result.throughput)
                run_results["latencies"].append(result.mean_latency)
                run_results["p95_latencies"].append(result.p95_latency)
                run_results["packet_losses"].append(result.packet_loss)
                run_results["jitters"].append(result.jitter)
                run_results["decision_times"].append(result.decision_time_ms)
                run_results["flow_updates"].append(result.flow_table_updates)
                run_results["rule_delays"].append(result.rule_install_delay_ms)
                if method == "gnnplus":
                    do_no_harm_fallback = float(bool(gnnplus_state.get("select_info", {}).get("do_no_harm_fallback", False)))
                    run_results["do_no_harm_fallbacks"].append(do_no_harm_fallback)
                else:
                    do_no_harm_fallback = 0.0
                    run_results["do_no_harm_fallbacks"].append(do_no_harm_fallback)
                if SAVE_PACKET_SDN_TIMESERIES:
                    ts_rows.append(
                        {
                            "topology": topo_key,
                            "status": "known" if topo_key in KNOWN_TOPOLOGIES else "unseen",
                            "method": method,
                            "scenario": "normal",
                            "run_id": int(run_id),
                            "timestep": int(t_idx),
                            "mlu": float(result.post_mlu),
                            "disturbance": float(result.disturbance),
                            "throughput": float(result.throughput),
                            "mean_latency_au": float(result.mean_latency),
                            "p95_latency_au": float(result.p95_latency),
                            "packet_loss": float(result.packet_loss),
                            "jitter_au": float(result.jitter),
                            "decision_time_ms": float(result.decision_time_ms),
                            "flow_table_updates": float(result.flow_table_updates),
                            "rule_install_delay_ms": float(result.rule_install_delay_ms),
                            "do_no_harm_fallback": float(do_no_harm_fallback),
                        }
                    )

        row = {
            "topology": topo_key,
            "status": "known" if topo_key in KNOWN_TOPOLOGIES else "unseen",
            "method": method,
            "scenario": "normal",
            "nodes": len(dataset.nodes),
            "edges": len(dataset.edges),
            "mean_mlu": float(np.mean(run_results["post_mlus"])),
            "mean_disturbance": float(np.mean(run_results["disturbances"])),
            "throughput": float(np.mean(run_results["throughputs"])),
            "mean_latency_au": float(np.mean(run_results["latencies"])),
            "p95_latency_au": float(np.mean(run_results["p95_latencies"])),
            "packet_loss": float(np.mean(run_results["packet_losses"])),
            "jitter_au": float(np.mean(run_results["jitters"])),
            "decision_time_ms": float(np.mean(run_results["decision_times"])),
            "flow_table_updates": float(np.mean(run_results["flow_updates"])),
            "rule_install_delay_ms": float(np.mean(run_results["rule_delays"])),
            "do_no_harm_fallback_rate": float(np.mean(run_results["do_no_harm_fallbacks"])) if run_results["do_no_harm_fallbacks"] else 0.0,
        }
        rows.append(row)
        print(f"[eval:normal] {topo_key} {method} mean_mlu={row['mean_mlu']:.4f}", flush=True)
    return rows, ts_rows


def benchmark_topology_failures_improved(runner, topo_key: str, gnn_cache: dict, gnnplus_model, inference_calibration=None) -> tuple[list[dict], list[dict]]:
    dataset, path_library = runner.load_dataset(topo_key)
    capacities = np.asarray(dataset.capacities, dtype=float)
    weights = np.asarray(dataset.weights, dtype=float)
    ecmp_base = ecmp_splits(path_library)
    gnn_model = gnn_cache.get(topo_key)
    if gnn_model is None:
        gnn_model = runner.load_gnn_model(dataset, path_library)
        gnn_cache[topo_key] = gnn_model
    test_indices = list(range(int(dataset.split["test_start"]), dataset.tm.shape[0]))

    rows = []
    ts_rows = []
    for scenario in runner.FAILURE_SCENARIOS:
        if scenario == "normal":
            continue
        sample_indices = test_indices[:: max(1, len(test_indices) // 10)]
        for method in CORE_METHODS:
            run_results = defaultdict(list)
            for t_idx in sample_indices:
                tm_vec = dataset.tm[t_idx]
                if method == "gnnplus":
                    recovery_ms, pre_mlu, post_mlu, _, select_info = run_failure_scenario_gnnplus_improved(
                        runner,
                        scenario=scenario,
                        tm_vector=tm_vec,
                        dataset=dataset,
                        path_library=path_library,
                        ecmp_base=ecmp_base,
                        capacities=capacities,
                        weights=weights,
                        gnnplus_model=gnnplus_model,
                        inference_calibration=inference_calibration,
                    )
                    do_no_harm_fallback = float(bool(select_info.get("do_no_harm_fallback", False)))
                    run_results["do_no_harm_fallbacks"].append(do_no_harm_fallback)
                else:
                    recovery_ms, pre_mlu, post_mlu, _ = runner.run_failure_scenario(
                        scenario=scenario,
                        tm_vector=tm_vec,
                        method=method,
                        dataset=dataset,
                        path_library=path_library,
                        ecmp_base=ecmp_base,
                        capacities=capacities,
                        weights=weights,
                        topo_mapping=None,
                        gnn_model=gnn_model,
                        gnnplus_model=None,
                    )
                    do_no_harm_fallback = 0.0
                    run_results["do_no_harm_fallbacks"].append(do_no_harm_fallback)
                run_results["recovery_times"].append(recovery_ms)
                run_results["pre_mlus"].append(pre_mlu)
                run_results["post_mlus"].append(post_mlu)
                if SAVE_PACKET_SDN_TIMESERIES:
                    ts_rows.append(
                        {
                            "topology": topo_key,
                            "status": "known" if topo_key in KNOWN_TOPOLOGIES else "unseen",
                            "method": method,
                            "scenario": scenario,
                            "timestep": int(t_idx),
                            "pre_failure_mlu": float(pre_mlu),
                            "post_recovery_mlu": float(post_mlu),
                            "failure_recovery_ms": float(recovery_ms),
                            "do_no_harm_fallback": float(do_no_harm_fallback),
                        }
                    )

            row = {
                "topology": topo_key,
                "status": "known" if topo_key in KNOWN_TOPOLOGIES else "unseen",
                "method": method,
                "scenario": scenario,
                "nodes": len(dataset.nodes),
                "edges": len(dataset.edges),
                "mean_mlu": float(np.mean(run_results["post_mlus"])),
                "pre_failure_mlu": float(np.mean(run_results["pre_mlus"])),
                "failure_recovery_ms": float(np.mean(run_results["recovery_times"])),
                "do_no_harm_fallback_rate": float(np.mean(run_results["do_no_harm_fallbacks"])) if run_results["do_no_harm_fallbacks"] else 0.0,
            }
            rows.append(row)
            print(f"[eval:failure] {topo_key} {scenario} {method} mean_mlu={row['mean_mlu']:.4f}", flush=True)
    return rows, ts_rows


def proportional_budget_k(tm_vector: np.ndarray, path_library, ratio: float) -> int:
    active_mask = (
        (np.asarray(tm_vector, dtype=np.float64) > 1e-12)
        & surviving_od_mask(path_library)
    ).astype(np.float32)
    num_active_od = int(np.sum(active_mask > 0.5))
    return max(0, int(float(ratio) * float(num_active_od)))


def run_proportional_budget_cycle(
    runner,
    *,
    method: str,
    ratio: float,
    tm_vector,
    dataset,
    path_library,
    ecmp_base,
    current_splits,
    current_groups,
    topo_mapping,
    capacities,
    weights,
    gnnplus_model,
    prev_latency_by_od,
    gnnplus_state,
    inference_calibration=None,
):
    budget_k = proportional_budget_k(tm_vector, path_library, ratio)
    t_total_start = time.perf_counter()
    routing_pre = apply_routing(tm_vector, current_splits, path_library, capacities)
    pre_telemetry = runner.compute_telemetry(
        tm_vector=tm_vector,
        splits=current_splits,
        path_library=path_library,
        routing=routing_pre,
        weights=weights,
        prev_latency_by_od=prev_latency_by_od,
        fast_links_only=True,
    )

    selected_ods: list[int] = []
    if method == "ecmp" or budget_k <= 0:
        new_splits = [s.copy() for s in ecmp_base]
        next_state = dict(gnnplus_state)
    elif method == "bottleneck":
        selected_ods = select_bottleneck_critical(tm_vector, ecmp_base, path_library, capacities, budget_k)
        lp_result = runner.solve_selected_path_lp_safe(
            tm_vector=tm_vector,
            selected_ods=selected_ods,
            base_splits=ecmp_base,
            path_library=path_library,
            capacities=capacities,
            warm_start_splits=current_splits,
            time_limit_sec=LP_TIME_LIMIT,
            context=f"{dataset.key}:bottleneck:ratio_{ratio:.2f}",
        )
        new_splits = [s.copy() for s in lp_result.splits]
        next_state = dict(gnnplus_state)
    elif method == "gnnplus":
        controls = resolve_inference_controls(inference_calibration, str(dataset.key))
        selected_ods, select_info = gnnplus_select_stateful(
            gnnplus_model,
            dataset=dataset,
            path_library=path_library,
            capacities=capacities,
            tm_vector=tm_vector,
            telemetry=pre_telemetry,
            prev_tm=gnnplus_state["prev_tm"],
            prev_util=gnnplus_state["prev_util"],
            prev_selected_indicator=gnnplus_state["prev_selected_indicator"],
            prev_disturbance=gnnplus_state["prev_disturbance"],
            k_crit=budget_k,
            gate_temperature=controls["gate_temperature"],
            tie_break_eps=controls["tie_break_eps"],
        )
        _prev_indicator_ratio = gnnplus_state.get("prev_selected_indicator")
        prev_selected_ods_ratio: list[int] = []
        if _prev_indicator_ratio is not None:
            _prev_arr_ratio = np.asarray(_prev_indicator_ratio, dtype=np.float32).reshape(-1)
            prev_selected_ods_ratio = [int(i) for i in np.where(_prev_arr_ratio > 0.5)[0].tolist()]
        selected_ods, lp_result, gate_info, updated_guard_cache, updated_guard_fallback_cooldown = apply_do_no_harm_gate(
            runner,
            tm_vector=np.asarray(tm_vector, dtype=float),
            selected_ods=selected_ods,
            base_splits=ecmp_base,
            warm_start_splits=current_splits,
            path_library=path_library,
            capacities=capacities,
            k_crit=budget_k,
            context=f"{dataset.key}:gnnplus:ratio_{ratio:.2f}",
            topology_key=str(dataset.key),
            guard_bottleneck_selected=select_info.get("prefilter_bottleneck_selected", []),
            guard_cache=gnnplus_state.get("guard_cache", {}),
            step_index=int(gnnplus_state.get("guard_cycle_index", 0)),
            guard_fallback_cooldown=int(gnnplus_state.get("guard_fallback_cooldown", 0)),
            prev_selected_ods=prev_selected_ods_ratio,
        )
        select_info.update(gate_info)
        new_splits = [s.copy() for s in lp_result.splits]
        next_state = {
            "prev_tm": np.asarray(tm_vector, dtype=float),
            "prev_util": None,
            "prev_selected_indicator": selection_indicator(len(dataset.od_pairs), selected_ods),
            "prev_disturbance": float(gnnplus_state["prev_disturbance"]),
            "select_info": dict(select_info),
            "guard_cache": dict(updated_guard_cache),
            "guard_cycle_index": int(gnnplus_state.get("guard_cycle_index", 0)) + 1,
            "guard_fallback_cooldown": int(updated_guard_fallback_cooldown),
        }
    else:
        raise ValueError(f"Unsupported proportional-budget method: {method}")

    decision_time_ms = (time.perf_counter() - t_total_start) * 1000.0
    if method == "ecmp" or budget_k <= 0:
        new_groups, _ = runner.build_ecmp_baseline_rules(path_library, topo_mapping, dataset.edges)
    else:
        new_groups, _ = runner.splits_to_openflow_rules(new_splits, selected_ods, path_library, topo_mapping, dataset.edges)

    flow_table_updates = len(runner.compute_rule_diff(current_groups, new_groups))
    routing_post = apply_routing(tm_vector, new_splits, path_library, capacities)
    post_mlu = float(routing_post.mlu)
    disturbance = float(compute_disturbance(current_splits, new_splits, tm_vector))
    telemetry_post = runner.compute_telemetry(
        tm_vector=tm_vector,
        splits=new_splits,
        path_library=path_library,
        routing=routing_post,
        weights=weights,
        prev_latency_by_od=prev_latency_by_od,
    )
    if method == "gnnplus" and budget_k > 0:
        next_state["prev_util"] = np.asarray(telemetry_post.utilization, dtype=float)
        next_state["prev_disturbance"] = disturbance

    return {
        "budget_k": int(budget_k),
        "mean_mlu": float(post_mlu),
        "mean_disturbance": disturbance,
        "decision_time_ms": float(decision_time_ms),
        "flow_table_updates": float(flow_table_updates),
    }, [s.copy() for s in new_splits], new_groups, telemetry_post.latency_by_od, next_state


def benchmark_proportional_budget_normal(runner, topo_key: str, gnnplus_model, ratio: float, inference_calibration=None) -> list[dict]:
    dataset, path_library = runner.load_dataset(topo_key)
    capacities = np.asarray(dataset.capacities, dtype=float)
    weights = np.asarray(dataset.weights, dtype=float)
    ecmp_base = ecmp_splits(path_library)
    topo_mapping = runner.SDNTopologyMapping.from_mininet(dataset.nodes, dataset.edges, dataset.od_pairs)
    test_indices = list(range(int(dataset.split["test_start"]), dataset.tm.shape[0]))

    rows = []
    for method in PROPORTIONAL_METHODS:
        run_results = defaultdict(list)
        for _ in range(PROPORTIONAL_NUM_RUNS):
            current_splits = [s.copy() for s in ecmp_base]
            current_groups, _ = runner.build_ecmp_baseline_rules(path_library, topo_mapping, dataset.edges)
            prev_latency = None
            gnnplus_state = {
                "prev_tm": None,
                "prev_util": None,
                "prev_selected_indicator": np.zeros(len(dataset.od_pairs), dtype=np.float32),
                "prev_disturbance": 0.0,
                "guard_cache": {},
                "guard_cycle_index": 0,
                "guard_fallback_cooldown": 0,
            }
            for t_idx in test_indices:
                tm_vec = np.asarray(dataset.tm[t_idx], dtype=float)
                result, current_splits, current_groups, prev_latency, gnnplus_state = run_proportional_budget_cycle(
                    runner,
                    method=method,
                    ratio=ratio,
                    tm_vector=tm_vec,
                    dataset=dataset,
                    path_library=path_library,
                    ecmp_base=ecmp_base,
                    current_splits=current_splits,
                    current_groups=current_groups,
                    topo_mapping=topo_mapping,
                    capacities=capacities,
                    weights=weights,
                    gnnplus_model=gnnplus_model,
                    prev_latency_by_od=prev_latency,
                    gnnplus_state=gnnplus_state,
                    inference_calibration=inference_calibration,
                )
                for key, value in result.items():
                    run_results[key].append(float(value))

        row = {
            "topology": topo_key,
            "status": "known" if topo_key in KNOWN_TOPOLOGIES else "unseen",
            "budget_ratio": float(ratio),
            "method": method,
            "nodes": len(dataset.nodes),
            "edges": len(dataset.edges),
            "mean_budget_k": float(np.mean(run_results["budget_k"])) if run_results["budget_k"] else 0.0,
            "mean_mlu": float(np.mean(run_results["mean_mlu"])) if run_results["mean_mlu"] else 0.0,
            "mean_disturbance": float(np.mean(run_results["mean_disturbance"])) if run_results["mean_disturbance"] else 0.0,
            "decision_time_ms": float(np.mean(run_results["decision_time_ms"])) if run_results["decision_time_ms"] else 0.0,
            "flow_table_updates": float(np.mean(run_results["flow_table_updates"])) if run_results["flow_table_updates"] else 0.0,
        }
        rows.append(row)
        print(
            f"[eval:ratio] {topo_key} ratio={ratio:.2f} {method} "
            f"mean_k={row['mean_budget_k']:.2f} mean_mlu={row['mean_mlu']:.4f}",
            flush=True,
        )
    return rows


def build_cascading_random_failure_state(
    runner,
    *,
    dataset,
    path_library,
    capacities,
    weights,
    tm_vector,
    num_failures: int,
    seed: int,
):
    rng = random.Random(int(seed))
    fail_count = min(max(int(num_failures), 1), len(capacities))
    failed_edges = sorted(rng.sample(range(len(capacities)), fail_count))
    failure_mask = np.ones(len(capacities), dtype=float)
    for idx in failed_edges:
        failure_mask[int(idx)] = 0.0
    keep = [idx for idx in range(len(dataset.edges)) if idx not in set(int(x) for x in failed_edges)]
    kept_edges = [dataset.edges[idx] for idx in keep]
    kept_caps = np.asarray([capacities[idx] for idx in keep], dtype=float)
    kept_weights = np.asarray([weights[idx] for idx in keep], dtype=float)
    effective_path_library = runner.build_modified_paths(
        dataset.nodes,
        kept_edges,
        kept_weights,
        dataset.od_pairs,
        k_paths=runner.K_PATHS,
    )
    effective_dataset = runner._build_dataset_view(
        dataset,
        edges=kept_edges,
        capacities=kept_caps,
        weights=kept_weights,
    )
    return {
        "effective_tm": np.asarray(tm_vector, dtype=float),
        "effective_caps": kept_caps,
        "effective_weights": kept_weights,
        "effective_path_library": effective_path_library,
        "effective_dataset": effective_dataset,
        "effective_ecmp": ecmp_splits(effective_path_library),
        "failure_mask": failure_mask,
        "failed_edges": failed_edges,
        "scenario": f"cascading_random_failure_{fail_count}",
    }


def benchmark_extreme_stress_vtlwavenet(runner, gnnplus_model, ratio: float, inference_calibration=None) -> list[dict]:
    topo_key = "vtlwavenet2011"
    dataset, path_library = runner.load_dataset(topo_key)
    capacities = np.asarray(dataset.capacities, dtype=float)
    weights = np.asarray(dataset.weights, dtype=float)
    test_indices = list(range(int(dataset.split["test_start"]), dataset.tm.shape[0]))
    sample_indices = test_indices[:: max(1, len(test_indices) // 10)]

    rows = []
    for method in PROPORTIONAL_METHODS:
        run_results = defaultdict(list)
        for t_idx in sample_indices:
            tm_vec = np.asarray(dataset.tm[t_idx], dtype=float)
            stress_state = build_cascading_random_failure_state(
                runner,
                dataset=dataset,
                path_library=path_library,
                capacities=capacities,
                weights=weights,
                tm_vector=tm_vec,
                num_failures=EXTREME_STRESS_FAIL_COUNT,
                seed=SEED + int(t_idx),
            )
            effective_tm = stress_state["effective_tm"]
            effective_caps = stress_state["effective_caps"]
            effective_weights = stress_state["effective_weights"]
            effective_path_library = stress_state["effective_path_library"]
            effective_dataset = stress_state["effective_dataset"]
            effective_ecmp = stress_state["effective_ecmp"]
            failure_mask = stress_state["failure_mask"]
            budget_k = proportional_budget_k(effective_tm, effective_path_library, ratio)

            pre_routing = apply_routing(effective_tm, effective_ecmp, effective_path_library, effective_caps)
            pre_telemetry = runner.compute_telemetry(
                tm_vector=effective_tm,
                splits=effective_ecmp,
                path_library=effective_path_library,
                routing=pre_routing,
                weights=effective_weights,
                prev_latency_by_od=None,
                fast_links_only=True,
            )
            t_start = time.perf_counter()
            if method == "ecmp" or budget_k <= 0:
                recovery_splits = [s.copy() for s in effective_ecmp]
            elif method == "bottleneck":
                selected = select_bottleneck_critical(effective_tm, effective_ecmp, effective_path_library, effective_caps, budget_k)
                lp_result = runner.solve_selected_path_lp_safe(
                    tm_vector=effective_tm,
                    selected_ods=selected,
                    base_splits=effective_ecmp,
                    path_library=effective_path_library,
                    capacities=effective_caps,
                    warm_start_splits=effective_ecmp,
                    time_limit_sec=LP_TIME_LIMIT,
                    context=f"{topo_key}:extreme_stress:bottleneck",
                )
                recovery_splits = [s.copy() for s in lp_result.splits]
            else:
                controls = resolve_inference_controls(inference_calibration, str(topo_key))
                selected, select_info = gnnplus_select_stateful(
                    gnnplus_model,
                    dataset=effective_dataset,
                    path_library=effective_path_library,
                    capacities=effective_caps,
                    tm_vector=effective_tm,
                    telemetry=pre_telemetry,
                    prev_tm=None,
                    prev_util=None,
                    prev_selected_indicator=np.zeros(len(effective_dataset.od_pairs), dtype=np.float32),
                    prev_disturbance=0.0,
                    k_crit=budget_k,
                    failure_mask=failure_mask,
                    gate_temperature=controls["gate_temperature"],
                    tie_break_eps=controls["tie_break_eps"],
                )
                _, lp_result, _, _, _ = apply_do_no_harm_gate(
                    runner,
                    topology_key=str(topo_key),
                    tm_vector=np.asarray(effective_tm, dtype=float),
                    selected_ods=selected,
                    base_splits=effective_ecmp,
                    warm_start_splits=effective_ecmp,
                    path_library=effective_path_library,
                    capacities=effective_caps,
                    k_crit=budget_k,
                    guard_bottleneck_selected=select_info.get("prefilter_bottleneck_selected", []),
                    guard_cache={},
                    step_index=0,
                    context=f"{topo_key}:extreme_stress:gnnplus",
                )
                recovery_splits = [s.copy() for s in lp_result.splits]

            recovery_ms = (time.perf_counter() - t_start) * 1000.0
            post_routing = apply_routing(effective_tm, recovery_splits, effective_path_library, effective_caps)
            run_results["budget_k"].append(float(budget_k))
            run_results["pre_failure_mlu"].append(float(pre_routing.mlu))
            run_results["post_recovery_mlu"].append(float(post_routing.mlu))
            run_results["failure_recovery_ms"].append(float(recovery_ms))

        row = {
            "topology": topo_key,
            "status": "unseen",
            "scenario": f"cascading_random_failure_{EXTREME_STRESS_FAIL_COUNT}",
            "budget_ratio": float(ratio),
            "method": method,
            "nodes": len(dataset.nodes),
            "edges": len(dataset.edges),
            "mean_budget_k": float(np.mean(run_results["budget_k"])) if run_results["budget_k"] else 0.0,
            "pre_failure_mlu": float(np.mean(run_results["pre_failure_mlu"])) if run_results["pre_failure_mlu"] else 0.0,
            "post_recovery_mlu": float(np.mean(run_results["post_recovery_mlu"])) if run_results["post_recovery_mlu"] else 0.0,
            "failure_recovery_ms": float(np.mean(run_results["failure_recovery_ms"])) if run_results["failure_recovery_ms"] else 0.0,
        }
        rows.append(row)
        print(
            f"[eval:stress] {topo_key} ratio={ratio:.2f} {method} "
            f"mean_k={row['mean_budget_k']:.2f} post_mlu={row['post_recovery_mlu']:.4f}",
            flush=True,
        )
    return rows


def prepare_sdn_metrics(summary_df: pd.DataFrame, failure_df: pd.DataFrame) -> pd.DataFrame:
    recovery = (
        failure_df.groupby(["topology", "method"], as_index=False)["failure_recovery_ms"]
        .mean()
        .rename(columns={"failure_recovery_ms": "avg_failure_recovery_ms"})
    )
    metrics = summary_df.merge(recovery, on=["topology", "method"], how="left")
    metrics["Method"] = metrics["method"].map(METHOD_LABELS)
    metrics["Topology"] = metrics["topology"].map(TOPOLOGY_DISPLAY)
    metrics["Status"] = metrics["status"].str.title()
    metrics = metrics.rename(
        columns={
            "mean_mlu": "Mean MLU",
            "throughput": "Throughput",
            "mean_disturbance": "Disturbance",
            "mean_latency_au": "Mean Delay",
            "p95_latency_au": "P95 Delay",
            "packet_loss": "Packet Loss",
            "jitter_au": "Jitter",
            "decision_time_ms": "Decision Time (ms)",
            "flow_table_updates": "Flow Table Updates",
            "rule_install_delay_ms": "Rule Install Delay (ms)",
            "avg_failure_recovery_ms": "Failure Recovery (ms)",
        }
    )
    cols = [
        "Method",
        "Topology",
        "Status",
        "Mean MLU",
        "Throughput",
        "Disturbance",
        "Mean Delay",
        "P95 Delay",
        "Packet Loss",
        "Jitter",
        "Decision Time (ms)",
        "Flow Table Updates",
        "Rule Install Delay (ms)",
        "Failure Recovery (ms)",
    ]
    return metrics[cols]


def build_split_manifest() -> dict:
    return {
        "base_objective": "zero-shot generalization",
        "main_reference_configuration": "step1to5_failgate",
        "methods": CORE_METHODS,
        "known_topologies": KNOWN_TOPOLOGIES,
        "unseen_topologies": UNSEEN_TOPOLOGIES,
        "train_topologies": KNOWN_TOPOLOGIES,
        "validation_topologies": KNOWN_TOPOLOGIES,
        "evaluation_topologies": ALL_TOPOLOGIES,
        "zero_shot_guards": {
            "bayesian_calibration_used": False,
            "few_shot_adaptation_used": False,
            "per_topology_tuning_used": False,
            "metagate_used": False,
            "stable_metagate_used": False,
        },
        "improvements_enabled": {
            "section1_failure_gated_features_and_cache": True,
            "section1_temporal_gate_during_failures": True,
            "section2_topology_balanced_supervised_training": True,
            "section2_gentle_rl_finetuning": True,
            "section3_physical_features": True,
            "section4_soft_teacher_targets": True,
            "section4_continuous_regression_targets": True,
            "section5_rl_reward_alignment": True,
            "section7_disturbance_aware_temporal_path": True,
            "candidate_prefiltering": bool(USE_CANDIDATE_PREFILTER),
            "gated_residual_with_learned_deferral": bool(USE_GATED_RESIDUAL),
            "cross_od_attention": bool(USE_CROSS_OD_ATTENTION),
            "lp_warm_start": True,
            "synthetic_failure_training": float(SYNTHETIC_FAILURE_PROB) > 0.0,
        },
        "fixed_k": K_CRIT,
        "proportional_budget_study": {
            "enabled": True,
            "main_ratio": float(PROPORTIONAL_BUDGET_RATIO),
            "extreme_stress_ratio": float(EXTREME_STRESS_RATIO),
            "extreme_stress_failure_count": int(EXTREME_STRESS_FAIL_COUNT),
            "methods": list(PROPORTIONAL_METHODS),
        },
        "feature_variant": FEATURE_VARIANT,
        "feature_profile": feature_profile_description(),
        "teacher_profile": teacher_profile_description(),
        "continuity_bonus": CONTINUITY_BONUS,
        "supervised_topology_weights": {str(k): float(v) for k, v in SUPERVISED_TOPOLOGY_WEIGHTS.items()},
        "reward_profile": {
            "w_reward_mlu": REWARD_MLU,
            "w_reward_mlu_normal": REWARD_MLU_NORMAL,
            "w_reward_mlu_failure": REWARD_MLU_FAILURE,
            "w_reward_improvement": REWARD_IMPROVEMENT,
            "w_reward_disturbance": REWARD_DISTURBANCE,
            "w_reward_vs_bottleneck": REWARD_VS_BOTTLENECK,
            "w_reward_vs_reference": REWARD_VS_REFERENCE,
            "w_reward_bottleneck_margin": REWARD_BOTTLENECK_MARGIN,
            "entropy_weight": RL_ENTROPY_WEIGHT,
            "rl_failure_rollout_probability": RL_FAILURE_ROLLOUT_PROB,
            "hinge_reward_enabled": bool(ENABLE_HINGE_REWARD),
            "hinge_threshold": float(HINGE_THRESHOLD),
            "hinge_multiplier": float(HINGE_MULTIPLIER),
        },
        "base_checkpoint": str(BASE_GNNPLUS_CKPT.relative_to(PROJECT_ROOT)),
    }


def available_reference_dirs() -> list[tuple[str, Path]]:
    ordered = [
        ("gnnplus_archfix_fulltrain", ARCHFIX_OUTPUT_DIR),
        ("gnnplus_step1to5_failgate_main_reference", MAIN_REFERENCE_OUTPUT_DIR),
    ]
    out: list[tuple[str, Path]] = []
    seen: set[str] = set()
    for name, path in ordered:
        if name in seen:
            continue
        if (path / "packet_sdn_summary.csv").exists() and (path / "packet_sdn_failure.csv").exists():
            out.append((name, path))
            seen.add(name)
    if not out:
        fallback_name = PREVIOUS_OUTPUT_DIR.name
        if (PREVIOUS_OUTPUT_DIR / "packet_sdn_summary.csv").exists() and (PREVIOUS_OUTPUT_DIR / "packet_sdn_failure.csv").exists():
            out.append((fallback_name, PREVIOUS_OUTPUT_DIR))
        elif (BASELINE_OUTPUT_DIR / "packet_sdn_summary.csv").exists() and (BASELINE_OUTPUT_DIR / "packet_sdn_failure.csv").exists():
            out.append((BASELINE_OUTPUT_DIR.name, BASELINE_OUTPUT_DIR))
    return out


def build_comparison_tables(summary_df: pd.DataFrame, failure_df: pd.DataFrame) -> dict[str, Path]:
    COMPARISON_DIR.mkdir(parents=True, exist_ok=True)
    reference_dirs = available_reference_dirs()
    comparison_paths: dict[str, Path] = {}
    overall_rows = []

    current_gnnplus = summary_df[summary_df["method"] == "gnnplus"].copy()
    current_failure = failure_df[failure_df["method"] == "gnnplus"].copy()

    for ref_name, reference_dir in reference_dirs:
        old_summary = pd.read_csv(reference_dir / "packet_sdn_summary.csv")
        old_failure = pd.read_csv(reference_dir / "packet_sdn_failure.csv")

        old_new = old_summary[old_summary["method"] == "gnnplus"].merge(
            current_gnnplus,
            on=["topology", "status", "method", "scenario", "nodes", "edges"],
            suffixes=("_old", "_new"),
        )
        for metric in ["mean_mlu", "throughput", "mean_disturbance", "decision_time_ms"]:
            old_new[f"{metric}_delta"] = old_new[f"{metric}_new"] - old_new[f"{metric}_old"]
            old_new[f"{metric}_pct_delta"] = np.where(
                np.abs(old_new[f"{metric}_old"]) > 1e-12,
                (old_new[f"{metric}_new"] - old_new[f"{metric}_old"]) / np.abs(old_new[f"{metric}_old"]) * 100.0,
                0.0,
            )
        normal_path = COMPARISON_DIR / f"gnnplus_normal_vs_{ref_name}.csv"
        old_new.to_csv(normal_path, index=False)
        comparison_paths[f"normal_vs_{ref_name}"] = normal_path

        old_new_failure = old_failure[old_failure["method"] == "gnnplus"].merge(
            current_failure,
            on=["topology", "status", "method", "scenario", "nodes", "edges"],
            suffixes=("_old", "_new"),
        )
        for metric in ["mean_mlu", "failure_recovery_ms"]:
            old_new_failure[f"{metric}_delta"] = old_new_failure[f"{metric}_new"] - old_new_failure[f"{metric}_old"]
            old_new_failure[f"{metric}_pct_delta"] = np.where(
                np.abs(old_new_failure[f"{metric}_old"]) > 1e-12,
                (old_new_failure[f"{metric}_new"] - old_new_failure[f"{metric}_old"]) / np.abs(old_new_failure[f"{metric}_old"]) * 100.0,
                0.0,
            )
        failure_path = COMPARISON_DIR / f"gnnplus_failure_vs_{ref_name}.csv"
        old_new_failure.to_csv(failure_path, index=False)
        comparison_paths[f"failure_vs_{ref_name}"] = failure_path

        overall_rows.append(
            {
                "bundle": ref_name,
                "gnnplus_mean_mlu": float(old_summary[old_summary["method"] == "gnnplus"]["mean_mlu"].mean()),
                "gnnplus_mean_throughput": float(old_summary[old_summary["method"] == "gnnplus"]["throughput"].mean()),
                "gnnplus_mean_disturbance": float(old_summary[old_summary["method"] == "gnnplus"]["mean_disturbance"].mean()),
                "gnnplus_mean_decision_time_ms": float(old_summary[old_summary["method"] == "gnnplus"]["decision_time_ms"].mean()),
            }
        )

    ranking_rows = []
    for topo in ALL_TOPOLOGIES:
        topo_df = summary_df[summary_df["topology"] == topo].sort_values("mean_mlu")
        for rank, (_, row) in enumerate(topo_df.iterrows(), start=1):
            ranking_rows.append(
                {
                    "topology": topo,
                    "method": str(row["method"]),
                    "rank_by_mlu": int(rank),
                    "mean_mlu": float(row["mean_mlu"]),
                    "throughput": float(row["throughput"]),
                    "mean_disturbance": float(row["mean_disturbance"]),
                    "decision_time_ms": float(row["decision_time_ms"]),
                }
    )
    path3 = COMPARISON_DIR / "new_run_method_ranking.csv"
    pd.DataFrame(ranking_rows).to_csv(path3, index=False)
    comparison_paths["method_ranking"] = path3

    overall_rows.append(
        {
            "bundle": EXPERIMENT_TAG,
            "gnnplus_mean_mlu": float(current_gnnplus["mean_mlu"].mean()),
            "gnnplus_mean_throughput": float(current_gnnplus["throughput"].mean()),
            "gnnplus_mean_disturbance": float(current_gnnplus["mean_disturbance"].mean()),
            "gnnplus_mean_decision_time_ms": float(current_gnnplus["decision_time_ms"].mean()),
        }
    )
    overall = pd.DataFrame(overall_rows)
    path4 = COMPARISON_DIR / "overall_bundle_comparison.csv"
    overall.to_csv(path4, index=False)
    comparison_paths["overall"] = path4
    return comparison_paths


def add_title_page(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(os.environ.get("GNNPLUS_REPORT_TITLE", "GNN+ Roadmap Step1-6 Zero-Shot Report"))
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x80)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "Scope: ECMP, Bottleneck, Original GNN, GNN+\n"
        f"Feature profile: {FEATURE_VARIANT}\n"
        "Main reference: Fixed-K40 failgate zero-shot branch\n"
        "No MetaGate, No Stable MetaGate, No Bayesian Calibration"
    )
    run.font.size = Pt(11)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "Important: packet-level SDN metrics are model-based analytical metrics, not live Mininet measurements."
    )
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    doc.add_page_break()


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_image(doc: Document, path: Path, caption: str, width: float = 6.2) -> None:
    if not path.exists():
        doc.add_paragraph(f"[Image missing: {path.name}]")
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def build_report(summary_df: pd.DataFrame, failure_df: pd.DataFrame, metrics_df: pd.DataFrame, training_summary: dict, split_manifest: dict):
    helper = load_helper()
    helper.PLOTS_DIR = PLOTS_DIR
    PLOTS_DIR.mkdir(parents=True, exist_ok=True)
    helper.create_plots(summary_df, failure_df)
    num_topologies = len(ALL_TOPOLOGIES)
    unseen_display = ", ".join(TOPOLOGY_DISPLAY.get(topo, topo) for topo in UNSEEN_TOPOLOGIES)

    overall_compare = pd.read_csv(COMPARISON_DIR / "overall_bundle_comparison.csv")
    reference_dirs = available_reference_dirs()
    comparison_normal_tables = {
        ref_name: pd.read_csv(COMPARISON_DIR / f"gnnplus_normal_vs_{ref_name}.csv")
        for ref_name, _ in reference_dirs
        if (COMPARISON_DIR / f"gnnplus_normal_vs_{ref_name}.csv").exists()
    }
    comparison_failure_tables = {
        ref_name: pd.read_csv(COMPARISON_DIR / f"gnnplus_failure_vs_{ref_name}.csv")
        for ref_name, _ in reference_dirs
        if (COMPARISON_DIR / f"gnnplus_failure_vs_{ref_name}.csv").exists()
    }
    proportional_df = pd.read_csv(PROPORTIONAL_SUMMARY_CSV) if PROPORTIONAL_SUMMARY_CSV.exists() else pd.DataFrame()
    proportional_stress_df = pd.read_csv(PROPORTIONAL_STRESS_CSV) if PROPORTIONAL_STRESS_CSV.exists() else pd.DataFrame()
    rl_used = str(training_summary.get("reinforce", {}).get("mode", "")).lower() == "reinforce"

    doc = Document()
    helper.set_default_style(doc)
    add_title_page(doc)

    doc.add_heading("1. Scope", level=1)
    doc.add_paragraph(
        "This experiment locks the fixed-K40 failgate branch as the main standalone GNN+ reference configuration, "
        "then retrains and reevaluates it under the clean zero-shot protocol. The method scope remains ECMP, "
        "Bottleneck, Original GNN, and GNN+ for the main benchmark."
    )
    add_bullet(doc, "Known training topologies remain the original six known topologies.")
    add_bullet(doc, f"Unseen zero-shot evaluation topologies in this run: {unseen_display}.")
    add_bullet(doc, "No MetaGate, Stable MetaGate, or unseen-topology adaptation is used.")
    add_bullet(doc, "Gate-temperature calibration is fit only on known-topology validation slices; unseen topologies use the global default.")
    add_bullet(doc, "Section 1 selector engine fixes are locked before training: failure-path generator check and temporal failure gate.")
    add_bullet(doc, f"Active feature profile: {feature_profile_description()}")

    doc.add_heading("2. Improvements Enabled", level=1)
    improvements = pd.DataFrame(
        [
            {"Section": "1", "Change": "Failure-gated OD feature path + cached topology tensors", "Enabled": "Yes"},
            {"Section": "1", "Change": "Failure-active temporal gate drops prev_selected and prev_disturbance", "Enabled": "Yes"},
            {"Section": "A2", "Change": "Synthetic failures reduced to p=0.08 with 1-2 failed edges and capacity perturbation stratification", "Enabled": "Yes"},
            {"Section": "A3", "Change": f"Training candidate window = {TRAIN_CANDIDATE_POOL_SIZE}; inference window = {INFER_CANDIDATE_POOL_SIZE}", "Enabled": "Yes"},
            {"Section": "A5", "Change": "Oracle-window auditing + prefilter recall auxiliary loss", "Enabled": "Yes"},
            {"Section": "2", "Change": "Topology-balanced supervised pretraining", "Enabled": "Yes"},
            {"Section": "A1", "Change": "KL-anchored RL fine-tuning against the supervised reference policy", "Enabled": "Yes" if rl_used else "No (Task 7 ships supervised checkpoint)"},
            {"Section": "A4", "Change": "Churn-aware disturbance penalty in RL reward", "Enabled": "Yes" if rl_used else "No (Task 7 ships supervised checkpoint)"},
            {"Section": "1", "Change": f"Feature pruning profile ({FEATURE_VARIANT})", "Enabled": "Yes"},
            {"Section": "3", "Change": "Physical/stress-aware features", "Enabled": "Yes"},
            {"Section": "4", "Change": "LP-strengthened soft teacher targets", "Enabled": "Yes"},
            {"Section": "C1", "Change": "Explicit Bottleneck-vs-GNN+ MoE floor in the gated residual path", "Enabled": "Yes"},
            {"Section": "C2", "Change": "Auxiliary bottleneck replacement head", "Enabled": "Yes"},
            {"Section": "C3", "Change": "Plackett-Luce set-level supervised loss", "Enabled": "Yes"},
            {"Section": "C4", "Change": "2-head cross-OD attention with edge-sharing bias", "Enabled": "Yes"},
            {"Section": "C5", "Change": "Supervised temporal consistency regularization", "Enabled": "Yes"},
            {"Section": "B1-B3", "Change": "Inference do-no-harm gate + tie-aware ranking + gate-temperature calibration", "Enabled": "Yes"},
            {"Section": "6", "Change": f"Fixed K = {K_CRIT} main thesis branch", "Enabled": "Yes"},
        ]
    )
    helper.add_dataframe_table(doc, improvements, font_size=9)

    doc.add_heading("3. Training Summary", level=1)
    supervised_losses = training_summary["supervised"].get("losses", {})
    pl_start = supervised_losses.get("plackett_luce_weight_start")
    pl_end = supervised_losses.get("plackett_luce_weight_end")
    pl_ramp = supervised_losses.get("plackett_luce_ramp_epochs")
    pl_schedule_value = "N/A" if (pl_start is None or pl_end is None or pl_ramp is None) else f"{pl_start} -> {pl_end} (epochs 1-{pl_ramp})"
    train_rows_list = [
        {"Field": "Base checkpoint", "Value": training_summary["base_checkpoint"]},
        {"Field": "Final checkpoint", "Value": training_summary["final_checkpoint"]},
        {"Field": "Final model source", "Value": training_summary.get("final_model_source", "reinforce")},
        {"Field": "Feature variant", "Value": training_summary["feature_variant"]},
        {"Field": "Feature profile", "Value": training_summary["supervised"]["feature_profile"]},
        {"Field": "Teacher profile", "Value": training_summary["supervised"]["teacher_profile"]},
        {"Field": "Supervised topology weights", "Value": json.dumps(training_summary["supervised"]["supervised_topology_weights"], sort_keys=True)},
        {"Field": "Continuity bonus", "Value": training_summary["continuity_bonus"]},
        {"Field": "Train candidate pool size", "Value": training_summary["supervised"]["candidate_prefilter"]["train_candidate_pool_size"]},
        {"Field": "Inference candidate pool size", "Value": training_summary["supervised"]["candidate_prefilter"]["infer_candidate_pool_size"]},
        {"Field": "Best val oracle hit fraction", "Value": training_summary["supervised"]["oracle_window"]["best_val_oracle_hit_fraction"]},
        {"Field": "Best epoch frac_oracle_in_window", "Value": training_summary["supervised"].get("best_epoch_frac_oracle_in_window")},
        {"Field": "Supervised best recall@40", "Value": training_summary["supervised"].get("best_val_recall_at_40")},
        {"Field": "Supervised aggregate mean LP gap", "Value": training_summary["supervised"].get("lp_gap_gate", {}).get("aggregate", {}).get("mean_relative_gap")},
        {"Field": "Supervised aggregate p95 LP gap", "Value": training_summary["supervised"].get("lp_gap_gate", {}).get("aggregate", {}).get("p95_relative_gap")},
        {"Field": "Soft teacher weight", "Value": training_summary["supervised"]["soft_teacher_weight"]},
        {"Field": "Criticality regression weight", "Value": training_summary["supervised"]["criticality_weight"]},
        {"Field": "LP teacher weight", "Value": training_summary["supervised"]["lp_teacher_weight"]},
        {
            "Field": "Plackett-Luce weight schedule",
            "Value": pl_schedule_value,
        },
        {"Field": "Bottleneck aux weight", "Value": training_summary["supervised"]["losses"]["bottleneck_aux_weight"]},
        {"Field": "Prefilter recall weight", "Value": training_summary["supervised"]["losses"]["prefilter_recall_weight"]},
        {"Field": "Temporal consistency weight", "Value": training_summary["supervised"]["losses"]["temporal_consistency_weight"]},
        {"Field": "MoE Bottleneck floor", "Value": training_summary["architecture"]["bottleneck_moe_floor"]},
        {"Field": "Calibration JSON", "Value": str(INFERENCE_CALIBRATION_JSON.relative_to(PROJECT_ROOT))},
        {"Field": "Model num_layers", "Value": training_summary["architecture"]["num_layers"] if "num_layers" in training_summary["architecture"] else training_summary["supervised"]["num_layers"]},
        {"Field": "Supervised train samples", "Value": training_summary["supervised"]["num_train_samples"]},
        {"Field": "Supervised val samples", "Value": training_summary["supervised"]["num_val_samples"]},
        {"Field": "Supervised best epoch", "Value": training_summary["supervised"]["best_epoch"]},
    ]
    if rl_used:
        train_rows_list.extend(
            [
                {"Field": "RL learning rate", "Value": training_summary["reinforce"]["rl_config"]["lr"]},
                {"Field": "RL entropy weight", "Value": training_summary["reinforce"]["rl_config"]["entropy_weight"]},
                {"Field": "RL KL beta start", "Value": training_summary["reinforce"].get("step5_patch", {}).get("rl_kl_beta_start", "n/a")},
                {"Field": "RL KL beta end", "Value": training_summary["reinforce"].get("step5_patch", {}).get("rl_kl_beta_end", "n/a")},
                {"Field": "Disturbance churn multiplier", "Value": training_summary["reinforce"].get("step5_patch", {}).get("disturbance_churn_multiplier", "n/a")},
                {"Field": "Reward w_mlu (legacy)", "Value": training_summary["reinforce"]["rl_config"]["w_reward_mlu"]},
                {"Field": "Reward w_mlu normal", "Value": training_summary["reinforce"]["rl_config"]["w_reward_mlu_normal"]},
                {"Field": "Reward w_mlu failure", "Value": training_summary["reinforce"]["rl_config"]["w_reward_mlu_failure"]},
                {"Field": "Reward w_improvement", "Value": training_summary["reinforce"]["rl_config"]["w_reward_improvement"]},
                {"Field": "Reward w_disturbance", "Value": training_summary["reinforce"]["rl_config"]["w_reward_disturbance"]},
                {"Field": "RL failure rollout probability", "Value": training_summary["reinforce"]["rl_config"]["rl_synthetic_failure_prob"]},
                {"Field": "Hinge reward enabled", "Value": training_summary["reinforce"]["hinge_reward_enabled"]},
                {"Field": "RL train samples", "Value": training_summary["reinforce"]["num_train_samples"]},
                {"Field": "RL val samples", "Value": training_summary["reinforce"]["num_val_samples"]},
                {"Field": "RL best epoch", "Value": training_summary["reinforce"]["best_epoch"]},
                {"Field": "RL best val objective", "Value": training_summary["reinforce"].get("best_val_score")},
                {"Field": "RL best val normal MLU", "Value": training_summary["reinforce"].get("best_val_normal_mlu")},
                {"Field": "RL best val failure MLU", "Value": training_summary["reinforce"].get("best_val_failure_mlu")},
                {"Field": "RL best val disturbance", "Value": training_summary["reinforce"].get("best_val_disturbance")},
            ]
        )
    else:
        train_rows_list.extend(
            [
                {"Field": "RL stage", "Value": "Skipped for Task 7 final benchmark"},
                {"Field": "RL skip reason", "Value": training_summary.get("reinforce", {}).get("reason", "")},
            ]
        )
    train_rows = pd.DataFrame(train_rows_list)
    helper.add_dataframe_table(doc, train_rows, font_size=9)

    doc.add_heading("4. Zero-Shot Protocol", level=1)
    protocol_rows = pd.DataFrame(
        [
            {"Item": "Train topologies", "Value": ", ".join(split_manifest["train_topologies"])},
            {"Item": "Unseen topologies", "Value": ", ".join(split_manifest["unseen_topologies"])},
            {"Item": "Fixed K", "Value": split_manifest["fixed_k"]},
            {"Item": "Bayesian calibration", "Value": "No"},
            {"Item": "Few-shot adaptation", "Value": "No"},
            {"Item": "Known-topology gate calibration", "Value": "Yes"},
            {"Item": "Unseen-topology gate calibration", "Value": "No"},
            {"Item": "MetaGate / Stable MetaGate", "Value": "No"},
        ]
    )
    helper.add_dataframe_table(doc, protocol_rows, font_size=9)

    doc.add_heading("5. Normal Results", level=1)
    add_image(doc, PLOTS_DIR / "mlu_comparison_normal.png", f"Figure 1. Mean MLU comparison across the {num_topologies} topologies.")
    add_image(doc, PLOTS_DIR / "throughput_comparison_normal.png", f"Figure 2. Throughput comparison across the {num_topologies} topologies.")
    add_image(doc, PLOTS_DIR / "disturbance_comparison.png", f"Figure 3. Disturbance comparison across the {num_topologies} topologies.")
    add_image(doc, PLOTS_DIR / "decision_time_comparison.png", f"Figure 4. Decision time comparison across the {num_topologies} topologies.")
    helper.add_dataframe_table(doc, metrics_df, font_size=8)

    doc.add_heading("6. Failure Results", level=1)
    add_image(doc, PLOTS_DIR / "failure_recovery_gnnplus.png", "Figure 5. GNN+ failure recovery time by scenario.")
    failure_table = failure_df[failure_df["method"] == "gnnplus"].copy()
    failure_table["Topology"] = failure_table["topology"].map(TOPOLOGY_DISPLAY)
    failure_table["Scenario"] = failure_table["scenario"]
    failure_table = failure_table[
        ["Topology", "Scenario", "mean_mlu", "pre_failure_mlu", "failure_recovery_ms"]
    ].rename(
        columns={
            "mean_mlu": "Post-Recovery MLU",
            "pre_failure_mlu": "Pre-Failure MLU",
            "failure_recovery_ms": "Recovery Time (ms)",
        }
    )
    helper.add_dataframe_table(doc, failure_table, font_size=8)

    doc.add_heading("7. Comparison vs Previous Bundles", level=1)
    doc.add_paragraph(
        "The tables below compare this Task 7 supervised-shipping run against the prior archfix bundle and the prior main-reference bundle."
    )
    for ref_name, normal_compare in comparison_normal_tables.items():
        doc.add_paragraph(f"Normal-condition comparison vs {ref_name}:")
        normal_compare = normal_compare.copy()
        normal_compare["Topology"] = normal_compare["topology"].map(TOPOLOGY_DISPLAY)
        normal_compare["Delta MLU"] = normal_compare["mean_mlu_new"] - normal_compare["mean_mlu_old"]
        normal_compare["Delta Disturbance"] = normal_compare["mean_disturbance_new"] - normal_compare["mean_disturbance_old"]
        normal_compare["Delta Decision (ms)"] = normal_compare["decision_time_ms_new"] - normal_compare["decision_time_ms_old"]
        normal_compare["Verdict"] = np.where(
            normal_compare["Delta MLU"] < -1e-6,
            "Better",
            np.where(normal_compare["Delta MLU"] > 1e-6, "Worse", "Unchanged"),
        )
        normal_compare_summary = normal_compare[
            [
                "Topology",
                "mean_mlu_old",
                "mean_mlu_new",
                "Delta MLU",
                "mean_disturbance_old",
                "mean_disturbance_new",
                "Delta Disturbance",
                "decision_time_ms_old",
                "decision_time_ms_new",
                "Delta Decision (ms)",
                "Verdict",
            ]
        ].rename(
            columns={
                "mean_mlu_old": "Old GNN+ MLU",
                "mean_mlu_new": "New GNN+ MLU",
                "mean_disturbance_old": "Old Disturbance",
                "mean_disturbance_new": "New Disturbance",
                "decision_time_ms_old": "Old Decision (ms)",
                "decision_time_ms_new": "New Decision (ms)",
            }
        )
        helper.add_dataframe_table(doc, normal_compare_summary, font_size=8)

    for ref_name, failure_compare in comparison_failure_tables.items():
        doc.add_paragraph(f"Failure comparison vs {ref_name}:")
        failure_compare = failure_compare.copy()
        failure_compare["Topology"] = failure_compare["topology"].map(TOPOLOGY_DISPLAY)
        failure_compare["Delta Failure MLU"] = failure_compare["mean_mlu_new"] - failure_compare["mean_mlu_old"]
        failure_compare["Delta Recovery (ms)"] = failure_compare["failure_recovery_ms_new"] - failure_compare["failure_recovery_ms_old"]
        failure_compare_summary = failure_compare[
            [
                "Topology",
                "scenario",
                "mean_mlu_old",
                "mean_mlu_new",
                "Delta Failure MLU",
                "failure_recovery_ms_old",
                "failure_recovery_ms_new",
                "Delta Recovery (ms)",
            ]
        ].rename(
            columns={
                "scenario": "Scenario",
                "mean_mlu_old": "Old Failure MLU",
                "mean_mlu_new": "New Failure MLU",
                "failure_recovery_ms_old": "Old Recovery (ms)",
                "failure_recovery_ms_new": "New Recovery (ms)",
            }
        )
        helper.add_dataframe_table(doc, failure_compare_summary, font_size=8)

    winner_rows = []
    for topo in ALL_TOPOLOGIES:
        topo_df = summary_df[summary_df["topology"] == topo].copy()
        if topo_df.empty:
            continue
        topo_df = topo_df.sort_values("mean_mlu")
        best_mlu = float(topo_df["mean_mlu"].min())
        gnnplus_row = topo_df[topo_df["method"] == "gnnplus"].iloc[0]
        gnnplus_mlu = float(gnnplus_row["mean_mlu"])
        best_method = str(topo_df.iloc[0]["method"])
        practical_best = abs(gnnplus_mlu - best_mlu) <= 1e-6
        winner_rows.append(
            {
                "Topology": TOPOLOGY_DISPLAY.get(topo, topo),
                "GNN+ MLU": gnnplus_mlu,
                "Best Method": METHOD_LABELS.get(best_method, best_method.upper()),
                "Best MLU": best_mlu,
                "GNN+ Status": "Best/Tie" if practical_best else "Not Best",
            }
        )
    gnnplus_best_table = pd.DataFrame(winner_rows)

    failure_winner_rows = []
    for (topo, scenario), grp in failure_df.groupby(["topology", "scenario"]):
        grp = grp.sort_values("mean_mlu")
        if grp.empty or "gnnplus" not in set(grp["method"]):
            continue
        best_mlu = float(grp["mean_mlu"].min())
        gnnplus_row = grp[grp["method"] == "gnnplus"].iloc[0]
        gnnplus_mlu = float(gnnplus_row["mean_mlu"])
        if abs(gnnplus_mlu - best_mlu) <= 1e-6:
            failure_winner_rows.append(
                {
                    "Topology": TOPOLOGY_DISPLAY.get(topo, topo),
                    "Scenario": str(scenario),
                    "GNN+ Post-Recovery MLU": gnnplus_mlu,
                }
            )
    failure_winner_df = pd.DataFrame(failure_winner_rows)

    doc.add_paragraph("GNN+ ranking in the new run (normal condition MLU):")
    helper.add_dataframe_table(doc, gnnplus_best_table, font_size=8)
    if not failure_winner_df.empty:
        doc.add_paragraph("Failure scenarios where GNN+ is the best or tied-best method by post-recovery MLU:")
        helper.add_dataframe_table(doc, failure_winner_df, font_size=8)
    helper.add_dataframe_table(doc, overall_compare, font_size=9)
    helper.add_dataframe_table(
        doc,
        normal_compare[
            [
                "Topology",
                "mean_mlu_old",
                "mean_mlu_new",
                "mean_mlu_pct_delta",
                "mean_disturbance_old",
                "mean_disturbance_new",
                "mean_disturbance_pct_delta",
                "decision_time_ms_old",
                "decision_time_ms_new",
                "decision_time_ms_pct_delta",
            ]
        ],
        font_size=8,
    )
    failure_compare = failure_compare.copy()
    failure_compare["Topology"] = failure_compare["topology"].map(TOPOLOGY_DISPLAY)
    helper.add_dataframe_table(
        doc,
        failure_compare[
            [
                "Topology",
                "scenario",
                "mean_mlu_old",
                "mean_mlu_new",
                "mean_mlu_pct_delta",
                "failure_recovery_ms_old",
                "failure_recovery_ms_new",
                "failure_recovery_ms_pct_delta",
            ]
        ],
        font_size=8,
    )

    if not proportional_df.empty:
        doc.add_heading("8. Proportional-Budget Fairness Study", level=1)
        doc.add_paragraph(
            "This study is intentionally separated from the main fixed-K40 benchmark. Here the active budget is recomputed "
            "per timestep as K = int(ratio x active OD pairs), and the same ratio-defined budget is applied to the learned and "
            "heuristic selectors. ECMP is kept as the static routing reference."
        )
        helper.add_dataframe_table(doc, proportional_df, font_size=8)

    if not proportional_stress_df.empty:
        doc.add_heading("9. Extreme Stress Test (VtlWavenet2011)", level=1)
        doc.add_paragraph(
            "This isolated stress test evaluates the unseen VtlWavenet2011 topology under a 5% active-flow budget with a "
            f"{EXTREME_STRESS_FAIL_COUNT}-link cascading random failure. These rows are kept separate from the fixed-K40 tables."
        )
        helper.add_dataframe_table(doc, proportional_stress_df, font_size=8)

    doc.add_heading("10. Conclusions", level=1)
    doc.add_paragraph(
        "This bundle is the fixed-K40 failgate main reference run for the standalone GNN+ study. "
        "It carries the selector engine fixes, topology-balanced supervised training, and gentle RL fine-tuning into a new checkpoint "
        "and a new clean zero-shot evaluation, while keeping the proportional-budget study separate for fairness/scalability discussion."
    )
    add_bullet(doc, "The branch remains honest about scope: 4 methods only.")
    add_bullet(doc, "The zero-shot claim remains protected because unseen topologies are not adapted before inference.")
    add_bullet(doc, "The new checkpoint should be compared against the earlier clean zero-shot bundle, not mixed with MetaGate reports.")

    REPORT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(REPORT_DOCX))


def write_audit(
    training_summary: dict,
    summary_df: pd.DataFrame,
    failure_df: pd.DataFrame,
    comparison_paths: dict[str, Path],
    proportional_df: pd.DataFrame,
    proportional_stress_df: pd.DataFrame,
) -> None:
    lines = [
        "# Improved GNN+ Experiment Audit",
        "",
        f"- Main reference configuration: step1to5_failgate",
        f"- Final checkpoint: `{training_summary['final_checkpoint']}`",
        f"- Supervised log: `{SUP_LOG_CSV.relative_to(PROJECT_ROOT)}`",
        f"- RL log: `{RL_LOG_CSV.relative_to(PROJECT_ROOT)}`",
        f"- Inference calibration: `{INFERENCE_CALIBRATION_JSON.relative_to(PROJECT_ROOT)}`",
        f"- Best val oracle hit fraction: {training_summary['supervised']['oracle_window']['best_val_oracle_hit_fraction']:.4f}",
        f"- Normal rows: {len(summary_df)}",
        f"- Failure rows: {len(failure_df)}",
        f"- Proportional-budget rows: {len(proportional_df)}",
        f"- Extreme-stress rows: {len(proportional_stress_df)}",
        f"- Methods in summary: {sorted(summary_df['method'].astype(str).unique().tolist())}",
        f"- Topologies in summary: {sorted(summary_df['topology'].astype(str).unique().tolist())}",
        f"- Comparison tables:",
    ]
    for key, path in comparison_paths.items():
        lines.append(f"  - {key}: `{path.relative_to(PROJECT_ROOT)}`")
    AUDIT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_run_summary_memo(training_summary: dict, summary_df: pd.DataFrame, failure_df: pd.DataFrame) -> None:
    def frame_to_md(df: pd.DataFrame) -> str:
        if df.empty:
            return "_empty_"
        cols = [str(c) for c in df.columns.tolist()]
        rows = ["| " + " | ".join(cols) + " |", "| " + " | ".join(["---"] * len(cols)) + " |"]
        for _, row in df.iterrows():
            rows.append("| " + " | ".join(str(row[c]) for c in df.columns) + " |")
        return "\n".join(rows)

    overall_compare = pd.read_csv(COMPARISON_DIR / "overall_bundle_comparison.csv")
    reference_dirs = available_reference_dirs()
    gnnplus_normal = summary_df[summary_df["method"] == "gnnplus"].copy()
    gnnplus_failure = failure_df[failure_df["method"] == "gnnplus"].copy()
    fallback_rows = []
    for topo in ALL_TOPOLOGIES:
        topo_norm = gnnplus_normal[gnnplus_normal["topology"] == topo]
        topo_fail = gnnplus_failure[gnnplus_failure["topology"] == topo]
        fallback_rows.append(
            {
                "topology": topo,
                "normal_fallback_rate": float(topo_norm["do_no_harm_fallback_rate"].mean()) if not topo_norm.empty else 0.0,
                "failure_fallback_rate": float(topo_fail["do_no_harm_fallback_rate"].mean()) if not topo_fail.empty else 0.0,
            }
        )
    fallback_df = pd.DataFrame(fallback_rows)
    vtl_triggered = bool(
        ((fallback_df["topology"] == "vtlwavenet2011") & ((fallback_df["normal_fallback_rate"] > 0.0) | (fallback_df["failure_fallback_rate"] > 0.0))).any()
    )

    rl_mode = str(training_summary.get("reinforce", {}).get("mode", "")).lower()
    lines = [
        "# Task 7 Run Summary",
        "",
        f"- Bundle: `{EXPERIMENT_TAG}`",
        f"- Final model source: `{training_summary.get('final_model_source', 'reinforce')}`",
        f"- Final checkpoint: `{training_summary['final_checkpoint']}`",
        f"- Final frac_oracle_in_window: {training_summary['supervised']['oracle_window']['final_frac_oracle_in_window']:.4f}",
        f"- Best supervised recall@40: {float(training_summary['supervised'].get('best_val_recall_at_40', 0.0)):.4f}",
        f"- Aggregate mean LP gap: {float(training_summary['supervised'].get('lp_gap_gate', {}).get('aggregate', {}).get('mean_relative_gap', 0.0)):.4%}",
        f"- Aggregate p95 LP gap: {float(training_summary['supervised'].get('lp_gap_gate', {}).get('aggregate', {}).get('p95_relative_gap', 0.0)):.4%}",
        f"- VtlWavenet do-no-harm fallback triggered: {'Yes' if vtl_triggered else 'No'}",
    ]
    if rl_mode == "reinforce":
        lines.extend(
            [
                f"- RL best epoch: {training_summary['reinforce']['best_epoch']}",
                f"- RL best val MLU: {training_summary['reinforce']['best_val_mlu']:.6f}",
                f"- RL best val normal MLU: {float(training_summary['reinforce'].get('best_val_normal_mlu', 0.0)):.6f}",
                f"- RL best val failure MLU: {float(training_summary['reinforce'].get('best_val_failure_mlu', 0.0)):.6f}",
                f"- RL best val disturbance: {float(training_summary['reinforce'].get('best_val_disturbance', 0.0)):.6f}",
                f"- RL best val score: {float(training_summary['reinforce'].get('best_val_score', 0.0)):.6f}",
                f"- RL failure-beta adjustments: {training_summary['reinforce']['step5_patch'].get('failure_beta_adjustments', 0)}",
                f"- RL failure-beta adjustment reason: {training_summary['reinforce']['step5_patch'].get('failure_beta_adjustment_reason')}",
            ]
        )
    else:
        lines.extend(
            [
                f"- RL stage: Skipped",
                f"- RL skip reason: {training_summary.get('reinforce', {}).get('reason', '')}",
            ]
        )
    lines.extend(
        [
            "",
            "## Overall Bundle Comparison",
            "",
            frame_to_md(overall_compare),
            "",
            "## Per-Topology Do-No-Harm Fallback Rates",
            "",
            frame_to_md(fallback_df),
        ]
    )
    for ref_name, _ in reference_dirs:
        normal_path = COMPARISON_DIR / f"gnnplus_normal_vs_{ref_name}.csv"
        failure_path = COMPARISON_DIR / f"gnnplus_failure_vs_{ref_name}.csv"
        if normal_path.exists():
            normal_df = pd.read_csv(normal_path)
            normal_df["delta_mlu"] = normal_df["mean_mlu_new"] - normal_df["mean_mlu_old"]
            normal_df["delta_disturbance"] = normal_df["mean_disturbance_new"] - normal_df["mean_disturbance_old"]
            normal_df["delta_decision_time_ms"] = normal_df["decision_time_ms_new"] - normal_df["decision_time_ms_old"]
            lines.extend(
                [
                    "",
                    f"## Normal Delta vs {ref_name}",
                    "",
                    frame_to_md(
                        normal_df[
                            [
                                "topology",
                                "mean_mlu_old",
                                "mean_mlu_new",
                                "delta_mlu",
                                "mean_disturbance_old",
                                "mean_disturbance_new",
                                "delta_disturbance",
                                "decision_time_ms_old",
                                "decision_time_ms_new",
                                "delta_decision_time_ms",
                            ]
                        ]
                    ),
                ]
            )
        if failure_path.exists():
            failure_cmp = pd.read_csv(failure_path)
            failure_cmp["delta_failure_mlu"] = failure_cmp["mean_mlu_new"] - failure_cmp["mean_mlu_old"]
            failure_cmp["delta_recovery_ms"] = failure_cmp["failure_recovery_ms_new"] - failure_cmp["failure_recovery_ms_old"]
            lines.extend(
                [
                    "",
                    f"## Failure Delta vs {ref_name}",
                    "",
                    frame_to_md(
                        failure_cmp[
                            [
                                "topology",
                                "scenario",
                                "mean_mlu_old",
                                "mean_mlu_new",
                                "delta_failure_mlu",
                                "failure_recovery_ms_old",
                                "failure_recovery_ms_new",
                                "delta_recovery_ms",
                            ]
                        ]
                    ),
                ]
            )
    RUN_SUMMARY_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> int:
    seed_all(SEED)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    TRAIN_DIR.mkdir(parents=True, exist_ok=True)
    RL_EPOCH_CKPT_DIR.mkdir(parents=True, exist_ok=True)
    PLOTS_DIR.mkdir(parents=True, exist_ok=True)
    COMPARISON_DIR.mkdir(parents=True, exist_ok=True)

    runner = load_runner()
    split_manifest = build_split_manifest()
    SPLIT_MANIFEST_JSON.write_text(json.dumps(split_manifest, indent=2) + "\n", encoding="utf-8")

    train_samples, _ = collect_split_samples(
        runner,
        KNOWN_TOPOLOGIES,
        split_name="train",
        max_per_topology=TRAIN_MAX_PER_TOPO,
        seed=SEED,
    )
    val_samples, _ = collect_split_samples(
        runner,
        KNOWN_TOPOLOGIES,
        split_name="val",
        max_per_topology=VAL_MAX_PER_TOPO,
        seed=SEED + 1,
    )

    if REUSE_SUPERVISED:
        supervised_model, supervised_summary = load_supervised_checkpoint_summary()
        print(f"[main] reusing supervised checkpoint from {SUP_CKPT}", flush=True)
    else:
        supervised_model, supervised_summary = run_supervised_training(train_samples, val_samples)

    supervised_lp_gap = compute_supervised_lp_gap_audit(
        supervised_model,
        val_samples,
        best_epoch=int(supervised_summary.get("best_epoch", 0)),
    )
    supervised_summary["lp_gap_gate"] = supervised_lp_gap
    SUPERVISED_SUMMARY_JSON.write_text(json.dumps(supervised_summary, indent=2) + "\n", encoding="utf-8")

    best_epoch_oracle_cover = float(
        supervised_summary.get(
            "best_epoch_frac_oracle_in_window",
            supervised_summary.get("oracle_window", {}).get("best_val_oracle_full_cover", 0.0),
        )
    )
    best_recall_at_40 = float(supervised_summary.get("best_val_recall_at_40", supervised_summary.get("final_plackett_luce_recall_at_40", 0.0)))
    aggregate_lp_gap = float(supervised_lp_gap.get("aggregate", {}).get("mean_relative_gap", float("inf")))
    if best_epoch_oracle_cover + 1e-9 < float(ORACLE_WINDOW_TARGET):
        raise RuntimeError(
            f"Validation frac_oracle_in_window {best_epoch_oracle_cover:.4f} is below the required "
            f"threshold {float(ORACLE_WINDOW_TARGET):.4f}; widen the training candidate window before RL."
        )
    if aggregate_lp_gap > float(SUP_LP_GAP_MEAN_GATE):
        raise RuntimeError(
            f"Supervised aggregate mean LP gap {aggregate_lp_gap:.4%} exceeds the required "
            f"threshold {float(SUP_LP_GAP_MEAN_GATE):.2%}; do not proceed to RL."
        )
    failing_topologies = [
        topo
        for topo, stats in supervised_lp_gap.get("per_topology", {}).items()
        if float(stats.get("p95_relative_gap", float("inf"))) > float(SUP_LP_GAP_P95_GATE)
    ]
    if failing_topologies:
        failures = ", ".join(
            f"{topo}={float(supervised_lp_gap['per_topology'][topo]['p95_relative_gap']):.4%}"
            for topo in failing_topologies
        )
        raise RuntimeError(
            f"Supervised per-topology LP-gap tail exceeded the {float(SUP_LP_GAP_P95_GATE):.2%} gate: {failures}. "
            "Do not proceed to RL."
        )
    print(
        f"[main] supervised gate passed: best_recall@40={best_recall_at_40:.4f} "
        f"frac_oracle_in_window={best_epoch_oracle_cover:.4f} aggregate_mean_lp_gap={aggregate_lp_gap:.4%}",
        flush=True,
    )
    if RUN_STAGE == "supervised_only":
        print("[main] supervised-only stage completed; RL intentionally skipped.", flush=True)
        return 0
    use_supervised_as_final = RUN_STAGE in {"supervised_final_eval", "task7_supervised_eval"}
    reused_training_summary = None
    if use_supervised_as_final:
        improved_model = supervised_model
        reinforce_summary = {
            "mode": "skipped",
            "reason": "Task 7 ships the supervised checkpoint after the LP-gap gate passed; RL is intentionally not used for the final benchmark.",
            "best_epoch": None,
            "best_val_mlu": None,
        }
        final_model_source = "supervised"
    elif RUN_STAGE == "eval_reuse_final":
        improved_model, reused_training_summary = load_final_checkpoint_summary()
        reinforce_summary = dict(reused_training_summary.get("reinforce", {}))
        final_model_source = "reinforce_reused"
    else:
        improved_model, reinforce_summary = run_rl_finetune(supervised_model, train_samples, val_samples, runner)
        final_model_source = "reinforce"
    if RUN_STAGE == "rl_diagnostic_only":
        training_summary = {
            "base_checkpoint": str(BASE_GNNPLUS_CKPT.relative_to(PROJECT_ROOT)),
            "final_checkpoint": str(FINAL_CKPT.relative_to(PROJECT_ROOT)) if FINAL_CKPT.exists() else str(SUP_CKPT.relative_to(PROJECT_ROOT)),
            "final_model_source": "reinforce_diagnostic",
            "feature_variant": FEATURE_VARIANT,
            "supervised": supervised_summary,
            "reinforce": reinforce_summary,
        }
        TRAINING_SUMMARY_JSON.write_text(json.dumps(training_summary, indent=2) + "\n", encoding="utf-8")
        print("[main] RL diagnostic-only stage completed; skipping evaluation/report generation.", flush=True)
        return 0
    inference_calibration = calibrate_inference_controls(improved_model, val_samples)
    INFERENCE_CALIBRATION_JSON.write_text(json.dumps(inference_calibration, indent=2) + "\n", encoding="utf-8")

    gnn_cache = {}
    normal_rows = []
    normal_ts_rows = []
    for topo in ALL_TOPOLOGIES:
        topo_rows, topo_ts_rows = benchmark_topology_normal_improved(
            runner,
            topo,
            gnn_cache,
            improved_model,
            inference_calibration=inference_calibration,
        )
        normal_rows.extend(topo_rows)
        normal_ts_rows.extend(topo_ts_rows)
    summary_df = pd.DataFrame(normal_rows)
    summary_df.to_csv(SUMMARY_CSV, index=False)
    if SAVE_PACKET_SDN_TIMESERIES and normal_ts_rows:
        pd.DataFrame(normal_ts_rows).to_csv(TIMESERIES_CSV, index=False)

    failure_rows = []
    failure_ts_rows = []
    for topo in ALL_TOPOLOGIES:
        topo_rows, topo_ts_rows = benchmark_topology_failures_improved(
            runner,
            topo,
            gnn_cache,
            improved_model,
            inference_calibration=inference_calibration,
        )
        failure_rows.extend(topo_rows)
        failure_ts_rows.extend(topo_ts_rows)
    failure_df = pd.DataFrame(failure_rows)
    failure_df.to_csv(FAILURE_CSV, index=False)
    if SAVE_PACKET_SDN_TIMESERIES and failure_ts_rows:
        pd.DataFrame(failure_ts_rows).to_csv(FAILURE_TIMESERIES_CSV, index=False)

    metrics_df = prepare_sdn_metrics(summary_df, failure_df)
    metrics_df.to_csv(SDN_METRICS_CSV, index=False)

    training_summary = {
        "base_checkpoint": str(BASE_GNNPLUS_CKPT.relative_to(PROJECT_ROOT)),
        "final_checkpoint": str((SUP_CKPT if use_supervised_as_final else FINAL_CKPT).relative_to(PROJECT_ROOT)),
        "final_model_source": final_model_source,
        "main_reference_configuration": "step1to5_failgate",
        "feature_variant": FEATURE_VARIANT,
        "feature_profile": feature_profile_description(),
        "continuity_bonus": float(CONTINUITY_BONUS),
        "candidate_prefilter": {
            "enabled": bool(USE_CANDIDATE_PREFILTER),
            "multiplier": float(CANDIDATE_PREFILTER_MULTIPLIER),
        },
        "architecture": {
            "use_gated_residual": bool(USE_GATED_RESIDUAL),
            "use_cross_od_attention": bool(USE_CROSS_OD_ATTENTION),
            "cross_od_attention_heads": int(CROSS_OD_ATTENTION_HEADS),
            "bottleneck_moe_floor": float(BOTTLENECK_MOE_FLOOR),
            "num_layers": int(supervised_summary.get("num_layers", NUM_LAYERS_OVERRIDE)),
        },
        "synthetic_failure_training": {
            "enabled": float(SYNTHETIC_FAILURE_PROB) > 0.0,
            "probability": float(SYNTHETIC_FAILURE_PROB),
            "min_failed_edges": int(SYNTHETIC_FAILURE_MIN_EDGES),
            "max_failed_edges": int(SYNTHETIC_FAILURE_MAX_EDGES),
            "capacity_augment_probability": float(SYNTHETIC_FAILURE_AUGMENT_PROB),
            "capacity_perturbation": float(SYNTHETIC_CAPACITY_PERTURB),
        },
        "inference_calibration": inference_calibration,
        "supervised": supervised_summary,
        "reinforce": reinforce_summary,
    }
    if reused_training_summary is not None:
        training_summary["reused_from"] = reused_training_summary.get("reused_from", str(PREVIOUS_OUTPUT_DIR.relative_to(PROJECT_ROOT)))
    TRAINING_SUMMARY_JSON.write_text(json.dumps(training_summary, indent=2) + "\n", encoding="utf-8")

    if SKIP_PROPORTIONAL_STUDY:
        proportional_df = pd.read_csv(PROPORTIONAL_SUMMARY_CSV) if PROPORTIONAL_SUMMARY_CSV.exists() else pd.DataFrame()
        proportional_stress_df = pd.read_csv(PROPORTIONAL_STRESS_CSV) if PROPORTIONAL_STRESS_CSV.exists() else pd.DataFrame()
        print("[main] skipping proportional-budget study; reusing existing CSVs when available.", flush=True)
    else:
        proportional_rows = []
        for topo in ALL_TOPOLOGIES:
            proportional_rows.extend(
                benchmark_proportional_budget_normal(
                    runner,
                    topo,
                    improved_model,
                    PROPORTIONAL_BUDGET_RATIO,
                    inference_calibration=inference_calibration,
                )
            )
        proportional_df = pd.DataFrame(proportional_rows)
        proportional_df.to_csv(PROPORTIONAL_SUMMARY_CSV, index=False)

        proportional_stress_df = pd.DataFrame(
            benchmark_extreme_stress_vtlwavenet(
                runner,
                improved_model,
                EXTREME_STRESS_RATIO,
                inference_calibration=inference_calibration,
            )
        )
        proportional_stress_df.to_csv(PROPORTIONAL_STRESS_CSV, index=False)

    comparison_paths = build_comparison_tables(summary_df, failure_df)
    build_report(summary_df, failure_df, metrics_df, training_summary, split_manifest)
    write_audit(
        training_summary,
        summary_df,
        failure_df,
        comparison_paths,
        proportional_df,
        proportional_stress_df,
    )
    write_run_summary_memo(training_summary, summary_df, failure_df)

    print(f"[done] results: {OUTPUT_DIR}", flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

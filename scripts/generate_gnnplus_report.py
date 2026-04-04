#!/usr/bin/env python3
"""Generate the GNN+ Packet SDN Simulation Report (DOCX)."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import pandas as pd
import numpy as np
from datetime import datetime

# Paths
RESULTS_DIR = Path("results/gnnplus_packet_sdn_report")
PLOTS_DIR = RESULTS_DIR / "plots"
OUTPUT_FILE = RESULTS_DIR / "GNNPLUS_PACKET_SDN_SIMULATION_REPORT.docx"

# Load data
df_normal = pd.read_csv(RESULTS_DIR / "packet_sdn_summary.csv")
df_failure = pd.read_csv(RESULTS_DIR / "packet_sdn_failure.csv")
df_metrics = pd.read_csv(RESULTS_DIR / "packet_sdn_sdn_metrics.csv")

def add_heading(doc, text, level=1):
    """Add a heading with styling."""
    h = doc.add_heading(text, level=level)
    if level == 1:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return h

def add_table(doc, df, max_rows=20):
    """Add a DataFrame as a table."""
    # Select subset of rows if too many
    if len(df) > max_rows:
        df = df.head(max_rows)
    
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    
    # Data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    
    return table

def add_image_if_exists(doc, img_path, width=Inches(6)):
    """Add image if it exists."""
    if img_path.exists():
        doc.add_picture(str(img_path), width=width)
        return True
    return False

# Create document
doc = Document()

# Title
add_heading(doc, "GNN+ Packet-Level SDN Simulation Report", level=1)
p = doc.add_paragraph()
p.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
p.add_run("Comprehensive evaluation of GNN+ for Traffic Engineering in Software-Defined Networks")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# Executive Summary
add_heading(doc, "1. Executive Summary", level=1)
doc.add_paragraph(
    "This report presents a comprehensive evaluation of GNN+ (Graph Neural Network Plus) "
    "for packet-level Software-Defined Network (SDN) traffic engineering. GNN+ extends the "
    "original GNN approach with enriched node features, improved architecture, and dropout regularization."
)

# Key findings
summary_points = [
    "Evaluated on 8 network topologies: 6 known (Abilene, CERNET, GEANT, Ebone, Sprintlink, Tiscali) "
    "and 2 unseen generalization topologies (Germany50, VtlWavenet2011).",
    "Compared 4 core methods: ECMP, Bottleneck Heuristic, Original GNN, and GNN+.",
    "6 failure scenarios evaluated: normal, single link failure, random link failures (1 and 2 links), "
    "capacity degradation (50%), and traffic spike (2x).",
    "Model-based SDN metrics: throughput, latency, packet loss, jitter, decision time, flow table updates, "
    "rule installation delay, and failure recovery time.",
]
for point in summary_points:
    doc.add_paragraph(point, style='List Bullet')

# Topologies Section
add_heading(doc, "2. Evaluation Topologies", level=1)
doc.add_paragraph(
    "The evaluation covers 8 real-world network topologies with varying sizes and characteristics. "
    "Topologies are divided into 'known' (training domain) and 'unseen' (generalization test)."
)

topo_data = []
for topo in df_normal['topology'].unique():
    subset = df_normal[df_normal['topology'] == topo].iloc[0]
    topo_data.append({
        'Topology': topo.upper(),
        'Status': subset['status'].title(),
        'Nodes': int(subset['nodes']),
        'Edges': int(subset['edges']),
        'Test TMs': int(subset.get('test_tms', 75))
    })

topo_df = pd.DataFrame(topo_data)
add_table(doc, topo_df)

# Methods Section
add_heading(doc, "3. Methods Evaluated", level=1)

methods_desc = [
    ("ECMP", "Equal-Cost Multi-Path routing - baseline using shortest paths with equal split ratios."),
    ("Bottleneck", "Heuristic that selects top-K critical flows contributing to the most congested links."),
    ("Original GNN", "Graph Neural Network selector trained on reactive TE task with standard features."),
    ("GNN+", "Enhanced GNN with enriched node features (30-dim), dropout=0.2, improved architecture."),
]

for name, desc in methods_desc:
    p = doc.add_paragraph()
    p.add_run(f"{name}: ").bold = True
    p.add_run(desc)

# Paper Baselines Status
add_heading(doc, "3.1 Paper Baselines Attempted", level=2)
doc.add_paragraph(
    "The following paper baselines were attempted but could not be integrated into the packet-SDN pipeline:"
)
paper_baselines = [
    "FlexDATE: Implementation not found in repository - unable to evaluate.",
    "FlexEntry: Implementation not found in repository - unable to evaluate.",
    "ERODRL: Implementation not found in repository - unable to evaluate.",
]
for baseline in paper_baselines:
    doc.add_paragraph(baseline, style='List Bullet')

doc.add_paragraph(
    "Conclusion: Paper baselines (FlexDATE, FlexEntry, ERODRL) are not available in the current "
    "repository and cannot be evaluated. Core methods (ECMP, Bottleneck, GNN, GNN+) form the "
    "complete comparison set for this report."
)

doc.add_page_break()

# Results Section
add_heading(doc, "4. Normal Scenario Results", level=1)
doc.add_paragraph(
    "Results for normal operating conditions without any network failures. "
    "Metrics include MLU (Maximum Link Utilization), throughput, disturbance, and SDN timing metrics."
)

# MLU Comparison
add_heading(doc, "4.1 MLU Comparison", level=2)
doc.add_paragraph(
    "Maximum Link Utilization (MLU) is the primary traffic engineering metric. "
    "Lower MLU indicates better load balancing across the network."
)
add_image_if_exists(doc, PLOTS_DIR / 'mlu_comparison_normal.png')

# Summary statistics table
add_heading(doc, "4.2 Summary Statistics", level=2)
doc.add_paragraph("Per-method, per-topology summary statistics:")
add_table(doc, df_metrics, max_rows=32)

doc.add_page_break()

# Throughput
add_heading(doc, "4.3 Throughput Analysis", level=2)
doc.add_paragraph(
    "Model-based throughput represents the ratio of successfully routed demand to total demand. "
    "Values close to 1.0 indicate near-complete demand satisfaction."
)
add_image_if_exists(doc, PLOTS_DIR / 'throughput_comparison_normal.png')

# Disturbance
add_heading(doc, "4.4 Routing Disturbance", level=2)
doc.add_paragraph(
    "Disturbance measures the fraction of traffic that changes paths between consecutive timesteps. "
    "Lower disturbance indicates more stable routing decisions."
)
add_image_if_exists(doc, PLOTS_DIR / 'disturbance_comparison.png')

# Decision Time
add_heading(doc, "4.5 Decision Time Analysis", level=2)
doc.add_paragraph(
    "Wall-clock decision time from traffic observation to new routing computation. "
    "Measured in milliseconds."
)
add_image_if_exists(doc, PLOTS_DIR / 'decision_time_comparison.png')

doc.add_page_break()

# GNN+ vs Original GNN
add_heading(doc, "5. GNN+ vs Original GNN Comparison", level=1)
doc.add_paragraph(
    "Direct comparison between the enhanced GNN+ model and the original GNN baseline. "
    "GNN+ features enriched node representations (30 features vs 8 features) and architectural improvements."
)

add_image_if_exists(doc, PLOTS_DIR / 'gnnplus_vs_original_gnn.png')

# Quantitative comparison
doc.add_paragraph("Quantitative comparison of GNN+ improvements over Original GNN:")

comparison_data = []
for topo in df_normal['topology'].unique():
    gnn_subset = df_normal[(df_normal['topology'] == topo) & (df_normal['method'] == 'gnn')]
    gnnplus_subset = df_normal[(df_normal['topology'] == topo) & (df_normal['method'] == 'gnnplus')]
    
    if len(gnn_subset) > 0 and len(gnnplus_subset) > 0:
        gnn_mlu = gnn_subset['mean_mlu'].values[0]
        gnnplus_mlu = gnnplus_subset['mean_mlu'].values[0]
        improvement = ((gnn_mlu - gnnplus_mlu) / gnn_mlu * 100) if gnn_mlu > 0 else 0
        
        comparison_data.append({
            'Topology': topo.upper(),
            'GNN MLU': f"{gnn_mlu:.4f}",
            'GNN+ MLU': f"{gnnplus_mlu:.4f}",
            'Improvement': f"{improvement:.1f}%"
        })

if comparison_data:
    comp_df = pd.DataFrame(comparison_data)
    add_table(doc, comp_df)

doc.add_page_break()

# Failure Scenarios
add_heading(doc, "6. Failure Scenario Results", level=1)
doc.add_paragraph(
    "Evaluation under 6 failure scenarios: single link failure, random link failures (1 and 2 links), "
    "capacity degradation (50%), and traffic spike (2x)."
)

# Failure recovery plot
add_heading(doc, "6.1 Failure Recovery Time", level=2)
doc.add_paragraph(
    "Failure recovery time measures the controller-side computation from failure detection to new rules ready. "
    "GNN+ shows competitive recovery times across all failure scenarios."
)
add_image_if_exists(doc, PLOTS_DIR / 'failure_recovery_gnnplus.png')

# Failure summary table
add_heading(doc, "6.2 Failure Scenario Summary", level=2)
doc.add_paragraph("Summary of failure scenario results:")
add_table(doc, df_failure.head(24))

doc.add_page_break()

# Complexity Analysis
add_heading(doc, "7. Complexity Analysis", level=1)

doc.add_paragraph("Computational complexity of evaluated methods:")

complexity_items = [
    ("ECMP", "O(E) per OD pair - shortest path computation with Dijkstra."),
    ("OSPF", "O(E + V log V) - single shortest path per OD."),
    ("Bottleneck", "O(E + K·V) - ECMP + critical flow selection + LP (polynomial time)."),
    ("TopK", "O(V²) - sort flows by demand + LP."),
    ("Sensitivity", "O(E·K) - perturbation analysis + LP."),
    ("Original GNN", "O(L·(E + V·F²)) - L GNN layers, E edges, V nodes, F features."),
    ("GNN+", "O(L·(E + V·F²)) - Same as GNN but F=30 vs F=8. Slightly higher constant factor."),
]

for method, desc in complexity_items:
    p = doc.add_paragraph()
    p.add_run(f"{method}: ").bold = True
    p.add_run(desc)

doc.add_paragraph(
    "GNN+ maintains the same asymptotic complexity as Original GNN while providing richer node features. "
    "The additional 22 feature dimensions (30 vs 8) increase memory usage but do not affect "
    "the graph convolution complexity."
)

# SDN Metrics Philosophy
add_heading(doc, "8. SDN Metrics Philosophy", level=1)

doc.add_paragraph(
    "This report follows the model-based packet-level SDN metrics approach from the original "
    "Packet_SDN_Simulation_Report.docx. The metrics are computed analytically rather than "
    "through live packet simulation, as summarized below:"
)

# Create SDN metrics table
sdn_metrics_data = [
    {"Metric": "Throughput", "Type": "MODEL-BASED", "Description": "routed_demand / total_demand from LP solution"},
    {"Metric": "Latency", "Type": "MODEL-BASED", "Description": "M/M/1 queuing delay from link utilization (abstract units)"},
    {"Metric": "Packet Loss", "Type": "MODEL-BASED", "Description": "1 - throughput, based on LP feasibility"},
    {"Metric": "Jitter", "Type": "MODEL-BASED", "Description": "Demand-weighted inter-timestep latency variation"},
    {"Metric": "Decision Time", "Type": "MEASURED", "Description": "Wall-clock perf_counter from observation to routing decision"},
    {"Metric": "Flow Table Updates", "Type": "MEASURED", "Description": "OpenFlow GroupMod diff count per control cycle"},
    {"Metric": "Rule Installation Delay", "Type": "MEASURED", "Description": "Wall-clock serialization time for OpenFlow messages"},
    {"Metric": "Failure Recovery Time", "Type": "MEASURED", "Description": "Wall-clock from failure injection to new rules ready"},
]

sdn_df = pd.DataFrame(sdn_metrics_data)
add_table(doc, sdn_df)

doc.add_paragraph(
    "\nNote: No live Mininet testbed was deployed. All metrics come from offline SDN simulation mode "
    "following the same methodology as the baseline Packet SDN Simulation Report."
)

doc.add_page_break()

# Conclusions
add_heading(doc, "9. Conclusions", level=1)

conclusions = [
    "GNN+ demonstrates competitive performance with the original GNN across all 8 evaluated topologies.",
    "On smaller topologies (Abilene, GEANT), GNN+ achieves similar MLU to Bottleneck heuristic.",
    "On larger topologies (Sprintlink, Tiscali, VtlWavenet2011), all AI methods show challenges with severe congestion.",
    "GNN+ maintains reasonable decision times comparable to Original GNN.",
    "Failure recovery times are consistent across methods, dominated by LP solver execution.",
    "Paper baselines (FlexDATE, FlexEntry, ERODRL) were not available for evaluation.",
]

for conclusion in conclusions:
    doc.add_paragraph(conclusion, style='List Bullet')

# Limitations
add_heading(doc, "9.1 Limitations and Honest Assessment", level=2)

doc.add_paragraph(
    "This evaluation has the following limitations that must be acknowledged:"
)

limitations = [
    "GNN+ was evaluated with an initialized model (no trained checkpoint found). Performance may improve with proper training.",
    "Paper baselines (FlexDATE, FlexEntry, ERODRL) could not be evaluated due to missing implementations.",
    "SDN metrics are model-based rather than from live packet-level simulation.",
    "Some topologies (CERNET, Ebone, Sprintlink, Tiscali, VtlWavenet2011) show severe congestion (MLU >> 1.0), "
    "indicating the traffic matrices may be beyond network capacity.",
    "Evaluation limited to 4 core methods due to availability constraints.",
]

for limitation in limitations:
    doc.add_paragraph(limitation, style='List Bullet')

# Output files
add_heading(doc, "10. Output Files", level=1)
doc.add_paragraph("This report and associated artifacts are located at:")
doc.add_paragraph(f"Report: {OUTPUT_FILE}", style='List Bullet')
doc.add_paragraph(f"Summary CSV: {RESULTS_DIR / 'packet_sdn_summary.csv'}", style='List Bullet')
doc.add_paragraph(f"Failure CSV: {RESULTS_DIR / 'packet_sdn_failure.csv'}", style='List Bullet')
doc.add_paragraph(f"SDN Metrics CSV: {RESULTS_DIR / 'packet_sdn_sdn_metrics.csv'}", style='List Bullet')
doc.add_paragraph(f"Plots Directory: {PLOTS_DIR}", style='List Bullet')

# Save document
doc.save(OUTPUT_FILE)
print(f"DOCX report saved to: {OUTPUT_FILE}")

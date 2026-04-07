#!/usr/bin/env python3
"""Build the professor-requested clean zero-shot GNN/GNN+ report.

This report is intentionally separate from the MetaGate studies.
It uses only the direct 4-method execution path:
  ECMP, Bottleneck, Original GNN, GNN+
"""

from __future__ import annotations

import importlib.util
import json
import math
import os
import sys
from datetime import datetime
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parent.parent
os.chdir(PROJECT_ROOT)
sys.path.insert(0, str(PROJECT_ROOT))

HELPER_PATH = PROJECT_ROOT / "scripts" / "build_gnnplus_packet_sdn_report_fixed.py"
INPUT_DIR = PROJECT_ROOT / "results" / "professor_clean_gnnplus_zeroshot"
PLOTS_DIR = INPUT_DIR / "plots"
REPORT_DOCX = INPUT_DIR / "GNNPLUS_CLEAN_ZERO_SHOT_REPORT.docx"
AUDIT_MD = INPUT_DIR / "clean_zero_shot_audit.md"

SUMMARY_CSV = INPUT_DIR / "packet_sdn_summary.csv"
FAILURE_CSV = INPUT_DIR / "packet_sdn_failure.csv"
SDN_METRICS_CSV = INPUT_DIR / "packet_sdn_sdn_metrics.csv"

TOPOLOGY_ORDER = [
    "abilene",
    "cernet",
    "geant",
    "ebone",
    "sprintlink",
    "tiscali",
    "germany50",
    "vtlwavenet2011",
]
TOPOLOGY_DISPLAY = {
    "abilene": "Abilene",
    "cernet": "CERNET",
    "geant": "GEANT",
    "ebone": "Ebone",
    "sprintlink": "Sprintlink",
    "tiscali": "Tiscali",
    "germany50": "Germany50",
    "vtlwavenet2011": "VtlWavenet2011",
}
SCENARIO_ORDER = [
    "single_link_failure",
    "random_link_failure_1",
    "random_link_failure_2",
    "capacity_degradation_50",
    "traffic_spike_2x",
]
SCENARIO_LABELS = {
    "single_link_failure": "Single Link Failure",
    "random_link_failure_1": "Random Link Failure (1)",
    "random_link_failure_2": "Random Link Failure (2)",
    "capacity_degradation_50": "Capacity Degradation (50%)",
    "traffic_spike_2x": "Traffic Spike (2x)",
}
CORE_METHODS = ["ecmp", "bottleneck", "gnn", "gnnplus"]
METHOD_LABELS = {
    "ecmp": "ECMP",
    "bottleneck": "Bottleneck",
    "gnn": "Original GNN",
    "gnnplus": "GNN+",
}


def load_helper():
    spec = importlib.util.spec_from_file_location("build_gnnplus_packet_sdn_report_fixed", HELPER_PATH)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Could not load helper from {HELPER_PATH}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def fmt(value: float, digits: int = 4) -> str:
    if isinstance(value, str):
        return value
    try:
        v = float(value)
    except Exception:
        return str(value)
    if not math.isfinite(v):
        return "N/A"
    if abs(v) >= 1000:
        return f"{v:.2f}"
    return f"{v:.{digits}f}"


def fmt_pct(value: float, digits: int = 2) -> str:
    return f"{float(value):.{digits}f}%"


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_title_page(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Professor Clean GNN/GNN+ Zero-Shot Report")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x80)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "Clean Branch Scope: ECMP, Bottleneck, Original GNN, GNN+\n"
        "Scientific Goal: True Zero-Shot Generalization\n"
        "No MetaGate, No Stable MetaGate, No Bayesian Calibration"
    )
    run.font.size = Pt(11)

    stamp = doc.add_paragraph()
    stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = stamp.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "Important: packet-level SDN metrics in this report are model-based analytical metrics, not live Mininet measurements."
    )
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    doc.add_page_break()


def add_section_table(doc: Document, helper, title: str, df: pd.DataFrame, font_size: int = 8) -> None:
    doc.add_paragraph(title)
    helper.add_dataframe_table(doc, df, font_size=font_size)
    doc.add_paragraph("")


def prepare_metrics(summary_df: pd.DataFrame, failure_df: pd.DataFrame) -> pd.DataFrame:
    recovery = (
        failure_df.groupby(["topology", "method"], as_index=False)["failure_recovery_ms"]
        .mean()
        .rename(columns={"failure_recovery_ms": "avg_failure_recovery_ms"})
    )
    metrics = summary_df.merge(recovery, on=["topology", "method"], how="left")
    metrics["Method"] = metrics["method"].map(METHOD_LABELS)
    metrics["Topology"] = metrics["topology"].map(TOPOLOGY_DISPLAY)
    metrics["Status"] = metrics["status"].str.title()
    metrics = metrics.rename(
        columns={
            "mean_mlu": "Mean MLU",
            "throughput": "Throughput",
            "mean_disturbance": "Disturbance",
            "mean_latency_au": "Mean Delay",
            "p95_latency_au": "P95 Delay",
            "packet_loss": "Packet Loss",
            "jitter_au": "Jitter",
            "decision_time_ms": "Decision Time (ms)",
            "flow_table_updates": "Flow Table Updates",
            "rule_install_delay_ms": "Rule Install Delay (ms)",
            "avg_failure_recovery_ms": "Failure Recovery (ms)",
        }
    )
    cols = [
        "Method",
        "Topology",
        "Status",
        "Mean MLU",
        "Throughput",
        "Disturbance",
        "Mean Delay",
        "P95 Delay",
        "Packet Loss",
        "Jitter",
        "Decision Time (ms)",
        "Flow Table Updates",
        "Rule Install Delay (ms)",
        "Failure Recovery (ms)",
    ]
    metrics = metrics[cols]
    metrics.to_csv(SDN_METRICS_CSV, index=False)
    return metrics


def build_report():
    helper = load_helper()
    helper.PLOTS_DIR = PLOTS_DIR
    PLOTS_DIR.mkdir(parents=True, exist_ok=True)

    if not SUMMARY_CSV.exists() or not FAILURE_CSV.exists():
        missing = [str(p.relative_to(PROJECT_ROOT)) for p in [SUMMARY_CSV, FAILURE_CSV] if not p.exists()]
        raise FileNotFoundError("Missing required clean rerun files:\n" + "\n".join(missing))

    summary_df = pd.read_csv(SUMMARY_CSV)
    failure_df = pd.read_csv(FAILURE_CSV)
    metrics_df = prepare_metrics(summary_df, failure_df)
    provenance = helper.load_checkpoint_provenance()

    helper.create_plots(summary_df, failure_df)

    doc = Document()
    helper.set_default_style(doc)
    add_title_page(doc)

    doc.add_heading("1. Clean Branch Objective", level=1)
    doc.add_paragraph(
        "This clean branch is separated from the MetaGate branch specifically to test true zero-shot generalization. "
        "Only four direct routing methods are evaluated: ECMP, Bottleneck, Original GNN, and GNN+."
    )
    add_bullet(doc, "No MetaGate is used in this branch.")
    add_bullet(doc, "No Stable MetaGate extension is used in this branch.")
    add_bullet(doc, "No Bayesian calibration or few-shot adaptation is applied to unseen topologies before inference.")
    add_bullet(doc, "Germany50 and VtlWavenet2011 are evaluated directly with no retraining and no calibration.")

    doc.add_heading("2. Zero-Shot Evaluation Design", level=1)
    doc.add_paragraph(
        "The learned models used in this branch come from the known-topology training path. "
        "At evaluation time, unseen topologies are passed directly into the already-trained GNN and GNN+ selectors. "
        "There is no per-topology tuning, no adaptation step, and no deployment-time prior estimation."
    )
    protocol_rows = pd.DataFrame(
        [
            {"Item": "Known-topology training", "Value": "Used for Original GNN and GNN+ checkpoint preparation"},
            {"Item": "Unseen topologies", "Value": "Germany50 and VtlWavenet2011"},
            {"Item": "Calibration before inference", "Value": "No"},
            {"Item": "Bayesian prior fusion", "Value": "No"},
            {"Item": "MetaGate / Stable MetaGate", "Value": "Not part of this branch"},
        ]
    )
    helper.add_dataframe_table(doc, protocol_rows, font_size=9)

    doc.add_heading("3. Checkpoint Provenance", level=1)
    provenance_rows = pd.DataFrame(
        [
            {"Field": "GNN+ checkpoint used", "Value": provenance["checkpoint_path"]},
            {"Field": "Source checkpoint", "Value": provenance["source_checkpoint_path"]},
            {"Field": "Dropout", "Value": provenance["dropout"]},
            {"Field": "learn_k_crit", "Value": provenance["learn_k_crit"]},
            {"Field": "Fixed K", "Value": provenance["fixed_k"]},
            {"Field": "Legacy packet-SDN scorer input dim", "Value": provenance["legacy_packet_sdn_in_channels"]},
        ]
    )
    helper.add_dataframe_table(doc, provenance_rows, font_size=9)

    doc.add_heading("4. Topologies and Methods", level=1)
    topo_rows = (
        summary_df[["topology", "status", "nodes", "edges"]]
        .drop_duplicates()
        .assign(Topology=lambda f: f["topology"].map(TOPOLOGY_DISPLAY))
        .assign(Status=lambda f: f["status"].str.title())
        [["Topology", "Status", "nodes", "edges"]]
        .rename(columns={"nodes": "Nodes", "edges": "Edges"})
    )
    helper.add_dataframe_table(doc, topo_rows, font_size=8)

    doc.add_heading("5. Normal Scenario Results", level=1)
    doc.add_paragraph(
        "The normal-scenario results below come directly from the clean zero-shot rerun bundle. "
        "All four methods are rendered from actual rows in the new summary CSV."
    )
    helper.add_image(doc, PLOTS_DIR / "mlu_comparison_normal.png", "Figure 1. Zero-shot MLU comparison across all 8 topologies.")
    helper.add_image(doc, PLOTS_DIR / "throughput_comparison_normal.png", "Figure 2. Zero-shot throughput comparison across all 8 topologies.")
    helper.add_image(doc, PLOTS_DIR / "disturbance_comparison.png", "Figure 3. Zero-shot routing disturbance comparison across all 8 topologies.")
    helper.add_image(doc, PLOTS_DIR / "decision_time_comparison.png", "Figure 4. Zero-shot decision-time comparison across all 8 topologies.")
    normal_table = summary_df.copy()
    normal_table["Method"] = normal_table["method"].map(METHOD_LABELS)
    normal_table["Topology"] = normal_table["topology"].map(TOPOLOGY_DISPLAY)
    normal_table["Status"] = normal_table["status"].str.title()
    normal_table = normal_table[
        ["Method", "Topology", "Status", "mean_mlu", "throughput", "mean_disturbance", "decision_time_ms"]
    ].rename(
        columns={
            "mean_mlu": "Mean MLU",
            "throughput": "Throughput",
            "mean_disturbance": "Disturbance",
            "decision_time_ms": "Decision Time (ms)",
        }
    )
    helper.add_dataframe_chunks(doc, normal_table, "Normal-scenario summary table", chunk_size=16)

    doc.add_heading("6. GNN+ vs Original GNN", level=1)
    doc.add_paragraph(
        "This comparison isolates whether GNN+ improves over the original GNN in the same zero-shot 4-method setting."
    )
    helper.add_image(doc, PLOTS_DIR / "gnnplus_vs_original_gnn.png", "Figure 5. GNN+ vs Original GNN on MLU and decision time.")
    comp_rows = []
    for topo in TOPOLOGY_ORDER:
        gnn_row = summary_df[(summary_df["topology"] == topo) & (summary_df["method"] == "gnn")]
        gnnplus_row = summary_df[(summary_df["topology"] == topo) & (summary_df["method"] == "gnnplus")]
        if gnn_row.empty or gnnplus_row.empty:
            continue
        g = gnn_row.iloc[0]
        gp = gnnplus_row.iloc[0]
        delta = ((float(g["mean_mlu"]) - float(gp["mean_mlu"])) / max(float(g["mean_mlu"]), 1e-12)) * 100.0
        comp_rows.append(
            {
                "Topology": TOPOLOGY_DISPLAY[topo],
                "GNN MLU": float(g["mean_mlu"]),
                "GNN+ MLU": float(gp["mean_mlu"]),
                "MLU Improvement %": delta,
                "GNN Time (ms)": float(g["decision_time_ms"]),
                "GNN+ Time (ms)": float(gp["decision_time_ms"]),
            }
        )
    helper.add_dataframe_table(doc, pd.DataFrame(comp_rows), font_size=8)

    doc.add_heading("7. Failure Scenario Results", level=1)
    doc.add_paragraph(
        "Failure results are also from the same clean 4-method zero-shot rerun. "
        "No failure-time adaptive gate is present in this branch."
    )
    helper.add_image(doc, PLOTS_DIR / "failure_recovery_gnnplus.png", "Figure 6. GNN+ failure recovery time by topology and scenario.")
    for idx, scenario in enumerate(SCENARIO_ORDER, start=1):
        doc.add_heading(f"7.{idx} {SCENARIO_LABELS[scenario]}", level=2)
        scenario_df = failure_df[failure_df["scenario"] == scenario].copy()
        scenario_df["Method"] = scenario_df["method"].map(METHOD_LABELS)
        scenario_df["Topology"] = scenario_df["topology"].map(TOPOLOGY_DISPLAY)
        scenario_df["Status"] = scenario_df["status"].str.title()
        scenario_df = scenario_df[
            ["Method", "Topology", "Status", "mean_mlu", "pre_failure_mlu", "failure_recovery_ms"]
        ].rename(
            columns={
                "mean_mlu": "Post-Failure MLU",
                "pre_failure_mlu": "Pre-Failure MLU",
                "failure_recovery_ms": "Recovery (ms)",
            }
        )
        helper.add_dataframe_chunks(doc, scenario_df, f"{SCENARIO_LABELS[scenario]} results", chunk_size=16)

    doc.add_heading("8. Model-Based SDN Metrics", level=1)
    doc.add_paragraph(
        "These SDN metrics come from the clean rerun bundle itself. They remain model-based analytical control-loop metrics rather than live Mininet dataplane measurements."
    )
    helper.add_dataframe_chunks(doc, metrics_df, "Integrated clean-branch SDN metrics", chunk_size=16)

    doc.add_heading("9. Complexity Analysis", level=1)
    helper.add_dataframe_table(doc, helper.build_complexity_rows(), font_size=8)
    helper.add_dataframe_table(doc, helper.build_complexity_numeric_rows(summary_df), font_size=8)

    doc.add_heading("10. Limitations", level=1)
    add_bullet(doc, "This branch is intentionally limited to ECMP, Bottleneck, Original GNN, and GNN+.")
    add_bullet(doc, "No MetaGate / Stable MetaGate results are part of this branch by design.")
    add_bullet(doc, "Paper baselines such as FlexDATE, FlexEntry, and ERODRL are not added unless they are actually runnable in the repo.")
    add_bullet(doc, "Packet-level SDN metrics are model-based analytical metrics, not live Mininet measurements.")

    REPORT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_DOCX)

    audit_lines = [
        "# Clean Zero-Shot GNN/GNN+ Report Audit",
        "",
        f"- Report: `{REPORT_DOCX.relative_to(PROJECT_ROOT)}`",
        f"- Results directory: `{INPUT_DIR.relative_to(PROJECT_ROOT)}`",
        "- Scope: ECMP, Bottleneck, Original GNN, GNN+",
        "- MetaGate used: no",
        "- Stable MetaGate used: no",
        "- Bayesian calibration used: no",
        "- Unseen-topology adaptation used: no",
    ]
    AUDIT_MD.write_text("\n".join(audit_lines), encoding="utf-8")
    print(f"Report saved to {REPORT_DOCX}")


if __name__ == "__main__":
    build_report()

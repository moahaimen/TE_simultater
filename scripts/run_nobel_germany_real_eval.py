#!/usr/bin/env python3
"""Run a real zero-shot Nobel-Germany evaluation with the expanded baseline set."""

from __future__ import annotations

import importlib.util
import json
import os
import sys
from collections import defaultdict
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
os.environ.setdefault("MPLCONFIGDIR", str(PROJECT_ROOT / ".mplconfig"))
os.environ.setdefault("XDG_CACHE_HOME", str(PROJECT_ROOT / ".cache"))
Path(os.environ["MPLCONFIGDIR"]).mkdir(parents=True, exist_ok=True)
Path(os.environ["XDG_CACHE_HOME"]).mkdir(parents=True, exist_ok=True)

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

sys.path.insert(0, str(PROJECT_ROOT))

RESULT_TAG = os.environ.get("GNNPLUS_NOBEL_RESULT_TAG", "gnnplus_nobel_germany_fullbaselines_realrun")
OUT_DIR = PROJECT_ROOT / "results" / RESULT_TAG
PLOTS_DIR = OUT_DIR / "plots"
SUMMARY_CSV = OUT_DIR / "packet_sdn_summary.csv"
FAILURE_CSV = OUT_DIR / "packet_sdn_failure.csv"
METRICS_CSV = OUT_DIR / "packet_sdn_sdn_metrics.csv"
TIMESERIES_CSV = OUT_DIR / "packet_sdn_timeseries.csv"
FAILURE_TIMESERIES_CSV = OUT_DIR / "packet_sdn_failure_timeseries.csv"
REPORT_DOCX = OUT_DIR / "GNNPLUS_NOBEL_GERMANY_FULL_BASELINES_REPORT.docx"
AUDIT_MD = OUT_DIR / "nobel_germany_run_audit.md"
SAVE_PACKET_SDN_TIMESERIES = os.environ.get("GNNPLUS_SAVE_PACKET_SDN_TIMESERIES", "0") == "1"

SOURCE_TAG = os.environ.get("GNNPLUS_NOBEL_SOURCE_TAG", "gnnplus_task17_tiebreak_eval")
SOURCE_DIR = PROJECT_ROOT / "results" / SOURCE_TAG
SOURCE_TRAIN_DIR = SOURCE_DIR / "training"
CHECKPOINT_PATH = SOURCE_TRAIN_DIR / "gnn_plus_improved_fixedk40.pt"
INFERENCE_CALIBRATION_JSON = SOURCE_TRAIN_DIR / "inference_calibration.json"

TOPO_KEY = "nobel_germany"
TOPOLOGY_DISPLAY = "Nobel-Germany"
TOPOLOGY_URL = "https://sndlib.put.poznan.pl/download/sndlib-networks-native/nobel-germany.txt"
DYNAMIC_ARCHIVE_URL = "https://sndlib.put.poznan.pl/download/directed-nobel-germany-DFN-aggregated-5min-over-1day-native.tgz"

METHOD_ORDER = ["ecmp", "ospf", "topk", "sensitivity", "bottleneck", "gnn", "gnnplus"]
METHOD_LABELS = {
    "ecmp": "ECMP",
    "ospf": "OSPF",
    "topk": "TopK",
    "sensitivity": "Sensitivity",
    "bottleneck": "Bottleneck",
    "gnn": "Original GNN",
    "gnnplus": "GNN+",
}
METHOD_COLORS = {
    "ecmp": "#1f77b4",
    "ospf": "#17becf",
    "topk": "#9467bd",
    "sensitivity": "#8c564b",
    "bottleneck": "#ff7f0e",
    "gnn": "#2ca02c",
    "gnnplus": "#d62728",
}
SCENARIO_ORDER = [
    "single_link_failure",
    "multiple_link_failure",
    "three_link_failure",
    "capacity_degradation_50",
    "traffic_spike_2x",
]
SCENARIO_LABELS = {
    "single_link_failure": "Single Link Failure",
    "multiple_link_failure": "Multiple Link Failure (2 Links)",
    "three_link_failure": "3-Link Failure",
    "capacity_degradation_50": "Capacity Degradation (50%)",
    "traffic_spike_2x": "Traffic Spike (2x)",
}


def load_module(path: Path, name: str):
    spec = importlib.util.spec_from_file_location(name, path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Could not load module from {path}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[name] = module
    spec.loader.exec_module(module)
    return module


def add_image(doc: Document, path: Path, caption: str, width: float = 6.0) -> None:
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def title_page(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Nobel-Germany Zero-Shot Real Evaluation")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Standalone extension on top of the final Task 17 GNN+ checkpoint\n"
        "Methods: ECMP, OSPF, TopK, Sensitivity, Bottleneck, Original GNN, GNN+\n"
        "Protocol: zero-shot unseen topology evaluation"
    )
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Important: SDN metrics in this report are model-based analytical metrics, not live Mininet measurements."
    )
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    doc.add_page_break()


def plot_method_metric(summary_df: pd.DataFrame, metric: str, title: str, ylabel: str, output_name: str, *, log_scale: bool = False) -> None:
    ordered = summary_df.copy()
    ordered["method"] = pd.Categorical(ordered["method"], categories=METHOD_ORDER, ordered=True)
    ordered = ordered.sort_values("method")

    fig, ax = plt.subplots(figsize=(9.0, 4.8))
    x = np.arange(len(ordered))
    colors = [METHOD_COLORS[str(method)] for method in ordered["method"].astype(str)]
    ax.bar(x, ordered[metric].astype(float).to_numpy(), color=colors, width=0.72)
    ax.set_xticks(x)
    ax.set_xticklabels([METHOD_LABELS[str(method)] for method in ordered["method"].astype(str)], rotation=30, ha="right")
    ax.set_ylabel(ylabel)
    ax.set_title(title)
    ax.grid(axis="y", alpha=0.25)
    if log_scale:
        ax.set_yscale("log")
    fig.tight_layout()
    fig.savefig(PLOTS_DIR / output_name, dpi=200, bbox_inches="tight")
    plt.close(fig)


def plot_failure_metric(
    failure_df: pd.DataFrame,
    metric: str,
    title: str,
    ylabel: str,
    output_name: str,
    *,
    log_scale: bool = False,
) -> None:
    methods = [method for method in METHOD_ORDER if method in set(failure_df["method"].astype(str))]
    scenarios = [scenario for scenario in SCENARIO_ORDER if scenario in set(failure_df["scenario"].astype(str))]
    pivot = (
        failure_df.pivot(index="scenario", columns="method", values=metric)
        .reindex(index=scenarios, columns=methods)
    )

    x = np.arange(len(scenarios), dtype=float)
    width = 0.11
    offsets = (np.arange(len(methods), dtype=float) - (len(methods) - 1) / 2.0) * width

    fig, ax = plt.subplots(figsize=(11.5, 5.2))
    for idx, method in enumerate(methods):
        values = pivot[method].astype(float).to_numpy()
        ax.bar(
            x + offsets[idx],
            values,
            width=width * 0.94,
            label=METHOD_LABELS[method],
            color=METHOD_COLORS[method],
        )
    ax.set_xticks(x)
    ax.set_xticklabels([SCENARIO_LABELS[scenario] for scenario in scenarios], rotation=18, ha="right")
    ax.set_ylabel(ylabel)
    ax.set_title(title)
    ax.grid(axis="y", alpha=0.25)
    if log_scale:
        ax.set_yscale("log")
    ax.legend(ncol=4, fontsize=8, frameon=False)
    fig.tight_layout()
    fig.savefig(PLOTS_DIR / output_name, dpi=200, bbox_inches="tight")
    plt.close(fig)


def create_plots(summary_df: pd.DataFrame, failure_df: pd.DataFrame) -> None:
    PLOTS_DIR.mkdir(parents=True, exist_ok=True)
    plot_method_metric(summary_df, "mean_mlu", "Nobel-Germany Normal Mean MLU", "Mean MLU", "mlu_comparison_normal.png")
    plot_method_metric(summary_df, "throughput", "Nobel-Germany Throughput", "Throughput", "throughput_comparison_normal.png")
    plot_method_metric(summary_df, "mean_disturbance", "Nobel-Germany Routing Disturbance", "Mean Disturbance", "disturbance_comparison.png")
    plot_method_metric(summary_df, "decision_time_ms", "Nobel-Germany Decision Time", "Decision Time (ms)", "decision_time_comparison.png")
    plot_failure_metric(
        failure_df,
        "mean_mlu",
        "Nobel-Germany Failure MLU by Scenario",
        "Post-Recovery MLU",
        "failure_mlu_comparison.png",
    )
    plot_failure_metric(
        failure_df,
        "failure_recovery_ms",
        "Nobel-Germany Failure Recovery Time",
        "Recovery Time (ms)",
        "failure_recovery_comparison.png",
    )


def prepare_sdn_metrics(summary_df: pd.DataFrame, failure_df: pd.DataFrame) -> pd.DataFrame:
    recovery = (
        failure_df.groupby(["topology", "method"], as_index=False)["failure_recovery_ms"]
        .mean()
        .rename(columns={"failure_recovery_ms": "Failure Recovery (ms)"})
    )
    metrics = summary_df.merge(recovery, on=["topology", "method"], how="left")
    metrics["Method"] = metrics["method"].map(METHOD_LABELS)
    metrics["Topology"] = TOPOLOGY_DISPLAY
    metrics["Status"] = metrics["status"].str.title()
    metrics = metrics.rename(
        columns={
            "mean_mlu": "Mean MLU",
            "throughput": "Throughput",
            "mean_disturbance": "Disturbance",
            "mean_latency_au": "Mean Delay",
            "p95_latency_au": "P95 Delay",
            "packet_loss": "Packet Loss",
            "jitter_au": "Jitter",
            "decision_time_ms": "Decision Time (ms)",
            "flow_table_updates": "Flow Table Updates",
            "rule_install_delay_ms": "Rule Install Delay (ms)",
            "do_no_harm_fallback_rate": "Fallback Rate",
        }
    )
    cols = [
        "Method",
        "Topology",
        "Status",
        "Mean MLU",
        "Throughput",
        "Disturbance",
        "Mean Delay",
        "P95 Delay",
        "Packet Loss",
        "Jitter",
        "Decision Time (ms)",
        "Flow Table Updates",
        "Rule Install Delay (ms)",
        "Failure Recovery (ms)",
        "Fallback Rate",
    ]
    return metrics[cols]


def benchmark_nobel_normal(runner, improved, *, gnnplus_model, inference_calibration: dict) -> tuple[pd.DataFrame, pd.DataFrame]:
    dataset, path_library = runner.load_dataset(TOPO_KEY)
    capacities = np.asarray(dataset.capacities, dtype=float)
    weights = np.asarray(dataset.weights, dtype=float)
    ecmp_base = runner.ecmp_splits(path_library)
    topo_mapping = runner.SDNTopologyMapping.from_mininet(dataset.nodes, dataset.edges, dataset.od_pairs)
    test_indices = list(range(int(dataset.split["test_start"]), dataset.tm.shape[0]))

    gnn_model = runner.load_gnn_model(dataset, path_library)
    rows: list[dict] = []
    ts_rows: list[dict] = []

    for method in METHOD_ORDER:
        run_results = defaultdict(list)
        for run_id in range(runner.NUM_RUNS):
            current_splits = [split.copy() for split in ecmp_base]
            current_groups, _ = runner.build_ecmp_baseline_rules(path_library, topo_mapping, dataset.edges)
            prev_latency = None
            gnnplus_state = {
                "prev_tm": None,
                "prev_util": None,
                "prev_selected_indicator": np.zeros(len(dataset.od_pairs), dtype=np.float32),
                "prev_disturbance": 0.0,
                "guard_cache": {},
                "guard_cycle_index": 0,
            }
            for t_idx in test_indices:
                tm_vec = dataset.tm[t_idx]
                if method == "gnnplus":
                    result, current_splits, current_groups, prev_latency, gnnplus_state = improved.run_sdn_cycle_gnnplus_improved(
                        runner,
                        tm_vector=tm_vec,
                        dataset=dataset,
                        path_library=path_library,
                        ecmp_base=ecmp_base,
                        current_splits=current_splits,
                        current_groups=current_groups,
                        topo_mapping=topo_mapping,
                        capacities=capacities,
                        weights=weights,
                        gnnplus_model=gnnplus_model,
                        prev_latency_by_od=prev_latency,
                        gnnplus_state=gnnplus_state,
                        inference_calibration=inference_calibration,
                    )
                    do_no_harm_fallback = float(bool(gnnplus_state.get("select_info", {}).get("do_no_harm_fallback", False)))
                    run_results["do_no_harm_fallbacks"].append(do_no_harm_fallback)
                else:
                    result, current_splits, current_groups, prev_latency = runner.run_sdn_cycle(
                        tm_vector=tm_vec,
                        method=method,
                        dataset=dataset,
                        path_library=path_library,
                        ecmp_base=ecmp_base,
                        current_splits=current_splits,
                        current_groups=current_groups,
                        topo_mapping=topo_mapping,
                        capacities=capacities,
                        weights=weights,
                        gnn_model=gnn_model,
                        gnnplus_model=None,
                        prev_latency_by_od=prev_latency,
                    )
                    do_no_harm_fallback = 0.0
                    run_results["do_no_harm_fallbacks"].append(do_no_harm_fallback)
                run_results["post_mlus"].append(result.post_mlu)
                run_results["disturbances"].append(result.disturbance)
                run_results["throughputs"].append(result.throughput)
                run_results["latencies"].append(result.mean_latency)
                run_results["p95_latencies"].append(result.p95_latency)
                run_results["packet_losses"].append(result.packet_loss)
                run_results["jitters"].append(result.jitter)
                run_results["decision_times"].append(result.decision_time_ms)
                run_results["flow_updates"].append(result.flow_table_updates)
                run_results["rule_delays"].append(result.rule_install_delay_ms)
                if SAVE_PACKET_SDN_TIMESERIES:
                    ts_rows.append(
                        {
                            "topology": TOPO_KEY,
                            "status": "unseen",
                            "method": method,
                            "scenario": "normal",
                            "run_id": int(run_id),
                            "timestep": int(t_idx),
                            "mlu": float(result.post_mlu),
                            "disturbance": float(result.disturbance),
                            "throughput": float(result.throughput),
                            "mean_latency_au": float(result.mean_latency),
                            "p95_latency_au": float(result.p95_latency),
                            "packet_loss": float(result.packet_loss),
                            "jitter_au": float(result.jitter),
                            "decision_time_ms": float(result.decision_time_ms),
                            "flow_table_updates": float(result.flow_table_updates),
                            "rule_install_delay_ms": float(result.rule_install_delay_ms),
                            "do_no_harm_fallback": float(do_no_harm_fallback),
                        }
                    )

        row = {
            "topology": TOPO_KEY,
            "status": "unseen",
            "method": method,
            "scenario": "normal",
            "nodes": len(dataset.nodes),
            "edges": len(dataset.edges),
            "mean_mlu": float(np.mean(run_results["post_mlus"])),
            "mean_disturbance": float(np.mean(run_results["disturbances"])),
            "throughput": float(np.mean(run_results["throughputs"])),
            "mean_latency_au": float(np.mean(run_results["latencies"])),
            "p95_latency_au": float(np.mean(run_results["p95_latencies"])),
            "packet_loss": float(np.mean(run_results["packet_losses"])),
            "jitter_au": float(np.mean(run_results["jitters"])),
            "decision_time_ms": float(np.mean(run_results["decision_times"])),
            "flow_table_updates": float(np.mean(run_results["flow_updates"])),
            "rule_install_delay_ms": float(np.mean(run_results["rule_delays"])),
            "do_no_harm_fallback_rate": float(np.mean(run_results["do_no_harm_fallbacks"])),
        }
        rows.append(row)
        print(f"[nobel:normal] {method} mean_mlu={row['mean_mlu']:.6f}", flush=True)

    return pd.DataFrame(rows), pd.DataFrame(ts_rows)


def benchmark_nobel_failures(runner, improved, *, gnnplus_model, inference_calibration: dict) -> tuple[pd.DataFrame, pd.DataFrame]:
    dataset, path_library = runner.load_dataset(TOPO_KEY)
    capacities = np.asarray(dataset.capacities, dtype=float)
    weights = np.asarray(dataset.weights, dtype=float)
    ecmp_base = runner.ecmp_splits(path_library)
    test_indices = list(range(int(dataset.split["test_start"]), dataset.tm.shape[0]))
    sample_indices = test_indices[:: max(1, len(test_indices) // 10)]
    gnn_model = runner.load_gnn_model(dataset, path_library)

    rows: list[dict] = []
    ts_rows: list[dict] = []
    for scenario in SCENARIO_ORDER:
        for method in METHOD_ORDER:
            run_results = defaultdict(list)
            for t_idx in sample_indices:
                tm_vec = dataset.tm[t_idx]
                if method == "gnnplus":
                    recovery_ms, pre_mlu, post_mlu, _, select_info = improved.run_failure_scenario_gnnplus_improved(
                        runner,
                        scenario=scenario,
                        tm_vector=tm_vec,
                        dataset=dataset,
                        path_library=path_library,
                        ecmp_base=ecmp_base,
                        capacities=capacities,
                        weights=weights,
                        gnnplus_model=gnnplus_model,
                        inference_calibration=inference_calibration,
                    )
                    do_no_harm_fallback = float(bool(select_info.get("do_no_harm_fallback", False)))
                    run_results["do_no_harm_fallbacks"].append(do_no_harm_fallback)
                else:
                    recovery_ms, pre_mlu, post_mlu, _ = runner.run_failure_scenario(
                        scenario=scenario,
                        tm_vector=tm_vec,
                        method=method,
                        dataset=dataset,
                        path_library=path_library,
                        ecmp_base=ecmp_base,
                        capacities=capacities,
                        weights=weights,
                        topo_mapping=None,
                        gnn_model=gnn_model,
                        gnnplus_model=None,
                    )
                    do_no_harm_fallback = 0.0
                    run_results["do_no_harm_fallbacks"].append(do_no_harm_fallback)
                run_results["recovery_times"].append(recovery_ms)
                run_results["pre_mlus"].append(pre_mlu)
                run_results["post_mlus"].append(post_mlu)
                if SAVE_PACKET_SDN_TIMESERIES:
                    ts_rows.append(
                        {
                            "topology": TOPO_KEY,
                            "status": "unseen",
                            "method": method,
                            "scenario": scenario,
                            "timestep": int(t_idx),
                            "pre_failure_mlu": float(pre_mlu),
                            "post_recovery_mlu": float(post_mlu),
                            "failure_recovery_ms": float(recovery_ms),
                            "do_no_harm_fallback": float(do_no_harm_fallback),
                        }
                    )

            row = {
                "topology": TOPO_KEY,
                "status": "unseen",
                "method": method,
                "scenario": scenario,
                "nodes": len(dataset.nodes),
                "edges": len(dataset.edges),
                "mean_mlu": float(np.mean(run_results["post_mlus"])),
                "pre_failure_mlu": float(np.mean(run_results["pre_mlus"])),
                "failure_recovery_ms": float(np.mean(run_results["recovery_times"])),
                "do_no_harm_fallback_rate": float(np.mean(run_results["do_no_harm_fallbacks"])),
            }
            rows.append(row)
            print(f"[nobel:failure] {scenario} {method} mean_mlu={row['mean_mlu']:.6f}", flush=True)

    return pd.DataFrame(rows), pd.DataFrame(ts_rows)


def build_report(helper, summary_df: pd.DataFrame, failure_df: pd.DataFrame, metrics_df: pd.DataFrame) -> None:
    doc = Document()
    helper.set_default_style(doc)
    title_page(doc)

    doc.add_heading("1. Scope", level=1)
    doc.add_paragraph(
        "This bundle evaluates the real SNDlib Nobel-Germany topology as an additional unseen zero-shot topology. "
        "The trained GNN+ checkpoint is reused from the final Task 17 branch; Nobel-Germany is not added to training or validation. "
        "This version widens the baseline set to include OSPF and the heuristic selectors TopK and Sensitivity."
    )
    helper.add_dataframe_table(
        doc,
        pd.DataFrame(
            [
                {"Field": "Topology", "Value": "Nobel-Germany"},
                {"Field": "Nodes", "Value": 17},
                {"Field": "Directed links", "Value": 52},
                {"Field": "Traffic type", "Value": "Real Measured TMs (SNDlib)"},
                {"Field": "Demand matrices", "Value": 288},
                {"Field": "Granularity / horizon", "Value": "5 minutes / 1 day"},
                {"Field": "Status in this run", "Value": "Unseen zero-shot evaluation"},
                {"Field": "Methods", "Value": "ECMP, OSPF, TopK, Sensitivity, Bottleneck, Original GNN, GNN+"},
                {"Field": "Checkpoint reused", "Value": str(CHECKPOINT_PATH.relative_to(PROJECT_ROOT))},
                {"Field": "Inference calibration reused", "Value": str(INFERENCE_CALIBRATION_JSON.relative_to(PROJECT_ROOT))},
                {"Field": "Topology source URL", "Value": TOPOLOGY_URL},
                {"Field": "Dynamic TM archive URL", "Value": DYNAMIC_ARCHIVE_URL},
            ]
        ),
        font_size=9,
    )

    doc.add_heading("2. Normal Results", level=1)
    add_image(doc, PLOTS_DIR / "mlu_comparison_normal.png", "Figure 1. Nobel-Germany normal-condition mean MLU with expanded baselines.")
    add_image(doc, PLOTS_DIR / "throughput_comparison_normal.png", "Figure 2. Nobel-Germany throughput comparison with expanded baselines.")
    add_image(doc, PLOTS_DIR / "disturbance_comparison.png", "Figure 3. Nobel-Germany routing disturbance comparison.")
    add_image(doc, PLOTS_DIR / "decision_time_comparison.png", "Figure 4. Nobel-Germany decision-time comparison.")

    normal_display = summary_df.copy()
    normal_display["Method"] = normal_display["method"].map(METHOD_LABELS)
    helper.add_dataframe_table(
        doc,
        normal_display[
            [
                "Method",
                "mean_mlu",
                "mean_disturbance",
                "throughput",
                "decision_time_ms",
                "flow_table_updates",
                "rule_install_delay_ms",
                "do_no_harm_fallback_rate",
            ]
        ].rename(
            columns={
                "mean_mlu": "Mean MLU",
                "mean_disturbance": "Mean Disturbance",
                "throughput": "Throughput",
                "decision_time_ms": "Decision Time (ms)",
                "flow_table_updates": "Flow Table Updates",
                "rule_install_delay_ms": "Rule Install Delay (ms)",
                "do_no_harm_fallback_rate": "Fallback Rate",
            }
        ),
        font_size=8,
    )

    doc.add_heading("3. Failure Results", level=1)
    add_image(doc, PLOTS_DIR / "failure_mlu_comparison.png", "Figure 5. Nobel-Germany post-recovery MLU across failure scenarios.")
    add_image(doc, PLOTS_DIR / "failure_recovery_comparison.png", "Figure 6. Nobel-Germany recovery time across failure scenarios.")

    failure_winners = (
        failure_df.sort_values(["scenario", "mean_mlu"])
        .groupby("scenario", as_index=False)
        .first()[["scenario", "method", "mean_mlu"]]
        .rename(columns={"method": "best_method", "mean_mlu": "best_post_recovery_mlu"})
    )
    failure_winners["Scenario"] = failure_winners["scenario"].map(SCENARIO_LABELS)
    failure_winners["Best Method"] = failure_winners["best_method"].map(METHOD_LABELS)
    helper.add_dataframe_table(
        doc,
        failure_winners[["Scenario", "Best Method", "best_post_recovery_mlu"]].rename(
            columns={"best_post_recovery_mlu": "Best Post-Recovery MLU"}
        ),
        title="Failure winners by scenario",
        font_size=8,
    )

    failure_display = failure_df.copy()
    failure_display["Method"] = failure_display["method"].map(METHOD_LABELS)
    failure_display["Scenario"] = failure_display["scenario"].map(SCENARIO_LABELS).fillna(failure_display["scenario"])
    helper.add_dataframe_chunks(
        doc,
        failure_display[
            [
                "Method",
                "Scenario",
                "mean_mlu",
                "pre_failure_mlu",
                "failure_recovery_ms",
                "do_no_harm_fallback_rate",
            ]
        ].rename(
            columns={
                "mean_mlu": "Post-Recovery MLU",
                "pre_failure_mlu": "Pre-Failure MLU",
                "failure_recovery_ms": "Recovery Time (ms)",
                "do_no_harm_fallback_rate": "Fallback Rate",
            }
        ),
        heading_prefix="Detailed failure results",
        chunk_size=18,
    )

    doc.add_heading("4. Model-Based Packet-SDN Metrics", level=1)
    helper.add_dataframe_table(doc, metrics_df, font_size=8)

    doc.add_heading("5. Conclusion", level=1)
    best_row = summary_df.sort_values("mean_mlu").iloc[0]
    gnnplus_row = summary_df[summary_df["method"] == "gnnplus"].iloc[0]
    bottleneck_row = summary_df[summary_df["method"] == "bottleneck"].iloc[0]
    delta_pct = 0.0
    if abs(float(bottleneck_row["mean_mlu"])) > 1e-12:
        delta_pct = (
            (float(gnnplus_row["mean_mlu"]) - float(bottleneck_row["mean_mlu"]))
            / float(bottleneck_row["mean_mlu"])
        ) * 100.0
    doc.add_paragraph(
        f"On Nobel-Germany, the best normal-condition MLU in this zero-shot extension run is achieved by "
        f"{METHOD_LABELS[str(best_row['method'])]}. GNN+ records mean MLU {float(gnnplus_row['mean_mlu']):.6f} "
        f"versus Bottleneck {float(bottleneck_row['mean_mlu']):.6f} ({delta_pct:+.2f}%). "
        "This expanded Nobel bundle now includes the added baselines requested by the student: OSPF, TopK, and Sensitivity, "
        "alongside ECMP, Bottleneck, Original GNN, and GNN+."
    )

    REPORT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(REPORT_DOCX))


def write_audit(summary_df: pd.DataFrame, failure_df: pd.DataFrame) -> None:
    normal_best = summary_df.sort_values("mean_mlu").iloc[0]
    failure_winners = (
        failure_df.sort_values(["scenario", "mean_mlu"])
        .groupby("scenario", as_index=False)
        .first()[["scenario", "method", "mean_mlu"]]
    )
    lines = [
        "# Nobel-Germany Real Evaluation Audit",
        "",
        f"- Source bundle reused: `{SOURCE_TAG}`",
        f"- Checkpoint reused: `{CHECKPOINT_PATH.relative_to(PROJECT_ROOT)}`",
        f"- Inference calibration reused: `{INFERENCE_CALIBRATION_JSON.relative_to(PROJECT_ROOT)}`",
        f"- Prepared dataset: `data/processed/nobel_germany.npz`",
        "- Topology: Nobel-Germany",
        "- Nodes: 17",
        "- Directed links: 52",
        "- Traffic: Real Measured TMs (SNDlib)",
        "- Timesteps: 288",
        "- Zero-shot status: unseen; not retrained or calibrated on Nobel-Germany",
        "- Methods included: ECMP, OSPF, TopK, Sensitivity, Bottleneck, Original GNN, GNN+",
        f"- Normal summary rows: {len(summary_df)}",
        f"- Failure rows: {len(failure_df)}",
        f"- Normal best method: {METHOD_LABELS[str(normal_best['method'])]} (`mean_mlu={float(normal_best['mean_mlu']):.6f}`)",
        "",
        "## Failure Winners",
        "",
    ]
    for row in failure_winners.itertuples(index=False):
        lines.append(
            f"- {SCENARIO_LABELS[str(row.scenario)]}: {METHOD_LABELS[str(row.method)]} (`mean_mlu={float(row.mean_mlu):.6f}`)"
        )
    lines.extend(
        [
            "",
            f"- Normal summary CSV: `{SUMMARY_CSV.relative_to(PROJECT_ROOT)}`",
            f"- Failure CSV: `{FAILURE_CSV.relative_to(PROJECT_ROOT)}`",
            f"- SDN metrics CSV: `{METRICS_CSV.relative_to(PROJECT_ROOT)}`",
            f"- Report DOCX: `{REPORT_DOCX.relative_to(PROJECT_ROOT)}`",
        ]
    )
    AUDIT_MD.write_text("\n".join(lines), encoding="utf-8")


def main() -> int:
    os.environ.setdefault("GNNPLUS_EXTRA_UNSEEN_TOPOLOGIES", TOPO_KEY)
    improved = load_module(PROJECT_ROOT / "scripts" / "run_gnnplus_improved_fixedk40_experiment.py", "nobel_eval_exp")
    helper = load_module(PROJECT_ROOT / "scripts" / "build_gnnplus_packet_sdn_report_fixed.py", "nobel_eval_helper")

    if not CHECKPOINT_PATH.exists():
        raise FileNotFoundError(f"Missing reused final checkpoint: {CHECKPOINT_PATH}")
    if not INFERENCE_CALIBRATION_JSON.exists():
        raise FileNotFoundError(f"Missing inference calibration: {INFERENCE_CALIBRATION_JSON}")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    PLOTS_DIR.mkdir(parents=True, exist_ok=True)

    runner = improved.load_runner()
    runner.TOPOLOGIES[TOPO_KEY] = {"status": "unseen", "key": TOPO_KEY}

    from phase1_reactive.drl.gnn_plus_selector import load_gnn_plus

    gnnplus_model, _ = load_gnn_plus(CHECKPOINT_PATH, device="cpu")
    inference_calibration = json.loads(INFERENCE_CALIBRATION_JSON.read_text(encoding="utf-8"))

    summary_df, summary_ts_df = benchmark_nobel_normal(
        runner,
        improved,
        gnnplus_model=gnnplus_model,
        inference_calibration=inference_calibration,
    )
    failure_df, failure_ts_df = benchmark_nobel_failures(
        runner,
        improved,
        gnnplus_model=gnnplus_model,
        inference_calibration=inference_calibration,
    )
    metrics_df = prepare_sdn_metrics(summary_df, failure_df)

    summary_df.to_csv(SUMMARY_CSV, index=False)
    failure_df.to_csv(FAILURE_CSV, index=False)
    metrics_df.to_csv(METRICS_CSV, index=False)
    if SAVE_PACKET_SDN_TIMESERIES and not summary_ts_df.empty:
        summary_ts_df.to_csv(TIMESERIES_CSV, index=False)
    if SAVE_PACKET_SDN_TIMESERIES and not failure_ts_df.empty:
        failure_ts_df.to_csv(FAILURE_TIMESERIES_CSV, index=False)

    create_plots(summary_df, failure_df)
    build_report(helper, summary_df, failure_df, metrics_df)
    write_audit(summary_df, failure_df)

    print(f"Wrote Nobel-Germany expanded-baseline evaluation bundle to: {OUT_DIR}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

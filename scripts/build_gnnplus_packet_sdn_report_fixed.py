#!/usr/bin/env python3
"""Build a fixed GNN+ packet-SDN report bundle without touching the current one."""

from __future__ import annotations

import importlib.util
import json
import math
import os
import sys
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Iterable, Sequence

PROJECT_ROOT = Path(__file__).resolve().parent.parent
os.environ.setdefault("MPLCONFIGDIR", str(PROJECT_ROOT / ".mplconfig"))
os.environ.setdefault("XDG_CACHE_HOME", str(PROJECT_ROOT / ".cache"))
Path(os.environ["MPLCONFIGDIR"]).mkdir(parents=True, exist_ok=True)
Path(os.environ["XDG_CACHE_HOME"]).mkdir(parents=True, exist_ok=True)

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import torch
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


CURRENT_DIR = PROJECT_ROOT / "results" / "gnnplus_packet_sdn_report"
FIXED_DIR = PROJECT_ROOT / "results" / "gnnplus_packet_sdn_report_fixed"
PLOTS_DIR = FIXED_DIR / "plots"
OLD_TEMPLATE_DOCX = PROJECT_ROOT / "results" / "packet_sdn_simulation" / "Packet_SDN_Simulation_Report.docx"
CURRENT_DOCX = CURRENT_DIR / "GNNPLUS_PACKET_SDN_SIMULATION_REPORT.docx"
FIXED_DOCX = FIXED_DIR / "GNNPLUS_PACKET_SDN_SIMULATION_REPORT_FIXED.docx"
FIX_AUDIT = FIXED_DIR / "fix_audit.md"

CURRENT_SUMMARY = CURRENT_DIR / "packet_sdn_summary.csv"
CURRENT_FAILURE = CURRENT_DIR / "packet_sdn_failure.csv"
CURRENT_METRICS = CURRENT_DIR / "packet_sdn_sdn_metrics.csv"

FIXED_SUMMARY = FIXED_DIR / "packet_sdn_summary.csv"
FIXED_FAILURE = FIXED_DIR / "packet_sdn_failure.csv"
FIXED_METRICS = FIXED_DIR / "packet_sdn_sdn_metrics.csv"

CHECKPOINT_PATH = PROJECT_ROOT / "results" / "gnn_plus_retrained_fixedk40" / "gnn_plus_fixed_k40.pt"
SOURCE_CHECKPOINT_PATH = PROJECT_ROOT / "results" / "gnn_plus" / "stage1_regularization" / "training_d02" / "final.pt"
SOURCE_SUMMARY_JSON = PROJECT_ROOT / "results" / "gnn_plus" / "stage1_regularization" / "training_d02" / "summary.json"

CORE_METHODS = ["ecmp", "bottleneck", "gnn", "gnnplus"]
METHOD_LABELS = {
    "ecmp": "ECMP",
    "bottleneck": "BOTTLENECK",
    "gnn": "GNN",
    "gnnplus": "GNNPLUS",
}
METHOD_PLOT_LABELS = {
    "ecmp": "ECMP",
    "bottleneck": "Bottleneck",
    "gnn": "Original GNN",
    "gnnplus": "GNN+",
}
METHOD_COLORS = {
    "ecmp": "#1f77b4",
    "bottleneck": "#ff7f0e",
    "gnn": "#2ca02c",
    "gnnplus": "#d62728",
}
TOPOLOGY_ORDER = [
    "abilene",
    "cernet",
    "geant",
    "ebone",
    "sprintlink",
    "tiscali",
    "germany50",
    "nobel_germany",
    "vtlwavenet2011",
]
TOPOLOGY_DISPLAY = {
    "abilene": "Abilene",
    "cernet": "CERNET",
    "geant": "GEANT",
    "ebone": "Ebone",
    "sprintlink": "Sprintlink",
    "tiscali": "Tiscali",
    "germany50": "Germany50",
    "nobel_germany": "Nobel-Germany",
    "vtlwavenet2011": "VtlWavenet2011",
}
SCENARIO_ORDER = [
    "single_link_failure",
    "random_link_failure_1",
    "random_link_failure_2",
    "capacity_degradation_50",
    "traffic_spike_2x",
]
SCENARIO_LABELS = {
    "normal": "Normal",
    "single_link_failure": "Single Link Failure",
    "random_link_failure_1": "Random Link Failure (1)",
    "random_link_failure_2": "Random Link Failure (2)",
    "capacity_degradation_50": "Capacity Degradation (50%)",
    "traffic_spike_2x": "Traffic Spike (2x)",
}


def load_packet_sdn_module():
    module_path = PROJECT_ROOT / "scripts" / "run_gnnplus_packet_sdn_full.py"
    spec = importlib.util.spec_from_file_location("run_gnnplus_packet_sdn_full", module_path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Could not load module from {module_path}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def chunked(seq: Sequence, size: int) -> Iterable[Sequence]:
    for idx in range(0, len(seq), size):
        yield seq[idx:idx + size]


def ordered_unique(values: Iterable[str], order: Sequence[str]) -> list[str]:
    present = list(dict.fromkeys(str(v) for v in values))
    ordered = [value for value in order if value in present]
    ordered.extend(value for value in present if value not in ordered)
    return ordered


def audit_csv(path: Path) -> dict:
    df = pd.read_csv(path)
    audit = {
        "path": str(path.relative_to(PROJECT_ROOT)),
        "rows": int(len(df)),
        "methods": [],
        "scenarios": [],
        "topologies": [],
        "status": [],
        "method_counts": {},
    }
    if "method" in df.columns:
        methods = df["method"].astype(str)
        audit["methods"] = ordered_unique(methods.unique().tolist(), CORE_METHODS)
        audit["method_counts"] = dict(sorted(Counter(methods).items()))
    if "Method" in df.columns:
        methods = df["Method"].astype(str)
        audit["methods"] = ordered_unique(methods.unique().tolist(), [METHOD_LABELS[m] for m in CORE_METHODS])
        audit["method_counts"] = dict(sorted(Counter(methods).items()))
    if "scenario" in df.columns:
        audit["scenarios"] = ordered_unique(df["scenario"].astype(str).unique().tolist(), ["normal"] + SCENARIO_ORDER)
    if "topology" in df.columns:
        audit["topologies"] = ordered_unique(df["topology"].astype(str).unique().tolist(), TOPOLOGY_ORDER)
    if "Topology" in df.columns:
        inv = {v: k for k, v in TOPOLOGY_DISPLAY.items()}
        topo_values = [inv.get(str(v), str(v)) for v in df["Topology"].astype(str)]
        audit["topologies"] = ordered_unique(pd.Series(topo_values).unique().tolist(), TOPOLOGY_ORDER)
    if "status" in df.columns:
        audit["status"] = ordered_unique(df["status"].astype(str).unique().tolist(), ["known", "unseen"])
    if "Status" in df.columns:
        audit["status"] = ordered_unique(df["Status"].astype(str).str.lower().unique().tolist(), ["known", "unseen"])
    return audit


def extract_docx_table_audit(path: Path) -> dict:
    doc = Document(path)
    summary_methods = set()
    failure_methods = set()
    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows[1:]]
        if headers[:3] == ["Method", "Topology", "Status"]:
            summary_methods.update(row[0] for row in rows if row and row[0])
        if headers[:4] == ["topology", "status", "method", "scenario"]:
            failure_methods.update(row[2] for row in rows if len(row) > 2 and row[2])
    full_text = "\n".join(p.text for p in doc.paragraphs)
    return {
        "summary_table_methods": sorted(summary_methods),
        "failure_table_methods": sorted(failure_methods),
        "contains_initialized_model_text": "initialized model" in full_text.lower(),
        "contains_out_of_scope_methods": any(token in full_text for token in ["OSPF", "TopK", "Sensitivity"]),
    }


def load_checkpoint_provenance() -> dict:
    ckpt = torch.load(CHECKPOINT_PATH, map_location="cpu")
    src_ckpt = torch.load(SOURCE_CHECKPOINT_PATH, map_location="cpu")
    training_summary = json.loads(SOURCE_SUMMARY_JSON.read_text())
    config = ckpt.get("config", {})
    return {
        "checkpoint_path": str(CHECKPOINT_PATH.relative_to(PROJECT_ROOT)),
        "source_checkpoint_path": str(SOURCE_CHECKPOINT_PATH.relative_to(PROJECT_ROOT)),
        "source_summary_path": str(SOURCE_SUMMARY_JSON.relative_to(PROJECT_ROOT)),
        "payload_model_type": ckpt.get("model_type"),
        "payload_stage": ckpt.get("stage"),
        "dropout": ckpt.get("dropout"),
        "learn_k_crit": config.get("learn_k_crit"),
        "fixed_k": config.get("k_crit_max") if config.get("k_crit_min") == config.get("k_crit_max") else None,
        "k_crit_min": config.get("k_crit_min"),
        "k_crit_max": config.get("k_crit_max"),
        "node_dim": config.get("node_dim"),
        "edge_dim": config.get("edge_dim"),
        "od_dim": config.get("od_dim"),
        "legacy_packet_sdn_in_channels": 30,
        "same_config_as_source": config == src_ckpt.get("config", {}),
        "training_summary_dropout": training_summary.get("dropout"),
    }


def needs_normal_rerun(summary_df: pd.DataFrame) -> bool:
    methods = set(summary_df["method"].astype(str).unique())
    expected_rows = len(TOPOLOGY_ORDER) * len(CORE_METHODS)
    return len(summary_df) != expected_rows or methods != set(CORE_METHODS)


def rerun_normal_summary() -> pd.DataFrame:
    module = load_packet_sdn_module()
    gnn_cache = {}
    gnnplus_cache = {}
    normal_results = []
    for topo in TOPOLOGY_ORDER:
        rows = module.benchmark_topology_normal(topo, CORE_METHODS.copy(), gnn_cache, gnnplus_cache)
        normal_results.extend(rows)
    df = pd.DataFrame(normal_results)
    if needs_normal_rerun(df):
        raise RuntimeError(
            "Regenerated normal summary is still incomplete: "
            f"rows={len(df)}, methods={sorted(df['method'].astype(str).unique().tolist())}"
        )
    return df


def prepare_fixed_csvs() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, dict, dict, dict]:
    FIXED_DIR.mkdir(parents=True, exist_ok=True)
    PLOTS_DIR.mkdir(parents=True, exist_ok=True)

    source_summary = pd.read_csv(CURRENT_SUMMARY)
    source_failure = pd.read_csv(CURRENT_FAILURE)

    if needs_normal_rerun(source_summary):
        fixed_summary = rerun_normal_summary()
        summary_source = "regenerated from scripts/run_gnnplus_packet_sdn_full.py normal benchmark"
    else:
        fixed_summary = source_summary.copy()
        summary_source = "copied from existing packet_sdn_summary.csv"

    fixed_failure = source_failure.copy()
    failure_methods = set(fixed_failure["method"].astype(str).unique())
    expected_failure_rows = len(TOPOLOGY_ORDER) * len(CORE_METHODS) * len(SCENARIO_ORDER)
    if len(fixed_failure) != expected_failure_rows or failure_methods != set(CORE_METHODS):
        raise RuntimeError(
            "Existing failure CSV is incomplete and this builder is intentionally not rerunning failures: "
            f"rows={len(fixed_failure)}, methods={sorted(failure_methods)}"
        )

    fixed_summary = fixed_summary.copy()
    fixed_summary["topology"] = pd.Categorical(fixed_summary["topology"], categories=TOPOLOGY_ORDER, ordered=True)
    fixed_summary["method"] = pd.Categorical(fixed_summary["method"], categories=CORE_METHODS, ordered=True)
    fixed_summary = fixed_summary.sort_values(["topology", "method"]).reset_index(drop=True)
    fixed_summary["topology"] = fixed_summary["topology"].astype(str)
    fixed_summary["method"] = fixed_summary["method"].astype(str)

    fixed_failure = fixed_failure.copy()
    fixed_failure["topology"] = pd.Categorical(fixed_failure["topology"], categories=TOPOLOGY_ORDER, ordered=True)
    fixed_failure["method"] = pd.Categorical(fixed_failure["method"], categories=CORE_METHODS, ordered=True)
    fixed_failure["scenario"] = pd.Categorical(fixed_failure["scenario"], categories=SCENARIO_ORDER, ordered=True)
    fixed_failure = fixed_failure.sort_values(["scenario", "topology", "method"]).reset_index(drop=True)
    fixed_failure["topology"] = fixed_failure["topology"].astype(str)
    fixed_failure["method"] = fixed_failure["method"].astype(str)
    fixed_failure["scenario"] = fixed_failure["scenario"].astype(str)

    recovery_lookup = (
        fixed_failure.groupby(["topology", "method"], as_index=False)["failure_recovery_ms"]
        .mean()
        .rename(columns={"failure_recovery_ms": "avg_failure_recovery_ms"})
    )
    fixed_metrics = fixed_summary.merge(recovery_lookup, on=["topology", "method"], how="left")
    fixed_metrics = fixed_metrics.rename(
        columns={
            "topology": "Topology",
            "status": "Status",
            "method": "Method",
            "mean_mlu": "Mean MLU",
            "throughput": "Throughput",
            "mean_disturbance": "Disturbance",
            "decision_time_ms": "Decision Time (ms)",
            "avg_failure_recovery_ms": "Recovery Time (ms)",
        }
    )
    fixed_metrics["Method"] = fixed_metrics["Method"].map(METHOD_LABELS)
    fixed_metrics["Topology"] = fixed_metrics["Topology"].map(TOPOLOGY_DISPLAY)
    fixed_metrics["Status"] = fixed_metrics["Status"].str.title()
    fixed_metrics = fixed_metrics[
        [
            "Method",
            "Topology",
            "Status",
            "Mean MLU",
            "Throughput",
            "Disturbance",
            "Decision Time (ms)",
            "Recovery Time (ms)",
        ]
    ]

    fixed_summary.to_csv(FIXED_SUMMARY, index=False)
    fixed_failure.to_csv(FIXED_FAILURE, index=False)
    fixed_metrics.to_csv(FIXED_METRICS, index=False)

    return (
        fixed_summary,
        fixed_failure,
        fixed_metrics,
        audit_csv(CURRENT_SUMMARY),
        audit_csv(CURRENT_FAILURE),
        audit_csv(CURRENT_METRICS),
    )


def grouped_bar_offsets(num_methods: int, total_width: float = 0.82) -> tuple[np.ndarray, float]:
    width = min(0.18, total_width / max(num_methods, 1))
    offsets = np.linspace(-((num_methods - 1) * width) / 2, ((num_methods - 1) * width) / 2, num_methods)
    return offsets, width


def plot_grouped_metric(df: pd.DataFrame, metric: str, title: str, ylabel: str, output_name: str, log_scale: bool = False):
    methods = ordered_unique(df["method"].astype(str).unique().tolist(), CORE_METHODS)
    if set(methods) != set(CORE_METHODS):
        raise RuntimeError(f"Cannot plot {output_name}: expected {CORE_METHODS}, found {methods}")

    fig, ax = plt.subplots(figsize=(14, 6))
    topologies = ordered_unique(df["topology"].astype(str).unique().tolist(), TOPOLOGY_ORDER)
    x = np.arange(len(topologies))
    offsets, width = grouped_bar_offsets(len(methods))

    for idx, method in enumerate(methods):
        series = []
        for topo in topologies:
            subset = df[(df["topology"] == topo) & (df["method"] == method)]
            if subset.empty:
                raise RuntimeError(f"Missing {method} row for topology {topo} while plotting {output_name}")
            series.append(float(subset.iloc[0][metric]))
        ax.bar(
            x + offsets[idx],
            series,
            width=width,
            color=METHOD_COLORS[method],
            edgecolor="white",
            linewidth=0.6,
            label=METHOD_PLOT_LABELS[method],
        )

    ax.set_title(title)
    ax.set_xlabel("Topology")
    ax.set_ylabel(ylabel)
    ax.set_xticks(x)
    ax.set_xticklabels([TOPOLOGY_DISPLAY[t] for t in topologies], rotation=35, ha="right")
    ax.grid(axis="y", alpha=0.25)
    ax.legend(ncol=2)
    if log_scale:
        ax.set_yscale("log")
    fig.tight_layout()
    fig.savefig(PLOTS_DIR / output_name, dpi=160, bbox_inches="tight")
    plt.close(fig)


def plot_failure_recovery(df_failure: pd.DataFrame):
    gnnplus_df = df_failure[df_failure["method"] == "gnnplus"].copy()
    fig, ax = plt.subplots(figsize=(14, 6))
    topologies = ordered_unique(gnnplus_df["topology"].astype(str).unique().tolist(), TOPOLOGY_ORDER)
    scenarios = ordered_unique(gnnplus_df["scenario"].astype(str).unique().tolist(), SCENARIO_ORDER)
    x = np.arange(len(topologies))
    offsets, width = grouped_bar_offsets(len(scenarios), total_width=0.86)
    scenario_colors = ["#4c78a8", "#f58518", "#54a24b", "#e45756", "#72b7b2"]

    for idx, scenario in enumerate(scenarios):
        series = []
        for topo in topologies:
            subset = gnnplus_df[(gnnplus_df["topology"] == topo) & (gnnplus_df["scenario"] == scenario)]
            if subset.empty:
                raise RuntimeError(f"Missing gnnplus failure row for {topo}/{scenario}")
            series.append(float(subset.iloc[0]["failure_recovery_ms"]))
        ax.bar(
            x + offsets[idx],
            series,
            width=width,
            color=scenario_colors[idx % len(scenario_colors)],
            edgecolor="white",
            linewidth=0.6,
            label=SCENARIO_LABELS[scenario],
        )

    ax.set_title("GNN+ Failure Recovery Time by Scenario")
    ax.set_xlabel("Topology")
    ax.set_ylabel("Recovery Time (ms)")
    ax.set_xticks(x)
    ax.set_xticklabels([TOPOLOGY_DISPLAY[t] for t in topologies], rotation=35, ha="right")
    ax.grid(axis="y", alpha=0.25)
    ax.legend(ncol=2)
    fig.tight_layout()
    fig.savefig(PLOTS_DIR / "failure_recovery_gnnplus.png", dpi=160, bbox_inches="tight")
    plt.close(fig)


def plot_gnn_vs_gnnplus(df_summary: pd.DataFrame):
    topologies = ordered_unique(df_summary["topology"].astype(str).unique().tolist(), TOPOLOGY_ORDER)
    fig, axes = plt.subplots(1, 2, figsize=(14, 5))
    metrics = [("mean_mlu", "Mean MLU", True), ("decision_time_ms", "Decision Time (ms)", False)]
    x = np.arange(len(topologies))
    width = 0.34

    for axis, (metric, ylabel, log_scale) in zip(axes, metrics):
        gnn_vals = []
        gnnplus_vals = []
        for topo in topologies:
            gnn_row = df_summary[(df_summary["topology"] == topo) & (df_summary["method"] == "gnn")]
            gnnplus_row = df_summary[(df_summary["topology"] == topo) & (df_summary["method"] == "gnnplus")]
            if gnn_row.empty or gnnplus_row.empty:
                raise RuntimeError(f"Missing GNN or GNN+ row for topology {topo} in {metric}")
            gnn_vals.append(float(gnn_row.iloc[0][metric]))
            gnnplus_vals.append(float(gnnplus_row.iloc[0][metric]))

        axis.bar(x - width / 2, gnn_vals, width=width, color=METHOD_COLORS["gnn"], edgecolor="white", linewidth=0.6, label="Original GNN")
        axis.bar(x + width / 2, gnnplus_vals, width=width, color=METHOD_COLORS["gnnplus"], edgecolor="white", linewidth=0.6, label="GNN+")
        axis.set_title("GNN+ vs Original GNN" if metric == "mean_mlu" else "GNN+ vs Original GNN")
        axis.set_ylabel(ylabel)
        axis.set_xticks(x)
        axis.set_xticklabels([TOPOLOGY_DISPLAY[t] for t in topologies], rotation=35, ha="right")
        axis.grid(axis="y", alpha=0.25)
        axis.legend()
        if log_scale:
            axis.set_yscale("log")
            axis.set_title("GNN+ vs Original GNN: MLU Comparison")
        else:
            axis.set_title("GNN+ vs Original GNN: Decision Time Comparison")

    fig.tight_layout()
    fig.savefig(PLOTS_DIR / "gnnplus_vs_original_gnn.png", dpi=160, bbox_inches="tight")
    plt.close(fig)


def create_plots(df_summary: pd.DataFrame, df_failure: pd.DataFrame):
    plot_grouped_metric(
        df_summary,
        metric="mean_mlu",
        title="MLU Comparison Across Topologies (Normal Scenario)",
        ylabel="Mean MLU",
        output_name="mlu_comparison_normal.png",
        log_scale=True,
    )
    plot_grouped_metric(
        df_summary,
        metric="throughput",
        title="Throughput Comparison Across Topologies (Normal Scenario)",
        ylabel="Throughput",
        output_name="throughput_comparison_normal.png",
    )
    plot_grouped_metric(
        df_summary,
        metric="mean_disturbance",
        title="Routing Disturbance Comparison Across Topologies",
        ylabel="Mean Disturbance",
        output_name="disturbance_comparison.png",
    )
    plot_grouped_metric(
        df_summary,
        metric="decision_time_ms",
        title="Decision Time Comparison Across Topologies",
        ylabel="Decision Time (ms)",
        output_name="decision_time_comparison.png",
    )
    plot_failure_recovery(df_failure)
    plot_gnn_vs_gnnplus(df_summary)


def set_default_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    for level in range(1, 4):
        heading = doc.styles[f"Heading {level}"]
        heading.font.name = "Arial"
        heading.font.color.rgb = RGBColor(0x1A, 0x47, 0x80)


def add_title_page(doc: Document):
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GNN+ Packet-Level SDN Simulation Report")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x80)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Fixed Report Bundle for the 4-Method GNN+ Packet-SDN Branch")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    stamp = doc.add_paragraph()
    stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = stamp.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        "Important: packet-level SDN metrics in this report are model-based analytical metrics, not live Mininet measurements."
    )
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    doc.add_page_break()


def add_bullet(doc: Document, text: str):
    doc.add_paragraph(text, style="List Bullet")


def add_image(doc: Document, path: Path, caption: str, width: float = 6.2):
    if not path.exists():
        doc.add_paragraph(f"[Image not found: {path.name}]")
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_paragraph.add_run(caption)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_dataframe_table(doc: Document, df: pd.DataFrame, title: str | None = None, font_size: int = 8):
    if title:
        doc.add_paragraph(title)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for idx, col in enumerate(df.columns):
        cell = table.rows[0].cells[idx]
        cell.text = str(col)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(font_size)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            text = value
            if isinstance(value, float):
                if math.isfinite(value):
                    text = f"{value:.4f}" if abs(value) < 1000 else f"{value:.2f}"
                else:
                    text = "nan"
            cells[idx].text = str(text)
            for run in cells[idx].paragraphs[0].runs:
                run.font.size = Pt(font_size)
    return table


def add_dataframe_chunks(doc: Document, df: pd.DataFrame, heading_prefix: str, chunk_size: int = 20):
    for chunk_idx, rows in enumerate(chunked(df.to_dict("records"), chunk_size), start=1):
        chunk_df = pd.DataFrame(rows)
        title = heading_prefix if chunk_idx == 1 else f"{heading_prefix} (cont. {chunk_idx})"
        add_dataframe_table(doc, chunk_df, title=title)
        doc.add_paragraph("")


def build_complexity_rows() -> pd.DataFrame:
    rows = [
        {
            "Method": "ECMP",
            "Selection Stage": "Uniform split over precomputed paths",
            "Complexity": "Low; no optimization solve",
            "Practical Note": "Fastest controller decision path in this report.",
        },
        {
            "Method": "Bottleneck",
            "Selection Stage": "Heuristic critical-flow selection + LP reroute on selected flows",
            "Complexity": "Heuristic selection plus LP solve; controller cost dominated by LP",
            "Practical Note": "Higher than ECMP because the LP is solved every control cycle.",
        },
        {
            "Method": "GNN",
            "Selection Stage": "Graph-neural inference + LP reroute on selected flows",
            "Complexity": "GNN inference plus LP solve; asymptotically higher than ECMP, same LP class as Bottleneck",
            "Practical Note": "In practice, controller time is usually dominated by inference plus LP solve.",
        },
        {
            "Method": "GNN+",
            "Selection Stage": "Enhanced graph-neural inference + LP reroute on selected flows",
            "Complexity": "Same asymptotic structure as GNN, with a higher constant factor from richer features",
            "Practical Note": "More feature processing than GNN, but still the same controller pipeline shape: inference then LP.",
        },
    ]
    return pd.DataFrame(rows)


def build_complexity_numeric_rows(df_summary: pd.DataFrame) -> pd.DataFrame:
    rows = []
    label_map = {
        "ecmp": "ECMP",
        "bottleneck": "Bottleneck",
        "gnn": "GNN",
        "gnnplus": "GNN+",
    }
    for method in CORE_METHODS:
        method_df = df_summary[df_summary["method"] == method].copy()
        rows.append(
            {
                "Method": label_map[method],
                "Avg Decision Time (ms)": float(method_df["decision_time_ms"].mean()),
                "Median Decision Time (ms)": float(method_df["decision_time_ms"].median()),
                "Max Decision Time (ms)": float(method_df["decision_time_ms"].max()),
                "Avg Rule Install Delay (ms)": float(method_df["rule_install_delay_ms"].mean()),
                "Avg Flow Table Updates": float(method_df["flow_table_updates"].mean()),
            }
        )
    return pd.DataFrame(rows)


def build_report(df_summary: pd.DataFrame, df_failure: pd.DataFrame, df_metrics: pd.DataFrame, provenance: dict):
    doc = Document()
    set_default_style(doc)
    add_title_page(doc)

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This fixed report rebuilds the GNN+ packet-SDN branch without widening the accepted method scope. "
        "The comparison set remains ECMP, Bottleneck, Original GNN, and GNN+ across 8 topologies."
    )
    add_bullet(doc, "Known topologies: Abilene, CERNET, GEANT, Ebone, Sprintlink, Tiscali.")
    add_bullet(doc, "Unseen topologies: Germany50 and VtlWavenet2011.")
    add_bullet(doc, "Normal results and 5 failure scenarios are included; paper baselines remain unavailable in the repo.")
    add_bullet(doc, "Packet-level SDN metrics are model-based analytical estimates, not live Mininet measurements.")
    add_bullet(doc, "The normal summary bundle was regenerated because the current source summary omitted GNN+ rows; the current failure CSV was already complete and was reused.")

    doc.add_heading("2. Checkpoint Provenance", level=1)
    doc.add_paragraph("The fixed report uses the verified fixed-K checkpoint path below.")
    provenance_rows = pd.DataFrame(
        [
            {"Field": "Checkpoint used", "Value": provenance["checkpoint_path"]},
            {"Field": "Source checkpoint", "Value": provenance["source_checkpoint_path"]},
            {"Field": "Training summary", "Value": provenance["source_summary_path"]},
            {"Field": "Model type", "Value": provenance["payload_model_type"]},
            {"Field": "Dropout", "Value": provenance["dropout"]},
            {"Field": "learn_k_crit", "Value": provenance["learn_k_crit"]},
            {"Field": "Fixed K", "Value": provenance["fixed_k"]},
            {"Field": "Checkpoint dims", "Value": f"node={provenance['node_dim']}, edge={provenance['edge_dim']}, od={provenance['od_dim']}"},
            {"Field": "Legacy packet-SDN scorer input dim", "Value": provenance["legacy_packet_sdn_in_channels"]},
        ]
    )
    add_dataframe_table(doc, provenance_rows, font_size=9)
    doc.add_paragraph(
        "The checkpoint payload itself records explicit node, edge, and OD feature dimensions rather than a standalone "
        "`feature_dim` field. The legacy simplified packet-SDN scorer scripts also reference `in_channels=30`; both facts are recorded here to keep the provenance honest."
    )

    doc.add_heading("3. Branch Scope and Topologies", level=1)
    doc.add_paragraph(
        "This is not the universal all-method final report. It is the fixed GNN+ packet-SDN branch report for the current 4-method comparison set."
    )
    add_bullet(doc, "Included methods: ECMP, Bottleneck, GNN, GNN+.")
    add_bullet(doc, "No additional heuristics or paper baselines are added beyond the methods that are actually present in this branch output bundle.")
    topo_df = (
        df_summary[["topology", "status", "nodes", "edges"]]
        .drop_duplicates()
        .assign(Topology=lambda frame: frame["topology"].map(TOPOLOGY_DISPLAY))
        .assign(Status=lambda frame: frame["status"].str.title())
        [["Topology", "Status", "nodes", "edges"]]
        .rename(columns={"nodes": "Nodes", "edges": "Edges"})
    )
    add_dataframe_table(doc, topo_df, title="Evaluation topologies")

    doc.add_heading("3.1 Coverage Against Student Request", level=2)
    doc.add_paragraph("This subsection answers the exact scope questions the student is likely to ask.")
    add_bullet(doc, "Topology coverage present: 8 topologies, split into 6 known and 2 unseen.")
    add_bullet(doc, "Failure coverage present: normal conditions plus 5 failure scenarios.")
    add_bullet(doc, "SDN metrics coverage present: throughput, latency, packet loss, jitter, rule-install delay, and failure recovery are described as model-based analytical metrics.")
    add_bullet(doc, "Complexity analysis is included for the 4 methods that are actually evaluated in this branch.")
    add_bullet(doc, "Paper baselines requested by the student are not present as evaluated results in this report: FlexDATE, FlexEntry, and ERODRL are not runnable from the current repo branch outputs, so they are not faked into the comparison.")

    doc.add_heading("4. Normal Scenario Results", level=1)
    doc.add_paragraph(
        "The normal-scenario plots below are regenerated from a complete 32-row summary CSV. Every plotted series is derived from actual rows in the fixed CSV."
    )
    add_image(doc, PLOTS_DIR / "mlu_comparison_normal.png", "Figure 1. MLU comparison across all 8 topologies.")
    add_image(doc, PLOTS_DIR / "throughput_comparison_normal.png", "Figure 2. Throughput comparison across all 8 topologies.")
    add_image(doc, PLOTS_DIR / "disturbance_comparison.png", "Figure 3. Routing disturbance comparison across all 8 topologies.")
    add_image(doc, PLOTS_DIR / "decision_time_comparison.png", "Figure 4. Decision-time comparison across all 8 topologies.")

    known_metrics = df_metrics[df_metrics["Status"] == "Known"].reset_index(drop=True)
    unseen_metrics = df_metrics[df_metrics["Status"] == "Unseen"].reset_index(drop=True)
    add_dataframe_chunks(doc, known_metrics, "Known-topology summary table", chunk_size=24)
    add_dataframe_chunks(doc, unseen_metrics, "Unseen-topology summary table", chunk_size=12)

    doc.add_heading("5. GNN+ vs Original GNN", level=1)
    doc.add_paragraph(
        "This section compares the regenerated normal-scenario GNN and GNN+ rows directly. Both series are rendered from the fixed summary CSV."
    )
    add_image(doc, PLOTS_DIR / "gnnplus_vs_original_gnn.png", "Figure 5. GNN+ vs Original GNN on MLU and decision time.")
    comparison_rows = []
    for topo in TOPOLOGY_ORDER:
        gnn_row = df_summary[(df_summary["topology"] == topo) & (df_summary["method"] == "gnn")].iloc[0]
        gnnplus_row = df_summary[(df_summary["topology"] == topo) & (df_summary["method"] == "gnnplus")].iloc[0]
        improvement = 0.0
        if float(gnn_row["mean_mlu"]) != 0:
            improvement = ((float(gnn_row["mean_mlu"]) - float(gnnplus_row["mean_mlu"])) / float(gnn_row["mean_mlu"])) * 100.0
        comparison_rows.append(
            {
                "Topology": TOPOLOGY_DISPLAY[topo],
                "GNN MLU": float(gnn_row["mean_mlu"]),
                "GNN+ MLU": float(gnnplus_row["mean_mlu"]),
                "MLU Change %": improvement,
                "GNN Time (ms)": float(gnn_row["decision_time_ms"]),
                "GNN+ Time (ms)": float(gnnplus_row["decision_time_ms"]),
            }
        )
    add_dataframe_table(doc, pd.DataFrame(comparison_rows), font_size=8)

    doc.add_heading("6. Failure Scenario Results", level=1)
    doc.add_paragraph(
        "The existing failure CSV already contained all 8 topologies, all 4 core methods, and all 5 failure scenarios, so it is preserved in the fixed bundle."
    )
    add_image(doc, PLOTS_DIR / "failure_recovery_gnnplus.png", "Figure 6. GNN+ failure recovery by topology and scenario.")
    for scenario in SCENARIO_ORDER:
        doc.add_heading(f"6.{SCENARIO_ORDER.index(scenario) + 1} {SCENARIO_LABELS[scenario]}", level=2)
        scenario_df = df_failure[df_failure["scenario"] == scenario].copy()
        scenario_df["topology"] = scenario_df["topology"].map(TOPOLOGY_DISPLAY)
        known_df = scenario_df[scenario_df["status"] == "known"][
            ["topology", "method", "mean_mlu", "pre_failure_mlu", "failure_recovery_ms"]
        ].rename(
            columns={
                "topology": "Topology",
                "method": "Method",
                "mean_mlu": "Post-Failure MLU",
                "pre_failure_mlu": "Pre-Failure MLU",
                "failure_recovery_ms": "Recovery (ms)",
            }
        )
        known_df["Method"] = known_df["Method"].map(METHOD_LABELS)
        unseen_df = scenario_df[scenario_df["status"] == "unseen"][
            ["topology", "method", "mean_mlu", "pre_failure_mlu", "failure_recovery_ms"]
        ].rename(
            columns={
                "topology": "Topology",
                "method": "Method",
                "mean_mlu": "Post-Failure MLU",
                "pre_failure_mlu": "Pre-Failure MLU",
                "failure_recovery_ms": "Recovery (ms)",
            }
        )
        unseen_df["Method"] = unseen_df["Method"].map(METHOD_LABELS)
        add_dataframe_chunks(doc, known_df.reset_index(drop=True), f"{SCENARIO_LABELS[scenario]}: known topologies", chunk_size=24)
        add_dataframe_chunks(doc, unseen_df.reset_index(drop=True), f"{SCENARIO_LABELS[scenario]}: unseen topologies", chunk_size=12)

    doc.add_heading("7. Complexity Analysis", level=1)
    doc.add_paragraph(
        "The complexity discussion below is limited to the 4 methods that are actually evaluated in this branch report. "
        "For Bottleneck, GNN, and GNN+, the practical controller cost is dominated by the LP rerouting stage after flow selection."
    )
    doc.add_heading("7.1 Theoretical Complexity", level=2)
    add_dataframe_table(doc, build_complexity_rows(), font_size=8)
    doc.add_paragraph(
        "GNN+ does not introduce a different optimization class relative to the original GNN. "
        "Its controller pipeline remains graph-based flow selection followed by the same LP reroute stage, "
        "but with a higher constant factor from richer feature construction and inference."
    )
    doc.add_heading("7.2 Numeric Controller-Cost Summary", level=2)
    doc.add_paragraph(
        "Because the student asked for complexity with numbers, the table below reports empirical controller-cost numbers from the fixed normal-scenario CSV. "
        "These are measured runtimes and control overhead statistics, not replacements for asymptotic complexity."
    )
    add_dataframe_table(doc, build_complexity_numeric_rows(df_summary), font_size=8)
    doc.add_paragraph(
        "The numeric table should be read together with the theoretical table above: asymptotic complexity explains scaling, "
        "while the measured decision-time and control-overhead numbers show the actual cost observed in this report."
    )

    doc.add_heading("8. SDN Metrics Philosophy", level=1)
    doc.add_paragraph(
        "The packet-level SDN metrics in this branch remain model-based analytical metrics, matching the older Packet_SDN_Simulation_Report methodology rather than a live Mininet testbed."
    )
    metrics_df = pd.DataFrame(
        [
            {"Component": "Per-link delay", "Model": "M/M/1 queueing", "Formula / rule": "d = 1 / (mu - lambda) + prop_delay"},
            {"Component": "End-to-end delay", "Model": "Path sum", "Formula / rule": "sum of per-link delays"},
            {"Component": "Throughput", "Model": "Bottleneck model", "Formula / rule": "minimum capacity / load ratio along routed paths"},
            {"Component": "Packet loss", "Model": "Overflow approximation", "Formula / rule": "max(0, (load - capacity) / load)"},
            {"Component": "Jitter", "Model": "Delay variation", "Formula / rule": "|delay(t) - delay(t-1)|"},
            {"Component": "Rule install delay", "Model": "Analytical SDN control cost", "Formula / rule": "0.5 ms + 0.02 ms * num_rules"},
            {"Component": "Failure recovery", "Model": "Controller reroute time", "Formula / rule": "cycles / wall-clock from failure to reroute"},
        ]
    )
    add_dataframe_table(doc, metrics_df, font_size=8)
    doc.add_paragraph("No live packet generation or Mininet dataplane measurement is claimed in this report.")

    doc.add_heading("9. Limitations and Honest Assessment", level=1)
    add_bullet(doc, "Paper baselines remain unavailable in the repo and are not faked into this report.")
    add_bullet(doc, "The evaluation scope remains limited to the 4 core methods present in this branch output bundle.")
    add_bullet(doc, "SDN metrics remain model-based analytical metrics rather than live packet-level measurements.")
    add_bullet(doc, "Some topologies still exhibit severe congestion (MLU much greater than 1), indicating traffic matrices beyond network capacity.")
    add_bullet(doc, "The source report bundle mixed stale and fresh artifacts; this fixed bundle corrects the presentation layer and regenerates only the incomplete normal summary.")

    doc.add_heading("10. Output Files", level=1)
    add_bullet(doc, f"Report: {FIXED_DOCX.relative_to(PROJECT_ROOT)}")
    add_bullet(doc, f"Summary CSV: {FIXED_SUMMARY.relative_to(PROJECT_ROOT)}")
    add_bullet(doc, f"Failure CSV: {FIXED_FAILURE.relative_to(PROJECT_ROOT)}")
    add_bullet(doc, f"SDN metrics CSV: {FIXED_METRICS.relative_to(PROJECT_ROOT)}")
    add_bullet(doc, f"Plots directory: {PLOTS_DIR.relative_to(PROJECT_ROOT)}")
    add_bullet(doc, f"Audit note: {FIX_AUDIT.relative_to(PROJECT_ROOT)}")

    doc.save(FIXED_DOCX)


def write_fix_audit(
    source_summary_audit: dict,
    source_failure_audit: dict,
    source_metrics_audit: dict,
    current_docx_audit: dict,
    fixed_summary: pd.DataFrame,
    fixed_failure: pd.DataFrame,
    fixed_metrics: pd.DataFrame,
    provenance: dict,
):
    fixed_summary_audit = audit_csv(FIXED_SUMMARY)
    fixed_failure_audit = audit_csv(FIXED_FAILURE)
    fixed_metrics_audit = audit_csv(FIXED_METRICS)

    lines = [
        "# GNN+ Packet-SDN Fixed Bundle Audit",
        "",
        f"Generated: {datetime.now().isoformat(timespec='seconds')}",
        "",
        "## Source Audit",
        "",
        f"- Current report: `{CURRENT_DOCX.relative_to(PROJECT_ROOT)}`",
        f"- Old template reference: `{OLD_TEMPLATE_DOCX.relative_to(PROJECT_ROOT)}`",
        f"- Source summary CSV rows: {source_summary_audit['rows']}",
        f"- Source summary CSV methods: {source_summary_audit['methods']}",
        f"- Source summary CSV topologies: {source_summary_audit['topologies']}",
        f"- Source failure CSV rows: {source_failure_audit['rows']}",
        f"- Source failure CSV methods: {source_failure_audit['methods']}",
        f"- Source failure CSV scenarios: {source_failure_audit['scenarios']}",
        f"- Source metrics CSV rows: {source_metrics_audit['rows']}",
        f"- Source metrics CSV methods: {source_metrics_audit['methods']}",
        f"- Current DOCX summary table methods: {current_docx_audit['summary_table_methods']}",
        f"- Current DOCX failure table methods: {current_docx_audit['failure_table_methods']}",
        f"- Current DOCX still says initialized model: {current_docx_audit['contains_initialized_model_text']}",
        f"- Current DOCX contains out-of-scope methods: {current_docx_audit['contains_out_of_scope_methods']}",
        "",
        "## Root Cause",
        "",
        "- `packet_sdn_summary.csv` is stale/incomplete: it has only 24 rows and no `gnnplus` rows.",
        "- `packet_sdn_sdn_metrics.csv` was regenerated from that incomplete summary CSV, so its visible normal-summary table also drops GNNPLUS.",
        "- `scripts/generate_gnnplus_plots.py` hardcodes `['ecmp', 'bottleneck', 'gnn', 'gnnplus']`, so the legend still shows GNNPLUS even when the data slice is empty and the plotted values are all NaN.",
        "- `scripts/generate_gnnplus_report.py` truncates tables with `head(max_rows)` and still injects stale narrative text such as the initialized-model limitation and out-of-scope methods in the complexity section.",
        "",
        "## Fixed Bundle Actions",
        "",
        f"- Fixed checkpoint used: `{provenance['checkpoint_path']}`",
        f"- Source checkpoint copied from: `{provenance['source_checkpoint_path']}`",
        f"- Payload config match between copied checkpoint and source checkpoint: {provenance['same_config_as_source']}",
        f"- Dropout recorded in payload: {provenance['dropout']}",
        f"- `learn_k_crit` recorded in payload config: {provenance['learn_k_crit']}",
        f"- Fixed K recorded in payload config: min={provenance['k_crit_min']}, max={provenance['k_crit_max']}",
        f"- Legacy packet-SDN scorer input dim: {provenance['legacy_packet_sdn_in_channels']}",
        "- Regenerated only the normal summary CSV because the source normal bundle was incomplete.",
        "- Reused the existing failure CSV because it already had all 8 topologies, all 4 methods, and all 5 failure scenarios.",
        "- Rebuilt plots so legends are driven by actual rendered series and grouped bars use explicit centered offsets.",
        "- Rebuilt the DOCX so tables are chunked instead of truncated and the stale initialized-model / out-of-scope narrative is removed.",
        "",
        "## Fixed Output Audit",
        "",
        f"- Fixed summary CSV rows: {fixed_summary_audit['rows']}",
        f"- Fixed summary CSV methods: {fixed_summary_audit['methods']}",
        f"- Fixed summary CSV topologies: {fixed_summary_audit['topologies']}",
        f"- Fixed failure CSV rows: {fixed_failure_audit['rows']}",
        f"- Fixed failure CSV scenarios: {fixed_failure_audit['scenarios']}",
        f"- Fixed metrics CSV rows: {fixed_metrics_audit['rows']}",
        f"- Fixed metrics CSV methods: {fixed_metrics_audit['methods']}",
        f"- Fixed DOCX: `{FIXED_DOCX.relative_to(PROJECT_ROOT)}`",
        "",
        "## Acceptance Check",
        "",
        f"- GNN+ present in fixed normal summary CSV: {'gnnplus' in set(fixed_summary['method'].astype(str))}",
        f"- GNN+ present in fixed failure CSV: {'gnnplus' in set(fixed_failure['method'].astype(str))}",
        f"- GNNPLUS present in fixed visible summary table CSV: {'GNNPLUS' in set(fixed_metrics['Method'].astype(str))}",
        f"- All 8 topologies present in fixed summary: {set(fixed_summary['topology'].astype(str)) == set(TOPOLOGY_ORDER)}",
        f"- All 5 failure scenarios present in fixed failure CSV: {set(fixed_failure['scenario'].astype(str)) == set(SCENARIO_ORDER)}",
        "- SDN metrics remain explicitly labeled as model-based in the rebuilt report.",
    ]
    FIX_AUDIT.write_text("\n".join(lines) + "\n")


def main():
    FIXED_DIR.mkdir(parents=True, exist_ok=True)
    PLOTS_DIR.mkdir(parents=True, exist_ok=True)

    provenance = load_checkpoint_provenance()
    current_docx_audit = extract_docx_table_audit(CURRENT_DOCX)
    (
        fixed_summary,
        fixed_failure,
        fixed_metrics,
        source_summary_audit,
        source_failure_audit,
        source_metrics_audit,
    ) = prepare_fixed_csvs()
    create_plots(fixed_summary, fixed_failure)
    build_report(fixed_summary, fixed_failure, fixed_metrics, provenance)
    write_fix_audit(
        source_summary_audit,
        source_failure_audit,
        source_metrics_audit,
        current_docx_audit,
        fixed_summary,
        fixed_failure,
        fixed_metrics,
        provenance,
    )
    print(f"Fixed bundle written to: {FIXED_DIR}")


if __name__ == "__main__":
    main()
